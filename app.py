
# app_unificado.py
# Arquivo ÚNICO: app.py + move_db.py + move_creator_ui.py (com cg_draft)


import streamlit as st
import pandas as pd
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import json
import requests
import unicodedata
import os
from textwrap import dedent
import io
import html
import re
import difflib
import uuid
from datetime import datetime
from PIL import Image, ImageDraw
from streamlit_image_coordinates import streamlit_image_coordinates
import random
import gzip
import base64
import urllib.parse
import streamlit.components.v1 as components
from advantages_engine import suggest_advantages
import queue
import threading
import time
import hashlib

import mimetypes
from PIL import Image, ImageDraw, ImageEnhance

from pathlib import Path
from biome_generator import BiomeGenerator, BIOME_CONFIG
from streamlit.runtime import scriptrunner # <--- IMPORTANTE: Importar o módulo inteiro
from streamlit.runtime.scriptrunner import add_script_run_ctx, get_script_run_ctx
from streamlit.errors import StreamlitAPIException

# ─── Golpe Builder MM3e (criador avançado de golpes) ─────────────────────────
try:
    from golpe_builder_ui import render_golpe_builder as _gb_render_builder
    from golpe_builder import (
        parse_build_string as _gb_parse_build,
        calculate_total_pp as _gb_calc_pp,
        GolpeDraft as _GolpeDraft,
    )
    _GB_AVAILABLE = True
except Exception:  # ImportError, ModuleNotFoundError, etc.
    _GB_AVAILABLE = False
    _gb_render_builder = None
    _gb_parse_build = None
    _gb_calc_pp = None
    _GolpeDraft = None

# ─── Encounter Generator (aba exclusiva Ezenek) ─────────────────────────────
try:
    from encounter_generator import (
        carregar_pokedex as _enc_carregar_pokedex,
        gerar_encontro as _enc_gerar_encontro,
        CLIMAS_POR_REGIAO as _ENC_CLIMAS_POR_REGIAO,
    )
    _ENC_AVAILABLE = True
except Exception:
    _ENC_AVAILABLE = False
    _enc_carregar_pokedex = None
    _enc_gerar_encontro = None
    _ENC_CLIMAS_POR_REGIAO = {}


# ================================
# MOVE DB + MOVE CREATOR (UNIFICADO)
# ================================
# (conteúdo original do move_db.py e move_creator_ui.py foi incorporado aqui para ficar em um único arquivo)


# move_db.py
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple
import unicodedata

import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# ----------------------------
# Utils
# ----------------------------
def _norm_pid(v) -> str:
    s = str(v or "").strip().replace("#", "")
    # caso "16.0", "16.00", etc
    if re.fullmatch(r"\d+\.0+", s):
        s = s.split(".", 1)[0]
    # remove zeros à esquerda para comparação consistente: "019" -> "19", "076" -> "76"
    if re.fullmatch(r"\d+", s):
        s = str(int(s))
    return s
    
def _norm(s: str) -> str:
    import unicodedata

    s = (s or "").strip().lower()
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"\s+", " ", s)
    return s


def _safe_str(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, float) and pd.isna(v):
        return ""
    return str(v).strip()


def _boolish(v: Any) -> bool:
    x = _norm(_safe_str(v))
    return x in {"sim", "true", "1", "yes", "y"}

# =========================
# UI helpers: Lista & ajustes (Golpes)
# =========================
def _mv_ui_id(m: dict) -> str:
    """ID estável por golpe para chaves de widgets, mesmo com filtro/ordenação."""
    if "_ui_id" not in m or not m["_ui_id"]:
        m["_ui_id"] = uuid.uuid4().hex
    return str(m["_ui_id"])

def _mv_tags_from_move(m: dict) -> list[str]:
    meta = (m.get("meta") or {})
    build = str(m.get("build") or "").lower()

    tags = []

    # Meta/flags
    if meta.get("is_reaction"):
        tags.append("⚡ Reaction")
    if meta.get("perception_area"):
        tags.append("👁️ Área")
    if meta.get("ranged"):
        tags.append("🎯 Ranged")
    if meta.get("area_type") and meta.get("area_type") != "—":
        ext_label = " Extended" if meta.get("area_extended") else ""
        tags.append(f"🔴 {meta['area_type']}{ext_label}")
    if meta.get("resist_stat") and meta.get("resist_stat") not in ("— (padrão)", "Thg"):
        tags.append(f"Resiste: {meta['resist_stat']}")

    # Heurística pelo build (não altera lógica; só UI)
    if "damage" in build:
        tags.append("Damage")
    if "affliction" in build:
        tags.append("Affliction")
    if "weaken" in build:
        tags.append("Weaken")
    if "linked" in build:
        tags.append("Linked")

    # Resistências (pra ajudar leitura)
    if "will" in build:
        tags.append("Will")
    if "dodge" in build:
        tags.append("Dodge")
    if "toughness" in build or "thg" in build:
        tags.append("Toughness")
    if "fortitude" in build:
        tags.append("Fortitude")

    # Remove duplicados preservando ordem
    seen = set()
    out = []
    for t in tags:
        if t not in seen:
            seen.add(t)
            out.append(t)
    return out

def _mv_desc_for(db_moves_guided, move_name: str) -> str | None:
    if not db_moves_guided:
        return None
    try:
        mv_desc = db_moves_guided.get_by_name(str(move_name or ""))
    except Exception:
        mv_desc = None
    if mv_desc and getattr(mv_desc, "descricao", None):
        return mv_desc.descricao
    return None

def _mv_preview(text: str | None, n: int = 180) -> str:
    t = (text or "").strip()
    if not t:
        return "Descrição não disponível."
    if len(t) <= n:
        return t
    return t[:n].rstrip() + "…"

def _mv_sort_key(sort_mode: str, name: str, pp: int, final_rank: int, acc: int, mod_acc: int):
    # retornos “compatíveis” com sorted()
    if sort_mode == "Nome (A→Z)":
        return (name.lower(),)
    if sort_mode == "PP (maior→menor)":
        return (-pp, name.lower())
    if sort_mode == "Rank final (maior→menor)":
        return (-final_rank, name.lower())
    if sort_mode == "Acerto (maior→menor)":
        return (-acc, name.lower())
    if sort_mode == "Mod. acerto (maior→menor)":
        return (-mod_acc, name.lower())
    return (name.lower(),)

def _draw_tactical_grid(img, grid, tile_size: int):
    """
    Desenha a grade por cima da imagem.

    Aceita:
      - grid como matriz (H x W): list[list[...]]
      - grid como int: tamanho (NxN)
      - grid como tuple/list (h, w)
    """
    draw = ImageDraw.Draw(img)

    # 1) Descobre dimensões (h, w)
    h = w = 0

    if isinstance(grid, int):
        h = w = int(grid)
    elif isinstance(grid, (tuple, list)) and len(grid) == 2 and all(isinstance(x, int) for x in grid):
        h, w = int(grid[0]), int(grid[1])
    else:
        # assume matriz
        try:
            h = len(grid)
            w = len(grid[0]) if h else 0
        except Exception:
            return  # fail-safe: não desenha grade se vier algo inesperado

    if h <= 0 or w <= 0:
        return

    # 2) linhas verticais
    for x in range(w + 1):
        px = x * tile_size
        draw.line([(px, 0), (px, h * tile_size)], fill=(0, 0, 0, 90), width=1)

    # 3) linhas horizontais
    for y in range(h + 1):
        py = y * tile_size
        draw.line([(0, py), (w * tile_size, py)], fill=(0, 0, 0, 90), width=1)


def _infer_based_from_text(*texts: str) -> str | None:
    blob = " ".join(str(t or "") for t in texts).lower()

    if not blob:
        return None

    if "status" in blob:
        return "—"

    int_tokens = (
        "intelect based", "intellect based", "int based", "special-based", "special based",
        "especial", "special", "intelect", "intellect"
    )
    stgr_tokens = (
        "stgr based", "strength based", "strength-based", "physical-based", "physical based",
        "físico", "fisico", "physical", "stgr", "strength"
    )

    if any(t in blob for t in int_tokens):
        return "Int"
    if any(t in blob for t in stgr_tokens):
        return "Stgr"

    return None


def _move_based_stat_from_meta(move_like: dict | None) -> str:
    move_like = move_like or {}
    move_meta = move_like.get("meta") if any(k in move_like for k in ("meta", "name", "build", "description")) else move_like
    move_meta = move_meta or {}
    cat_meta = str(move_meta.get("category", "") or "").strip().lower()

    if move_meta.get("is_special") is True:
        return "Int"
    if move_meta.get("is_special") is False:
        return "Stgr"

    if "status" in cat_meta:
        return "—"
    if "especial" in cat_meta or "special" in cat_meta:
        return "Int"
    if "físico" in cat_meta or "fisico" in cat_meta or "physical" in cat_meta:
        return "Stgr"

    text_based = _infer_based_from_text(
        move_like.get("name", ""),
        move_like.get("build", ""),
        move_like.get("description", ""),
        move_meta.get("raw_power_name", ""),
        move_meta.get("category", ""),
    )
    if text_based:
        return text_based

    return "Stgr"


def _move_stat_value(move_like: dict | None, stats: dict) -> tuple[str, int]:
    based = _move_based_stat_from_meta(move_like)
    if based == "Int":
        return based, int(stats.get("int", 0) or 0)
    if based == "Stgr":
        return based, int(stats.get("stgr", 0) or 0)
    return based, 0





def _move_accuracy_limit(move: dict, np_value: int, stats: dict) -> int:
    rank = int(move.get("rank", 0) or 0)
    _, stat_val = _move_stat_value(move, stats)
    return max(0, (2 * int(np_value)) - rank - stat_val)


def _move_accuracy_pp(move: dict) -> float:
    return float(int(move.get("accuracy", 0) or 0)) / 2

# COLOQUE NO TOPO DO ARQUIVO
st.markdown("""
<style>
/* Estilo para os slots da BOX (Grama) */
.box-slot-grass {
    background: #55a64b !important;
    background-image: 
        linear-gradient(rgba(255,255,255,0.1) 1px, transparent 1px),
        linear-gradient(90deg, rgba(255,255,255,0.1) 1px, transparent 1px) !important;
    background-size: 8px 8px !important;
    border: 3px solid #2d5a27 !important;
    border-radius: 12px !important;
    padding: 10px !important;
    display: flex !important;
    justify-content: center !important;
    align-items: center !important;
    box-shadow: inset 0 0 15px rgba(0,0,0,0.4) !important;
    margin-bottom: 8px !important;
}

/* Estilo para os cards da Equipe (Menu Azul) */
.gba-party-slot {
    background: linear-gradient(180deg, #4d88ff 0%, #2e5cb8 100%) !important;
    border: 4px solid #f8fafc !important;
    border-radius: 15px !important;
    padding: 15px !important;
    margin-bottom: 20px !important;
    box-shadow: 6px 6px 0px rgba(0,0,0,0.3) !important;
    display: flex !important;
    flex-direction: column !important;
    align-items: center !important;
}

.gba-party-slot img {
    filter: drop-shadow(0 4px 6px rgba(0,0,0,0.4)) !important;
}

/* Fix para remover margens extras do Streamlit dentro dos slots */
[data-testid="stVerticalBlock"] > div:has(.box-slot-grass) {
    padding: 0px !important;
}
</style>
""", unsafe_allow_html=True)


    
@st.cache_resource(show_spinner=False)
def load_effect_icon_rgba(path: str):
    try:
        if not path or not os.path.exists(path):
            return None
        return Image.open(path).convert("RGBA")
    except Exception:
        return None

# ----------------------------
# Data model
# ----------------------------
@dataclass
class Move:
    name: str
    tipo: str
    categoria: str
    descricao: str
    build: str
    how_it_works: str
    resist_stat: str
    ranged: bool
    perception_area: bool
    tags: List[str]
    raw: Dict[str, Any]

    def render_build(self, rank: int, sub_ranks: dict | None = None) -> str:
        b = (self.build or "").strip()
        if not b:
            return ""
    
        # 1) Trocar Rank = PL por Rank = <rank>
        b = re.sub(r"Rank\s*=\s*PL", f"Rank = {rank}", b, flags=re.IGNORECASE)
        b = re.sub(r"Rank\s*=\s*X",  f"Rank = {rank}", b, flags=re.IGNORECASE)
    
        # 2) Escalar efeitos numéricos: cada efeito pode ter rank próprio via sub_ranks
        def _scale(m):
            effect = m.group(1)               # "Damage", "Weaken"...
            eff_key = effect.strip().lower()  # "damage", "weaken"...
            if sub_ranks and eff_key in sub_ranks and int(sub_ranks[eff_key]) > 0:
                r = int(sub_ranks[eff_key])
            else:
                r = int(rank)
            return f"{effect} {r}"
    
        b = re.sub(
            r"\b(Damage|Weaken|Affliction|Healing|Nullify|Create|Environment)\s+\d+\b",
            _scale,
            b,
            flags=re.IGNORECASE
        )
    
        # 3) Deduplicar segmentos "Linked ..." idênticos (exatos)
        parts = [p.strip() for p in b.split(";") if p.strip()]
        seen = set()
        uniq = []
        for p in parts:
            key = re.sub(r"\s+", " ", p.lower()).strip()
            if key not in seen:
                seen.add(key)
                uniq.append(p)
    
        # ==========================================
        # DEFINIÇÃO DE RESISTÊNCIA DO DANO (FINAL)
        # ==========================================
        name_desc = f"{self.name} {self.descricao or ''}".lower()
        tipo = (self.tipo or "").lower()
    
        damage_resist = "Thg"  # padrão absoluto
    
        # 1) DODGE — prioridade máxima
        if any(k in name_desc for k in [
            "ohko", "one-hit", "hit kill",
            "guillotine", "horn drill", "sheer cold", "fissure",
            "diferença de velocidade", "speed difference",
            "diferença de peso", "weight difference"
        ]):
            damage_resist = "Dodge"
    
        # 2) WILL — psíquico / fantasma / redução de will/spdef/spatk
        elif tipo in {"psychic", "psíquico", "ghost", "fantasma"}:
            damage_resist = "Will"
    
        elif any(k in name_desc for k in [
            "reduce will", "reduz will",
            "special defense down", "spdef down",
            "special attack down", "spatk down"
        ]):
            damage_resist = "Will"
    
        # 3) Aplica no primeiro Damage (no resultado final)
        out = "; ".join(uniq)
        out = re.sub(
            r"(Damage\s+\d+)(?![^\[]*\])",
            rf"\1 (Resisted by {damage_resist})",
            out,
            count=1,
            flags=re.IGNORECASE
        )
    
        return out



    def pp_cost(self, rank: int) -> Tuple[Optional[float], str]:
        """
        Retorna (pp_cost, explicacao). Aqui "PP" = Power Points do M&M.
    
        - Se existir coluna PP_Custo no Excel, usa direto.
        - Caso contrário, usa um estimador simples (você pode trocar depois pela sua tabela oficial).
        Regra base do M&M: custo final por graduação = custo básico + extras - falhas.
        """
    
        # (A0) Novo: "PP por Rank" vindo do Excel -> PP total = (pp_por_rank * rank)
        if "PP por Rank" in self.raw and _safe_str(self.raw.get("PP por Rank")):
            try:
                ppr = float(str(self.raw["PP por Rank"]).replace(",", "."))
                return ppr * float(rank), 'PP total = ("PP por Rank" do Excel) × rank.'
            except Exception:
                pass
    
        # (A) override do Excel
        for key in ("PP_Custo", "PP", "Custo_PP"):
            if key in self.raw and _safe_str(self.raw.get(key)):
                try:
                    val = float(str(self.raw[key]).replace(",", "."))
                    return val, f"PP_Custo vindo do Excel ({key})."
                except Exception:
                    pass


        # (B) estimador bem simples (trocável!)
        # Ideia: golpes que têm Damage tendem a escalar mais caro; Linked/Área etc aumentam.
        build = _norm(self.build)
        base_per_rank = 0.0

        if "damage" in build:
            base_per_rank += 1.0
        if "affliction" in build:
            base_per_rank += 1.0
        if "weaken" in build:
            base_per_rank += 1.0
        if "healing" in build:
            base_per_rank += 2.0
        if "create" in build:
            base_per_rank += 1.0
        if "environment" in build:
            base_per_rank += 1.0
        if "nullify" in build:
            base_per_rank += 1.0

        # extras comuns no texto
        if "linked" in build:
            base_per_rank += 0.5
        if "area" in build:
            base_per_rank += 0.5
        if "perception" in build:
            base_per_rank += 0.5
        if "extended range" in build or "ranged" in build:
            base_per_rank += 0.25

        if base_per_rank <= 0:
            return None, "Sem Custo no Excel e não foi possível estimar por falta de palavras-chave na build."
        pp_est = base_per_rank * int(rank)
        return float(pp_est), f"PP estimado ({base_per_rank:.2f} por rank)"



class MoveDB:
    def __init__(self, df: pd.DataFrame):
        self.df = df.copy()

        # normaliza colunas esperadas
        self.df["__name_norm"] = self.df["Nome"].apply(lambda x: _norm(_safe_str(x)))

        # texto para busca por descrição
        self.df["__text"] = (
            self.df["Nome"].fillna("").astype(str)
            + " | " + self.df["Tipo"].fillna("").astype(str)
            + " | " + self.df["Categoria"].fillna("").astype(str)
            + " | " + self.df["Descricao"].fillna("").astype(str)
            + " | " + self.df.get("Como funciona (regras/condições)", pd.Series([""] * len(self.df))).fillna("").astype(str)
            + " | " + self.df.get("Build M&M (adaptado)", pd.Series([""] * len(self.df))).fillna("").astype(str)
        )

        self._vectorizer = TfidfVectorizer(
            lowercase=True,
            ngram_range=(1, 2),
            max_features=25000,
        )
        self._tfidf = self._vectorizer.fit_transform(self.df["__text"].tolist())

    @staticmethod
    def from_excel(path: str, sheet_name: str = "Golpes_MM") -> "MoveDB":
        df = pd.read_excel(path, sheet_name=sheet_name)
        # garante colunas mínimas
        required = ["Nome", "Tipo", "Categoria", "Descricao", "Build M&M (adaptado)"]
        missing = [c for c in required if c not in df.columns]
        if missing:
            raise ValueError(f"Excel sem colunas obrigatórias: {missing}")
        return MoveDB(df)

    def _row_to_move(self, row: Dict[str, Any]) -> Move:
        return Move(
            name=_safe_str(row.get("Nome")),
            tipo=_safe_str(row.get("Tipo")),
            categoria=_safe_str(row.get("Categoria")),
            descricao=_safe_str(row.get("Descricao")),
            build=_safe_str(row.get("Build M&M (adaptado)")),
            how_it_works=_safe_str(row.get("Como funciona (regras/condições)")),
            resist_stat=_safe_str(row.get("Resist Stat")),
            ranged=_boolish(row.get("Ranged")),
            perception_area=_boolish(row.get("Perception Area")),
            tags=self._infer_tags(row),
            raw=row,
        )

    def get_by_name(self, name: str) -> Optional[Move]:
        key = _norm(name)
        hit = self.df[self.df["__name_norm"] == key]
        if hit.empty:
            return None
        row = hit.iloc[0].to_dict()
        return self._row_to_move(row)

    def get_all_by_name(self, name: str) -> List[Move]:
        key = _norm(name)
        hit = self.df[self.df["__name_norm"] == key]
        if hit.empty:
            return []
        return [self._row_to_move(row.to_dict()) for _, row in hit.iterrows()]

    def search_by_name_prefix(self, name: str) -> List[Move]:
        key = _norm(name)
        if not key:
            return []
        hit = self.df[self.df["__name_norm"].str.startswith(key)]
        if hit.empty:
            return []
        return [self._row_to_move(row.to_dict()) for _, row in hit.iterrows()]

    def suggest_by_description(self, description: str, top_k: int = 5) -> List[Tuple[Move, float]]:
        q = _safe_str(description)
        if not q:
            return []
        q_vec = self._vectorizer.transform([q])
        sims = cosine_similarity(q_vec, self._tfidf).flatten()
        idxs = sims.argsort()[::-1][:top_k]

        out: List[Tuple[Move, float]] = []
        for i in idxs:
            row = self.df.iloc[i].to_dict()
            mv = self._row_to_move(row)
            out.append((mv, float(sims[i])))
        return out

    def _infer_tags(self, row: Dict[str, Any]) -> List[str]:
        tags: List[str] = []

        # sinalizadores do seu Excel
        if _boolish(row.get("Dano")):
            tags.append("damage")
        if _boolish(row.get("Cura")):
            tags.append("healing")
        if _boolish(row.get("Teleport")):
            tags.append("teleport")
        if _boolish(row.get("Create")):
            tags.append("create")
        if _boolish(row.get("Clima")):
            tags.append("environment")
        if _boolish(row.get("Prioridade")):
            tags.append("priority")
        if _boolish(row.get("Perception Area")):
            tags.append("perception")
        if _boolish(row.get("Ranged")):
            tags.append("ranged")

        # detecta por texto
        b = _norm(_safe_str(row.get("Build M&M (adaptado)")))
        for kw, tag in [
            ("affliction", "affliction"),
            ("weaken", "weaken"),
            ("nullify", "nullify"),
            ("immunity", "immunity"),
            ("enhanced", "enhance"),
            ("linked", "linked"),
            ("area", "area"),
        ]:
            if kw in b and tag not in tags:
                tags.append(tag)

        return tags

@st.cache_data(show_spinner=False)
def _audio_data_uri(path_str: str) -> str:
    p = Path(path_str)
    if not p.exists():
        return ""
    data = p.read_bytes()
    mime, _ = mimetypes.guess_type(str(p))
    if not mime:
        mime = "audio/mpeg"
    b64 = base64.b64encode(data).decode("utf-8")
    return f"data:{mime};base64,{b64}"
    
@st.cache_data(show_spinner=False)
def comp_img_data_uri(path_str: str) -> str:
    p = Path(path_str)
    if not p.exists():
        return ""

    data = p.read_bytes()
    mime, _ = mimetypes.guess_type(str(p))
    if not mime:
        mime = "image/png"
    b64 = base64.b64encode(data).decode("utf-8")
    return f"data:{mime};base64,{b64}"

def apply_npc_overrides(base_npcs: dict, overrides: dict) -> dict:
    base_npcs = base_npcs or {}
    overrides = overrides or {}

    for npc_name, over in overrides.items():
        if not isinstance(over, dict):
            continue

        base = base_npcs.get(npc_name) or {"name": npc_name, "sections": {}}
        if not isinstance(base, dict):
            base = {"name": npc_name, "sections": {}}

        merged = dict(base)  # shallow copy

        # nunca deixar perder sections (lore)
        merged_sections = {}
        if isinstance(base.get("sections"), dict):
            merged_sections.update(base["sections"])
        if isinstance(over.get("sections"), dict):
            # override só adiciona/atualiza; não apaga o que existe
            for k, v in over["sections"].items():
                if (str(v).strip() if v is not None else ""):
                    merged_sections[k] = v
        merged["sections"] = merged_sections

        # campos "seguros" que você pode sobrescrever
        for k in ["idade", "origem", "ocupacao", "status", "img", "name"]:
            if k in over and over[k] not in (None, ""):
                merged[k] = over[k]

        # pokémons: dedupe preservando ordem
        def _dedupe_list(xs):
            out, seen = [], set()
            for x in xs or []:
                s = str(x).strip()
                if not s:
                    continue
                key = s.lower()
                if key in seen:
                    continue
                seen.add(key)
                out.append(s)
            return out

        if "pokemons" in over:
            merged["pokemons"] = _dedupe_list(over.get("pokemons"))
        if "pokemons_conhecidos" in over:
            merged["pokemons_conhecidos"] = _dedupe_list(over.get("pokemons_conhecidos"))

        base_npcs[npc_name] = merged

    return base_npcs


def render_bgm(track_path: str, volume: float = 0.35) -> None:
    src = _audio_data_uri(track_path)
    vol = max(0.0, min(1.0, float(volume)))

    components.html(
        """
<div style="height:0; overflow:hidden;">
  <audio id="ds-bgm" preload="auto" playsinline></audio>
</div>

<script>
(function () {
  const SRC = __SRC__;
  const VOL = __VOL__;

  const a = document.getElementById("ds-bgm");
  if (!a) return;

  a.loop = true;
  a.volume = VOL;

  if (a.dataset.src !== SRC) {
      // PARA a faixa anterior
      try { a.pause(); } catch(e) {}
      try { a.currentTime = 0; } catch(e) {}
    
      // limpa fontes antigas e descarrega
      a.innerHTML = "";
      a.removeAttribute("src");
      a.load();
    
      // carrega a nova faixa
      a.dataset.src = SRC;
      const s = document.createElement("source");
      s.src = SRC;
      a.appendChild(s);
      a.load();
    }

  async function tryPlay() {
    try {
      await a.play();
      localStorage.setItem("ds_bgm_unlocked", "1");
    } catch (e) {
      if (a.dataset.clickHooked === "1") return;
      a.dataset.clickHooked = "1";
      const parentDoc = window.parent && window.parent.document ? window.parent.document : document;
      const handler = async () => {
        parentDoc.removeEventListener("click", handler, true);
        try { await a.play(); } catch (e2) {}
        localStorage.setItem("ds_bgm_unlocked", "1");
      };
      parentDoc.addEventListener("click", handler, true);
    }
  }

  tryPlay();
})();
</script>
        """.replace("__SRC__", json.dumps(src)).replace("__VOL__", str(vol)),
        height=0,
        width=0,
        scrolling=False,
    )

PVP_RERUN_COOLDOWN_SEC = 0.05


def pvp_in_action() -> bool:
    """Retorna True quando o usuário está no meio de uma ação de mapa/PvP."""
    return bool(
        st.session_state.get("moving_piece_id")
        or st.session_state.get("placing_pid")
        or st.session_state.get("placing_trainer")
        or st.session_state.get("placing_effect")
    )


def request_rerun(reason: str, *, force: bool = False) -> bool:
    """Gate único de rerun para evitar cascata/tremedeira no PvP/mapa."""
    now = time.time()
    st.session_state["rerun_reason"] = reason

    if not force:
        if pvp_in_action() or float(st.session_state.get("arena_pause_until", 0) or 0) > now:
            st.session_state["pvp_sync_pending"] = True
            return False

        last_rerun_ts = float(st.session_state.get("last_rerun_ts", 0.0) or 0.0)
        if (now - last_rerun_ts) < PVP_RERUN_COOLDOWN_SEC:
            st.session_state["pvp_sync_pending"] = True
            return False

    st.session_state["last_rerun_ts"] = now
    st.session_state["pvp_sync_pending"] = False
    st.rerun()
    return True


# --- PLANO B: VIGIA DE SINCRONIZAÇÃO ---
@st.fragment(run_every=2) # Roda esta função sozinha a cada 2 segundos
def sync_watchdog(db, rid):
    if not rid:
        return

    try:
        doc_ref = db.collection("rooms").document(rid).collection("public_state").document("state")
        snapshot = doc_ref.get()
        if not snapshot.exists:
            return

        server_data = snapshot.to_dict() or {}
        server_time = server_data.get("updatedAt")

        if "last_map_update" not in st.session_state:
            st.session_state["last_map_update"] = server_time
            st.session_state["pvp_sync_pending"] = False
            return

        if server_time != st.session_state.get("last_map_update"):
            st.session_state["last_map_update"] = server_time

            # Se o usuário está no meio de uma ação, marcamos como pendente e não rerunamos agora.
            if pvp_in_action() or float(st.session_state.get("arena_pause_until", 0) or 0) > time.time():
                st.session_state["pvp_sync_pending"] = True
                return

            request_rerun("map_update")
            return

        if st.session_state.get("pvp_sync_pending"):
            request_rerun("map_update_pending")

    except Exception:
        # Se der erro de conexão, ignora e tenta na próxima
        return

try:
    from move_interpreter import interpret_effects_to_build
except Exception:
    interpret_effects_to_build = None




@st.fragment(run_every=2)
def battle_watchdog(db, rid):
    if not rid:
        return
    try:
        doc_ref = db.collection("rooms").document(rid).collection("public_state").document("battle")
        snapshot = doc_ref.get()
        if not snapshot.exists:
            return

        server_data = snapshot.to_dict() or {}
        server_time = server_data.get("updatedAt")

        if "last_battle_update" not in st.session_state:
            st.session_state["last_battle_update"] = server_time
            return

        if server_time != st.session_state.get("last_battle_update"):
            st.session_state["last_battle_update"] = server_time

            if pvp_in_action() or float(st.session_state.get("arena_pause_until", 0) or 0) > time.time():
                st.session_state["pvp_sync_pending"] = True
                return

            request_rerun("battle_update")
            return

    except Exception:
        return



@st.cache_resource
def load_move_db(excel_path: str) -> "MoveDB":
    return MoveDB.from_excel(excel_path, sheet_name="Golpes_MM")

def _comp_mtime(p: str) -> float:
    try:
        return os.path.getmtime(p) if p and os.path.exists(p) else 0.0
    except Exception:
        return 0.0


def render_move_creator(
    excel_path: str,
    state_key_prefix: str = "mc",
    return_to_view: str | None = None,
):
    """
    Tela do criador de golpes (UNIFICADO):
    - Busca por nome (persistente: não some no rerun)
    - Sugestão por descrição (persistente)
    - Criar do zero
    - Confirmação adiciona em st.session_state["cg_moves"] (alias do cg_draft["moves"])
    """
    db = load_move_db(excel_path)

    # garante draft / alias
    cg_init()
    st.session_state.setdefault("cg_moves", st.session_state["cg_draft"]["moves"])
    st.session_state["cg_draft"]["moves"] = st.session_state["cg_moves"]

    # estados persistentes da tela
    last_name_key = f"{state_key_prefix}_last_name_found"
    last_desc_key = f"{state_key_prefix}_last_desc_suggestions"
    st.session_state.setdefault(last_name_key, None)
    st.session_state.setdefault(last_desc_key, [])
    st.session_state.setdefault("cg_manual_moves", [])

    st.subheader("⚔️ Criação de Golpes (M&M 3e)")
    tab1, tab2, tab3, tab4 = st.tabs(
        ["🔎 Buscar por nome", "🧩 Procurar por Descrição", "🛠️ Criar Semi Manual", "✍️ Entrada Manual"]
    )

    def _move_default_accuracy(mv) -> int:
        raw = getattr(mv, "raw", {}) or {}
        return int(raw.get("Accuracy") or raw.get("Acerto") or raw.get("acerto") or 0)

    def _confirm_move(mv, rank: int, build: str, pp):
        st.session_state["cg_moves"].append({
            "name": mv.name,
            "rank": int(rank),
            "build": build,
            "pp_cost": pp,
            "accuracy": _move_default_accuracy(mv),
            "meta": {
                "ranged": bool(getattr(mv, "ranged", False)),
                "perception_area": bool(getattr(mv, "perception_area", False)),
                "category": str(getattr(mv, "categoria", "") or ""),
            }
        })

      
    
    def _render_move_card(mv, rank: int):
        st.markdown(f"### 🌀 {mv.name}  ({getattr(mv,'tipo','—')} / {getattr(mv,'categoria','—')})")
        c1, c2, c3 = st.columns(3)
        c1.metric("Ranged", "SIM" if getattr(mv, "ranged", False) else "NÃO")
        c2.metric("Perception Area", "SIM" if getattr(mv, "perception_area", False) else "NÃO")
        c3.metric("Resist Stat", getattr(mv, "resist_stat", None) or "—")

        st.write("**Descrição:**")
        st.write(getattr(mv, "descricao", None) or "—")

        st.write("**Build M&M (rank escolhido):**")
        
        # ✅ opção de escolher ranks por sub-efeito
        custom_sub = st.checkbox(
            "Quero escolher rank por sub-efeito (Damage/Affliction/Weaken etc.)",
            key=f"{state_key_prefix}_customsub_{mv.name}_{rank}"
        )
        
        sub_ranks = None
        manual_pp = None
        
        if custom_sub:
            st.caption("Defina o rank de cada sub-efeito. Se deixar 0, ele não entra / não altera.")
            cA, cB, cC = st.columns(3)
        
            with cA:
                r_damage = st.number_input(
                    "Rank Damage", min_value=0, max_value=30, value=int(rank),
                    key=f"{state_key_prefix}_r_damage_{mv.name}_{rank}"
                )
                r_aff = st.number_input(
                    "Rank Affliction", min_value=0, max_value=30, value=int(rank),
                    key=f"{state_key_prefix}_r_aff_{mv.name}_{rank}"
                )
        
            with cB:
                r_weaken = st.number_input(
                    "Rank Weaken", min_value=0, max_value=30, value=int(rank),
                    key=f"{state_key_prefix}_r_weaken_{mv.name}_{rank}"
                )
                r_heal = st.number_input(
                    "Rank Healing", min_value=0, max_value=30, value=0,
                    key=f"{state_key_prefix}_r_heal_{mv.name}_{rank}"
                )
        
            with cC:
                r_create = st.number_input(
                    "Rank Create", min_value=0, max_value=30, value=0,
                    key=f"{state_key_prefix}_r_create_{mv.name}_{rank}"
                )
                r_env = st.number_input(
                    "Rank Environment", min_value=0, max_value=30, value=0,
                    key=f"{state_key_prefix}_r_env_{mv.name}_{rank}"
                )
        
            sub_ranks = {
                "damage": int(r_damage),
                "affliction": int(r_aff),
                "weaken": int(r_weaken),
                "healing": int(r_heal),
                "create": int(r_create),
                "environment": int(r_env),
            }
        
            manual_pp = st.number_input(
                "PP total do golpe (obrigatório quando você customiza ranks)",
                min_value=0,
                value=0,
                step=1,
                key=f"{state_key_prefix}_manualpp_{mv.name}_{rank}"
            )
        
        # ✅ build: normal ou customizado
        if sub_ranks:
            build = mv.render_build(rank, sub_ranks=sub_ranks)
        else:
            build = mv.render_build(rank)
        
        st.code(build, language="text")
        
        # --- FIX: sempre inicializa antes de ramificar ---
        pp_final = None
        pp_auto = None
        why_auto = ""
        why = ""
        
        # ✅ PP: se customizou, usa o manual; se não, usa o Excel
        if sub_ranks:
            pp_final = int(manual_pp or 0)
            why = "PP informado manualmente (porque você escolheu ranks por sub-efeito)."
        else:
            tmp = mv.pp_cost(rank)
            if tmp is None:
                pp_auto, why_auto = None, "pp_cost() retornou None (erro interno)."
            else:
                pp_auto, why_auto = tmp  # pode ser None
        
            if pp_auto is None:
                # PP obrigatório manual (não definido no banco)
                pp_manual = st.number_input(
                    "PP total do golpe (obrigatório)",
                    min_value=1,
                    value=1,
                    step=1,
                    key=f"{state_key_prefix}_pp_required_{mv.name}_{rank}"
                )
                pp_final = int(pp_manual)
                why = "PP manual obrigatório (não definido no banco)."
            else:
                pp_final = pp_auto
                why = why_auto
        
        st.info(f"PP: **{pp_final}** — {why}")


        # define o PP final a ser exibido
        if pp_final is not None:
            pp = pp_final
            why = "PP definido manualmente"
        elif pp_auto is not None:
            pp = pp_auto
            why = why_auto
        else:
            pp = None
            why = "PP não definido"

        if pp is not None:
            st.info(f"PP (estimado ou do Excel): **{pp}** — {why}")
        else:
            st.warning(f"PP: não definido — {why}")

        how = getattr(mv, "how_it_works", None)
        if how:
            st.write("**Como funciona:**")
            st.write(how)

        col_confirm, col_add = st.columns(2)
        with col_confirm:
            if st.button(f"✅ Confirmar {mv.name}", key=f"{state_key_prefix}_confirm_{mv.name}_{rank}"):
                _confirm_move(mv, rank, build, pp)
                st.success(f"Adicionado: {mv.name} (Rank {rank})")

        if return_to_view:
            with col_add:
                if st.button("➕ Adicionar golpe à ficha", key=f"{state_key_prefix}_add_to_sheet_{mv.name}_{rank}"):
                    _confirm_move(mv, rank, build, pp)
                    st.success(f"Adicionado: {mv.name} (Rank {rank})")
                    st.session_state["cg_view"] = return_to_view
                    st.rerun()

    def _manual_move_to_move(manual: dict) -> Move:
        return Move(
            name=_safe_str(manual.get("Nome")),
            tipo=_safe_str(manual.get("Tipo")),
            categoria=_safe_str(manual.get("Categoria") or "Manual"),
            descricao=_safe_str(manual.get("Descricao")),
            build=_safe_str(manual.get("Formula")),
            how_it_works="",
            # ✅ AGORA LÊ A RESISTÊNCIA DO DICIONÁRIO
            resist_stat=_safe_str(manual.get("Resist Stat") or "—"),
            ranged=bool(manual.get("Ranged", False)),
            perception_area=bool(manual.get("Area", False)),
            tags=[],
            raw=manual,
        )

    def _search_moves_by_name(name: str) -> List[Move]:
        hits = db.search_by_name_prefix(name)
        manual_hits = [
            _manual_move_to_move(mv)
            for mv in st.session_state.get("cg_manual_moves", [])
            if _norm(mv.get("Nome")).startswith(_norm(name))
        ]
        return hits + manual_hits

    with tab1:
        name = st.text_input("Nome do golpe", key=f"{state_key_prefix}_name")
        rank = st.slider("Rank", 1, 20, 10, key=f"{state_key_prefix}_rank")

        if st.button("Buscar", key=f"{state_key_prefix}_search", type="primary"):
            matches = _search_moves_by_name(name)
            if not matches:
                st.session_state[last_name_key] = None
                st.error("Não achei pelo nome. Use a aba 'Criar por descrição'.")
            else:
                st.session_state[last_name_key] = name

        if st.session_state.get(last_name_key):
            for mv in _search_moves_by_name(st.session_state[last_name_key]):
                _render_move_card(mv, rank)

    with tab2:
        desc = st.text_area("Descrição do golpe", height=120, key=f"{state_key_prefix}_desc")
        rank2 = st.slider("Rank para renderizar sugestões", 1, 20, 10, key=f"{state_key_prefix}_rank2")
        top_k = st.slider("Sugestões", 3, 10, 5, key=f"{state_key_prefix}_topk")

        if st.button("Sugerir", key=f"{state_key_prefix}_suggest"):
            sugg = db.suggest_by_description(desc, top_k=top_k)
            if not sugg:
                st.session_state[last_desc_key] = []
                st.warning("Digite uma descrição.")
            else:
                st.session_state[last_desc_key] = [mv.name for (mv, _score) in sugg]

        for i, mv_name in enumerate(st.session_state.get(last_desc_key, []), start=1):
            mv = db.get_by_name(mv_name)
            if not mv:
                continue
            with st.expander(f"{i}) {mv.name}", expanded=(i == 1)):
                _render_move_card(mv, rank2)

    with tab3:
        st.subheader("🛠️ Criar Semi Manual")

        rank3 = st.slider("Rank do golpe", 1, 20, 10, key=f"{state_key_prefix}_z_rank")
        is_special = st.checkbox("Golpe Especial (Intelect Based)", value=True, key=f"{state_key_prefix}_z_special")

        st.markdown("### Efeitos")
        effects = {
            "damage": st.checkbox("Causar Dano", key=f"{state_key_prefix}_z_damage"),
            "affliction": st.checkbox("Causar Affliction", key=f"{state_key_prefix}_z_affliction"),
            "weaken": st.checkbox("Causar Weaken", key=f"{state_key_prefix}_z_weaken"),
            "healing": st.checkbox("Cura", key=f"{state_key_prefix}_z_healing"),
            "create": st.checkbox("Create", key=f"{state_key_prefix}_z_create"),
            "environment": st.checkbox("Environment", key=f"{state_key_prefix}_z_environment"),
        }

        st.markdown("### Detalhes do Weaken")
        effects["weaken_stgr"] = st.checkbox("Weaken Strength", key=f"{state_key_prefix}_z_w_stgr")
        effects["weaken_int"] = st.checkbox("Weaken Intellect", key=f"{state_key_prefix}_z_w_int")
        effects["weaken_dodge"] = st.checkbox("Weaken Dodge", key=f"{state_key_prefix}_z_w_dodge")
        effects["weaken_will"] = st.checkbox("Weaken Will", key=f"{state_key_prefix}_z_w_will")

        st.markdown("### Ranks por sub-efeito (opcional)")
        st.caption("Se você mudar algum rank aqui, você terá que informar o PP total do golpe.")

        c1, c2, c3 = st.columns(3)
        with c1:
            r_damage = st.number_input("Rank Damage", 0, 30, int(rank3), key=f"{state_key_prefix}_z_r_damage")
            r_aff = st.number_input("Rank Affliction", 0, 30, 0, key=f"{state_key_prefix}_z_r_aff")
        with c2:
            r_weaken = st.number_input("Rank Weaken", 0, 30, 0, key=f"{state_key_prefix}_z_r_weaken")
            r_heal = st.number_input("Rank Healing", 0, 30, 0, key=f"{state_key_prefix}_z_r_heal")
        with c3:
            r_create = st.number_input("Rank Create", 0, 30, 0, key=f"{state_key_prefix}_z_r_create")
            r_env = st.number_input("Rank Environment", 0, 30, 0, key=f"{state_key_prefix}_z_r_env")

        sub_ranks = {
            "damage": int(r_damage),
            "affliction": int(r_aff),
            "weaken": int(r_weaken),
            "healing": int(r_heal),
            "create": int(r_create),
            "environment": int(r_env),
        }

        # considera "custom" se qualquer coisa ficar diferente do rank3 (ou se algum efeito ganhar rank >0 diferente)
        custom_sub = any(
            (k == "damage" and sub_ranks[k] != int(rank3)) or (k != "damage" and sub_ranks[k] > 0)
            for k in sub_ranks
        )

        manual_pp = None
        if custom_sub:
            manual_pp = st.number_input(
                "PP total do golpe (obrigatório quando customiza sub-ranks)",
                min_value=0, value=0, step=1,
                key=f"{state_key_prefix}_z_manual_pp"
            )

        st.markdown("### Modificadores")
        area = st.selectbox("Área", ["Nenhuma", "Burst", "Cone", "Line"], key=f"{state_key_prefix}_z_area")
        perception = st.checkbox("Perception Area", key=f"{state_key_prefix}_z_perception")
        ranged = st.checkbox("Ranged", key=f"{state_key_prefix}_z_ranged")

        if interpret_effects_to_build:
            build = interpret_effects_to_build(
                rank=rank3,
                is_special=is_special,
                effects=effects,
                sub_ranks=sub_ranks,
                area=None if area == "Nenhuma" else area,
                perception=perception,
                ranged=ranged,
            )
        else:
            build = "Erro: move_interpreter.py não encontrado (interpret_effects_to_build)."

        st.markdown("### Build Gerada")
        st.code(build, language="text")
                # =========================
        # PP do golpe criado do zero
        # =========================
        def _estimate_pp_from_build(build_txt: str, rank: int) -> int | None:
            b = (build_txt or "").lower()
            base = 0.0

            if "damage" in b: base += 1.0
            if "affliction" in b: base += 1.0
            if "weaken" in b: base += 1.0
            if "healing" in b: base += 2.0
            if "create" in b: base += 1.0
            if "environment" in b: base += 1.0
            if "nullify" in b: base += 1.0

            # extras comuns
            if "linked" in b: base += 0.5
            if "area" in b: base += 0.5
            if "perception" in b: base += 0.5
            if "ranged" in b: base += 0.25

            if base <= 0:
                return None

            return int(round(base * int(rank)))

        pp_auto = _estimate_pp_from_build(build, rank3)

        # Se customizou ranks, PP manual é obrigatório.
        # Se não der pra estimar, também obriga manual.
        if custom_sub or pp_auto is None:
            default_pp = 1
            if custom_sub and manual_pp:
                default_pp = max(1, int(manual_pp))
            pp_final = int(st.number_input(
                "PP total do golpe (obrigatório)",
                min_value=1,
                value=int(default_pp),
                step=1,
                key=f"{state_key_prefix}_z_pp_required"
            ))
            st.info("PP manual obrigatório." if custom_sub else "PP manual obrigatório (não foi possível estimar).")
        else:
            # Mesmo estimando, deixo você ajustar
            pp_final = int(st.number_input(
                "PP total do golpe",
                min_value=1,
                value=int(pp_auto),
                step=1,
                key=f"{state_key_prefix}_z_pp"
            ))
            st.info(f"PP sugerido: **{pp_auto}** (você pode ajustar).")


        st.markdown("### 🏷️ Nome do Golpe")
        custom_name_input = st.text_input(
            "Defina o nome do seu golpe",
            value="Golpe Customizado",
            key=f"{state_key_prefix}_z_custom_name"
        )
        # Garante que não fique vazio (se o usuário apagar tudo, volta para o padrão)
        final_custom_name = custom_name_input.strip() or "Golpe Customizado"

        col_confirm_zero, col_add_zero = st.columns(2)
        with col_confirm_zero:
            can_confirm = (pp_final is not None) and (int(pp_final) > 0)
            if st.button("✅ Confirmar golpe criado do zero", key=f"{state_key_prefix}_z_confirm", disabled=not can_confirm):
                st.session_state["cg_moves"].append({
                    "name": final_custom_name,
                    "rank": int(rank3),
                    "build": build,
                    "pp_cost": int(pp_final),
                    "accuracy": 0,
                    "meta": {
                        "custom": True,
                        "sub_ranks": sub_ranks,
                        "pp_manual": bool(custom_sub),
                        "is_special": bool(is_special),
                    }
                })
                st.success(f"Golpe '{final_custom_name}' adicionado à ficha.")

        if return_to_view:
            with col_add_zero:
                if st.button("➕ Adicionar golpe à ficha", key=f"{state_key_prefix}_z_add_sheet"):
                    st.session_state["cg_moves"].append({
                        "name": final_custom_name,
                        "rank": int(rank3),
                        "build": build,
                        "pp_cost": int(pp_final),
                        "accuracy": 0,
                        "meta": {
                            "custom": True,
                            "sub_ranks": sub_ranks,
                            "pp_manual": bool(custom_sub),
                            "is_special": bool(is_special),
                        }                    })
                    st.success(f"Golpe '{final_custom_name}' adicionado à ficha.")
                    st.session_state["cg_view"] = return_to_view
                    st.rerun()

    with tab4:
        st.subheader("✍️ Entrada Manual")
        st.caption("Preencha os campos para golpes homebrew ou que não estão no Excel.")

        c_man_1, c_man_2 = st.columns([2, 1])
        with c_man_1:
            manual_name = st.text_input("Nome do Golpe", key=f"{state_key_prefix}_m_name")
        with c_man_2:
            # ✅ CAMPO NOVO: CUSTO DE PP
            manual_pp = st.number_input("Custo (PP)", min_value=1, value=1, key=f"{state_key_prefix}_m_pp")

        c_man_3, c_man_4 = st.columns(2)
        with c_man_3:
            manual_rank = st.number_input("Rank do Efeito", min_value=1, max_value=20, value=10, key=f"{state_key_prefix}_m_rank")
        with c_man_4:
            manual_accuracy = st.number_input("Bônus de Acerto", min_value=0, max_value=30, value=0, key=f"{state_key_prefix}_m_accuracy")

        manual_formula = st.text_area("Fórmula / Build (Opcional)", height=80, key=f"{state_key_prefix}_m_formula", placeholder="Ex: Damage 10, Ranged")
        manual_desc = st.text_area("Descrição", height=100, key=f"{state_key_prefix}_m_desc")
        
        # --- Configurações Técnicas ---
        st.markdown("**Configurações:**")
        c_cfg1, c_cfg2 = st.columns(2)
        with c_cfg1:
            manual_type = st.text_input("Tipo (ex: Fogo, Psíquico)", key=f"{state_key_prefix}_m_type")
            # ✅ CAMPO NOVO: RESISTÊNCIA
            manual_resist = st.selectbox(
                "Resistência (Save do Alvo)", 
                ["Toughness", "Fortitude", "Will", "Dodge", "Parry", "—"],
                index=0, # Padrão Toughness
                key=f"{state_key_prefix}_m_resist"
            )
        with c_cfg2:
            st.caption("Flags:")
            manual_area = st.checkbox("É Área?", key=f"{state_key_prefix}_m_area")
            manual_ranged = st.checkbox("É Ranged?", key=f"{state_key_prefix}_m_ranged")
            manual_is_special = st.checkbox("É Especial (Int)?", key=f"{state_key_prefix}_m_special")

        # Define categoria baseada na flag
        cat_manual = "Manual Special" if manual_is_special else "Manual Physical"

        manual_data = {
            "Nome": manual_name.strip(),
            "Rank": int(manual_rank),
            "Formula": manual_formula.strip(),
            "Acerto": int(manual_accuracy),
            "Tipo": manual_type.strip(),
            "Descricao": manual_desc.strip(),
            "Area": bool(manual_area),
            "Ranged": bool(manual_ranged),
            "Categoria": cat_manual,
            # ✅ SALVANDO NOVOS DADOS
            "Resist Stat": manual_resist, 
            "PP_Custo": int(manual_pp)
        }

        can_save_manual = bool(manual_data["Nome"])

        col_save, col_save_add = st.columns(2)
        with col_save:
            if st.button("💾 Salvar no repertório", key=f"{state_key_prefix}_m_save", disabled=not can_save_manual):
                st.session_state["cg_manual_moves"].append(manual_data)
                st.success("Golpe manual salvo no repertório (disponível na busca).")

        with col_save_add:
            if st.button("✅ Salvar e adicionar à ficha", key=f"{state_key_prefix}_m_save_add", disabled=not can_save_manual):
                st.session_state["cg_manual_moves"].append(manual_data)
                
                # Converte para objeto Move
                mv = _manual_move_to_move(manual_data)
                
                # ✅ Passa o PP manual diretamente
                _confirm_move(mv, rank=int(manual_data["Rank"]), build=mv.build, pp=int(manual_data["PP_Custo"]))
                
                st.success("Golpe adicionado à ficha com sucesso!")

    st.divider()
    st.subheader("📦 Golpes confirmados nesta ficha")

    if not st.session_state["cg_moves"]:
        st.info("Nenhum golpe confirmado ainda.")
    else:
        for i, m in enumerate(list(st.session_state["cg_moves"])):
            c1, c2, c3 = st.columns([6, 2, 2])
            with c1:
                accuracy = int(m.get("accuracy", 0) or 0)
                st.write(f"**{m['name']}** (Rank {m['rank']}) — PP: {m.get('pp_cost')} | Acerto {accuracy}")
                build_txt = (m.get("build") or "").strip()
                if build_txt:
                    with st.expander("Ingredientes do golpe"):
                        st.code(build_txt, language="text")
            with c2:
                stats = st.session_state.get("cg_draft", {}).get("stats", {})
                np_value = int(st.session_state.get("cg_np", 0) or 0)
                
                # --- CÓDIGO NOVO (CORREÇÃO) ---
                acc_limit = _move_accuracy_limit(m, np_value, stats)
                current_acc = int(m.get("accuracy", 0) or 0)
                
                # O "pulo do gato": o máximo é o limite OU o valor atual (o que for maior)
                # Isso impede o erro se você digitou um valor alto manualmente antes
                safe_max = max(acc_limit, current_acc)

                new_acc = st.number_input(
                    "Acerto",
                    min_value=0,
                    max_value=safe_max,  # <--- Aqui usamos o safe_max
                    value=current_acc,
                    step=1,
                    key=f"{state_key_prefix}_acc_{i}",
                )
                
                st.caption(f"Limite: {acc_limit}")
                if current_acc > acc_limit:
                    st.warning("⚠️ Acima do limite!")

                if st.button("Definir acerto", key=f"{state_key_prefix}_set_acc_{i}"):
                    m["accuracy"] = int(new_acc)
                    st.rerun()
                # --- FIM DO CÓDIGO NOVO ---
            with c3:
                if st.button("❌ Remover", key=f"{state_key_prefix}_remove_{i}"):
                    st.session_state["cg_moves"].pop(i)
                    st.rerun()






from io import BytesIO
from PIL import ImageFont
if "carousel_click" not in st.session_state:
    st.session_state["carousel_click"] = None


SKILLS_MM3 = [
    "Acrobatics",
    "Athletics",
    "Close Combat (especialidade)",
    "Deception",
    "Expertise (especialidade)",
    "Insight",
    "Intimidation",
    "Investigation",
    "Perception",
    "Persuasion",
    "Ranged Combat (especialidade)",
    "Sleight of Hand",
    "Stealth",
    "Technology",
    "Treatment",
    "Vehicles",
]
# ================================
# CRIAÇÃO GUIADA - DRAFT ÚNICO
# ================================
def cg_init(reset: bool = False):
    """Inicializa (ou reseta) o rascunho único da ficha guiada."""
    if reset or ("cg_draft" not in st.session_state):
        st.session_state["cg_draft"] = {
            "pname": "",
            "stats": {
                "stgr": 0,
                "int": 0,
                "dodge": 0,
                "parry": 0,
                "thg": 0,
                "fortitude": 0,
                "will": 0,
            },
            "moves": [],
            "notas": "",
        }
    # Garante que notas existe em drafts antigos sem esse campo
    st.session_state["cg_draft"].setdefault("notas", "")
    # 🔗 alias: cg_moves aponta para a MESMA lista do draft
    st.session_state.setdefault("cg_moves", st.session_state["cg_draft"]["moves"])
    st.session_state["cg_draft"]["moves"] = st.session_state["cg_moves"]

def cg_sync_from_widgets():
    """Sincroniza valores dos widgets (keys) para o dicionário cg_draft."""
    d = st.session_state.get("cg_draft")
    if not d:
        return
    if "cg_pname" in st.session_state:
        d["pname"] = st.session_state["cg_pname"]
    if "cg_notas" in st.session_state:
        d["notas"] = st.session_state["cg_notas"]

    for k_widget, k_stat in [
        ("cg_stgr", "stgr"),
        ("cg_int", "int"),
        ("cg_dodge", "dodge"),
        ("cg_parry", "parry"),
        ("cg_thg", "thg"),
        ("cg_fortitude", "fortitude"),
        ("cg_will", "will"),
    ]:
        if k_widget in st.session_state:
            d["stats"][k_stat] = st.session_state[k_widget]


def cg_reset_for_new_pokemon(pname: str):
    """Limpa o estado da criação guiada quando o usuário troca o Pokémon."""
    cg_init(reset=True)
    st.session_state["cg_draft"]["pname"] = pname

    # Widgets/campos da ficha
    st.session_state["cg_np"] = 0
    st.session_state["cg_abilities"] = []
    st.session_state["cg_advantages"] = []
    st.session_state["cg_moves"] = []
    st.session_state["cg_draft"]["moves"] = st.session_state["cg_moves"]

    for k in [
        "cg_stgr",
        "cg_int",
        "cg_dodge",
        "cg_parry",
        "cg_thg",
        "cg_fortitude",
        "cg_will",
    ]:
        st.session_state[k] = 0

    # auxiliares da UI que podem carregar informações da ficha anterior
    for k in [
        "cg_quick_pick",
        "cg_viab_sel_idx",
        "cg_viab_raw_view",
        "cg_viab_rank_default",
        "cg_viab_selected",
        "cg_viab_mode",
        "cg_imported_name",
        "cg_imported_types",
        "cg_imported_abilities",
        "cg_loaded_sheet_id",
    ]:
        st.session_state.pop(k, None)





# ================================
# FIREBASE - TESTE DE CONEXÃO
# ================================
import firebase_admin
from firebase_admin import credentials, firestore, storage



def init_firebase():
    if not firebase_admin._apps:
        raw = st.secrets["firebase_service_account"]
        cred_dict = {k: raw[k] for k in raw.keys()}  # <-- dict puro

        cred = credentials.Certificate(cred_dict)
        firebase_admin.initialize_app(cred, {
            "projectId": cred_dict["project_id"],
            "storageBucket": "batalhas-de-gaal.firebasestorage.app",
        })

    db = firestore.client()
    bucket = storage.bucket()
    return db, bucket
    
# ==========================
# FIREBASE SAVE/LOAD (Fichas)
# ==========================
from datetime import timezone

def _utc_now_iso():
    return datetime.now(timezone.utc).isoformat()

def upload_pdf_to_storage(bucket, pdf_bytes: bytes, storage_path: str):
    blob = bucket.blob(storage_path)
    blob.upload_from_string(pdf_bytes, content_type="application/pdf")
    return storage_path

def upload_png_to_storage(bucket, png_bytes: bytes, storage_path: str):
    blob = bucket.blob(storage_path)
    blob.upload_from_string(png_bytes, content_type="image/png")
    return storage_path
def upload_png_to_storage_with_token(bucket, png_bytes: bytes, storage_path: str) -> dict:
    """Upload PNG ao Firebase Storage com download token e retorna {storage_path, token, url}."""
    blob = bucket.blob(storage_path)
    token = uuid.uuid4().hex
    try:
        blob.metadata = (blob.metadata or {})
        blob.metadata["firebaseStorageDownloadTokens"] = token
    except Exception:
        pass
    blob.upload_from_string(png_bytes, content_type="image/png")
    try:
        blob.patch()
    except Exception:
        pass

    bucket_name = getattr(bucket, "name", None) or "batalhas-de-gaal.firebasestorage.app"
    enc_path = urllib.parse.quote(storage_path, safe="")
    url = f"https://firebasestorage.googleapis.com/v0/b/{bucket_name}/o/{enc_path}?alt=media&token={token}"
    return {"storage_path": storage_path, "token": token, "url": url}



def upload_json_to_storage_with_token(bucket, json_bytes: bytes, storage_path: str) -> dict:
    """Upload JSON ao Firebase Storage com download token e retorna {storage_path, token, url}."""
    blob = bucket.blob(storage_path)
    token = uuid.uuid4().hex
    try:
        blob.metadata = (blob.metadata or {})
        blob.metadata["firebaseStorageDownloadTokens"] = token
    except Exception:
        pass
    blob.upload_from_string(json_bytes, content_type="application/json")
    try:
        blob.patch()
    except Exception:
        pass

    bucket_name = getattr(bucket, "name", None) or "batalhas-de-gaal.firebasestorage.app"
    enc_path = urllib.parse.quote(storage_path, safe="")
    url = f"https://firebasestorage.googleapis.com/v0/b/{bucket_name}/o/{enc_path}?alt=media&token={token}"
    return {"storage_path": storage_path, "token": token, "url": url}

def upload_avatar_choice_to_storage(bucket, trainer_name: str, avatar_choice: str) -> dict | None:
    """Publica o avatar escolhido do treinador no Storage e devolve metadados de acesso."""
    if not avatar_choice:
        return None

    avatar_lookup = build_trainer_avatar_lookup()
    avatar_info = avatar_lookup.get(avatar_choice) or {}
    avatar_path = avatar_info.get("path")
    if not avatar_path or not os.path.exists(avatar_path):
        return None

    try:
        with open(avatar_path, "rb") as fp:
            avatar_bytes = fp.read()
    except Exception:
        return None

    storage_path = f"trainer_avatars/{safe_doc_id(trainer_name)}/avatar_{safe_doc_id(avatar_choice)}.png"
    return upload_png_to_storage_with_token(bucket, avatar_bytes, storage_path)

def ensure_room_map_published(db, bucket, rid: str, grid: int, theme_key: str, seed: int, no_water: bool, show_grid: bool = True):
    """Gera o mapa base (sem peças) e publica no Storage, salvando mapUrl/mapStoragePath em public_state/state.

    Além do PNG, também publica um JSON com a saída do BiomeGenerator (terrain_grid + decorations),
    para que outros clientes (ex.: battle-site) possam consumir os dados estruturados do mapa.
    """
    try:
        # 1) Gera PNG + dados estruturados num único passo cacheado
        png_bytes, map_data = _generate_map_cached(grid, theme_key, seed, no_water, show_grid=show_grid)

        # 2) Define paths estáveis no Storage (PNG + JSON)
        water_tag = "nowater" if no_water else "water"
        png_path = f"rooms/{rid}/maps/{grid}x{grid}_{theme_key}_{seed}_{water_tag}{'_grid' if show_grid else ''}.png"
        json_path = f"rooms/{rid}/maps/{grid}x{grid}_{theme_key}_{seed}_{water_tag}.json"

        # 3) Upload do PNG
        out_png = upload_png_to_storage_with_token(bucket, png_bytes, png_path)

        # 4) Upload do JSON estruturado (terrain_grid + decorations)
        try:
            map_data = dict(map_data)  # cópia para não mutar o cache
            map_data.setdefault("theme_key", str(theme_key))
            map_data.setdefault("no_water", bool(no_water))
            map_data.setdefault("rid", str(rid))
            map_data_bytes = json.dumps(map_data, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
            out_json = upload_json_to_storage_with_token(bucket, map_data_bytes, json_path)
        except Exception:
            out_json = None

        # 5) Salva URLs no Firestore (mantém campos antigos; adiciona mapDataUrl para novos clientes)
        state_ref = db.collection("rooms").document(str(rid)).collection("public_state").document("state")
        payload = {
            "mapStoragePath": out_png["storage_path"],
            "mapUrl": out_png["url"],
            "mapUpdatedAt": firestore.SERVER_TIMESTAMP,
        }
        if out_json:
            payload.update({
                "mapDataStoragePath": out_json["storage_path"],
                "mapDataUrl": out_json["url"],
                "mapDataUpdatedAt": firestore.SERVER_TIMESTAMP,
            })
        state_ref.set(payload, merge=True)

        return {"png": out_png, "json": out_json}
    except Exception as e:
        try:
            st.warning(f"Falha ao publicar mapa no Storage: {e}")
        except Exception:
            pass
        return None

@st.cache_data(show_spinner=False, ttl=3600)
def download_storage_bytes(storage_path: str) -> bytes | None:
    if not storage_path:
        return None
    try:
        _, bucket = init_firebase()
        return bucket.blob(storage_path).download_as_bytes()
    except Exception:
        return None


def save_sheet_to_firestore(db, trainer_name: str, sheet_payload: dict, sheet_id=None):
    trainer_id = safe_doc_id(trainer_name)

    if not sheet_id:
        pname = sheet_payload.get("pokemon", {}).get("name", "pokemon")
        pid = sheet_payload.get("pokemon", {}).get("id", "0")
        sheet_id = safe_doc_id(f"{pname}_{pid}_{uuid.uuid4().hex[:8]}")

    ref = (
        db.collection("trainers")
        .document(trainer_id)
        .collection("sheets")
        .document(sheet_id)
    )

    pokemon_block = sheet_payload.get("pokemon")
    if isinstance(pokemon_block, dict):
        pokemon_name = pokemon_block.get("name", "")
        pokemon_block["id"] = normalize_sheet_pokemon_id(pokemon_block.get("id"), pokemon_name)

    now = _utc_now_iso()
    sheet_payload.setdefault("created_at", now)
    sheet_payload["updated_at"] = now
    sheet_payload["trainer_name"] = trainer_name

    ref.set(sheet_payload, merge=True)
    return sheet_id
# ==========================
# IMPORTADOR DE FICHA M&M (PDF) — Hero Lab / Wolf Lair
# ==========================
from typing import Any, Dict, List, Tuple, Optional
import re

def _mm_extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extrai texto do PDF (sem OCR)."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception:
        return ""

def _mm_parse_powers_block(block: str) -> List[Dict[str, Any]]:
    """
    Parser mais robusto pro Hero Lab:

    Problemas comuns que ele resolve:
    - Bullet do Hero Lab às vezes vem como "\uf0fc" ou como "ü" SEM espaço (ex.: "üDrain Punch").
      Se você fizer line[2:], você perde a 1ª letra do nome ("rain Punch").
    - Quebras de linha e linhas de detalhe ("DC", "Limited", "Affects"...)

    Regras:
    - Começa um novo power quando a linha "parece header" e contém "(X PP)" (com ou sem bullet).
    - Também considera "(alternate)" como header (pp_cost=None).
    - As demais linhas entram como detalhes do power atual.
    """
    block = (block or "").replace("\uf0fc", "ü")

    # bullets comuns em PDFs
    _BULLET_RE = re.compile(r"^[\s\u2022\u25cf\u00b7\uf0b7\uf0fcü•\-]+", flags=re.U)

    def strip_bullet(s: str) -> str:
        s = (s or "").strip()
        s = re.sub(_BULLET_RE, "", s)
        return s.lstrip()

    lines = [l.rstrip() for l in block.splitlines()]

    def is_header(line: str) -> bool:
        s = strip_bullet(line)
        if not s:
            return False
        # Headers típicos:
        if re.search(r"\(\s*\d+\s*PP\s*\)\s*$", s, flags=re.I):
            return True
        if re.search(r"\(\s*alternate\s*\)\s*$", s, flags=re.I):
            return True
        return False

    powers: List[Dict[str, Any]] = []
    cur: Optional[Dict[str, Any]] = None

    for raw in lines:
        line = strip_bullet(raw)
        if not line:
            continue

        if is_header(line):
            # fecha anterior
            if cur:
                powers.append(cur)

            header = line
            pp_m = re.search(r"\(\s*(\d+)\s*PP\s*\)\s*$", header, flags=re.I)
            alt_m = re.search(r"\(\s*alternate\s*\)\s*$", header, flags=re.I)

            if pp_m:
                pp_cost = int(pp_m.group(1))
                name = re.sub(r"\(\s*\d+\s*PP\s*\)\s*$", "", header, flags=re.I).strip()
            elif alt_m:
                pp_cost = None
                name = re.sub(r"\(\s*alternate\s*\)\s*$", "", header, flags=re.I).strip()
            else:
                pp_cost = None
                name = header.strip()

            cur = {"name": name or "SemNome", "pp_cost": pp_cost, "lines": []}
            continue

        # linha “normal” (detalhe)
        if not cur:
            cur = {"name": "(sem nome)", "pp_cost": None, "lines": []}

        cur["lines"].append(line)

    if cur:
        powers.append(cur)

    return powers
def _mm_parse_skills_block(full_text: str) -> Tuple[List[Dict[str, Any]], Dict[str, int]]:
    """
    A tabela de skills do Hero Lab às vezes vem como:
      'Ability Total OtherSkills' / 'OtherSkills Ranks' etc.
    Vamos capturar o bloco até 'Validation Report'.
    """
    skills: List[Dict[str, Any]] = []
    combat_bonus: Dict[str, int] = {}

    m = re.search(
        r"(?:Total Ability Other Skills|Ability Total OtherSkills).*?\nRanks\n(.*?)(?=\nValidation Report|\nSettings:|\Z)",
        full_text,
        flags=re.I | re.S,
    )
    if not m:
        return skills, combat_bonus

    block = m.group(1)
    for raw in block.splitlines():
        line = re.sub(r"\s+", " ", (raw or "").strip())
        if not line:
            continue

        # Exemplos:
        # "- Acrobatics -"
        m0 = re.match(r"^-\s+(.*?)\s+-$", line)
        if m0:
            sname = m0.group(1).strip()
            skills.append({"name": sname, "total": 0, "ranks": 0})
            continue

        # "+0 Persuasion -"
        m1 = re.match(r"^([+\-]?\d+)\s+(.*?)\s+-$", line)
        if m1:
            total = int(m1.group(1))
            sname = m1.group(2).strip()
            skills.append({"name": sname, "total": total, "ranks": 0})
        else:
            # "3 +3 Athletics -"
            m2 = re.match(r"^(\d+)\s+([+\-]?\d+)\s+(.*?)\s+-$", line)
            if m2:
                ranks = int(m2.group(1))
                total = int(m2.group(2))
                sname = m2.group(3).strip()
                skills.append({"name": sname, "total": total, "ranks": ranks})
            else:
                # "+4 Close Combat: Crunch 4"
                m3 = re.match(r"^([+\-]?\d+)\s+(.*?)\s+(\d+)$", line)
                if m3:
                    total = int(m3.group(1))
                    sname = m3.group(2).strip()
                    ranks = int(m3.group(3))
                    skills.append({"name": sname, "total": total, "ranks": ranks})
                else:
                    continue

        # Mapear bônus de acerto por golpe
        sname_l = skills[-1]["name"].lower()
        if sname_l.startswith("close combat:"):
            mv = skills[-1]["name"].split(":", 1)[1].strip().lower()
            combat_bonus[mv] = int(skills[-1]["total"])
        elif sname_l.startswith("ranged combat:"):
            mv = skills[-1]["name"].split(":", 1)[1].strip().lower()
            combat_bonus[mv] = int(skills[-1]["total"])

    return skills, combat_bonus

def _mm_infer_rank_from_lines(lines: List[str]) -> int:
    """
    Para golpes, o jeito mais robusto é:
      - Pegar o maior 'DC XX' que aparece nas linhas do golpe
      - Rank_total ~= DC - 15
    (Serve para Damage, Affliction, Weaken etc.)
    """
    joined = " ".join(lines or [])
    dcs = [int(x) for x in re.findall(r"\bDC\s*(\d+)\b", joined, flags=re.I)]
    if dcs:
        dc = max(dcs)
        return max(0, dc - 15)

    m = re.search(r"\bDamage\s*(\d+)\b", joined, flags=re.I)
    if m:
        return int(m.group(1))
    return 0

def _mm_clean_move_name(raw_name: str) -> str:
    """
    Tenta converter headers do Hero Lab em nomes “de golpe”.

    Exemplos:
      "Surf (TM): Line Area Damage 14" -> "Surf"
      "Knock Off: Strength-based Damage 4" -> "Knock Off"
      "Swords Dance (TM): Enhanced Strength 2" -> "Swords Dance"
      "üDrain Punch" -> "Drain Punch"
    """
    s = (raw_name or "").strip().replace("\uf0fc", "ü")

    # remove bullets comuns
    s = re.sub(r"^[\s\u2022\u25cf\u00b7\uf0b7\uf0fcü•\-]+", "", s).lstrip()

    # se tiver ":" e depois vier claramente “descrição mecânica”, pega só o lado esquerdo
    if ":" in s:
        left, right = s.split(":", 1)
        right_l = right.lower()
        if any(k in right_l for k in ["damage", "weaken", "affliction", "healing", "enhanced", "senses", "teleport", "leaping", "swimming", "area", "perception"]):
            s = left.strip()

    # remove (TM)
    s = re.sub(r"\(\s*tm\s*\)", "", s, flags=re.I).strip()
    # normaliza espaços
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _mm_is_probably_internal_power(name: str) -> bool:
    """
    Evita importar sub-powers genéricos do Hero Lab como se fossem golpes.
    (Ex.: "Dano", "Cura", "Weaken", "Pulo"...)
    """
    s = (name or "").strip().lower()
    bad_prefixes = (
        "dano", "cura", "weaken", "retorno a pokebola", "pulo", "nada bem",
        "panda da noite", "senses", "teleport", "leaping", "swimming"
    )
    return s.startswith(bad_prefixes)

def _mm_is_combat_power(power: Dict[str, Any]) -> bool:
    # 🔧 Inclui também o header, porque no Hero Lab muitas vezes “Damage X” fica no NOME do power
    # (ex.: "Surf (TM): Line Area Damage 14 (29 PP)") e as linhas só trazem "DC 29..."
    header = str(power.get("name", "") or "")
    txt = (header + " " + " ".join(power.get("lines", []) or [])).lower()
    return any(k in txt for k in [
        "damage", "weaken", "affliction", "healing", "nullify", "create", "environment", "enhanced"
    ])

def import_mm_pdf_to_sheet_payload(pdf_bytes: bytes) -> Dict[str, Any]:
    """
    Retorna um payload no formato "amigável pro site":
      - pokemon: {name, id}
      - stats: stgr/int/dodge/parry/thg/fortitude/will (o que seu front já usa)
      - moves: [{name, rank, build, pp_cost, accuracy, meta}]
      - mm: dados crus (pl/pp_total/abilities/defenses/skills/advantages/powers)
    """
    txt = _mm_extract_text_from_pdf(pdf_bytes)
    if not txt.strip():
        raise ValueError("Não consegui extrair texto do PDF (talvez ele seja imagem/scan).")

    # Detect simples do formato
    is_herolab = ("Hero Lab" in txt) and ("Mutants & Masterminds" in txt)

    # Nome: no Hero Lab, a linha antes de "Power Level ..." costuma ser o nome
    name = ""
    lines = [l.strip() for l in txt.splitlines() if l.strip()]
    pl_idx = next((i for i, l in enumerate(lines) if l.lower().startswith("power level")), None)
    if pl_idx is not None:
        for j in range(pl_idx - 1, -1, -1):
            cand = lines[j]
            if any(k in cand.lower() for k in ["hero lab", "mutants", "registered trademarks", "free download"]):
                continue
            if re.search(r"\b(male|female|age|height|weight)\b", cand.lower()):
                continue
            if len(cand) <= 50 and not any(ch.isdigit() for ch in cand):
                name = cand
                break
    if not name:
        name = lines[0] if lines else "SemNome"

    # PL / PP total
    pl = None
    pp_total = None
    m = re.search(r"Power Level\s+(\d+),\s*([0-9]+)\s*PP", txt, flags=re.I)
    if m:
        pl = int(m.group(1))
        pp_total = int(m.group(2))

    # Abilities
    abil_map = {
        "Strength": "str", "Agility": "agi", "Fighting": "fgt", "Stamina": "sta",
        "Dexterity": "dex", "Intellect": "int", "Awareness": "awe", "Presence": "pre",
    }
    abilities: Dict[str, int] = {}
    for ab, key in abil_map.items():
        m = re.search(rf"\b{ab}\s+(-?\d+)", txt, flags=re.I)
        if m:
            abilities[key] = int(m.group(1))

    # Defenses
    def_map = {"Dodge": "dodge", "Parry": "parry", "Fortitude": "fortitude", "Toughness": "toughness", "Will": "will"}
    defenses: Dict[str, int] = {}
    for df, key in def_map.items():
        m = re.search(rf"\b{df}\s+(\d+)", txt, flags=re.I)
        if m:
            defenses[key] = int(m.group(1))

    # Advantages (somente nomes)
    adv: List[str] = []
    m = re.search(r"\nAdvantages\n(.*?)(?=\nMovement\n|\nDefenses\n|\nBackground Information\n)", txt, flags=re.I | re.S)
    if m:
        for raw in m.group(1).splitlines():
            line = raw.strip(" .")
            if not line:
                continue
            adv_name = re.split(r"[:(]", line, maxsplit=1)[0].strip()
            if adv_name and adv_name not in adv:
                adv.append(adv_name)

    # Powers: pegar o bloco real "Powers\n ... \nAdvantages"
    powers: List[Dict[str, Any]] = []
    m = re.search(r"\nPowers\n(.*?)(?=\nAdvantages\n|\nMovement\n|\nDefenses\n|\nBackground Information\n)", txt, flags=re.I | re.S)
    if m:
        powers = _mm_parse_powers_block(m.group(1))

    # Skills + combate por golpe
    skills, combat_bonus = _mm_parse_skills_block(txt)

    # Converter para o “formato do site”
    stats = {
        "stgr": int(abilities.get("str", 0)),
        "int": int(abilities.get("int", 0)),
        "dodge": int(defenses.get("dodge", 0)),
        "parry": int(defenses.get("parry", 0)),
        "thg": int(defenses.get("toughness", 0)),
        "fortitude": int(defenses.get("fortitude", 0)),
        "will": int(defenses.get("will", 0)),
    }

    moves: List[Dict[str, Any]] = []
    # Para ficar robusto:
    # - limpa nomes (bullet sem espaço, sufixos mecânicos após ":")
    # - não depende de "damage" estar nas linhas; também olha o header
    # - evita importar sub-powers genéricos (Dano/Cura/Weaken/Pulo...)
    # - tenta mapear bônus de acerto mesmo quando o nome do power vem com descrição
    import difflib

    def _best_accuracy(move_name: str) -> int:
        key = (move_name or "").strip().lower()
        if not key:
            return 0
        if key in combat_bonus:
            return int(combat_bonus[key])

        # tenta match aproximado
        options = list(combat_bonus.keys())
        close = difflib.get_close_matches(key, options, n=1, cutoff=0.84)
        if close:
            return int(combat_bonus.get(close[0], 0))
        return 0

    for p in powers:
        if not _mm_is_combat_power(p):
            continue

        raw_header = p.get("name") or "SemNome"
        mv_name = _mm_clean_move_name(raw_header)

        if _mm_is_probably_internal_power(mv_name):
            continue

        mv_lines = p.get("lines") or []

        # accuracy: usa o nome limpo
        acc = _best_accuracy(mv_name)

        # rank: considera também o header (às vezes vem "Damage 14" nele)
        rank = int(_mm_infer_rank_from_lines(list(mv_lines) + [str(raw_header)]))

        # build: inclui o header se ele carregar "ingredientes" (Damage/Area/etc)
        header_has_rules = bool(re.search(r"\b(damage|weaken|affliction|healing|environment|enhanced|area|perception)\b", str(raw_header), flags=re.I))
        build_parts: List[str] = []
        if header_has_rules and str(raw_header).strip() and str(raw_header).strip() != mv_name:
            build_parts.append(str(raw_header).strip())
        build_parts.extend([l.strip() for l in mv_lines if (l or "").strip()])
        build = "; ".join(build_parts)

        moves.append({
            "name": mv_name,
            "rank": rank,
            "build": build,
            "pp_cost": int(p["pp_cost"]) if p.get("pp_cost") is not None else None,
            "accuracy": acc,
            "meta": {
                "source": "mm_pdf_herolab",
                "format_detected": "herolab" if is_herolab else "unknown",
                "raw_power_name": raw_header,
            },
        })
# id "estável" (você pode trocar depois)
    pokemon_id = normalize_sheet_pokemon_id(f"mm_{safe_doc_id(name)}", name)

    payload = {
        "pokemon": {
            "name": name,
            "id": pokemon_id,
            "source": "mm_pdf",
        },
        "np": pl,  # se seu site usa NP, dá pra mapear PL -> NP
        "stats": stats,
        "moves": moves,
        "mm": {
            "format": "herolab_pdf" if is_herolab else "unknown_pdf",
            "pl": pl,
            "pp_total": pp_total,
            "abilities": abilities,
            "defenses": defenses,
            "advantages": adv,
            "skills": skills,
            "powers_raw": powers,
        }
    }
    return payload
def save_sheet_with_pdf(db, bucket, trainer_name: str, sheet_payload: dict, pdf_bytes=None, sheet_id=None):
    storage_path = None

    if pdf_bytes:
        pname = sheet_payload.get("pokemon", {}).get("name", "pokemon")
        pid = sheet_payload.get("pokemon", {}).get("id", "0")
        storage_path = (
            f"fichas/{safe_doc_id(trainer_name)}/"
            f"{safe_doc_id(pname)}_{pid}_{uuid.uuid4().hex[:8]}.pdf"
        )
        upload_pdf_to_storage(bucket, pdf_bytes, storage_path)

        sheet_payload.setdefault("pdf", {})
        sheet_payload["pdf"].update({
            "storage_path": storage_path,
            "updated_at": _utc_now_iso(),
            "version": int(sheet_payload.get("pdf", {}).get("version", 0)) + 1,
        })

    sheet_id = save_sheet_to_firestore(db, trainer_name, sheet_payload, sheet_id=sheet_id)
    return sheet_id, storage_path

def build_sheet_pdf(
    pname: str,
    np_: int,
    types: list[str],
    abilities: list[str],
    stats: dict,
    chosen_adv: list[str],
    moves: list[dict],
) -> bytes:
    from reportlab.pdfgen import canvas

    buffer = BytesIO()
    c = canvas.Canvas(buffer)
    c.setFont("Helvetica", 12)
    c.drawString(40, 800, f"Ficha Pokémon - {pname} (NP {np_})")
    c.drawString(40, 780, f"Tipos: {', '.join(types)}")
    c.drawString(40, 760, f"Abilities: {', '.join(abilities)}")
    c.drawString(
        40,
        730,
        "Stgr {stgr} | Int {intellect} | Dodge {dodge} | Parry {parry} | "
        "Fort {fortitude} | Will {will}".format(**stats),
    )
    c.drawString(40, 710, f"Advantages: {', '.join(chosen_adv) if chosen_adv else '(nenhuma)'}")

    y = 680
    c.drawString(40, y, "Golpes:")
    y -= 18
    def _draw_line(text: str, indent: int = 0):
        nonlocal y
        c.drawString(50 + indent, y, text)
        y -= 14
        if y < 80:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = 800
    for m in moves:
        accuracy = int(m.get("accuracy", 0) or 0)
        _draw_line(f"- {m['name']} (Rank {m['rank']}) | PP {m.get('pp_cost')} | Acerto {accuracy}")
        build_txt = (m.get("build") or "").strip()
        if build_txt:
            _draw_line("Ingredientes:", indent=10)
            for line in build_txt.splitlines():
                _draw_line(line, indent=20)


    c.showPage()
    c.save()
    return buffer.getvalue()


def parse_sheet_pdf(pdf_bytes: bytes) -> dict:
    from PyPDF2 import PdfReader

    reader = PdfReader(BytesIO(pdf_bytes))
    raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]

    header_re = re.compile(r"^Ficha Pokémon - (.+) \(NP (\d+)\)$")
    stats_re = re.compile(
        r"Stgr\s+(\d+)\s*\|\s*Int\s+(\d+)\s*\|\s*Dodge\s+(\d+)\s*\|\s*Parry\s+(\d+)\s*\|\s*Fort\s+(\d+)\s*\|\s*Will\s+(\d+)"
    )
    move_re = re.compile(r"^- (.+) \(Rank (\d+)\) \| PP ([^|]+) \| Acerto (\d+)$")

    pname = ""
    np_value = 0
    types: list[str] = []
    abilities: list[str] = []
    stats: dict[str, int] = {}
    advantages: list[str] = []
    moves: list[dict] = []

    for line in lines:
        header_match = header_re.match(line)
        if header_match:
            pname = header_match.group(1).strip()
            np_value = int(header_match.group(2))
            continue
        if line.startswith("Tipos:"):
            raw_types = line.replace("Tipos:", "", 1).strip()
            types = [t.strip() for t in raw_types.split(",") if t.strip()]
            continue
        if line.startswith("Abilities:"):
            raw_abilities = line.replace("Abilities:", "", 1).strip()
            abilities = [a.strip() for a in raw_abilities.split(",") if a.strip()]
            continue
        if line.startswith("Advantages:"):
            raw_adv = line.replace("Advantages:", "", 1).strip()
            if raw_adv and raw_adv != "(nenhuma)":
                advantages = [a.strip() for a in raw_adv.split(",") if a.strip()]
            continue
        stats_match = stats_re.search(line)
        if stats_match:
            stats = {
                "stgr": int(stats_match.group(1)),
                "int": int(stats_match.group(2)),
                "dodge": int(stats_match.group(3)),
                "parry": int(stats_match.group(4)),
                "fortitude": int(stats_match.group(5)),
                "will": int(stats_match.group(6)),
            }

    in_moves = False
    current_move = None
    build_lines: list[str] = []
    collecting_build = False

    for line in lines:
        if line == "Golpes:":
            in_moves = True
            continue
        if not in_moves:
            continue

        move_match = move_re.match(line)
        if move_match:
            if current_move:
                current_move["build"] = "\n".join(build_lines).strip()
                moves.append(current_move)
            move_name = move_match.group(1).strip()
            move_rank = int(move_match.group(2))
            pp_raw = move_match.group(3).strip()
            pp_value = int(pp_raw) if pp_raw.isdigit() else None
            acc_value = int(move_match.group(4))
            current_move = {
                "name": move_name,
                "rank": move_rank,
                "pp_cost": pp_value,
                "accuracy": acc_value,
                "build": "",
            }
            build_lines = []
            collecting_build = False
            continue

        if line.startswith("Ingredientes"):
            collecting_build = True
            continue

        if collecting_build and current_move:
            build_lines.append(line)

    if current_move:
        current_move["build"] = "\n".join(build_lines).strip()
        moves.append(current_move)

    if not pname or np_value <= 0:
        raise ValueError("Não foi possível identificar nome ou NP no PDF.")

    return {
        "pokemon": {"name": pname, "id": normalize_sheet_pokemon_id(f"pdf_{safe_doc_id(pname)}", pname), "types": types, "abilities": abilities},
        "np": np_value,
        "stats": stats,
        "advantages": advantages,
        "moves": moves,
        "meta": {"source": "site_pdf"},
    }


def apply_imported_sheet_to_session(sheet: dict):
    pokemon = sheet.get("pokemon", {}) if isinstance(sheet, dict) else {}
    pname = str(pokemon.get("name", "") or "")
    np_ = int(sheet.get("np", 0) or 0)
    stats = sheet.get("stats") or {}
    moves = sheet.get("moves") or []
    advantages = sheet.get("advantages") or []
    abilities = pokemon.get("abilities") or []
    types = pokemon.get("types") or []

    cap = 2 * np_
    dodge_val = int(stats.get("dodge", 0) or 0)

    for m in moves:
        if isinstance(m, dict) and "accuracy" not in m:
            m["accuracy"] = 0

    st.session_state["cg_edit_sheet_id"] = None
    st.session_state["cg_imported_name"] = pname
    st.session_state["cg_imported_types"] = list(types)
    st.session_state["cg_imported_abilities"] = list(abilities)

    st.session_state["cg_draft"] = {
        "pname": pname,
        "np": np_,
        "stats": {
            "stgr": int(stats.get("stgr", 0) or 0),
            "int": int(stats.get("int", 0) or 0),
            "dodge": dodge_val,
            "parry": int(stats.get("parry", dodge_val) or 0),
            "thg": int(stats.get("thg") or max(0, cap - dodge_val)),
            "fortitude": int(stats.get("fortitude", 0) or 0),
            "will": int(stats.get("will", 0) or 0),
        },
        "moves": list(moves),
    }
    st.session_state["cg_moves"] = st.session_state["cg_draft"]["moves"]
    st.session_state["cg_pname"] = pname
    st.session_state["cg_np"] = np_
    st.session_state["cg_stgr"] = st.session_state["cg_draft"]["stats"]["stgr"]
    st.session_state["cg_int"] = st.session_state["cg_draft"]["stats"]["int"]
    st.session_state["cg_dodge"] = st.session_state["cg_draft"]["stats"]["dodge"]
    st.session_state["cg_parry"] = st.session_state["cg_draft"]["stats"]["parry"]
    st.session_state["cg_thg"] = st.session_state["cg_draft"]["stats"]["thg"]
    st.session_state["cg_fortitude"] = st.session_state["cg_draft"]["stats"]["fortitude"]
    st.session_state["cg_will"] = st.session_state["cg_draft"]["stats"]["will"]
    st.session_state["cg_advantages"] = list(advantages)
    st.session_state["cg_abilities"] = list(abilities)

def list_sheets(db, trainer_name: str, limit: int = 50):
    trainer_id = safe_doc_id(trainer_name)
    docs = (
        db.collection("trainers")
        .document(trainer_id)
        .collection("sheets")
        .order_by("updated_at", direction="DESCENDING")
        .limit(limit)
        .stream()
    )
    out = []
    for d in docs:
        item = d.to_dict() or {}
        item["_sheet_id"] = d.id
        out.append(item)
    return out

def load_sheet(db, trainer_name: str, sheet_id: str):
    trainer_id = safe_doc_id(trainer_name)
    ref = (
        db.collection("trainers")
        .document(trainer_id)
        .collection("sheets")
        .document(sheet_id)
    )
    snap = ref.get()
    return snap.to_dict() if snap.exists else None
def delete_sheet(db, bucket, trainer_name: str, sheet_id: str, storage_path: str | None = None):
    trainer_id = safe_doc_id(trainer_name)
    ref = (
        db.collection("trainers")
        .document(trainer_id)
        .collection("sheets")
        .document(sheet_id)
    )
    ref.delete()

    if storage_path:
        try:
            bucket.blob(storage_path).delete()
        except Exception:
            pass

def apply_sheet_to_session(sheet: dict, sheet_id: str | None = None):
    pokemon = sheet.get("pokemon", {}) if isinstance(sheet, dict) else {}
    pname = str(pokemon.get("name", "") or "")
    stats = sheet.get("stats") or {}
    moves = sheet.get("moves") or []
    advantages = sheet.get("advantages") or []
    skills = sheet.get("skills") or []
    np_ = int(sheet.get("np", 0) or 0)
    cap = 2 * np_
    abilities = pokemon.get("abilities") or []

    st.session_state["cg_edit_sheet_id"] = sheet_id
    for m in moves:
        if isinstance(m, dict) and "accuracy" not in m:
            m["accuracy"] = 0
    st.session_state["cg_draft"] = {
        "pname": pname,
        "stats": {
            "stgr": int(stats.get("stgr", 0) or 0),
            "int": int(stats.get("int", 0) or 0),
            "dodge": int(stats.get("dodge", 0) or 0),
            "parry": int(stats.get("parry", stats.get("dodge", 0)) or 0),
            "thg": int(stats.get("thg") or max(0, cap - int(stats.get("dodge", 0) or 0))),
            "fortitude": int(stats.get("fortitude", 0) or 0),
            "will": int(stats.get("will", 0) or 0),
        },
        "moves": list(moves),
    }
    st.session_state["cg_moves"] = st.session_state["cg_draft"]["moves"]
    st.session_state["cg_pname"] = pname
    st.session_state["cg_np"] = np_
    st.session_state["cg_stgr"] = st.session_state["cg_draft"]["stats"]["stgr"]
    st.session_state["cg_int"] = st.session_state["cg_draft"]["stats"]["int"]
    st.session_state["cg_dodge"] = st.session_state["cg_draft"]["stats"]["dodge"]
    st.session_state["cg_parry"] = st.session_state["cg_draft"]["stats"]["parry"]
    st.session_state["cg_thg"] = st.session_state["cg_draft"]["stats"]["thg"]
    st.session_state["cg_fortitude"] = st.session_state["cg_draft"]["stats"]["fortitude"]
    st.session_state["cg_will"] = st.session_state["cg_draft"]["stats"]["will"]
    st.session_state["cg_advantages"] = list(advantages)
    st.session_state["cg_abilities"] = list(abilities)

    base_skills = {k: 0 for k in SKILLS_MM3}
    custom_skills = []
    for row in skills:
        if not isinstance(row, dict):
            continue
        name = str(row.get("name", "")).strip()
        try:
            ranks = int(row.get("ranks", 0))
        except Exception:
            ranks = 0
        if not name:
            continue
        if name in base_skills:
            base_skills[name] = ranks
        else:
            custom_skills.append({"name": name, "ranks": ranks})

    st.session_state["cg_skills"] = base_skills
    st.session_state["cg_skill_custom"] = custom_skills



def _cg_cap():
    np_ = int(st.session_state.get("cg_np", 0) or 0)
    return 2 * np_

def _cg_sync_from_dodge():
    cap = _cg_cap()
    dodge = int(st.session_state.get("cg_dodge", 0) or 0)

    # Parry espelha Dodge (como no seu código original)
    st.session_state["cg_parry"] = dodge

    # Thg fecha o cap automaticamente
    st.session_state["cg_thg"] = max(0, min(99, cap - dodge))

def _cg_sync_from_fortitude():
    cap = _cg_cap()
    fort = int(st.session_state.get("cg_fortitude", 0) or 0)

    st.session_state["cg_will"] = max(0, min(99, cap - fort))

def _cg_sync_from_np():
    # Quando NP muda, recalcula Thg e Will para manter cap
    _cg_sync_from_dodge()
    _cg_sync_from_fortitude()

def _cg_init_defenses_if_missing(dodge_base, fort_base):
    cap = _cg_cap()

    if "cg_dodge" not in st.session_state:
        st.session_state["cg_dodge"] = int(dodge_base)
    if "cg_parry" not in st.session_state:
        st.session_state["cg_parry"] = int(st.session_state["cg_dodge"])
    if "cg_fortitude" not in st.session_state:
        st.session_state["cg_fortitude"] = int(fort_base)
    if "cg_will" not in st.session_state:
        st.session_state["cg_will"] = max(0, min(99, cap - int(st.session_state["cg_fortitude"])))
    if "cg_thg" not in st.session_state:
        st.session_state["cg_thg"] = max(0, min(99, cap - int(st.session_state["cg_dodge"])))

def _normalize_hub_pid(pid_value) -> str:
    s = str(pid_value or "").strip()
    if not s:
        return ""
    if s.startswith("EXT:"):
        return s
    
    # Strip prefixos residuais PID:
    if s.startswith("PID:"):
        s = s.replace("PID:", "", 1).strip()
    
    # Aceita "283", "0283", "283.0", int, float etc.
    try:
        if re.fullmatch(r"\d+(\.0+)?", s):
            return str(int(float(s)))
    except Exception:
        pass
    
    return s


def normalize_sheet_pokemon_id(pid_value, pokemon_name: str = ""):
    """Normaliza o ID da ficha para o mesmo padrão usado nas fichas geradas no sistema.

    Regras:
    - IDs numéricos viram int (ex.: "0283", "283.0" -> 283)
    - IDs externos ficam como "EXT:<nome>"
    - placeholders legados (MM_/mm_/pdf_) tentam resolver pelo nome
    """
    raw = str(pid_value or "").strip()
    pname = str(pokemon_name or "").strip()

    # IDs legados/importados (ex.: MM_Pikachu, mm_pikachu, pdf_pikachu)
    legacy_placeholder = False
    if raw:
        low = raw.lower()
        legacy_placeholder = low.startswith("mm_") or low.startswith("pdf_")

    # Prioriza o valor informado quando ele já está em formato válido.
    if raw and not legacy_placeholder:
        if raw.startswith("EXT:"):
            ext_name = raw.replace("EXT:", "", 1).strip()
            return f"EXT:{ext_name}" if ext_name else (f"EXT:{pname}" if pname else "")

        if raw.startswith("PID:"):
            raw = raw.replace("PID:", "", 1).strip()

        try:
            if re.fullmatch(r"\d+(\.0+)?", raw):
                return int(float(raw))
        except Exception:
            pass

    # Resolve por nome para alinhar com fichas nativas (id numérico da Pokédex).
    if pname:
        try:
            resolved = resolve_pokemon_pid(df, pname)
            if resolved is not None and str(resolved).isdigit():
                return int(resolved)
        except Exception:
            pass
        return f"EXT:{pname}"

    # Último fallback: mantém texto limpo para não perder referência.
    return raw

# ----------------------------
# Helpers UX (Criação Guiada)
# ----------------------------
def _resolve_asset_path(fname: str) -> str:
    """Resolve caminhos para assets (excel, etc.) sem quebrar em Streamlit Cloud/local.

    Também tolera nomes com sufixos comuns (ex: "(2)") quando o arquivo foi enviado/baixado.
    """
    try:
        base = os.path.dirname(__file__)
    except Exception:
        base = os.getcwd()

    candidates = [
        fname,
        os.path.join(os.getcwd(), fname),
        os.path.join(base, fname),
        os.path.join(base, "assets", fname),
        os.path.join(os.getcwd(), "assets", fname),
        os.path.join(os.getcwd(), "data", fname),
        os.path.join(base, "data", fname),
    ]
    for c in candidates:
        if c and os.path.exists(c):
            return c

    # fallback inteligente: procura variações com sufixo (ex: "arquivo (5).xlsx")
    try:
        import glob
        root_dirs = [os.getcwd(), base, os.path.join(base, "assets"), os.path.join(os.getcwd(), "assets"),
                     os.path.join(base, "data"), os.path.join(os.getcwd(), "data")]
        stem, ext = os.path.splitext(fname)
        if ext:
            patterns = [f"{stem}*{ext}"]
        else:
            patterns = [f"{stem}*"]
        for rd in root_dirs:
            for pat in patterns:
                hits = sorted(glob.glob(os.path.join(rd, pat)))
                if hits:
                    return hits[0]
    except Exception:
        pass

    return fname  # fallback (deixa o erro explícito se não achar)
    
    
        
def _pokeapi_parse_move_names(pjson: dict) -> list[str]:
    out: list[str] = []
    for it in (pjson or {}).get("moves", []) or []:
        mv = (it or {}).get("move", {}) or {}
        nm = mv.get("name")
        if nm:
            out.append(str(nm))
    # remove duplicados preservando ordem
    seen = set()
    uniq = []
    for nm in out:
        if nm not in seen:
            seen.add(nm)
            uniq.append(nm)
    return uniq

def _try_match_move_in_db(db, api_name: str):
    """Tenta casar nomes do PokeAPI (geralmente com hífen) com os nomes do Excel (geralmente com espaço)."""
    if not api_name:
        return None

    raw = str(api_name).strip()
    cands = [
        raw,
        raw.replace("-", " "),
        raw.replace("-", " ").replace("’", "'"),
        raw.replace("-", ""),
    ]
    # algumas exceções comuns (seu excel costuma usar abreviações / pontuação diferente)
    # você pode ir adicionando aqui conforme aparecerem casos novos
    EXCEPTIONS = {
        "u turn": "u-turn",
        "v create": "v-create",
    }
    for k, v in EXCEPTIONS.items():
        if _norm(cands[1]) == _norm(k):
            cands.append(v)

    for c in cands:
        try:
            mv = db.get_by_name(c)
            if mv:
                return mv
        except Exception:
            continue

    # fallback: busca por prefixo / melhor candidato
    try:
        hits = db.search_by_name_prefix(raw.replace("-", " "))
        if hits:
            return hits[0]
    except Exception:
        pass

    return None

def _summarize_build(build_txt: str) -> list[str]:
    b = (build_txt or "").strip()
    if not b:
        return []
    low = b.lower()
    bullets: list[str] = []

    def _has(p: str) -> bool:
        return p in low

    if _has("damage"):
        bullets.append("🗡️ Dano (Damage)")
    if _has("affliction"):
        bullets.append("🌀 Status (Affliction)")
    if _has("weaken"):
        bullets.append("📉 Debuff (Weaken)")
    if _has("healing"):
        bullets.append("💚 Cura (Healing)")
    if _has("create"):
        bullets.append("🧱 Barreira/Criação (Create)")
    if _has("environment"):
        bullets.append("🌦️ Ambiente/Clima (Environment)")
    if _has("nullify"):
        bullets.append("🚫 Nullify")

    if _has("area"):
        # tenta capturar tipo de área
        m = re.search(r"\[Area:\s*([^\]]+)\]", b, flags=re.IGNORECASE)
        if m:
            bullets.append(f"🧨 Área: {m.group(1).strip()}")
        else:
            bullets.append("🧨 Área")
    if _has("perception"):
        bullets.append("🎯 Perception")
    if _has("ranged") or _has("extended range"):
        bullets.append("🏹 Ranged")

    if _has("linked"):
        bullets.append("🔗 Linked (multi-efeito)")

    # deixa no máximo 6 para não poluir
    return bullets[:6]

def _default_accuracy_from_raw(mv) -> int:
    raw = getattr(mv, "raw", {}) or {}
    try:
        return int(raw.get("Accuracy") or raw.get("Acerto") or raw.get("acerto") or 0)
    except Exception:
        return 0

def _cg_confirm_move(mv, rank: int, pp_override: int | None = None, accuracy: int | None = None) -> dict:
    # PP
    pp_auto = None
    try:
        tmp = mv.pp_cost(int(rank))
        if isinstance(tmp, tuple):
            pp_auto = tmp[0]
    except Exception:
        pp_auto = None
    pp = int(pp_override) if pp_override is not None else (int(pp_auto) if pp_auto is not None else None)

    # accuracy
    acc = int(accuracy) if accuracy is not None else _default_accuracy_from_raw(mv)

    return {
        "name": mv.name,
        "rank": int(rank),
        "build": mv.render_build(int(rank)),
        "pp_cost": pp,
        "accuracy": acc,
        "meta": {
            "ranged": bool(getattr(mv, "ranged", False)),
            "perception_area": bool(getattr(mv, "perception_area", False)),
            "category": str(getattr(mv, "categoria", "") or ""),
        },
    }


def _cg_confirm_move_with_engine(mv, rank: int, accuracy: int | None = None) -> dict:
    """Como _cg_confirm_move, mas usa parse_build_string para calcular PP com o motor MM3e."""
    build_str = mv.render_build(int(rank))
    pp: int | None = None
    if _GB_AVAILABLE and _gb_parse_build is not None:
        try:
            comps = _gb_parse_build(build_str, default_rank=int(rank))
            draft = _GolpeDraft(name=mv.name, components=comps)
            pp_raw, _ = _gb_calc_pp(draft)
            pp = int(pp_raw) if pp_raw is not None else None
        except Exception:
            pp = None
    if pp is None:
        try:
            tmp = mv.pp_cost(int(rank))
            pp = int(tmp[0]) if isinstance(tmp, tuple) else (int(tmp) if tmp is not None else None)
        except Exception:
            pp = None
    acc = int(accuracy) if accuracy is not None else _default_accuracy_from_raw(mv)
    return {
        "name": mv.name,
        "rank": int(rank),
        "build": build_str,
        "pp_cost": pp,
        "accuracy": acc,
        "meta": {
            "ranged": bool(getattr(mv, "ranged", False)),
            "perception_area": bool(getattr(mv, "perception_area", False)),
            "category": str(getattr(mv, "categoria", "") or ""),
            "engine_pp": True,
        },
    }


def _cg_recalculate_pp(move_data: dict, rank: int, db_moves: Optional["MoveDB"]) -> Tuple[Optional[int], Optional[str]]:
    mv_name = str(move_data.get("name") or "").strip()
    if db_moves and mv_name:
        try:
            mv = db_moves.get_by_name(mv_name)
        except Exception:
            mv = None
        if mv:
            try:
                pp_auto, why = mv.pp_cost(int(rank))
                if pp_auto is not None:
                    return int(pp_auto), why
            except Exception:
                pass

    tmp_build = str(move_data.get("build") or "").strip()
    if tmp_build:
        tmp_meta = move_data.get("meta") or {}
        tmp_mv = Move(
            name=(mv_name or "Poder Personalizado"),
            tipo="—",
            categoria=str(tmp_meta.get("category") or ""),
            descricao="",
            build=tmp_build,
            how_it_works="",
            resist_stat="",
            ranged=bool(tmp_meta.get("ranged", False)),
            perception_area=bool(tmp_meta.get("perception_area", False)),
            tags=[],
            raw={},
        )
        try:
            pp_auto, why = tmp_mv.pp_cost(int(rank))
            if pp_auto is not None:
                return int(pp_auto), why
        except Exception:
            pass

    return None, None



import math
from io import BytesIO


REGION_ALIASES = {
    "alola": "alola", "alolan": "alola", "a": "alola",
    "galar": "galar", "galarian": "galar", "g": "galar",
    "hisui": "hisui", "hisuian": "hisui", "h": "hisui",
    "paldea": "paldea", "paldean": "paldea", "p": "paldea",
}

def to_pokeapi_name(user_text: str) -> str:
    s = (user_text or "").strip().lower()

    # símbolos comuns
    s = s.replace("♀", " f").replace("♂", " m")

    # troca espaços/underscore por hífen
    s = re.sub(r"[\s_]+", "-", s)

    # remove duplo hífen
    s = re.sub(r"-{2,}", "-", s).strip("-")

    # nidoran: vira nidoran-f / nidoran-m
    if s in ("nidoran", "nidoran-"):
        return "nidoran"  # deixa ambíguo e você força escolha na UI
    if s in ("nidoran-f", "nidoranf", "nidoran-female", "nidoran-fem", "nidoran-f."):
        return "nidoran-f"
    if s in ("nidoran-m", "nidoranm", "nidoran-male", "nidoran-masc", "nidoran-m."):
        return "nidoran-m"

    # formatos tipo "sandslash-a" / "weezing-g" / "g-weezing"
    if re.match(r"^[aghp]-", s):  # g-weezing
        tag, base = s.split("-", 1)
        region = REGION_ALIASES.get(tag)
        if region and base:
            return f"{base}-{region}"

    m = re.match(r"^(.+)-([aghp])$", s)  # sandslash-a
    if m:
        base, tag = m.group(1), m.group(2)
        region = REGION_ALIASES.get(tag)
        if region:
            return f"{base}-{region}"

    # formatos tipo "sandslash-alolan" / "weezing-galarian"
    parts = s.split("-")
    if len(parts) >= 2:
        last = parts[-1]
        region = REGION_ALIASES.get(last)
        if region:
            base = "-".join(parts[:-1])
            return f"{base}-{region}"

    return s


POKEAPI_BASE = "https://pokeapi.co/api/v2"

@st.cache_data(ttl=60*60)
def pokeapi_get_pokemon(name_or_id: str) -> dict:
    q = to_pokeapi_name(name_or_id)
    url = f"{POKEAPI_BASE}/pokemon/{q}"
    r = requests.get(url, timeout=15)
    r.raise_for_status()
    return r.json()


def pokeapi_parse_stats(p: dict) -> dict:
    # base stats: hp, attack, defense, special-attack, special-defense, speed
    out = {}
    for s in p.get("stats", []):
        out[s["stat"]["name"]] = int(s["base_stat"])
    return out

def pokeapi_parse_types(p: dict) -> list[str]:
    return [t["type"]["name"] for t in p.get("types", [])]

def pokeapi_parse_abilities(p: dict) -> list[str]:
    # retorna nomes das abilities
    return [a["ability"]["name"] for a in p.get("abilities", [])]

def get_np_for_pokemon(df_pokedex: pd.DataFrame, pid: str, fallback_np: int = 6) -> int:
    """
    Tenta achar NP/PL no seu DF da pokedex.
    Se não achar coluna, retorna fallback.
    """
    pid = str(pid)
    row = df_pokedex[df_pokedex["Nº"].astype(str) == pid]
    if row.empty:
        return fallback_np

    # tenta colunas comuns
    for col in ["NP", "PL", "Nivel de Poder", "Nível de Poder", "Power Level"]:
        if col in row.columns:
            try:
                return int(row.iloc[0][col])
            except:
                pass
    return fallback_np

def calc_pp_budget(np_: int) -> int:
    # sua regra: NP x 2 = PP
    return int(np_) * 15

def can_add_more_attack_points(np_: int, spent_attack_points: int) -> bool:
    # trava: quando atingir limite de 20 pontos a mais do NP (você descreveu assim)
    # ✅ limite = NP + 20
    return spent_attack_points < (int(np_) + 20)

def upload_pdf_to_bucket(bucket, pdf_bytes: bytes, dest_path: str) -> str:
    """
    Faz upload no Firebase Storage (bucket do init_firebase).
    Retorna o caminho salvo.
    """
    blob = bucket.blob(dest_path)
    blob.upload_from_string(pdf_bytes, content_type="application/pdf")
    return dest_path


# Configuração da Página
st.set_page_config(
    page_title="Pokedex RPG Cloud",
    page_icon="🔒",
    layout="wide"
)
# ==========================================
# 🎨 ESTILO VISUAL GLOBAL (POKÉMON RETRÔ)
# ==========================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Press+Start+2P&family=Inter:wght@400;500;600;700&display=swap');

:root{
  --gba-bg: #0f172a;
  --gba-panel: #f8fafc;
  --gba-border: #334155;
  --gba-ink: #0f172a;
  --gba-ink2: #475569;
  --gba-accent: #38bdf8;
  --gba-accent-border: #0ea5e9;
}

/* =========================================================
   1) FUNDO (pega o “miolo” do Streamlit também)
   ========================================================= */
html, body{
  background: var(--gba-bg) !important;
}
.stApp{
  background:
    radial-gradient(circle at 20% 10%, rgba(255,255,255,0.06), transparent 35%),
    radial-gradient(circle at 90% 30%, rgba(255,255,255,0.04), transparent 40%),
    var(--gba-bg) !important;
}
div[data-testid="stAppViewContainer"]{{
  background: radial-gradient(ellipse at center, #0a0a0a 0%, #000 62%, #000 100%) !important;
}}
div[data-testid="stAppViewContainer"] > .main{{ background: transparent !important; }}
div[data-testid="stSidebar"]{{ background: #000 !important; }}
div[data-testid="stHeader"]{{ background: transparent !important; }}
/* =========================================================
   IFRAME / COMPONENTES (deixar fundo transparente)
   ========================================================= */
div[data-testid="stIFrame"],
div[data-testid="stIFrame"] > iframe {
  background: transparent !important;
  border: none !important;
}

/* Alguns builds do Streamlit usam wrappers adicionais */
iframe {
  background: transparent !important;
}

/* =========================================================
   IFRAME DO st_click_detector (Custom Component)
   ========================================================= */
div[data-testid="stCustomComponentV1"],
div[data-testid="stCustomComponentV1"] > iframe,
div[data-testid="stCustomComponentV1"] iframe {
  background: transparent !important;
  border: none !important;
}

/* fallback: se Streamlit usar outro testid */
div[data-testid="stCustomComponent"],
div[data-testid="stCustomComponent"] > iframe,
div[data-testid="stCustomComponent"] iframe {
  background: transparent !important;
  border: none !important;
}

/* fallback geral */
iframe {
  background: transparent !important;
}

/* =========================================================
   2) FONTE RETRÔ E CONTRASTE GLOBAL (FORÇADO)
   ========================================================= */
/* Aplica a fonte e a cor clara em TODOS os elementos possíveis */
.stApp, .stMarkdown, .stMarkdown p, .stMarkdown span, li,
h1, h2, h3, h4, label, .stWidgetLabel,
.stTextInput input, .stNumberInput input, .stTextArea textarea,
[data-baseweb="tab"], [data-baseweb="tab-list"],
.pokedex-info-value, .pokedex-info-title, .section-title {
  font-family: "Press Start 2P", cursive !important;
  font-size: 13px !important;
  line-height: 1.6 !important;
  color: #f8fafc !important; /* Branco fosco para máximo contraste */
  text-shadow: 1px 1px 0px #000; /* Sombra leve para destacar no fundo escuro */
}

/* Força o texto claro especificamente dentro das abas e containers verticais */
div[data-testid="stVerticalBlock"] div, 
div[role="tabpanel"] p, 
div[role="tabpanel"] span,
div[data-testid="stExpander"] p {
    color: #f8fafc !important;
}

/* Ajuste para que o texto digitado nos campos de busca também seja visível */
.stTextInput input {
    color: #ffffff !important;
    background-color: rgba(0, 0, 0, 0.5) !important;
}

/* =========================================================
   3) PROTEÇÃO CONTRA BUG DE ÍCONE (keyboard_arrow_right etc.)
   ========================================================= */
[data-testid="stExpander"] summary,
[data-testid="stExpander"] svg,
[data-testid="stHeader"] svg,
.stSelectbox svg,
.stMultiSelect svg,
div[data-baseweb="icon"],
span[class*="icon"],
i[class*="icon"]{
  font-family: sans-serif !important;
}

/* =========================================================
   4) Sprites pixelados
   ========================================================= */
img{ image-rendering: pixelated !important; image-rendering: crisp-edges !important; }

/* =========================================================
   5) “Janelas” GBA
   ========================================================= */
.gba-window{
  background: rgba(255,255,255,0.08);
  border: 2px solid rgba(148,163,184,0.35);
  border-radius: 16px;
  padding: 12px;
  box-shadow: 0 10px 28px rgba(0,0,0,0.25);
  margin-bottom: 14px;
}
.gba-window.party{ background: rgba(59,130,246,0.10); }
.gba-window.box{ background: rgba(34,197,94,0.08); }
.gba-window.summary{ background: rgba(255,255,255,0.08); }

/* Cabeçalho */
.gba-header{
  display:flex;
  justify-content:space-between;
  align-items:center;
  gap:10px;
  margin-bottom: 10px;
}
.gba-chip{
  display:inline-flex;
  align-items:center;
  gap:6px;
  white-space:nowrap;
  background: rgba(15,23,42,0.70);
  color: rgba(255,255,255,0.95);
  padding: 6px 10px;
  border-radius: 999px;
  font-weight: 900;
  font-size: 11px;
}
.gba-title{
  font-family: "Press Start 2P", cursive !important;
  color: rgba(255,255,255,0.92) !important;
  font-size: 12px !important;
  line-height: 1.2 !important;
}

/* =========================================================
   6) Botões (contraste garantido)
   ========================================================= */
div.stButton > button, div.stDownloadButton > button{
  background: var(--gba-accent) !important;
  color: #0b1220 !important;
  border: 2px solid var(--gba-accent-border) !important;
  border-radius: 12px !important;
  padding: 10px 12px !important;
  font-weight: 800 !important;
  box-shadow: 0 6px 12px rgba(0,0,0,0.18) !important;
}
div.stButton > button:hover{ filter: brightness(1.05); transform: translateY(-1px); }
div.stButton > button:active{ transform: translateY(0px); }

/* =========================================================
   7) Tabs (cartucho) + remove underline
   ========================================================= */
.stTabs [data-baseweb="tab-list"]{
  background: linear-gradient(135deg, rgba(15, 23, 42, 0.75), rgba(37, 99, 235, 0.35)) !important;
  border: 2px solid rgba(148, 163, 184, 0.35) !important;
  border-radius: 12px !important;
  padding: 6px !important;
  gap: 6px !important;
  border-bottom: 0 !important;
  box-shadow: none !important;
}
div[data-baseweb="tab-border"]{ display:none !important; }
.stTabs [data-baseweb="tab"]{
  background: rgba(30, 64, 175, 0.22) !important;
  color: rgba(255,255,255,0.92) !important;
  border: 2px solid rgba(59, 130, 246, 0.40) !important;
  border-radius: 10px !important;
  font-weight: 900 !important;
}
.stTabs [aria-selected="true"]{
  background: var(--gba-accent) !important;
  color: #0b1220 !important;
  border-color: var(--gba-accent-border) !important;
}

/* =========================================================
   8) BOX com “grama” (AGORA SIM)
   ========================================================= */
.grass-box{
  border: 2px solid rgba(34,197,94,0.75);
  border-radius: 16px;
  padding: 12px;
  position: relative;
  overflow: hidden;
  background:
    linear-gradient(180deg, rgba(16,185,129,0.20), rgba(34,197,94,0.10)),
    repeating-linear-gradient(
      135deg,
      rgba(34,197,94,0.16) 0px,
      rgba(34,197,94,0.16) 6px,
      rgba(22,163,74,0.12) 6px,
      rgba(22,163,74,0.12) 12px
    );
  box-shadow: 0 10px 28px rgba(0,0,0,0.28);
}
.grass-box::after{
  content:"";
  position:absolute; inset:0;
  background-image: radial-gradient(rgba(255,255,255,0.06) 1px, transparent 1px);
  background-size: 10px 10px;
  opacity: .55;
  pointer-events:none;
}

/* =========================================================
   9) PARTY com slots quadrados
   ========================================================= */
.team-box{
  border: 2px solid rgba(59,130,246,0.55);
  border-radius: 16px;
  padding: 12px;
  background: rgba(15,23,42,0.35);
  box-shadow: 0 10px 28px rgba(0,0,0,0.28);
}
.team-slots{
  display:grid;
  grid-template-columns: repeat(4, minmax(54px, 72px));
  gap: 10px;
  align-items:center;
}
.team-slot{
  aspect-ratio: 1 / 1;
  border-radius: 12px;
  border: 2px solid rgba(255,255,255,0.18);
  background: linear-gradient(180deg, rgba(255,255,255,0.09), rgba(255,255,255,0.03));
  box-shadow: inset 0 0 0 2px rgba(0,0,0,0.12);
  display:flex;
  align-items:center;
  justify-content:center;
  overflow:hidden;
}
.team-slot img{
  width: 85%;
  height: 85%;
  object-fit: contain;
  image-rendering: pixelated;
  filter: drop-shadow(0 4px 8px rgba(0,0,0,0.25));
}
.box-slot-grass {
  background: #55a64b; /* Cor base da grama */
  background-image: 
    linear-gradient(rgba(255,255,255,0.1) 1px, transparent 1px),
    linear-gradient(90deg, rgba(255,255,255,0.1) 1px, transparent 1px);
  background-size: 8px 8px; /* Efeito de pixels/quadriculado */
  border: 2px solid #2d5a27;
  border-radius: 8px;
  padding: 5px;
  display: flex;
  justify-content: center;
  align-items: center;
  margin-bottom: 5px;
  box-shadow: inset 0 0 10px rgba(0,0,0,0.3);
}
/* Fundo de grama individual para cada Pokémon na BOX */
.box-slot-grass {
  background: #55a64b;
  background-image: 
    linear-gradient(rgba(255,255,255,0.1) 1px, transparent 1px),
    linear-gradient(90deg, rgba(255,255,255,0.1) 1px, transparent 1px);
  background-size: 8px 8px;
  border: 2px solid #2d5a27;
  border-radius: 8px;
  padding: 5px;
  display: flex;
  justify-content: center;
  align-items: center;
  box-shadow: inset 0 0 10px rgba(0,0,0,0.3);
  margin-bottom: 5px;
}


/* Card estilo GBA para a Equipe Ativa */
.gba-party-slot {
  background: linear-gradient(180deg, #4d88ff 0%, #2e5cb8 100%);
  border: 3px solid #f8fafc;
  border-radius: 12px;
  padding: 12px;
  margin-bottom: 15px;
  box-shadow: 4px 4px 0px rgba(0,0,0,0.2);
  display: flex;
  flex-direction: column;
  align-items: center;
  position: relative;
  overflow: hidden;
}

/* Detalhe de luz no card da equipe */
.gba-party-slot::before {
  content: "";
  position: absolute;
  top: 0; left: 0; right: 0; height: 50%;
  background: rgba(255,255,255,0.1);
  pointer-events: none;
}
</style>
""", unsafe_allow_html=True)




# --- CONEXÃO COM GOOGLE SHEETS ---
def get_google_sheet():
    try:
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds_dict = st.secrets["gcp_service_account"]
        creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
        client = gspread.authorize(creds)
        sheet = client.open("SaveData_RPG").sheet1
        return sheet
    except Exception as e:
        st.error(f"Erro de Conexão: {e}")
        st.stop()

def pack_tiles(tiles: list[list[str]]) -> str:
    # list -> json -> gzip -> base64 (string)
    raw = json.dumps(tiles, separators=(",", ":")).encode("utf-8")
    gz = gzip.compress(raw, compresslevel=9)
    return base64.b64encode(gz).decode("ascii")

def unpack_tiles(packed: str) -> list[list[str]]:
    gz = base64.b64decode(packed.encode("ascii"))
    raw = gzip.decompress(gz).decode("utf-8")
    return json.loads(raw)


# --- SISTEMA DE LOGIN SEGURO (CORRIGIDO) ---

def find_user_row(sheet, name):
    """
    Procura o usuário APENAS na Coluna 1 (Coluna A).
    Isso evita bugs de achar o nome em outros lugares.
    Retorna o número da linha ou None.
    """
    try:
        # Pega todos os valores da primeira coluna
        all_names = sheet.col_values(1)
        # Tenta achar o índice na lista (Python começa em 0, Sheets em 1)
        # O index lança erro se não achar, então usamos try/except
        row_index = all_names.index(name) + 1 
        return row_index
    except ValueError:
        return None
        
def coord_to_notation(row, col):
    # Converte coluna em letra (0=A, 1=B...) e linha em número (0=1, 1=2...)
    letter = chr(65 + int(col)) 
    number = int(row) + 1
    return f"{letter}{number}"

def stop_pvp_sync_listener():
    sync_data = st.session_state.get("pvp_sync_listener")
    if not sync_data:
        return
    for unsubscribe in sync_data.get("unsubscribers", []):
        try:
            unsubscribe()
        except Exception:
            pass
    stop_event = sync_data.get("stop_event")
    event_queue = sync_data.get("queue")
    if stop_event:
        stop_event.set()
    if event_queue:
        event_queue.put({"tag": "stop"})
    st.session_state.pop("pvp_sync_listener", None)


def render_ds_tools_nav(selected_view: str):
    opts = ["Menu", "NPCs", "Ginásios", "Locais", "Sessões", "Sair"]

    to_view = {
        "Menu": "home",
        "NPCs": "npcs",
        "Ginásios": "ginasios",
        "Locais": "locais",
        "Sessões": "sessoes",
        "Sair": "sair",
    }
    from_view = {v: k for k, v in to_view.items()}

    default_label = from_view.get(selected_view, "NPCs")
    idx = opts.index(default_label)

    st.markdown("""
    <style>
      /* ====== RESGATE (caso você tenha CSS que esconde widgets) ====== */
      .ds-topnav div[data-testid="stRadio"],
      .ds-topnav div[data-testid="stRadio"] *{
        visibility: visible !important;
        opacity: 1 !important;
        display: revert !important;
      }
      /* garante que o radio não seja escondido pelo seu CSS global */
      .ds-topnav div[data-testid="stRadio"]{
        visibility: visible !important;
        opacity: 1 !important;
      }
        
      /* radiogroup horizontal e centralizado */
      .ds-topnav div[data-testid="stRadio"] [role="radiogroup"]{
        display: flex !important;
        justify-content: center !important;
        align-items: center !important;
        gap: 160px !important;
      }

      /* cada opção (label) */
      .ds-topnav div[data-testid="stRadio"] [role="radiogroup"] > label{
        position: relative;
        margin: 0 !important;
        padding: 8px 42px !important;
        cursor: pointer !important;
        display: flex !important;
        align-items: center !important;
        color: rgba(255,255,255,0.70) !important;
        background: transparent !important;
        box-shadow: none !important;
        text-transform: uppercase !important;
        letter-spacing: 0.24em !important;
        font-size: 12px !important;
        user-select: none !important;
        transition: all 120ms ease !important;
        z-index: 0 !important;
      }

      .ds-topnav div[data-testid="stRadio"] [role="radiogroup"] > label::before{
        content: none !important;
        display: none !important;
      }

      /* o texto do radio no Streamlit geralmente fica em p */
      .ds-topnav div[data-testid="stRadio"] label p{
        position: relative;
        padding-left: 0;
        margin: 0 !important;
        user-select: none;
        font-family: "DarkSouls", serif;
        font-size: 12px;
        letter-spacing: 0.24em;
        line-height: 1;
        text-transform: uppercase;
        color: rgba(255,255,255,0.70);
      }

      .ds-topnav div[data-testid="stRadio"] [role="radiogroup"] > label:hover{
        color: #FFD700 !important;
        text-shadow: 0 0 10px rgba(255,215,0,0.65) !important;
      }

      .ds-topnav div[data-testid="stRadio"] [role="radiogroup"] > label[data-checked="true"]{
        color: #FFD700 !important;
        text-shadow: 0 0 10px rgba(255,215,0,0.65) !important;
      }

      .ds-topnav div[data-testid="stRadio"] [role="radiogroup"] > label:hover::before,
      .ds-topnav div[data-testid="stRadio"] [role="radiogroup"] > label[data-checked="true"]::before{
        display: none !important;
      }

      /* esconde o circulinho padrão do Streamlit */
      .ds-topnav div[data-testid="stRadio"] [data-baseweb="radio"]{
        background: transparent !important;
        box-shadow: none !important;
        border: 0 !important;
        padding: 0 !important;
      }
      .ds-topnav div[data-testid="stRadio"] [data-baseweb="radio"] > div:first-child{
        display: none !important;
      }
      .ds-topnav div[data-testid="stRadio"] [data-baseweb="radio"] > div:last-child{
        padding-left: 0 !important;
      }
      .ds-topnav div[data-testid="stRadio"] input[type="radio"]{
        display: none !important;
      }
      .ds-topnav div[data-testid="stRadio"] [role="radiogroup"] > label > div:first-child{
        display: none !important;
      }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("<div class='ds-topnav'>", unsafe_allow_html=True)
    choice = st.radio(
        "",
        opts,
        index=idx,
        horizontal=True,
        key="comp_topnav_radio",
        label_visibility="collapsed",
    )
    st.markdown("</div>", unsafe_allow_html=True)

    new_view = to_view[choice]

    if new_view == "sair":
        st.session_state["nav_to"] = "Pokédex (Busca)"
        st.rerun()
    elif new_view != st.session_state.get("comp_view"):
        st.session_state["comp_view"] = new_view
        if new_view != "npcs":
            st.session_state["comp_selected_npc"] = None
        st.rerun()

def _resolve_badge_path_assets_insignias(gym_key: str, g: dict) -> str:
    """
    Procura badge em Assets/insignias.
    Regra:
      - ginásio "por cidade" -> badge_<cidade>.png
      - ginásio "por líder"  -> badge_<lider>.png
    Fallbacks: gym_key -> meta.cidade -> meta.lider
    """

    def _norm(s: str) -> str:
        try:
            return _stem_key(str(s))
        except Exception:
            x = str(s).strip().lower()
            x = re.sub(r"\s+", " ", x)
            return x.replace(" ", "_")

    def _try(name: str) -> str:
        if not name:
            return ""
        base = f"badge_{_norm(name)}.png"
        p = os.path.join("Assets", "insignias", base)
        return p if os.path.exists(p) else ""

    meta = (g.get("meta") or {})
    staff = (g.get("staff") or {})

    city = (meta.get("cidade") or meta.get("city") or g.get("city") or "").strip()
    lider = (meta.get("lider") or meta.get("líder") or meta.get("leader") or staff.get("lider") or "").strip()

    # 1) tenta pela chave atual (pode ser cidade ou líder)
    p = _try(gym_key)
    if p: return p

    # 2) tenta pela cidade
    p = _try(city)
    if p: return p

    # 3) tenta pelo líder
    p = _try(lider)
    if p: return p

    return ""

def render_compendium_ginasios() -> None:

    comp_data = comp_load()
    gyms: dict = (comp_data.get("gyms") or {})
    npcs: dict = (comp_data.get("npcs") or {})
    cities: dict = (comp_data.get("cities") or {})

    if not gyms:
        st.markdown(
            "<div class='ds-frame'><div class='ds-name'>GINÁSIOS</div><div class='ds-meta'>Nenhum ginásio encontrado.</div></div>",
            unsafe_allow_html=True,
        )
        return
    st.markdown(
        """
        <style>
          .ds-gym-shell{
            margin-top: 6px;
            padding: 0 24px 12px 24px;
            box-sizing: border-box;
          }
          html:has(.ds-gym-shell),
          body:has(.ds-gym-shell),
          div[data-testid="stAppViewContainer"]:has(.ds-gym-shell){
            overflow: hidden !important;
          }
          .ds-frame-marker{
            display: none;
          }
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-left),
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-center),
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right),
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-left),
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-center),
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right){
            background: rgba(0,0,0,0.55);
            border: 2px solid rgba(176,143,60,0.55);
            box-shadow: 0 0 45px rgba(0,0,0,0.9);
            border-radius: 12px;
            padding: 18px 18px 14px 18px !important;
            position: relative;
            box-sizing: border-box;
          }
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-left)::after,
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-center)::after,
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right)::after,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-left)::after,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-center)::after,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right)::after{
            content: "";
            position: absolute;
            top: 10px; left: 10px; right: 10px; bottom: 10px;
            border: 1px solid rgba(255,255,255,0.10);
            border-radius: 10px;
            pointer-events: none;
          }
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-left) .comp-divider,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-left) .comp-divider{
            margin: 14px 0 14px 0 !important;
          }
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-left) div[data-testid="stSelectbox"],
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-left) div[data-testid="stSelectbox"]{
            width: 100% !important;
          }
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-left) div[data-testid="stSelectbox"] > div,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-left) div[data-testid="stSelectbox"] > div{
            background: rgba(0,0,0,0.25) !important;
            border: 1px solid rgba(176,143,60,0.45) !important;
            border-radius: 12px !important;
            padding: 6px 10px !important;
            box-shadow: 0 0 18px rgba(255,215,0,0.06) !important;
          }
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-left) div[data-testid="stSelectbox"] *,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-left) div[data-testid="stSelectbox"] *{
            font-family: "DarkSouls", serif !important;
            letter-spacing: 0.18em !important;
            text-transform: uppercase !important;
            color: rgba(255,255,255,0.82) !important;
          }
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right),
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right){
            display: flex;
            flex-direction: column;
            height: 78vh;
            min-height: 0;
            overflow: hidden;
          }

          /* 🔧 O st.markdown cria um stMarkdownContainer intermediário.
             Ele precisa ocupar o “resto” do painel para a lore poder rolar. */
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right) div[data-testid="stMarkdownContainer"]:has(.ds-lore-scroll),
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right) div[data-testid="stMarkdownContainer"]:has(.ds-lore-scroll){
            flex: 1 1 auto;
            min-height: 0;
          }

          /* wrapper interno do markdown precisa ter altura */
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right) div[data-testid="stMarkdownContainer"]:has(.ds-lore-scroll) > div,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right) div[data-testid="stMarkdownContainer"]:has(.ds-lore-scroll) > div{
            height: 100%;
            min-height: 0;
            display: flex;
            flex-direction: column;
          }

            /* agora sim: a lore vira a área rolável (modo “bruto”, não depende do flex do Streamlit) */
            div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right) .ds-lore-scroll,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right) .ds-lore-scroll{
              max-height: calc(78vh - 210px);
              overflow-y: auto !important;
            
              padding-right: 8px;
              overscroll-behavior: contain;
              scrollbar-gutter: stable;
            }


          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right) .ds-lore-scroll::-webkit-scrollbar,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right) .ds-lore-scroll::-webkit-scrollbar{ width: 8px; }
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right) .ds-lore-scroll::-webkit-scrollbar-thumb,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right) .ds-lore-scroll::-webkit-scrollbar-thumb{
            background: rgba(255,215,0,0.18);
            border-radius: 10px;
          }
          /* 🔧 Streamlit coloca um wrapper interno dentro da coluna.
           Ele precisa virar flex e ter 100% de altura, senão a .ds-lore-scroll não ganha altura real. */
        div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right) > div,
        div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right) > div{
          display: flex !important;
          flex-direction: column !important;
          height: 100% !important;
          min-height: 0 !important;
        }

          
          div[data-testid="column"]:has(.ds-frame-marker.ds-gym-right) .ds-lore-scroll::-webkit-scrollbar-track,
          div[data-testid="stColumn"]:has(.ds-frame-marker.ds-gym-right) .ds-lore-scroll::-webkit-scrollbar-track{
            background: rgba(255,255,255,0.06);
          }
        </style>
        """,
        unsafe_allow_html=True,
    )

    def _norm(s: str) -> str:
        try:
            return _stem_key(str(s))
        except Exception:
            x = str(s).strip().lower()
            x = re.sub(r"\s+", " ", x)
            return x.replace(" ", "_")

    def _pick(d: dict, *keys: str) -> str:
        for k in keys:
            v = d.get(k)
            if isinstance(v, str) and v.strip():
                return v.strip()
        return ""

    @st.cache_data(show_spinner=False)
    def _img_data_uri(path: str, max_w: int = 900) -> str:
        try:
            import base64, io
            from PIL import Image
            if not path or not os.path.exists(path):
                return ""
            img = Image.open(path).convert("RGBA")
            img.thumbnail((max_w, 2400))
            buf = io.BytesIO()
            img.save(buf, format="PNG", optimize=True)
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            return f"data:image/png;base64,{b64}"
        except Exception:
            return ""

    def _npc_obj(name: str, g: dict) -> dict:
        staff_npcs = g.get("staff_npcs") or {}
        return (npcs.get(name) or staff_npcs.get(name) or {})

    def _gym_staff(g: dict) -> tuple[str, str]:
        meta = g.get("meta") or {}
        staff = g.get("staff") or {}
        lider = (meta.get("lider") or staff.get("lider") or "").strip()
        vice  = (meta.get("vice_lider") or meta.get("vice-lider") or staff.get("vice_lider") or staff.get("vice") or "").strip()
        return lider, vice

    def _collect_gym_pokemons(lider_nm: str, vice_nm: str, g: dict) -> list[str]:
        seen = set()
        out: list[str] = []

        def add_from(nm: str):
            if not nm:
                return
            npc = _npc_obj(nm, g)
            pokes = npc.get("pokemons") or npc.get("pokemons_conhecidos") or []
            if not isinstance(pokes, list):
                return
            for p in pokes:
                ps = str(p).strip()
                if not ps:
                    continue
                key = _norm(ps)
                if key in seen:
                    continue
                seen.add(key)
                out.append(ps)

        add_from(lider_nm)
        add_from(vice_nm)
        return out

    def _section_html(title: str, text) -> str:
        # retorna HTML (string) para ser renderizado de uma vez só dentro do scroll
        if text is None:
            return ""
        txt = str(text).strip()
        if not txt:
            return ""
        safe_title = html.escape(str(title).strip())
        safe_body  = html.escape(txt).replace("\n", "<br>")
        return (
            f"<div class='ds-subtitle'>{safe_title}</div>"
            f"<div class='ds-history'>{safe_body}</div>"
            f"<div class='comp-divider'></div>"
        )

    # ---- lista de ginásios (chaves do seu bundle gyms) ----
    gym_keys = list(gyms.keys())

    def _sort_key(k: str):
        g = gyms.get(k) or {}
        meta = g.get("meta") or {}
        city = (meta.get("cidade") or meta.get("city") or "").strip() or k
        reg = ""
        cobj = cities.get(city) or {}
        if isinstance(cobj, dict):
            reg = (cobj.get("region") or "").strip()
        return (_norm(reg), _norm(city), _norm(k))

    gym_keys = sorted(gym_keys, key=_sort_key)

    # estado
    st.session_state.setdefault("comp_gym_key", gym_keys[0] if gym_keys else "")
    if st.session_state["comp_gym_key"] not in gyms:
        st.session_state["comp_gym_key"] = gym_keys[0] if gym_keys else ""

    st.session_state.setdefault("comp_gym_focus", "__visao__")  # "__visao__" | "lider" | "vice"

    gym_now = st.session_state["comp_gym_key"]
    focus_now = st.session_state["comp_gym_focus"]
    g = gyms.get(gym_now) or {}
    meta = g.get("meta") or {}
    lider_nm, vice_nm = _gym_staff(g)
    
    # --- EX-LÍDERES (não deixa quebrar + tenta puxar do próprio ginásio e do JSON de NPCs) ---
    staff = g.get("staff") or {}
    raw_ex = (
        meta.get("ex_lideres")
        or meta.get("ex_lider")
        or staff.get("ex_lideres")
        or staff.get("ex_lider")
        or []
    )
    
    # normaliza formatos
    if isinstance(raw_ex, str):
        ex_list = [s.strip() for s in re.split(r"[;,/|\n]+", raw_ex) if s.strip()]
    elif isinstance(raw_ex, list):
        ex_list = [str(s).strip() for s in raw_ex if str(s).strip()]
    else:
        ex_list = []
    
    # fallback: varre NPCs procurando "Ex-Líder de Ginásio de <cidade/ginásio>"
    if not ex_list:
        city = (meta.get("cidade") or meta.get("city") or "").strip() or gym_now
        targets = {_norm(city), _norm(gym_now)}
        for nm, npc in (npcs or {}).items():
            if not isinstance(npc, dict):
                continue
            occ = str(npc.get("ocupacao") or "")
            secs = npc.get("sections") or {}
            sec_gym = ""
            if isinstance(secs, dict):
                sec_gym = str(secs.get("Ginásio") or secs.get("Ginasio") or "")
            blob = f"{occ}\n{sec_gym}"
    
            m = re.search(r"Ex-?L[ií]der de Gin[aá]sio de\s*(.+)", blob, flags=re.IGNORECASE)
            if m:
                place = m.group(1).strip()
                if _norm(place) in targets:
                    ex_list.append(nm)
    
    # limpa duplicatas e remove líder/vice atuais
    seen = set()
    clean = []
    for nm in ex_list:
        if not nm:
            continue
        nk = _norm(nm)
        if nk in seen:
            continue
        if nk in {_norm(lider_nm), _norm(vice_nm)}:
            continue
        seen.add(nk)
        clean.append(nm)
    ex_list = clean
    
    # 3 colunas igual Locais
    st.markdown("<div class='ds-gym-shell'>", unsafe_allow_html=True)

    col_left, col_center, col_right = st.columns([1.05, 1.35, 2.15], gap="large")

    # ============================
    # ESQUERDA: seletor + pokémons
    # ============================
    with col_left:
        st.markdown("<div class='ds-frame-marker ds-gym-left'></div>", unsafe_allow_html=True)
        st.markdown("<div class='ds-name'>GINÁSIOS</div>", unsafe_allow_html=True)

        # label: se a chave for cidade -> mostra só cidade; senão "Líder — Cidade"
        label_map = {}
        for k in gym_keys:
            gg = gyms.get(k) or {}
            mm = gg.get("meta") or {}
            city = (mm.get("cidade") or mm.get("city") or "").strip() or k
            if _norm(k) == _norm(city):
                label = f"{city}"
            else:
                label = f"{k} — {city}"
            label_map[label] = k

        labels = list(label_map.keys())
        cur_label = next((lbl for lbl, kk in label_map.items() if kk == gym_now), labels[0] if labels else "")

        st.markdown("<div class='ds-meta'>GINÁSIO</div>", unsafe_allow_html=True)
        pick = st.selectbox(
            "Ginásio",
            labels,
            index=labels.index(cur_label) if cur_label in labels else 0,
            key="comp_gym_pick",
            label_visibility="collapsed",
        )
        if pick and label_map.get(pick) != st.session_state["comp_gym_key"]:
            st.session_state["comp_gym_key"] = label_map[pick]
            st.session_state["comp_gym_focus"] = "__visao__"
            st.rerun()
        
        st.markdown("<div class='comp-divider'></div>", unsafe_allow_html=True)
        
        # ---- FOCO ----
        focus_map = {
            "Visão": "__visao__",
            "Líder": "lider",
            "Vice-líder": "vice",
            **{f"Ex-líder: {nm}": f"ex::{nm}" for nm in ex_list},
        }
        focus_options = list(focus_map.keys())
        reverse_focus_map = {v: k for k, v in focus_map.items()}
        
        current_label = reverse_focus_map.get(focus_now, "Visão")
        
        st.markdown("<div class='ds-meta'>FOCO</div>", unsafe_allow_html=True)
        focus_label = st.selectbox(
            "Foco",
            options=focus_options,
            index=focus_options.index(current_label) if current_label in focus_options else 0,
            key="comp_gym_focus_sel",
            label_visibility="collapsed",
        )
        
        new_focus = focus_map.get(focus_label, "__visao__")
        if new_focus != st.session_state.get("comp_gym_focus"):
            st.session_state["comp_gym_focus"] = new_focus
            st.rerun()

        
        st.markdown("<div class='comp-divider'></div>", unsafe_allow_html=True)


        pokes = _collect_gym_pokemons(lider_nm, vice_nm, g)
        st.markdown("<div class='ds-subtitle'>Pokémons disponíveis</div>", unsafe_allow_html=True)

        if not pokes:
            st.markdown("<div class='ds-history'>(Nenhum pokémon cadastrado)</div>", unsafe_allow_html=True)
        else:
            try:
                cols = st.columns(4, gap="small")
                for i, pnm in enumerate(pokes[:48]):
                    with cols[i % 4]:
                        spr = ""
                        try:
                            spr = _poke_sprite_cached(pnm)
                        except Exception:
                            spr = ""
                        if spr:
                            st.image(spr, width=56)
                        st.caption(pnm)
            except Exception:
                st.markdown("<div class='ds-history'>" + "<br>".join(pokes[:48]) + "</div>", unsafe_allow_html=True)

    # ============================
    # CENTRO: badge + retratos
    # ============================
    with col_center:
        st.markdown("<div class='ds-frame-marker ds-gym-center'></div>", unsafe_allow_html=True)
        city = (meta.get("cidade") or meta.get("city") or "").strip()
        title = city if city else gym_now
        st.markdown(f"<div class='ds-name'>{title}</div>", unsafe_allow_html=True)

        chips = []
        if meta.get("tipo"): chips.append(f"Tipo: {meta.get('tipo')}")
        if meta.get("status"): chips.append(f"Status: {meta.get('status')}")
        if meta.get("localizacao"): chips.append(f"Local: {meta.get('localizacao')}")
        st.markdown(f"<div class='ds-meta'>{' | '.join(chips) if chips else ''}</div>", unsafe_allow_html=True)

        st.markdown("<div class='comp-divider'></div>", unsafe_allow_html=True)

        # BADGE (Assets/insignias)
        badge_path = _resolve_badge_path_assets_insignias(gym_now, g)
        if badge_path:
            badge_uri = _img_data_uri(badge_path, max_w=520)
            st.markdown(
                f"<div style='display:flex;justify-content:center;margin:10px 0 14px 0;'><img src='{badge_uri}' style='max-width:260px;width:72%;height:auto;opacity:0.95;'/></div>",
                unsafe_allow_html=True,
            )
        else:
            st.markdown("<div class='ds-history'>(Sem insígnia encontrada em Assets/insignias)</div>", unsafe_allow_html=True)

        def _portrait(role: str, nm: str):
            if not nm:
                st.markdown(f"<div class='ds-history'>({role}: não definido)</div>", unsafe_allow_html=True)
                return
        
            p = comp_find_image(nm)
            if p and os.path.exists(p):
                uri = _img_data_uri(p, max_w=520)
        
                st.markdown(f"<div class='ds-subtitle'>{role}</div>", unsafe_allow_html=True)
        
                st.markdown(
                    f"<div style='display:flex;justify-content:center;'>"
                    f"<img src='{uri}' style='max-width:340px;width:100%;height:auto;border-radius:10px;opacity:0.98;'/>"
                    f"</div>",
                    unsafe_allow_html=True,
                )
        
                st.markdown(
                    f"<div class='ds-meta' style='text-align:center;margin-top:8px;'>{nm}</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(f"<div class='ds-subtitle'>{role}</div>", unsafe_allow_html=True)
                st.markdown(f"<div class='ds-history'>{nm} (sem retrato)</div>", unsafe_allow_html=True)


        _portrait("Líder", lider_nm)
        _portrait("Vice-líder", vice_nm)

    # ============================
    # DIREITA: lore
    # ============================
    with col_right:
        st.markdown("<div class='ds-frame-marker ds-gym-right'></div>", unsafe_allow_html=True)
        st.markdown("<div class='ds-name'>LORE</div>", unsafe_allow_html=True)
    
        if isinstance(focus_now, str) and focus_now.startswith("ex::"):
            ex_nm = focus_now.split("::", 1)[1].strip()
            crumb = f"{title} — Ex-líder ({ex_nm})"
        else:
            crumb = title + (" — Visão" if focus_now == "__visao__" else (" — Líder" if focus_now == "lider" else " — Vice-líder"))
        st.markdown(f"<div class='ds-meta'>{crumb}</div>", unsafe_allow_html=True)
        st.markdown("<div class='comp-divider'></div>", unsafe_allow_html=True)

        active_sid = st.session_state.get("comp_session_active_id")
        if active_sid and gym_now:
            if st.button("➕ Adicionar à sessão atual", key=f"add_gym_session_{_stem_key(gym_now)}"):
                if add_entity_to_active_session(
                    "gyms",
                    gym_now,
                    "travel",
                    f"Ginásio visitado: {gym_now}",
                    {"gym": gym_now},
                ):
                    st.success(f"{gym_now} adicionado à sessão {active_sid}.")
    
        # Lore com scroll: render em 1 markdown só (não pode “abrir div” num markdown e renderizar depois)
        # ... você já está dentro do with col_right:

        lore_html_parts = []
        
        def _push_gym_section(title_txt: str, body_txt: str):
            if not isinstance(body_txt, str) or not body_txt.strip():
                return
            ttl = html.escape(str(title_txt))
            body = html.escape(body_txt.strip()).replace("\n", "<br>")
            lore_html_parts.append(f"<div class='ds-subtitle'>{ttl}</div>")
            lore_html_parts.append(f"<div class='ds-history'>{body}</div>")
            lore_html_parts.append("<div class='comp-divider'></div>")
        
        # --------- AQUI SÓ MONTA (APPEND) ---------
        if focus_now == "__visao__":
            narrative = (g.get("narrative") or "").strip()
            if narrative:
                _push_gym_section("Ginásio", narrative)
        
            if lider_nm:
                npcL = _npc_obj(lider_nm, g)
                secs = npcL.get("sections") or {}
                if isinstance(secs, dict):
                    _push_gym_section(f"Líder — {lider_nm}", _pick(secs, "Ginásio", "Historia", "História"))
        
            if vice_nm:
                npcV = _npc_obj(vice_nm, g)
                secs = npcV.get("sections") or {}
                if isinstance(secs, dict):
                    _push_gym_section(f"Vice-líder — {vice_nm}", _pick(secs, "Ginásio", "Historia", "História"))
        
        elif focus_now == "lider":
            npcL = _npc_obj(lider_nm, g) if lider_nm else {}
            secs = npcL.get("sections") or {}
        
            if not lider_nm:
                lore_html_parts.append("<div class='ds-history'>(Líder não definido)</div>")
            elif not isinstance(secs, dict) or not secs:
                lore_html_parts.append("<div class='ds-history'>(Sem seções para o líder)</div>")
            else:
                _push_gym_section("Ginásio", _pick(secs, "Ginásio"))
                _push_gym_section("História", _pick(secs, "História", "Historia", "Histórico", "Historico"))
                for k, v in secs.items():
                    if k in ("Ginásio", "História", "Historia", "Histórico", "Historico"):
                        continue
                    _push_gym_section(k, v)
        
        elif isinstance(focus_now, str) and focus_now.startswith("ex::"):
            ex_nm = focus_now.split("::", 1)[1].strip()
            npcX = _npc_obj(ex_nm, g) if ex_nm else {}
            secs = npcX.get("sections") or {}
        
            if not ex_nm:
                lore_html_parts.append("<div class='ds-history'>(Ex-líder não definido)</div>")
            elif not isinstance(secs, dict) or not secs:
                lore_html_parts.append("<div class='ds-history'>(Sem seções para o ex-líder)</div>")
            else:
                _push_gym_section("Ginásio", _pick(secs, "Ginásio"))
                _push_gym_section("História", _pick(secs, "História", "Historia", "Histórico", "Historico"))
                for k, v in secs.items():
                    if k in ("Ginásio", "História", "Historia", "Histórico", "Historico"):
                        continue
                    _push_gym_section(k, v)
        
        else:  # vice
            npcV = _npc_obj(vice_nm, g) if vice_nm else {}
            secs = npcV.get("sections") or {}
        
            if not vice_nm:
                lore_html_parts.append("<div class='ds-history'>(Vice-líder não definido)</div>")
            elif not isinstance(secs, dict) or not secs:
                lore_html_parts.append("<div class='ds-history'>(Sem seções para o vice-líder)</div>")
            else:
                _push_gym_section("Ginásio", _pick(secs, "Ginásio"))
                _push_gym_section("História", _pick(secs, "História", "Historia", "Histórico", "Historico"))
                for k, v in secs.items():
                    if k in ("Ginásio", "História", "Historia", "Histórico", "Historico"):
                        continue
                    _push_gym_section(k, v)
        
        # --------- AQUI SÓ RENDERIZA (UMA VEZ) ---------
        lore_html = "".join(lore_html_parts) or "<div class='ds-history'>(Sem lore cadastrada)</div>"
        st.markdown(f"<div class='ds-lore-scroll'>{lore_html}</div>", unsafe_allow_html=True)
        



def render_compendium_page() -> None:
    if "comp_view" not in st.session_state:
        st.session_state["comp_view"] = "home"

    # garante view inicial
    st.session_state.setdefault("comp_view", "home")
    
    # nav: aparece em TODAS as views do compendium, exceto home
    if st.session_state["comp_view"] != "home":
        render_ds_tools_nav(st.session_state["comp_view"])


    if st.session_state["comp_view"] == "home":
        render_compendium_home()
        return

    if st.session_state["comp_view"] == "npcs":
        render_compendium_npcs()
        return

    if st.session_state["comp_view"] == "ginasios":
        render_compendium_ginasios()
        return

    if st.session_state["comp_view"] == "locais":
        render_compendium_locais()
        return




def ensure_pvp_sync_listener(db, rid):
    """
    Inicia listeners do Firestore (Real-Time) para sincronizar o jogo 
    automaticamente APENAS quando houver interação de qualquer jogador.
    """
    if not rid:
        return
    
    # Se já existe um listener ativo para essa sala, não recria
    active = st.session_state.get("pvp_sync_listener")
    if active and active.get("rid") == rid:
        return

    # Se trocou de sala, mata o anterior
    stop_pvp_sync_listener()

    # Prepara filas e eventos para thread
    event_queue: queue.Queue = queue.Queue()
    stop_event = threading.Event()
    
    # Pega o contexto da sessão atual do usuário ANTES de entrar na thread
    ctx = get_script_run_ctx() 

    # Função auxiliar para colocar na fila de processamento
    def enqueue(tag: str, source: str):
        if stop_event.is_set():
            return
        # Coloca na fila para a thread principal processar
        event_queue.put({"tag": tag, "source": source, "ts": time.time()})

    # --- CALLBACKS DO FIRESTORE (Executam quando o banco muda) ---
    
    # 1. OUVINTE DE MAPA/PEÇAS
    def on_state_snapshot(doc_snapshot, changes, read_time):
        enqueue("state_change", "map")

    # 2. OUVINTE DE HP/STATUS (Party)
    def on_party_snapshot(col_snapshot, changes, read_time):
        enqueue("state_change", "party")

    # 3. OUVINTE DE EVENTOS (Dados, Logs, Ataques)
    def on_events_snapshot(col_snapshot, changes, read_time):
        for change in changes:
            if change.type.name == "ADDED": # Só queremos novos eventos
                enqueue("new_event", "log")

    # --- WORKER: A Thread que força o RERUN do Streamlit ---
    # --- WORKER: A Thread que força o RERUN do Streamlit ---
    def rerun_worker():
        if ctx:
            add_script_run_ctx(threading.current_thread(), ctx)
    
        last_req = 0.0
        min_gap = 0.06  # 60ms: bem responsivo e ainda evita tremedeira
    
        while not stop_event.is_set():
            try:
                item = event_queue.get(timeout=3)
                if item.get("tag") == "stop":
                    break
    
                # Drena rapidamente a fila pra juntar bursts em 1 rerun só
                drain_until = time.time() + 0.02
                while time.time() < drain_until:
                    try:
                        nxt = event_queue.get_nowait()
                        if nxt.get("tag") == "stop":
                            stop_event.set()
                            break
                    except queue.Empty:
                        break
    
                if ctx and ctx.script_requests:
                    now = time.time()
                    if (now - last_req) < min_gap:
                        time.sleep(min_gap - (now - last_req))
    
                    rerun_data = scriptrunner.RerunData(
                        query_string="",  # <<< limpa params da URL pra não “teleportar” de page
                        page_script_hash=ctx.page_script_hash,
                        cached_message_hashes=ctx.cached_message_hashes,
                    )

                    ctx.script_requests.request_rerun(rerun_data)
                    last_req = time.time()
                else:
                    break
    
            except queue.Empty:
                continue
            except Exception as e:
                print(f"Erro no worker de sync: {e}")
                time.sleep(0.2)


    # --- REGISTRO DOS LISTENERS NO FIREBASE ---
    try:
        # Listener 1: Documento de Estado (Mapa)
        state_unsub = db.collection("rooms").document(rid).collection("public_state").document("state").on_snapshot(on_state_snapshot)
        
        # Listener 2: Coleção de Eventos (Últimos 5 eventos)
        events_query = db.collection("rooms").document(rid).collection("public_events").order_by("ts", direction=firestore.Query.DESCENDING).limit(5)
        events_unsub = events_query.on_snapshot(on_events_snapshot)

        # Listener 3: Documento de Party (HP/Status)
        party_unsub = db.collection("rooms").document(rid).collection("public_state").document("party_states").on_snapshot(on_party_snapshot)

        # Inicia a Thread Worker
        worker = threading.Thread(target=rerun_worker, daemon=True)
        worker.start()

        # Salva na sessão para podermos cancelar depois
        st.session_state["pvp_sync_listener"] = {
            "rid": rid,
            "queue": event_queue,
            "stop_event": stop_event,
            "unsubscribers": [state_unsub, events_unsub, party_unsub], # Guarda funções para desligar
        }
    except Exception as e:
        st.error(f"Erro ao conectar no sync: {e}")


def render_public_log_fragment(db, rid, *, title: str | None = "📜 Log de Batalha (Tempo Real)", height: int = 280, show_divider: bool = True, limit: int = 25):
    """
    Renderiza um log enxuto baseado na coleção public_events.
    - title=None => não imprime título
    - show_divider=False => não imprime separador
    """
    if show_divider:
        st.markdown("---")
    if title:
        st.subheader(title)

    try:
        events = list_public_events(db, rid, limit=limit) or []
    except Exception:
        events = []

    with st.container(height=height):
        if not events:
            st.caption("Aguardando ações na arena...")
            return

        for ev in events:
            et = ev.get("type", "?")
            by = ev.get("by", "?")
            pl = ev.get("payload", {}) or {}
            ts = ev.get("ts")

            # Timestamp (quando existir)
            ts_txt = ""
            try:
                if ts:
                    # ts pode vir como float/epoch ou string; mantemos simples
                    if isinstance(ts, (int, float)):
                        ts_txt = datetime.fromtimestamp(ts).strftime("%H:%M:%S") + " · "
                    else:
                        ts_txt = ""
            except Exception:
                ts_txt = ""

            # Resolve nome do Pokémon quando possível
            p_id = pl.get("pid")
            p_name = by
            if p_id:
                try:
                    # get_poke_display_name é definido na view de battle e fica no escopo do módulo
                    p_name = get_poke_display_name(p_id)
                except Exception:
                    p_name = str(p_id)

            if et == "move":
                f_coord = coord_to_notation(*pl.get("from", [0, 0]))
                t_coord = coord_to_notation(*pl.get("to", [0, 0]))
                st.write(f"{ts_txt}👣 **{p_name}** se moveu de **({f_coord})** para **({t_coord})**")

            elif et == "dice":
                st.write(f"{ts_txt}🎲 **{by}** rolou d{pl.get('sides')}: **{pl.get('result')}**")

            elif et == "hit_confirmed":
                st.success(f"{ts_txt}⚔️ **{p_name}** ({by}) ACERTOU o ataque!")

            elif et == "missed":
                st.error(f"{ts_txt}🛡️ **{p_name}** ({by}) ERROU o ataque!")

            elif et == "finished":
                log_msg = ""
                logs = pl.get("logs")
                if isinstance(logs, list) and logs:
                    log_msg = str(logs[0])
                elif isinstance(logs, str):
                    log_msg = logs
                if log_msg:
                    st.info(f"{ts_txt}🩸 {log_msg}")
                else:
                    st.info(f"{ts_txt}🩸 Ação finalizada.")

            elif et == "effect":
                ic = pl.get("icon", "✨")
                t_coord = coord_to_notation(*pl.get("to", [0, 0]))
                st.write(f"{ts_txt}{ic} **{by}** aplicou efeito em **({t_coord})**")

            elif et == "effect_removed":
                t_coord = coord_to_notation(*pl.get("to", [0, 0]))
                st.write(f"{ts_txt}🧽 **{by}** removeu efeito em **({t_coord})**")

            elif et == "piece_placed":
                t_coord = coord_to_notation(*pl.get("to", [0, 0]))
                st.write(f"{ts_txt}📍 **{by}** colocou **{p_name}** em **({t_coord})**")

            elif et == "pokemon_removed":
                st.write(f"{ts_txt}❌ **{by}** removeu **{p_name}** do mapa")

            elif et == "trainer_placed":
                t_coord = coord_to_notation(*pl.get("to", [0, 0]))
                st.write(f"{ts_txt}🧍 **{by}** posicionou o avatar em **({t_coord})**")

            else:
                # fallback
                st.write(f"{ts_txt}🔹 **{by}** ({et}): {pl}")


import os, json, math
from PIL import Image
import streamlit as st

FLOOR_PREFIXES = ("agua", "areia", "grama", "pedra", "terra", "slope", "neve", "cave")

def _is_blueish(r, g, b):
    # "água" tende a ter B alto e G >= R
    return (b > 90) and (b > r + 20) and (b > g + 10)

def _brightness(r, g, b):
    return (0.2126*r + 0.7152*g + 0.0722*b)

def _tile_stats(tile: Image.Image):
    # retorna métricas úteis p/ classificar
    px = tile.convert("RGBA").getdata()
    total = 0
    blue = 0
    br_sum = 0.0
    r_sum = g_sum = b_sum = 0.0

    for r, g, b, a in px:
        if a < 10:
            continue
        total += 1
        r_sum += r; g_sum += g; b_sum += b
        br_sum += _brightness(r, g, b)
        if _is_blueish(r, g, b):
            blue += 1

    if total == 0:
        return {"blue_frac": 0.0, "avg": (0,0,0), "bright": 0.0}

    return {
        "blue_frac": blue / total,
        "avg": (r_sum/total, g_sum/total, b_sum/total),
        "bright": br_sum / total
    }

def _edge_blue_score(tile: Image.Image, band=4):
    """
    Mede "quão água" é cada borda (N,E,S,W) olhando uma faixa de pixels.
    Retorna 4 floats em [0..1].
    """
    im = tile.convert("RGBA")
    w, h = im.size
    pix = im.load()

    def score_region(x0, y0, x1, y1):
        total = 0
        blue = 0
        for y in range(y0, y1):
            for x in range(x0, x1):
                r, g, b, a = pix[x, y]
                if a < 10:
                    continue
                total += 1
                if _is_blueish(r, g, b):
                    blue += 1
        return (blue / total) if total else 0.0

    top    = score_region(0, 0, w, min(band, h))
    bottom = score_region(0, max(0, h-band), w, h)
    left   = score_region(0, 0, min(band, w), h)
    right  = score_region(max(0, w-band), 0, w, h)
    return (top, right, bottom, left)  # N, E, S, W

def _classify(tile: Image.Image):
    stt = _tile_stats(tile)
    blue = stt["blue_frac"]
    r, g, b = stt["avg"]
    bright = stt["bright"]

    # Heurística simples, mas funciona bem com atlas desse tipo:
    if blue > 0.55:
        # água: separa deep vs shallow pelo brilho
        return "water_shallow" if bright > 120 else "water_deep"

    # não-água: separa por cor média
    if g > r + 20 and g > b:
        return "grass"
    if r > g + 15 and r > b + 10:
        return "dirt"
    if r > 140 and g > 120 and b < 140:
        return "sand"
    if abs(r-g) < 15 and abs(g-b) < 15 and bright < 140:
        return "rock"

    return "detail"

def _pick_best_shore_tiles(all_tiles):
    """
    Cria um dicionário shore_variants[mask] = [tile_img, tile_img, ...]
    mask usa bits N,E,S,W (1 = vizinho é terra; 0 = vizinho é água)
    """
    shore_candidates = []
    deep_candidates = []
    shallow_candidates = []

    for t in all_tiles:
        cat = t["cat"]
        edges = t["edges"]  # N,E,S,W blue scores

        if cat.startswith("water"):
            # é água; se todas bordas bem "água", é candidato a miolo
            if all(e > 0.75 for e in edges):
                (shallow_candidates if cat == "water_shallow" else deep_candidates).append(t)
            else:
                shore_candidates.append(t)

    # se o atlas não tiver muitos shore tiles, usa os "não tão perfeitos"
    if len(shore_candidates) < 8:
        shore_candidates = [t for t in all_tiles if t["cat"].startswith("water")]

    shore_variants = {m: [] for m in range(16)}

    # para cada máscara, escolhe os melhores tiles por "match" das bordas
    for m in range(16):
        scored = []
        for t in shore_candidates:
            n,e,s,w = t["edges"]
            # Queremos:
            #  - onde m tem bit=1 (terra), borda deve ser MENOS água (1-score)
            #  - onde m tem bit=0 (água), borda deve ser MAIS água (score)
            target = [
                (1-n) if (m & 1) else n,         # N
                (1-e) if (m & 2) else e,         # E
                (1-s) if (m & 4) else s,         # S
                (1-w) if (m & 8) else w,         # W
            ]
            match = sum(target) / 4.0
            scored.append((match, t))

        scored.sort(key=lambda x: x[0], reverse=True)
        # pega top 3 para variação
        shore_variants[m] = [x[1] for x in scored[:3]]

    return deep_candidates, shallow_candidates, shore_variants


# =========================================================
# CORE_OUTDOOR (TSX) -> assets dict (rápido e com variações)
# =========================================================

def _resolve_asset_path_local(rel_path: str) -> str:
    """Resolve caminhos do Assets tanto em dev local quanto em Streamlit Cloud."""
    # 1) relativo ao CWD (raiz do repo, comum no Streamlit Cloud)
    if os.path.exists(rel_path):
        return rel_path
    # 2) relativo ao diretório deste arquivo
    try:
        here = os.path.dirname(__file__)
        cand = os.path.join(here, rel_path)
        if os.path.exists(cand):
            return cand
    except Exception:
        pass
    return rel_path  # deixa falhar lá embaixo (vai cair no fallback)

def _parse_tsx(tsx_path: str) -> dict:
    import xml.etree.ElementTree as ET
    tsx_path = _resolve_asset_path_local(tsx_path)
    root = ET.parse(tsx_path).getroot()
    img = root.find("image")
    return {
        "tilewidth": int(root.attrib.get("tilewidth", 16)),
        "tileheight": int(root.attrib.get("tileheight", 16)),
        "columns": int(root.attrib.get("columns", 1)),
        "tilecount": int(root.attrib.get("tilecount", 0)),
        "image_source": img.attrib.get("source", "") if img is not None else "",
        "tsx_dir": os.path.dirname(tsx_path),
    }

def _alpha_coverage(tile: Image.Image) -> float:
    im = tile.convert("RGBA")
    a = im.getchannel("A")
    # cobertura aproximada (bem rápida p/ tiles 16x16)
    total = im.size[0] * im.size[1]
    non0 = sum(1 for v in a.getdata() if v > 10)
    return non0 / total if total else 0.0

def _classify_core(tile: Image.Image) -> str:
    """Classificação estendida: grass/sand/dirt/rock + water + snow + cave."""
    stt = _tile_stats(tile)
    blue = stt["blue_frac"]
    r, g, b = stt["avg"]
    bright = stt["bright"]

    # água
    if blue > 0.55:
        return "water_shallow" if bright > 120 else "water_deep"

    # neve/gelo (muito claro e pouco saturado)
    if bright > 200 and abs(r - g) < 18 and abs(g - b) < 18:
        return "snow"

    # caverna (escuro e pouco saturado)
    if bright < 90 and abs(r - g) < 18 and abs(g - b) < 18:
        return "cave"

    # fallback "terra/areia/grama/pedra"
    cat = _classify(tile)
    return cat

def _border_contrast_score(tile: Image.Image, border: int = 2) -> float:
    """0..1 (maior = mais 'flat'). Compara bordas vs centro; tiles de cliff/edge tendem a ter contraste alto."""
    im = tile.convert("RGBA")
    w, h = im.size
    b = max(1, min(border, w // 4, h // 4))
    px = list(im.getdata())
    def avg_region(x0, y0, x1, y1):
        s = [0, 0, 0]
        n = 0
        for y in range(y0, y1):
            for x in range(x0, x1):
                r, g, bb, a = px[y * w + x]
                if a < 10:
                    continue
                s[0] += r; s[1] += g; s[2] += bb
                n += 1
        if n == 0:
            return (0, 0, 0)
        return (s[0] / n, s[1] / n, s[2] / n)

    center = avg_region(b, b, w - b, h - b)

    # bordas: top, bottom, left, right
    borders = [
        avg_region(0, 0, w, b),
        avg_region(0, h - b, w, h),
        avg_region(0, 0, b, h),
        avg_region(w - b, 0, w, h),
    ]

    # delta médio normalizado (0..~1)
    def dist(c1, c2):
        return (abs(c1[0] - c2[0]) + abs(c1[1] - c2[1]) + abs(c1[2] - c2[2])) / (3 * 255)

    d = sum(dist(center, br) for br in borders) / 4.0
    # flatness = 1 - contraste (clamp)
    return max(0.0, min(1.0, 1.0 - d))


def _pick_best_land_floors(candidates: list[Image.Image], want: int, border: int = 2) -> list[Image.Image]:
    """Escolhe pisos 'planos' primeiro para evitar tiles de borda/cliff virarem chão."""
    scored = []
    for t in candidates:
        cov = _alpha_coverage(t)
        if cov < 0.92:
            continue
        flat = _border_contrast_score(t, border=border)
        # penaliza tiles com buracos escuros (muito preto) -> geralmente 'pits' / recortes
        stt = _tile_stats(t)
        dark = stt.get("dark_frac", 0.0)
        score = flat - 0.75 * dark
        scored.append((score, t))
    scored.sort(key=lambda x: x[0], reverse=True)
    out = [t for _, t in scored[:want]]
    # fallback: se não teve o suficiente, completa sem score
    if len(out) < want:
        for t in candidates:
            if t not in out:
                out.append(t)
            if len(out) >= want:
                break
    return out


def _build_assets_from_core_tilesets(tile_size_out: int) -> dict[str, Image.Image]:
    """Constrói assets a partir de core_outdoor*.tsx evitando tiles de borda como piso."""
    # procura nos caminhos mais comuns
    tsx_triplet = [
        "Assets/core_outdoor.tsx",
        "Assets/core_outdoor_nature.tsx",
        "Assets/core_outdoor_water.tsx",
        "core_outdoor.tsx",
        "core_outdoor_nature.tsx",
        "core_outdoor_water.tsx",
    ]

    # pega um conjunto válido (3 arquivos)
    candidates = []
    for i in range(0, len(tsx_triplet), 3):
        trip = tsx_triplet[i:i+3]
        if len(trip) < 3:
            continue
        if all(os.path.exists(_resolve_asset_path_local(p)) for p in trip):
            candidates.append(trip)
    if not candidates:
        return {}

    tsx_trip = candidates[0]
    parsed = [_parse_tsx(p) for p in tsx_trip]

    sheets = []
    for meta in parsed:
        img_path = os.path.join(meta["tsx_dir"], meta["image_source"])
        img_path = _resolve_asset_path_local(img_path)
        if not os.path.exists(img_path):
            return {}
        sheets.append(Image.open(img_path).convert("RGBA"))

    # pools brutos
    floors_raw = {"grass": [], "sand": [], "dirt": [], "rock": [], "snow": [], "cave": []}
    objs = {"tree": [], "bush": [], "flower": [], "rock_prop": [], "wall": [], "stalagmite": [], "peak": []}
    water_tiles = []

    # quotas finais (o que você usa no render)
    floor_quota = {"grass": 28, "sand": 26, "dirt": 26, "rock": 28, "snow": 22, "cave": 22}
    obj_quota = {"tree": 14, "bush": 10, "flower": 16, "rock_prop": 10, "wall": 10, "stalagmite": 8, "peak": 6}
    water_quota_total = 1200

    def want_any_obj():
        return any(len(objs[k]) < obj_quota[k] for k in objs)

    # --------
    # 1) pisos + objetos (core_outdoor + nature)
    # --------
    for sheet, meta in [(sheets[0], parsed[0]), (sheets[1], parsed[1])]:
        tw, th = meta["tilewidth"], meta["tileheight"]
        cols = meta["columns"]
        count = meta["tilecount"]

        for tid in range(count):
            if (all(len(floors_raw[k]) >= floor_quota[k] * 5 for k in floors_raw) and not want_any_obj()):
                break

            x = (tid % cols) * tw
            y = (tid // cols) * th
            tile = sheet.crop((x, y, x + tw, y + th))
            cov = _alpha_coverage(tile)
            if cov < 0.20:
                continue

            cat = _classify_core(tile)

            # pisos sólidos (quase opacos): guardamos bruto para filtrar depois
            if cov > 0.92 and cat in floors_raw:
                floors_raw[cat].append(tile)
                continue

            if not want_any_obj():
                continue

            # objetos: heurística por bbox/altura + cor
            stt = _tile_stats(tile)
            r, g, b = stt["avg"]
            bright = stt["bright"]
            bbox = tile.getchannel("A").getbbox()
            if not bbox:
                continue
            bw, bh = bbox[2] - bbox[0], bbox[3] - bbox[1]
            rel_h = bh / max(1, th)
            rel_w = bw / max(1, tw)

            # trees: altos, bastante verde, cobertura moderada
            if (
                len(objs["tree"]) < obj_quota["tree"]
                and g > r + 12 and g > b + 8
                and 0.25 < cov < 0.90
                and rel_h > 0.65
                and rel_w > 0.40
            ):
                objs["tree"].append(tile)
                continue

            # bush: médio/baixo, verde, mais compacto
            if (
                len(objs["bush"]) < obj_quota["bush"]
                and g > r + 10 and g > b + 6
                and 0.18 < cov < 0.75
                and 0.25 < rel_h < 0.70
            ):
                objs["bush"].append(tile)
                continue

            # flower: bem claro/colorido e pequeno
            if (
                len(objs["flower"]) < obj_quota["flower"]
                and bright > 150
                and 0.10 < cov < 0.65
                and rel_h < 0.55
            ):
                objs["flower"].append(tile)
                continue

            # rochas de prop / picos / cave stuff
            if cat in ["rock", "cave"] and len(objs["rock_prop"]) < obj_quota["rock_prop"] and cov < 0.95:
                objs["rock_prop"].append(tile)
                continue
            if cat == "cave" and bright < 120 and len(objs["stalagmite"]) < obj_quota["stalagmite"]:
                objs["stalagmite"].append(tile)
                continue
            if cat == "rock" and bright < 140 and len(objs["peak"]) < obj_quota["peak"] and cov < 0.95:
                objs["peak"].append(tile)
                continue
            if cat == "cave" and cov > 0.55 and bright < 110 and len(objs["wall"]) < obj_quota["wall"]:
                objs["wall"].append(tile)
                continue

    # --------
    # 2) água (core_outdoor_water) -> autotile shore por máscara
    # --------
    sheet = sheets[2]
    meta = parsed[2]
    tw, th = meta["tilewidth"], meta["tileheight"]
    cols = meta["columns"]
    count = meta["tilecount"]
    for tid in range(count):
        if len(water_tiles) >= water_quota_total:
            break
        x = (tid % cols) * tw
        y = (tid // cols) * th
        tile = sheet.crop((x, y, x + tw, y + th))
        cov = _alpha_coverage(tile)
        if cov < 0.35:
            continue
        cat = _classify_core(tile)
        if not cat.startswith("water"):
            continue
        edges = _edge_blue_score(tile, band=max(2, tw // 4))
        water_tiles.append({"id": f"w{tid}", "img": tile, "cat": cat, "edges": edges})

    deep_tiles, shallow_tiles, shore_by_mask = _pick_best_shore_tiles(water_tiles)

    # --------
    # 3) filtra pisos: pega os 'mais planos' primeiro
    # --------
    floors = {}
    for k in floors_raw:
        floors[k] = _pick_best_land_floors(floors_raw[k], want=floor_quota[k], border=max(1, tw // 8))

    # --------
    # 4) monta assets no formato do render atual
    # --------
    assets: dict[str, Image.Image] = {}

    def _resize(tile: Image.Image) -> Image.Image:
        if tile.size != (tile_size_out, tile_size_out):
            return tile.resize((tile_size_out, tile_size_out), Image.Resampling.NEAREST)
        return tile

    def add_floor(base_key: str, pool: list[Image.Image], limit: int):
        if not pool:
            return
        assets[base_key] = _resize(pool[0])
        for i, t in enumerate(pool[:limit]):
            assets[f"{base_key}__v{i:02d}"] = _resize(t)

    add_floor("grama_1", floors["grass"], 24)
    add_floor("areia_1", floors["sand"], 24)
    add_floor("terra_1", floors["dirt"], 24)
    add_floor("pedra_1", floors["rock"], 24)
    add_floor("neve_1", floors["snow"], 20)
    add_floor("cave_1", floors["cave"], 20)

    # água miolo
    if deep_tiles:
        assets["agua_deep"] = _resize(deep_tiles[0]["img"])
        for i, t in enumerate(deep_tiles[:8]):
            assets[f"agua_deep__v{i:02d}"] = _resize(t["img"])
    if shallow_tiles:
        assets["agua_shallow"] = _resize(shallow_tiles[0]["img"])
        for i, t in enumerate(shallow_tiles[:8]):
            assets[f"agua_shallow__v{i:02d}"] = _resize(t["img"])

    for m in range(16):
        variants = shore_by_mask.get(m, [])
        if not variants:
            continue
        assets[f"agua_shore_{m:02d}"] = _resize(variants[0]["img"])
        for i, t in enumerate(variants[:3]):
            assets[f"agua_shore_{m:02d}__v{i:02d}"] = _resize(t["img"])

    def add_obj_keys(keys: list[str], pool: list[Image.Image]):
        for k, t in zip(keys, pool):
            assets[k] = _resize(t)

    add_obj_keys(["tree_1", "tree_2", "tree_3"], objs["tree"][:3])
    add_obj_keys(["brush_1", "brush_2"], objs["bush"][:2])
    if objs["flower"]:
        assets["flower"] = _resize(objs["flower"][0])
    add_obj_keys(["rochas", "rochas_2"], objs["rock_prop"][:2])
    if objs["wall"]:
        assets["wall_1"] = _resize(objs["wall"][0])
    if objs["stalagmite"]:
        assets["estalagmite_1"] = _resize(objs["stalagmite"][0])
    if objs["peak"]:
        assets["pico_1"] = _resize(objs["peak"][0])

    return assets

def _pick_variant(assets: dict, base_key: str, rng):
    """Escolhe base_key ou uma variante base_key__vXX se existir."""
    opts = [k for k in assets.keys() if k == base_key or k.startswith(base_key + "__v")]
    if not opts:
        return None
    return rng.choice(sorted(opts))

def _get_rotated_asset(assets: dict, key: str, angle: int):
    """
    Retorna uma versão rotacionada (cacheada) do asset.
    angle: 0, 90, 180, 270 (graus, sentido horário)
    """
    angle = angle % 360
    if angle == 0:
        return assets.get(key)

    cache_key = f"{key}__rot{angle}"
    if cache_key in assets:
        return assets[cache_key]

    img = assets.get(key)
    if img is None:
        return None

    # PIL rotate é anti-horário; por isso usamos -angle
    rot = img.rotate(-angle, expand=False)
    assets[cache_key] = rot
    return rot

def _water_tile_from_landmask(assets: dict, rng, landmask: int):
    """
    landmask bits (padrão):
      N=1, E=2, S=4, W=8  (vizinho é TERRA)
    Retorna (key_base, angle) para usar water_grass coerente.
    """
    N, E, S, W = 1, 2, 4, 8

    # miolo (sem terra encostando): water core
    if landmask == 0:
        # prefira shallow no lago; deep opcional
        k = _pick_variant(assets, "agua_shallow", rng) or _pick_variant(assets, "agua_deep", rng) or "agua_1"
        return (k, 0)

    # --- 1) UMA lateral de terra (borda reta) ---
    # usamos como "template" a borda reta com terra em CIMA: water_grass_m02_*
    # (no seu pack, a m02 é a borda reta horizontal perfeita)
    EDGE = "agua_shore_02"  # vindo de water_grass_m02_*

    if landmask in (N, E, S, W):
        angle = {N: 0, E: 90, S: 180, W: 270}[landmask]
        k = _pick_variant(assets, EDGE, rng) or EDGE
        return (k, angle)

    # --- 2) DUAS laterais adjacentes (canto externo) ---
    # templates ótimos no seu pack:
    # - m10 parece "canto" com terra em CIMA+ESQUERDA (NW)
    # - m11 parece "canto" com terra em CIMA+DIREITA (NE)
    CORNER_NW = "agua_shore_10"
    CORNER_NE = "agua_shore_11"

    if landmask == (N | W):   # NW
        k = _pick_variant(assets, CORNER_NW, rng) or CORNER_NW
        return (k, 0)
    if landmask == (N | E):   # NE
        k = _pick_variant(assets, CORNER_NE, rng) or CORNER_NE
        return (k, 0)
    if landmask == (S | E):   # SE = NW rot 180
        k = _pick_variant(assets, CORNER_NW, rng) or CORNER_NW
        return (k, 180)
    if landmask == (S | W):   # SW = NE rot 180
        k = _pick_variant(assets, CORNER_NE, rng) or CORNER_NE
        return (k, 180)

    # --- 3) TRÊS laterais de terra (baía/cova) ---
    # para lago ficar bonito, use tiles "cap" arredondados:
    # m00/m01 são excelentes (terra em cima e arredondamento lateral)
    CAP_A = "agua_shore_00"
    CAP_B = "agua_shore_01"

    if landmask == (N | E | W):     # aberto para S
        k = _pick_variant(assets, rng.choice([CAP_A, CAP_B]), rng) or CAP_A
        return (k, 0)
    if landmask == (S | E | W):     # aberto para N
        k = _pick_variant(assets, rng.choice([CAP_A, CAP_B]), rng) or CAP_A
        return (k, 180)
    if landmask == (N | S | W):     # aberto para E
        k = _pick_variant(assets, rng.choice([CAP_A, CAP_B]), rng) or CAP_A
        return (k, 270)
    if landmask == (N | S | E):     # aberto para W
        k = _pick_variant(assets, rng.choice([CAP_A, CAP_B]), rng) or CAP_A
        return (k, 90)

    # --- 4) casos raros (opostos / cruz) ---
    # Para evitar borda “nonsense”, cai pra água core.
    k = _pick_variant(assets, "agua_shallow", rng) or _pick_variant(assets, "agua_deep", rng) or "agua_1"
    return (k, 0)


@st.cache_resource
def load_map_assets():
    """
    Carrega assets para o gerador de mapas.

    - Base: Assets/Texturas (compatível com seu modelo atual)
    - Extra: Assets/map (usa TODOS os PNGs: pisos, água deep/shallow, shores por máscara, overlays)
    - Normaliza variações: *_2, *_3 => viram __v para o sistema de variantes funcionar
    - Garante chaves esperadas pelo render (agua_deep/shore_00..15 etc.)
    """
    base_dir = _resolve_asset_path_local("Assets/Texturas")
    extra_dir = _resolve_asset_path_local("Assets/map")

    assets: dict[str, Image.Image] = {}

    # compatibilidade Pillow
    _resampling = getattr(Image, "Resampling", Image)
    resample_nearest = getattr(_resampling, "NEAREST", Image.NEAREST)

    def _open_rgba(path: str) -> Image.Image | None:
        try:
            img = Image.open(path).convert("RGBA")
            if img.size != (TILE_SIZE, TILE_SIZE):
                img = img.resize((TILE_SIZE, TILE_SIZE), resample_nearest)
            return img
        except Exception:
            return None

    def pick_solid_color(img: Image.Image) -> tuple[int, int, int]:
        counts: dict[tuple[int,int,int], int] = {}
        for r, g, b, a in img.getdata():
            if a > 0:
                counts[(r, g, b)] = counts.get((r, g, b), 0) + 1
        return max(counts, key=counts.get) if counts else (0, 0, 0)

    def crop_to_alpha(img: Image.Image) -> Image.Image:
        bbox = img.getchannel("A").getbbox()
        return img.crop(bbox) if bbox else img

    def normalize_floor(img: Image.Image) -> Image.Image:
        alpha = img.getchannel("A")
        if alpha.getextrema()[0] < 255:
            solid = pick_solid_color(img)
            base = Image.new("RGBA", img.size, (*solid, 255))
            base.alpha_composite(img)
            img = base
        return img

    def add_asset(key: str, img: Image.Image):
        # normaliza pisos
        if key.startswith(FLOOR_PREFIXES):
            img = crop_to_alpha(img)
            if img.size != (TILE_SIZE, TILE_SIZE):
                img = img.resize((TILE_SIZE, TILE_SIZE), resample_nearest)
            img = normalize_floor(img)
        assets[key] = img

    # -----------------------------
    # 1) Carrega Assets/Texturas
    # -----------------------------
    if os.path.isdir(base_dir):
        for fn in os.listdir(base_dir):
            if not fn.lower().endswith(".png"):
                continue
            stem = os.path.splitext(fn)[0]
            img = _open_rgba(os.path.join(base_dir, fn))
            if img:
                add_asset(stem, img)

    # ----------------------------------------------------------
    # 2) Conserta o seu “sistema de variantes” ( *_2/_3 -> __v )
    # ----------------------------------------------------------
    # Ex: grama_2 vira grama_1__v01, grama_3 vira grama_1__v02
    def promote_numbered_variants(base_key: str, numbered_keys: list[str]):
        if base_key not in assets:
            return
        v = 1
        for nk in sorted(numbered_keys):
            assets[f"{base_key}__v{v:02d}"] = assets[nk]
            v += 1

    # Para cada família, se existir *_2/_3, promove pra variantes do *_1
    for fam in ["grama", "areia", "terra", "pedra", "agua"]:
        numbered = [k for k in list(assets.keys()) if re.fullmatch(rf"{fam}_[2-9]\d*", k)]
        promote_numbered_variants(f"{fam}_1", numbered)

    # -----------------------------
    # 3) Carrega TODOS Assets/map
    # -----------------------------
    # A ideia: mapear pastas conhecidas como variantes dos pisos/objetos
    def add_variants_from_folder(folder: str, base_key: str, max_take: int | None = None):
        path = os.path.join(extra_dir, folder)
        if not os.path.isdir(path):
            return
        files = [f for f in os.listdir(path) if f.lower().endswith(".png")]
        files.sort()
        if max_take:
            files = files[:max_take]
        v = 0
        for fn in files:
            img = _open_rgba(os.path.join(path, fn))
            if not img:
                continue
            assets[f"{base_key}__v{v:03d}"] = img
            v += 1

    # pisos (vira variedade enorme do seu “grama/areia/terra/pedra”)
    add_variants_from_folder("grass", "grama_1")
    add_variants_from_folder("dark_grass", "grama_1")
    add_variants_from_folder("dark_grass_light_grass", "grama_1")
    add_variants_from_folder("sand", "areia_1")
    add_variants_from_folder("wetdry_sand", "areia_1")
    add_variants_from_folder("light_dirt", "terra_1")
    # NOTE: "dirt_rock_edge" não entra como variação de terra_1; é usado como "muro" mascarado na arena.
    add_variants_from_folder("rocks", "pedra_1")

    # -----------------------------------------
    # 3.1) Água “deep/shallow” (water_core_tiles)
    # -----------------------------------------
    wc = os.path.join(extra_dir, "water_core_tiles")
    if os.path.isdir(wc):
        deep = sorted([f for f in os.listdir(wc) if f.lower().startswith("water_core_deep_") and f.lower().endswith(".png")])
        shallow = sorted([f for f in os.listdir(wc) if f.lower().startswith("water_core_shallow_") and f.lower().endswith(".png")])

        if deep:
            img0 = _open_rgba(os.path.join(wc, deep[0]))
            if img0:
                assets["agua_deep"] = img0
            for i, fn in enumerate(deep[:16]):
                im = _open_rgba(os.path.join(wc, fn))
                if im:
                    assets[f"agua_deep__v{i:02d}"] = im

        if shallow:
            img0 = _open_rgba(os.path.join(wc, shallow[0]))
            if img0:
                assets["agua_shallow"] = img0
            for i, fn in enumerate(shallow[:16]):
                im = _open_rgba(os.path.join(wc, fn))
                if im:
                    assets[f"agua_shallow__v{i:02d}"] = im

    # ---------------------------------------------------
    # 3.2) Shores por máscara (16 máscaras) -> agua_shore_XX
    # Pastas: water_grass / water_sand / water_overlay
    # Arquivo padrão: water_sand_m00_v0.png etc.
    # ---------------------------------------------------
    def load_shores(folder: str):
        p = os.path.join(extra_dir, folder)
        if not os.path.isdir(p):
            return

        family_map = {
            "water_grass": "grass",
            "water_sand": "sand",
            "water_overlay": "overlay",
        }
        family = family_map.get(folder, folder)

        files = sorted([f for f in os.listdir(p) if f.lower().endswith(".png")])
        for fn in files:
            m = re.match(r".*_m(\d{2})_v(\d+)\.png$", fn, re.IGNORECASE)
            if not m:
                continue
            mask = int(m.group(1))
            vid = int(m.group(2))
            img = _open_rgba(os.path.join(p, fn))
            if not img:
                continue

            base = f"agua_shore_{family}_{mask:02d}"
            # primeira que aparecer vira base; as demais viram variantes
            if base not in assets:
                assets[base] = img
            assets[f"{base}__v{vid:02d}"] = img

    load_shores("water_grass")
    load_shores("water_sand")
    load_shores("water_overlay")

    # ---------------------------------------------------
    # 3.2b) Tiles mascarados (m00..m15) para paredes/contornos
    # Ex.: dirt_rock_edge_m02_v0.png -> dirt_rock_edge_02 (+ variantes)
    # ---------------------------------------------------
    def load_masked_tiles(folder: str, key_prefix: str):
        p = os.path.join(extra_dir, folder)
        if not os.path.isdir(p):
            return

        files = sorted([f for f in os.listdir(p) if f.lower().endswith(".png")])
        for fn in files:
            m = re.match(r".*_m(\d{2})_v(\d+)\.png$", fn, re.IGNORECASE)
            if not m:
                continue
            mask = int(m.group(1))
            vid = int(m.group(2))
            img = _open_rgba(os.path.join(p, fn))
            if not img:
                continue

            base = f"{key_prefix}_{mask:02d}"
            if base not in assets:
                assets[base] = img
            assets[f"{base}__v{vid:02d}"] = img

    # Dirt+rock edge (usado como muro/parede na arena da caverna)
    load_masked_tiles("dirt_rock_edge", "dirt_rock_edge")

    # -----------------------------------------
    # 3.3) Objetos/overlays (usar tudo)
    # -----------------------------------------
    add_variants_from_folder("foliage", "brush_1")
    add_variants_from_folder("forest_overlays", "tree_1")
    add_variants_from_folder("overlays_and_objects", "brush_1")
    add_variants_from_folder("flower", "flower")
    add_variants_from_folder("cave_overlay", "cave_overlay")

    # -----------------------------
    # 4) Garantias / fallbacks
    # -----------------------------
    # Se não existir flower, tenta cair para brush
    if "flower" not in assets:
        if "brush_1" in assets:
            assets["flower"] = assets["brush_1"]

    # cave/neve não existem no seu zip -> fallback seguro
    if "cave_1" not in assets and "pedra_1" in assets:
        assets["cave_1"] = assets["pedra_1"]
    if "neve_1" not in assets and "pedra_1" in assets:
        assets["neve_1"] = assets["pedra_1"]

    # agua_deep/agua_shore pode faltar se pasta não existir
    if "agua_deep" not in assets and "agua_1" in assets:
        assets["agua_deep"] = assets["agua_1"]
    for m in range(16):
        k = f"agua_shore_{m:02d}"
        if k not in assets and "agua_1" in assets:
            assets[k] = assets["agua_1"]

    return assets



def authenticate_user(name, password):
    try:
        sheet = get_google_sheet()
        row_num = find_user_row(sheet, name)
        
        if row_num is None:
            return "NOT_FOUND"
        
        # Pega a linha exata
        row_values = sheet.row_values(row_num)
        
        # Validação de segurança se a linha estiver quebrada
        if len(row_values) < 3:
            return "WRONG_PASS"
            
        stored_password = str(row_values[2]) # Coluna C
        stored_data = row_values[1]          # Coluna B
        
        if stored_password == str(password):
            return json.loads(stored_data)
        else:
            return "WRONG_PASS"
            
    except Exception as e:
        st.error(f"Erro na autenticação: {e}")
        return None

def register_new_user(name, password):
    try:
        sheet = get_google_sheet()
        
        # Verifica APENAS na coluna 1
        if find_user_row(sheet, name) is not None:
            return "EXISTS"
            
        empty_data = json.dumps(get_empty_user_data())
        # Adiciona: Coluna A (Nome), Coluna B (Dados), Coluna C (Senha)
        sheet.append_row([name, empty_data, str(password)])
        return "SUCCESS"
    except Exception as e:
        st.error(f"Erro ao criar usuário: {e}")
        return "ERROR"

def save_data_cloud(trainer_name, data):
    try:
        sheet = get_google_sheet()

        # ✅ limpeza preventiva: remove campos que estouram a célula
        prof = (data or {}).get("trainer_profile")
        if isinstance(prof, dict):
            prof.pop("photo_b64", None)

        # ✅ reduz tamanho do JSON
        json_str = json.dumps(data, ensure_ascii=False, separators=(",", ":"))

        row_num = find_user_row(sheet, trainer_name)
        if not row_num:
            st.error("Erro crítico: Usuário sumiu da planilha enquanto salvava.")
            return False

        sheet.update_cell(row_num, 2, json_str)

        # ✅ Auto-sync NPC → Compendium (mesma sessão)
        try:
            if isinstance(data, dict) and data.get("npc_user"):
                npc_name = (data.get("npc_name") or trainer_name or "").strip() or trainer_name

                # --- pega NPC "base" do compendium pra não perder lore ---
                base_obj = {}
                try:
                    cd = st.session_state.get("comp_data")
                    if isinstance(cd, dict):
                        npcs = cd.get("npcs") or {}
                        if isinstance(npcs, dict):
                            base_obj = npcs.get(npc_name) or {}
                except Exception:
                    base_obj = {}

                overrides = st.session_state.setdefault("npc_sync_overrides", {})

                # usa prioridade: override anterior -> base do compendium -> default
                npc_obj = overrides.get(npc_name) or base_obj or {"name": npc_name, "sections": {}}
                npc_obj["name"] = npc_name  # garante

                # ✅ preserva lore: nunca zere sections
                if not isinstance(npc_obj.get("sections"), dict) or not npc_obj.get("sections"):
                    if isinstance(base_obj, dict) and isinstance(base_obj.get("sections"), dict) and base_obj.get("sections"):
                        npc_obj["sections"] = base_obj.get("sections")
                    else:
                        npc_obj["sections"] = npc_obj.get("sections") if isinstance(npc_obj.get("sections"), dict) else {}

                # 1) fontes possíveis (do user NPC)
                raw = []
                if isinstance(data.get("npc_pokemons"), list):
                    raw += (data.get("npc_pokemons") or [])
                if isinstance(data.get("party"), list):
                    raw += (data.get("party") or [])
                if isinstance(data.get("caught"), list):
                    raw += (data.get("caught") or [])

                # 2) converte para nomes + dedupe mantendo ordem
                pokes = []
                seen = set()
                for pid in raw:
                    pid_str = str(pid).strip()
                    if not pid_str:
                        continue

                    if pid_str.startswith("EXT:"):
                        nm = pid_str.replace("EXT:", "").strip()
                    else:
                        try:
                            nm = _get_pokemon_name(pid_str)
                        except Exception:
                            nm = pid_str

                    nm = (nm or "").strip()
                    key = nm.lower()
                    if nm and key not in seen:
                        seen.add(key)
                        pokes.append(nm)

                # 3) merge com pokemons já existentes no NPC base/override
                existing = npc_obj.get("pokemons") or npc_obj.get("pokemons_conhecidos") or []
                if isinstance(existing, list) and existing:
                    for x in existing:
                        xs = str(x).strip()
                        k = xs.lower()
                        if xs and k not in seen:
                            seen.add(k)
                            pokes.append(xs)

                npc_obj["pokemons"] = list(pokes)
                npc_obj["pokemons_conhecidos"] = list(pokes)
                overrides[npc_name] = npc_obj

                # ✅ IMPORTANTÍSSIMO: invalida cache do compendium na sessão
                st.session_state.pop("comp_data", None)

        except Exception:
            # falha no auto-sync não deve impedir salvar o usuário
            pass

        return True  # ✅ salva ok

    except Exception as e:
        st.error(f"Erro ao salvar dados: {e}")
        return False


def get_empty_user_data():
    return {
        "seen": [],
        "caught": [],
        "party": [],
        "wishlist": [],
        "shinies": [],
        "favorite_moves": {},
        "forms": {},
        "stats": {},
        "notes": {},
    }


# ============================================================
# LIMPEZA ÚNICA: remove resquícios do antigo sistema __dex_uid
# Roda uma vez por sessão e depois nunca mais interfere.
# ============================================================

def _dex_cleanup_once(data: dict, df_current: pd.DataFrame) -> None:
    """Remove __dex_uid e converte quaisquer entradas UID:/PID: residuais."""
    if not isinstance(data, dict):
        return
    if st.session_state.get("_dex_cleanup_done"):
        return

    # Mapa UID -> PID atual (para resolver entradas UID: que sobraram)
    uid_to_pid: dict[str, str] = {}
    if df_current is not None and not df_current.empty:
        cols = set(map(str, df_current.columns))
        has_region = "Região" in cols or "Regiao" in cols
        for _, r in df_current.iterrows():
            try:
                pid = _norm_pid(r.get("Nº", ""))
            except Exception:
                pid = ""
            if not pid:
                continue
            name = str(r.get("Nome", "")).strip()
            region = None
            if has_region:
                region = r.get("Região", r.get("Regiao", None))
            base = _norm(str(name))
            reg = _norm(str(region)) if region not in (None, "", "-", "—") else ""
            if reg:
                base = f"{base}|{reg}"
            uid_key = f"UID:{base}"
            if uid_key not in uid_to_pid:
                uid_to_pid[uid_key] = pid

    def _resolve(s: str) -> str | None:
        s = str(s).strip()
        if not s:
            return None
        if s.startswith("UID:"):
            return uid_to_pid.get(s)
        if s.startswith("PID:"):
            return s.replace("PID:", "", 1)
        return s

    # Limpa listas
    for k in ("seen", "caught", "party", "shinies", "wishlist"):
        cur = data.get(k)
        if not isinstance(cur, list):
            continue
        cleaned = []
        for item in cur:
            resolved = _resolve(item)
            if resolved:
                cleaned.append(resolved)
        data[k] = cleaned

    # Limpa dicts
    for dk in ("stats", "favorite_moves", "forms"):
        curd = data.get(dk)
        if not isinstance(curd, dict):
            continue
        newd = {}
        for key, val in curd.items():
            resolved = _resolve(key)
            if resolved:
                newd[resolved] = val
        data[dk] = newd

    # Remove o store __dex_uid
    data.pop("__dex_uid", None)

    st.session_state["_dex_cleanup_done"] = True


def _npc_pokemon_list(npc_obj: dict) -> list[str]:
    pokes = npc_obj.get("pokemons") or npc_obj.get("pokemons_conhecidos") or []
    if isinstance(pokes, str):
        pokes = [p.strip() for p in re.split(r"[;,/|\n]+", pokes) if p.strip()]
    if not isinstance(pokes, list):
        return []
    out = []
    for p in pokes:
        ps = str(p).strip()
        if ps:
            out.append(ps)
    return out

def _npc_pokemon_names_to_ids(pokes: list[str]) -> list[str]:
    seen = set()
    out: list[str] = []
    local_df = st.session_state.get("df_data") if "df_data" in st.session_state else df
    for p in pokes:
        pid = None
        if local_df is not None and "Nome" in local_df.columns and "Nº" in local_df.columns:
            pname = str(p).strip()
            if pname:
                lowered = pname.lower()
                hit = local_df[local_df["Nome"].astype(str).str.strip().str.lower() == lowered]
                if hit.empty:
                    norm = normalize_text(pname)
                    hit = local_df[local_df["Nome"].astype(str).apply(normalize_text) == norm]
                if not hit.empty:
                    pid = str(hit.iloc[0]["Nº"]).strip().replace("#", "")
        if pid:
            key = str(pid)
        else:
            key = f"EXT:{p}"
        if key in seen:
            continue
        seen.add(key)
        out.append(key)
    return out

def _npc_pokemon_ids_to_names(ids: list[str]) -> list[str]:
    if not ids:
        return []
    names = []
    for pid in ids:
        pid_str = str(pid).strip()
        if not pid_str:
            continue
        if pid_str.startswith("EXT:"):
            names.append(pid_str.replace("EXT:", "").strip())
            continue
        if df is not None:
            pid_norm = _norm_pid(pid_str)
            row = df[df["Nº"].apply(_norm_pid) == pid_norm]
            if not row.empty:
                names.append(str(row.iloc[0]["Nome"]))
                continue
        else:
            names.append(pid_str)
            continue
        names.append(pid_str)
    return names

def _npc_user_pokemon_names(user_data: dict) -> list[str]:
    caught = user_data.get("caught") or []
    if not isinstance(caught, list):
        caught = []

    names: list[str] = []
    if caught:
        for pid in caught:
            pid_str = str(pid).strip()
            if not pid_str:
                continue
            if pid_str.startswith("EXT:"):
                names.append(pid_str.replace("EXT:", "").strip())
                continue
            if pid_str.isdigit() or re.fullmatch(r"\d{3}", pid_str):
                if df is not None:
                    pid_norm = _norm_pid(pid_str)
                    row = df[df["Nº"].apply(_norm_pid) == pid_norm]
                    if not row.empty:
                        names.append(str(row.iloc[0]["Nome"]))
                        continue
                names.append(pid_str)
                continue
            names.append(pid_str)
        return names

    raw = user_data.get("npc_pokemons")
    if isinstance(raw, list):
        cleaned = []
        for p in raw:
            ps = str(p).strip()
            if ps:
                cleaned.append(ps)
        if cleaned:
            return cleaned

    return names

def _ensure_npc_user_payload(npc_name: str, npc_obj: dict) -> dict:
    npc_names = _npc_pokemon_list(npc_obj)
    npc_ids = _npc_pokemon_names_to_ids(npc_names)
    payload = get_empty_user_data()
    payload["caught"] = npc_ids
    payload["seen"] = npc_ids.copy()
    payload["npc_pokemons"] = npc_names
    payload["npc_user"] = True
    payload["npc_name"] = npc_name
    return payload

def _sync_npc_users_and_overrides(npc_map: dict[str, dict]) -> dict[str, dict]:
    try:
        sheet = get_google_sheet()
    except Exception:
        return npc_map

    try:
        rows = sheet.get_all_values()
    except Exception:
        return npc_map

    user_rows: dict[str, dict] = {}
    for idx, row in enumerate(rows, start=1):
        if not row:
            continue
        name = (row[0] or "").strip()
        if not name:
            continue
        data = {}
        if len(row) > 1 and row[1]:
            try:
                data = json.loads(row[1])
            except Exception:
                data = {}
        user_rows[name] = {"row": idx, "data": data, "password": row[2] if len(row) > 2 else ""}

    for npc_name, npc_obj in (npc_map or {}).items():
        entry = user_rows.get(npc_name)
        if entry is None:
            payload = _ensure_npc_user_payload(npc_name, npc_obj)
            try:
                sheet.append_row([npc_name, json.dumps(payload), "1"])
                user_rows[npc_name] = {"row": None, "data": payload, "password": "1"}
            except Exception:
                continue
            entry = user_rows.get(npc_name)

        if not entry:
            continue

        data = entry.get("data") or {}
        if not isinstance(data, dict) or not data.get("npc_user"):
            continue

        npc_names = _npc_user_pokemon_names(data)
        if not npc_names:
            fallback_names = _npc_pokemon_list(npc_obj)
            fallback_ids = _npc_pokemon_names_to_ids(fallback_names)
            if fallback_ids:
                data["caught"] = fallback_ids
                data["seen"] = list(dict.fromkeys((data.get("seen") or []) + fallback_ids))
                data["npc_pokemons"] = fallback_names
                try:
                    if entry.get("row"):
                        sheet.update_cell(entry["row"], 2, json.dumps(data))
                except Exception:
                    pass
                npc_names = fallback_names

        if npc_names:
            if data.get("npc_pokemons") != npc_names:
                data["npc_pokemons"] = npc_names
                try:
                    if entry.get("row"):
                        sheet.update_cell(entry["row"], 2, json.dumps(data))
                except Exception:
                    pass
            npc_obj["pokemons"] = npc_names
            npc_obj["pokemons_conhecidos"] = npc_names

    return npc_map
def render_login_menu(trainer_name: str, user_data: dict):
    caught_count = len(user_data.get("caught", []) or [])
    badge_count = int(user_data.get("badges", 0) or 0)
    profile = user_data.get("trainer_profile", {})
    avatar_choice = profile.get("avatar_choice")
    avatar_src = ""
    if avatar_choice:
        avatar_path = Path("trainer") / f"{avatar_choice}.png"
        if avatar_path.exists():
            avatar_src = comp_img_data_uri(str(avatar_path))
    badge_dir = os.path.join("Assets", "insignias")
    badge_files = []
    if os.path.isdir(badge_dir):
        badge_files = [
            f
            for f in os.listdir(badge_dir)
            if f.lower().endswith((".png", ".jpg", ".jpeg", ".webp"))
        ]
    badge_paths = {os.path.splitext(f)[0]: os.path.join(badge_dir, f) for f in badge_files}
    selected_badges = profile.get("badges", [])
    badge_srcs = [
        comp_img_data_uri(badge_paths[badge_key])
        for badge_key in selected_badges
        if badge_paths.get(badge_key)
    ]

    # Party (sprites) — usa a MESMA lógica do Trainer Hub:
    # - respeita shiny
    # - respeita forma salva (user_data["forms"][pid] => ex: "ponyta-galar", "lycanroc-midnight")
    raw_party = user_data.get("party") or []
    party_ids: list[str] = [str(p).strip() for p in raw_party if str(p).strip()]

    party_sprites: list[str] = []
    shinies = set(str(x).strip() for x in (user_data.get("shinies") or []) if str(x).strip())
    forms = user_data.get("forms", {}) if isinstance(user_data.get("forms", {}), dict) else {}

    for pid in party_ids:
        pid = str(pid).strip()
        if not pid:
            continue

        is_shiny = pid in shinies

        if pid.startswith("EXT:"):
            party_sprites.append(pokemon_pid_to_image(pid, mode="sprite", shiny=is_shiny))
            continue

        saved_form = forms.get(pid) or forms.get(str(pid))
        if saved_form:
            # Força lookup por NOME de forma (evita depender do Nº do Excel)
            party_sprites.append(pokemon_pid_to_image(f"EXT:{saved_form}", mode="sprite", shiny=is_shiny))
        else:
            party_sprites.append(pokemon_pid_to_image(pid, mode="sprite", shiny=is_shiny))

    st.markdown(
        """
        <style>
        .fr-login-wrap {
            max-width: 760px;
            margin: 32px auto 0 auto;
            font-family: "Press Start 2P", "Trebuchet MS", sans-serif;
        }
        .fr-login-layout {
            display: flex;
            flex-direction: column;
            align-items: stretch;
            gap: 18px;
        }
        .fr-login-card {
            flex: 1;
            position: relative;
            border: 4px solid #1f2a6b;
            border-radius: 14px;
            background: linear-gradient(180deg, #f8fbff 0%, #dbe7ff 100%);
            padding: 14px 16px;
            box-shadow: 0 8px 0 rgba(11, 23, 64, 0.4);
        }
        .fr-login-title {
            text-transform: uppercase;
            font-size: 12px;
            letter-spacing: 1px;
            margin-bottom: 10px;
        }
        .fr-login-info {
            border: 3px solid #3b5bd8;
            border-radius: 10px;
            background: linear-gradient(180deg, #f7fbff 0%, #d7e6ff 55%, #c8dcff 100%);
            padding: 10px 12px;
            box-shadow: inset 0 0 0 2px rgba(255,255,255,0.7), inset 0 -6px 10px rgba(25,45,120,0.15);
        }
        .fr-login-info-empty {
            min-height: 150px;
        }
        .fr-login-grid {
            display: grid;
            grid-template-columns: 1.1fr 1fr;
            gap: 8px 20px;
            font-size: 11px;
        }
        .fr-login-label {
            color: #1b3d7a;
            letter-spacing: 0.8px;
        }
        .fr-login-value {
            color: #0b2d6b;
            font-weight: 700;
            background: rgba(255,255,255,0.65);
            border: 1px solid rgba(59,91,216,0.35);
            border-radius: 6px;
            padding: 2px 6px;
            text-align: right;
            box-shadow: inset 0 -1px 0 rgba(0,0,0,0.1);
        }
        .fr-login-actions {
            display: flex;
            flex-direction: column;
            gap: 12px;
            padding-top: 26px;
            align-items: flex-start;
        }
        .fr-login-actions-row {
            display: flex;
            gap: 10px;
            align-items: center;
        }
        .fr-login-action {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            padding: 8px 12px;
            text-decoration: none;
            text-transform: uppercase;
            font-size: 9px;
            font-weight: 700;
            border: 3px solid #293356;
            color: #0f172a;
            background: linear-gradient(180deg, #eef3ff 0%, #c8d7ff 100%);
            box-shadow: 0 4px 0 rgba(0,0,0,0.35), inset 0 -3px 0 rgba(0,0,0,0.18);
            border-radius: 12px;
        }
        .fr-login-action-confirm {
            padding: 6px 10px;
            font-size: 8px;
            border-radius: 10px;
            background: linear-gradient(180deg, #f8fafc 0%, #e2e8f0 100%);
        }
        .fr-login-action-new-game {
            width: 150px;
            justify-content: center;
            font-size: 10px;
            padding: 10px 14px;
            border-radius: 6px;
            background: linear-gradient(180deg, #f3f4ff 0%, #a7b7ff 100%);
            border: 3px solid #3c4a88;
            box-shadow: 0 4px 0 rgba(0,0,0,0.35), inset 0 -3px 0 rgba(0,0,0,0.25);
        }
        .fr-login-title {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 12px;
        }
        .fr-login-continue-card,
        .fr-login-continue-card .fr-login-title,
        .fr-login-continue-card .fr-login-grid,
        .fr-login-continue-card .fr-login-label,
        .fr-login-continue-card .fr-login-value {
            color: #0b2d6b;
        }
        .fr-login-new-game-card,
        .fr-login-new-game-card .fr-login-title {
            color: #10203f;
        }
        .fr-login-continue-badge {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 26px;
            height: 26px;
            border-radius: 6px;
            background: #16a34a;
            color: #ffffff;
            font-weight: 800;
            border: 2px solid #166534;
            box-shadow: inset 0 -2px 0 rgba(0,0,0,0.2);
        }
        .fr-login-profile {
            display: flex;
            gap: 12px;
            margin-top: 12px;
            align-items: flex-start;
        }
        .fr-login-profile-col {
            flex: 1;
        }
        .fr-login-profile-title {
            font-size: 9px;
            text-transform: uppercase;
            letter-spacing: 1px;
            margin-bottom: 6px;
            color: #1f3b73;
        }
        .fr-login-avatar {
            width: 72px;
            height: 72px;
            border: 2px solid #2f3a87;
            border-radius: 10px;
            background: #ffffff;
            display: flex;
            align-items: center;
            justify-content: center;
            box-shadow: inset 0 -2px 0 rgba(0,0,0,0.15);
        }
        .fr-login-avatar img {
            width: 100%;
            height: 100%;
            object-fit: contain;
            image-rendering: pixelated;
        }
        .fr-login-badges,
        .fr-login-party {
            display: flex;
            flex-wrap: wrap;
            gap: 6px;
        }
        .fr-login-badges img {
            width: 30px;
            height: 30px;
            border-radius: 6px;
            border: 1px solid rgba(15,23,42,0.35);
            background: #ffffff;
            padding: 2px;
        }
        .fr-login-party img {
            width: 40px;
            height: 40px;
            image-rendering: pixelated;
            border-radius: 8px;
            border: 2px solid #1f3b73;
            background: #ffffff;
            padding: 2px;
        }
        .fr-login-empty {
            font-size: 9px;
            color: #475569;
        }
        .fr-login-card-link {
            position: absolute;
            inset: 0;
            z-index: 4;
            opacity: 0;
            background: transparent;
            border: none;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    params = st.query_params
    action_param = params.get("action")
    if action_param == ["continue"]:
        action_param = "continue"
    if action_param == ["new_game"]:
        action_param = "new_game"
    if action_param == "continue":
        st.session_state["show_login_menu"] = False
        st.session_state["nav_to"] = "Pokédex (Busca)"
        st.query_params.clear()
        st.rerun()
    if action_param == "new_game":
        st.session_state["confirm_new_game"] = True
        st.query_params.clear()
        st.rerun()

    st.markdown("<div class='fr-login-wrap'><div class='fr-login-layout'>", unsafe_allow_html=True)
    badges_html = (
        "".join(
            f"<img src='{badge_src}' alt='Insígnia'>"
            for badge_src in badge_srcs
            if badge_src
        )
        if badge_srcs
        else "<div class='fr-login-empty'>Sem insígnias</div>"
    )
    avatar_html = (
        f"<div class='fr-login-avatar'><img src='{avatar_src}' alt='Avatar do treinador'></div>"
        if avatar_src
        else "<div class='fr-login-empty'>Sem avatar</div>"
    )
    party_html = (
        "".join(
            f"<img src='{sprite}' alt='Pokémon da party'>"
            for sprite in party_sprites
            if sprite
        )
        if party_sprites
        else "<div class='fr-login-empty'>Party vazia</div>"
    )
    st.markdown(
        f"""
        <div id='fr_continue_card' class='fr-login-card fr-login-continue-card'>
            <div class='fr-login-title'>
                <span>Continue</span>
                <span class='fr-login-continue-badge'>V</span>
            </div>
            <div class='fr-login-info'>
                <div class='fr-login-grid'>
                    <div class='fr-login-label'>PLAYER</div>
                    <div class='fr-login-value'>{html.escape(trainer_name)}</div>
                    <div class='fr-login-label'>POKÉDEX</div>
                    <div class='fr-login-value'>{caught_count}</div>
                    <div class='fr-login-label'>VISTOS</div>
                    <div class='fr-login-value'>{len(user_data.get("seen", []))}</div>
                    <div class='fr-login-label'>BADGES</div>
                    <div class='fr-login-value'>{badge_count}</div>
                </div>
                <div class='fr-login-profile'>
                    <div class='fr-login-profile-col'>
                        <div class='fr-login-profile-title'>Avatar</div>
                        {avatar_html}
                    </div>
                    <div class='fr-login-profile-col'>
                        <div class='fr-login-profile-title'>Insígnias</div>
                        <div class='fr-login-badges'>
                            {badges_html}
                        </div>
                    </div>
                </div>
                <div class='fr-login-profile-title' style='margin-top: 10px;'>Party</div>
                <div class='fr-login-party'>
                    {party_html}
                </div>
            </div>

        </div>
        <div id='fr_new_game_card' class='fr-login-card fr-login-new-game-card'>
            <div class='fr-login-title'>
                <span>New Game</span>
            </div>
            <div class='fr-login-info fr-login-info-empty'></div>

        </div>
        """,
        unsafe_allow_html=True,
    )
    continue_clicked = st.button("Continuar jogo", key="fr_continue_card_button")
    new_game_clicked = st.button("Novo jogo", key="fr_new_game_card_button")

    if continue_clicked:
        st.session_state["show_login_menu"] = False
        st.session_state["nav_to"] = "Pokédex (Busca)"
        st.rerun()

    if new_game_clicked:
        st.session_state["confirm_new_game"] = True
        st.rerun()

    components.html(
        """
        <script>
        (function () {
            const parentDoc = window.parent && window.parent.document ? window.parent.document : document;

            function hookCard(cardId, buttonText) {
                const card = parentDoc.getElementById(cardId);
                if (!card) return;

                const buttons = Array.from(parentDoc.querySelectorAll("button"));
                const targetButton = buttons.find((btn) => btn.innerText.trim() === buttonText);
                if (!targetButton) return;

                targetButton.style.display = "none";
                card.style.cursor = "pointer";
                card.setAttribute("role", "button");
                card.setAttribute("tabindex", "0");

                const trigger = () => targetButton.click();
                card.addEventListener("click", trigger);
                card.addEventListener("keydown", (event) => {
                    if (event.key === "Enter" || event.key === " ") {
                        event.preventDefault();
                        trigger();
                    }
                });
            }

            hookCard("fr_continue_card", "Continuar jogo");
            hookCard("fr_new_game_card", "Novo jogo");
        })();
        </script>
        """,
        height=0,
    )


    # --- MENU + BOTÃO CONFIRMAR AO LADO (SEM click-detector) ---
    menu_col, confirm_col = st.columns([7, 1], gap="small")

    with menu_col:
        action = st.radio(
            "Ação",
            ["Continue", "Ok", "New Game"],
            index=0,
            key="fr_login_action",
            label_visibility="collapsed",
        )
        if st.button("Mystery Gift", key="fr_mystery_gift"):
            st.session_state["mystery_gift_prompt"] = True
            st.session_state["mystery_gift_confirmed"] = False

    with confirm_col:
        if st.button("Confirmar", type="primary", key="fr_login_confirm"):
            if action in ("Continue", "Ok"):
                st.session_state["show_login_menu"] = False
                st.session_state["nav_to"] = "Pokédex (Busca)"
                st.rerun()

            if action == "New Game":
                st.session_state["confirm_new_game"] = True
                st.rerun()

    if st.session_state.get("mystery_gift_prompt"):
        st.info("Você aceita o presente de Ezenek?")
        gift_cols = st.columns([1, 1], gap="small")
        with gift_cols[0]:
            if st.button("Confirmar presente", type="primary", key="fr_mystery_gift_confirm"):
                st.session_state["mystery_gift_confirmed"] = True
                st.session_state["mystery_gift_prompt"] = False
                st.rerun()
        with gift_cols[1]:
            if st.button("Cancelar", key="fr_mystery_gift_cancel"):
                st.session_state["mystery_gift_prompt"] = False
                st.rerun()

    if st.session_state.get("mystery_gift_confirmed"):
        st.success("Ezenek virá até você, aguarde.")


    if st.session_state.get("confirm_new_game"):
        st.warning("Isso vai apagar todos os seus dados salvos.")
        if st.button("Confirmar novo jogo", type="primary", key="fr_confirm_new_game"):
            st.session_state["user_data"] = get_empty_user_data()
            save_data_cloud(trainer_name, st.session_state["user_data"])
            st.session_state["confirm_new_game"] = False
            st.rerun()
        if st.button("Cancelar", key="fr_cancel_new_game"):
            st.session_state["confirm_new_game"] = False
            st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)





def render_intro_screen() -> None:
    title_src = comp_img_data_uri("Assets/titulo.png") or comp_img_data_uri("Assets/inicio.png")
    start_src = comp_img_data_uri("Assets/start.png")
    render_bgm("music/menu.mp3", volume=0.25)

    css = "\n".join([
        "<style>",
        ".gaal-intro{",
        "  position: fixed;",
        "  inset: 0;",
        "  width: 100vw;",
        "  height: 100vh;",
        "  background: #000;",
        "  overflow: hidden;",
        "  margin: 0;",
        "  padding: 0;",
        "  z-index: 9999;",
        "}",
        ".gaal-intro-title{",
        "  position: absolute;",
        "  inset: 0;",
        "  width: 100%;",
        "  height: 100%;",
        "  object-fit: cover;",
        "  object-position: center;",
        "  display: block;",
        "  z-index: 1;",
        "  filter: drop-shadow(0 16px 26px rgba(0,0,0,0.75));",
        "}",
        ".gaal-intro-start{",
        "  position: absolute;",
        "  left: 50%;",
        "  bottom: 8vh;",
        "  transform: translateX(-50%);",
        "  width: min(70vw, 720px);",
        "  height: auto;",
        "  display: block;",
        "  z-index: 2;",
        "  animation: gaalIntroBlink 1.05s ease-in-out infinite;",
        "  pointer-events: none;",
        "}",
        ".gaal-intro-skip-link{",
        "  position: fixed;",
        "  top: 24px;",
        "  right: 24px;",
        "  z-index: 10001;",
        "  display: inline-flex;",
        "  align-items: center;",
        "  justify-content: center;",
        "  font-family: \"Press Start 2P\", \"Trebuchet MS\", sans-serif !important;",
        "  background: linear-gradient(180deg, #ffd85c 0%, #ff9f1a 60%, #f57a00 100%) !important;",
        "  border: 3px solid #ffffff !important;",
        "  color: #2a1a00 !important;",
        "  font-weight: 700 !important;",
        "  letter-spacing: 1px !important;",
        "  text-transform: uppercase !important;",
        "  padding: 12px 18px !important;",
        "  border-radius: 999px !important;",
        "  box-shadow: 0 8px 0 rgba(0,0,0,0.45), inset 0 0 0 2px rgba(255,255,255,0.4) !important;",
        "  text-shadow: 2px 2px 0 rgba(0,0,0,0.35) !important;",
        "  transition: transform 0.12s ease, box-shadow 0.12s ease, filter 0.12s ease !important;",
        "  text-decoration: none !important;",
        "}",
        ".gaal-intro-skip-link:hover{",
        "  filter: brightness(1.08);",
        "}",
        ".gaal-intro-skip-link:active{",
        "  transform: translateY(2px);",
        "  box-shadow: 0 5px 0 rgba(0,0,0,0.45) !important;",
        "}",
        ".gaal-intro-press-link{",
        "  position: fixed;",
        "  left: 50%;",
        "  bottom: 8vh;",
        "  transform: translateX(-50%);",
        "  width: min(70vw, 720px);",
        "  height: 110px;",
        "  z-index: 10001;",
        "  background: transparent;",
        "  border: none;",
        "  color: transparent;",
        "  box-shadow: none;",
        "  opacity: 0;",
        "  cursor: pointer;",
        "  display: block;",
        "}",
        "@keyframes gaalIntroBlink{",
        "  0%, 45% { opacity: 0.1; }",
        "  55%, 100% { opacity: 0.95; }",
        "}",
        "</style>",
    ])
    st.markdown(css, unsafe_allow_html=True)

    html = (
        '<div class="gaal-intro">'
        f'<img class="gaal-intro-title" src="{title_src}" alt="Ga&#39;Al" />'
        f'<img class="gaal-intro-start" src="{start_src}" alt="Press Start" />'
        '<a class="gaal-intro-press-link" href="?intro=start" target="_self" aria-label="Press Start"></a>'
        '<a class="gaal-intro-skip-link" href="?intro=skip" target="_self">Pular</a>'
        "</div>"
    )
    st.markdown(html, unsafe_allow_html=True)


# --- INTRO / LOGIN ---
# --- INTRO / LOGIN ---
if not st.session_state.get("intro_done", False):
    action_param = st.query_params.get("action")

    if action_param:
        st.session_state["intro_done"] = True
        st.rerun()

    intro_action = st.query_params.get("intro")
    if isinstance(intro_action, list):
        intro_action = intro_action[0] if intro_action else None

    if intro_action in {"skip", "start"}:
        st.session_state["intro_done"] = True
        st.query_params.clear()
        st.rerun()

    # mostra a intro
    render_intro_screen()
    st.stop()  # <<< IMPORTANTÍSSIMO: impede cair no login e tocar login.mp3


# --- TELA DE LOGIN ---
if "trainer_name" not in st.session_state:
    render_bgm("music/login.mp3", volume=0.25)

    st.title("Bem-vindo(a) ao Ga'Al")

    tab_login, tab_register = st.tabs(["🔑 Acessar", "📝 Cadastrar"])

    with tab_login:
        l_user = st.text_input("Nome do treinador", key="l_user")
        l_pass = st.text_input("Senha", type="password", key="l_pass")

        if st.button("Acessar", type="primary"):
            if l_user and l_pass:
                with st.spinner("Verificando credenciais..."):
                    result = authenticate_user(l_user, l_pass)

                    if result == "WRONG_PASS":
                        st.error("🚫 Senha incorreta. Verifique e tente novamente.")
                    elif result == "NOT_FOUND":
                        st.warning("Treinador não encontrado. Cadastre-se na aba ao lado.")
                    elif isinstance(result, dict):
                        st.session_state["trainer_name"] = l_user
                        st.session_state["user_data"] = result
                        st.session_state["show_login_menu"] = True
                        st.rerun()

    st.stop()  # opcional: trava aqui até logar, se o resto do app vem abaixo

    
    # ABA DE REGISTRO
    with tab_register:
        st.info("Crie sua conta para salvar seu progresso e sincronizar seus dados.")
        r_user = st.text_input("Nome do treinador", key="r_user")
        r_pass = st.text_input("Senha", type="password", key="r_pass")
        
        if st.button("Criar Conta"):
            if r_user and r_pass:
                with st.spinner("Registrando..."):
                    res = register_new_user(r_user, r_pass)
                    if res == "SUCCESS":
                        st.success("Conta criada! Vá na aba 'Entrar' para fazer login.")
                    elif res == "EXISTS":
                        st.error("Esse nome de treinador já está em uso.")
                    else:
                        st.error("Não foi possível criar a conta. Tente novamente.")
            else:
                st.warning("Preencha nome e senha para continuar.")
                
    st.stop() 

# --- APP PRINCIPAL ---

user_data = st.session_state['user_data']
trainer_name = st.session_state['trainer_name']

# --- FUNÇÕES DO APP ---
import random

def roll_die(db, rid: str, by: str, sides: int = 20):
    result = random.randint(1, int(sides))
    add_public_event(db, rid, "dice", by, {"sides": int(sides), "result": int(result)})
    return result
    
def get_role(room: dict, trainer_name: str) -> str:
    owner = (room.get("owner") or {}).get("name")
    
    # CORREÇÃO: Pega a lista de desafiantes (plural)
    challengers = room.get("challengers") or []
    # Cria uma lista apenas com os nomes dos desafiantes
    challenger_names = [c.get("name") for c in challengers]

    if trainer_name == owner:
        return "owner"
    if trainer_name in challenger_names:
        return "challenger"
    return "spectator"

def safe_doc_id(name: str) -> str:
    # Evita caracteres problemáticos no Firestore doc id
    if not isinstance(name, str):
        name = str(name)
    return re.sub(r"[^a-zA-Z0-9_\-\.]", "_", name).strip("_")[:80] or "user"


# ================================
# FIRESTORE SYNC (Perfil → Sala)
# - espelha dados do jogador (avatar/party/fichas) para:
#   users/<uid>
#   rooms/<rid>/players/<uid>
# Isso é o "contrato" para o site exclusivo da batalha.
# ================================

def _uid_from_trainer(trainer_name: str) -> str:
    # UID estável (server-side). No site (frontend) você pode trocar por Firebase Auth UID.
    return safe_doc_id(trainer_name or "user")

def _build_party_snapshot_for_room(db, trainer_name: str, user_data: dict, limit_sheets: int = 120) -> list[dict]:
    party_ids = []
    if isinstance(user_data, dict):
        raw = user_data.get("party")
        if isinstance(raw, list):
            party_ids = [str(x) for x in raw if str(x).strip()]

    # puxa fichas mais recentes e tenta casar por pokemon.id (pid)
    by_pid: dict[str, dict] = {}
    try:
        sheets = list_sheets(db, trainer_name, limit=limit_sheets)  # usa sua coleção trainers/<id>/sheets
        for sh in sheets:
            p = (sh.get("pokemon") or {})
            pid = str(p.get("id") or "").strip()
            if pid and pid not in by_pid:
                by_pid[pid] = {
                    "sheet_id": sh.get("_sheet_id"),
                    "pokemon": {"id": p.get("id"), "name": p.get("name"), "types": p.get("types")},
                    "np": sh.get("np"),
                    "updated_at": sh.get("updated_at"),
                }
    except Exception:
        pass

    out = []
    for pid in party_ids:
        base = {"pid": pid}
        if pid in by_pid:
            base.update(by_pid[pid])
        out.append(base)
    return out

def _build_avatar_snapshot(user_data: dict) -> dict:
    if not isinstance(user_data, dict):
        return {}
    prof = user_data.get("trainer_profile") if isinstance(user_data.get("trainer_profile"), dict) else {}
    # thumb é pequeno; não mande photo_b64 inteira
    return {
        "avatar_choice": prof.get("avatar_choice"),
        "photo_thumb_b64": prof.get("photo_thumb_b64"),
        "avatar_storage_path": prof.get("avatar_storage_path"),
        "avatar_url": prof.get("avatar_url"),
    }

def sync_user_profile_to_firestore(db, trainer_name: str, user_data: dict):
    uid = _uid_from_trainer(trainer_name)
    payload = {
        "uid": uid,
        "displayName": trainer_name,
        "avatar": _build_avatar_snapshot(user_data),
        "party": (user_data or {}).get("party", []) if isinstance(user_data, dict) else [],
        "updatedAt": firestore.SERVER_TIMESTAMP,
    }
    db.collection("users").document(uid).set(payload, merge=True)
    return uid

def sync_room_player_to_firestore(db, rid: str, trainer_name: str, role: str, user_data: dict):
    uid = _uid_from_trainer(trainer_name)
    party_snapshot = _build_party_snapshot_for_room(db, trainer_name, user_data)

    db.collection("rooms").document(rid).collection("players").document(uid).set(
        {
            "uid": uid,
            "trainer_name": trainer_name,
            "role": role,  # owner | challenger | spectator
            "avatar": _build_avatar_snapshot(user_data),
            "party_snapshot": party_snapshot,   # pronto pro site de batalha
            "joinedAt": firestore.SERVER_TIMESTAMP,
            "lastSeenAt": firestore.SERVER_TIMESTAMP,
            "updatedAt": firestore.SERVER_TIMESTAMP,
        },
        merge=True,
    )
    return uid

def sync_me_into_room(db, rid: str, trainer_name: str, role: str):
    # wrapper seguro pra chamar no lobby (onde user_data está em session_state)
    user_data = st.session_state.get("user_data") or {}
    sync_user_profile_to_firestore(db, trainer_name, user_data)
    sync_room_player_to_firestore(db, rid, trainer_name, role, user_data)

def get_item_image_url(item_name):
    if not item_name:
        return "https://raw.githubusercontent.com/PokeAPI/sprites/master/sprites/items/question-mark.png"
    
    # Normalização para Pokébolas e TMs
    name = normalize_text(item_name).replace(" ", "-")
    if "ball" in name and "-" not in name:
        name = name.replace("ball", "-ball")
    if "tm" in name and "-" not in name:
        name = name.replace("tm", "tm-")

    return f"https://raw.githubusercontent.com/PokeAPI/sprites/master/sprites/items/{name}.png"

def render_item_row(category_key, index, item_data, show_image=True):
    # Layout das colunas
    cols = st.columns([2.5, 1, 1.5]) if show_image else st.columns([3, 1])
    
    with cols[0]:
        new_name = st.text_input("Item", value=item_data.get("name", ""), 
                                key=f"it_nm_{category_key}_{index}", label_visibility="collapsed")
    
    with cols[-1]:
        new_qty = st.number_input("Qtd", min_value=0, value=item_data.get("qty", 0), 
                                 key=f"it_qt_{category_key}_{index}", label_visibility="collapsed")

    if show_image:
        with cols[1]:
            if new_name:
                img_url = get_item_image_url(new_name)
                # Fallback se a imagem não carregar
                st.markdown(f'<img src="{img_url}" width="35" style="image-rendering: pixelated;">', unsafe_allow_html=True)

    return {"name": new_name, "qty": new_qty}


def room_id_new() -> str:
    # curto e fácil de digitar
    return str(random.randint(100, 999))

def mark_pid_seen(db, rid, pid):
    # Adiciona o ID do Pokémon à lista de "vistos" no banco de dados
    # Usa ArrayUnion para não duplicar se já estiver lá
    ref = db.collection("rooms").document(rid).collection("public_state").document("state")
    # Tenta atualizar, se o documento não tiver o campo 'seen', o firestore cria na hora se usarmos set com merge, 
    # mas aqui vamos assumir que o state existe.
    try:
        ref.update({"seen": firestore.ArrayUnion([str(pid)])})
    except:
        # Fallback caso o campo não exista ainda
        ref.set({"seen": [str(pid)]}, merge=True)
def get_user_doc_ref(db, trainer_name: str):
    return db.collection("users").document(safe_doc_id(trainer_name))

def list_my_rooms(db, trainer_name: str):
    uref = get_user_doc_ref(db, trainer_name)
    udoc = uref.get()
    if not udoc.exists:
        return []
    data = udoc.to_dict() or {}
    return data.get("active_rooms", []) or []

def add_room_to_user(db, trainer_name: str, rid: str):
    uref = get_user_doc_ref(db, trainer_name)
    uref.set(
        {"active_rooms": firestore.ArrayUnion([rid]),
         "updatedAt": firestore.SERVER_TIMESTAMP},
        merge=True
    )


def get_room_player_names(room: dict) -> set:
    names = set()
    owner = (room.get("owner") or {}).get("name")
    if owner:
        names.add(owner)
    challengers = room.get("challengers") or []
    names.update([c.get("name") for c in challengers if c.get("name")])
    legacy_challenger = room.get("challenger")
    if isinstance(legacy_challenger, dict):
        legacy_name = legacy_challenger.get("name")
    else:
        legacy_name = legacy_challenger
    if legacy_name:
        names.add(legacy_name)
    return names

def delete_room(db, rid: str):
    room_ref = db.collection("rooms").document(rid)
    for doc in room_ref.collection("public_state").list_documents():
        doc.delete()
    for doc in room_ref.collection("public_events").list_documents():
        doc.delete()
    room_ref.delete()

def cleanup_room_if_orphaned(db, rid: str):
    room = get_room(db, rid)
    if not room:
        return
    player_names = get_room_player_names(room)
    if not player_names:
        return
    for name in player_names:
        if rid in list_my_rooms(db, name):
            return
    delete_room(db, rid)
            
def remove_room_from_user(db, trainer_name: str, rid: str):
    uref = get_user_doc_ref(db, trainer_name)
    uref.set(
        {"active_rooms": firestore.ArrayRemove([rid]),
         "updatedAt": firestore.SERVER_TIMESTAMP},
        merge=True
    )


def _strip_html_if_any(s: str) -> str:
    s = html.unescape(s or "")
    # se parece HTML, remove tags
    if "<" in s and ">" in s:
        s = re.sub(r"(?i)<br\s*/?>", "\n", s)
        s = re.sub(r"(?is)<[^>]+>", "", s)
    return s.strip()


# --- FUNÇÃO DE CALLBACK CORRIGIDA (CORREÇÃO DO BUG DE STATS 0) ---
def update_poke_state_callback(db, rid, trainer_name, pid, index):
    # CHAVE CORRIGIDA: Agora inclui o nome do treinador e o índice da party
    key_hp = f"hp_{trainer_name}_{pid}_{index}"
    key_cond = f"cond_{trainer_name}_{pid}_{index}"
    
    new_hp = st.session_state.get(key_hp)
    new_cond = st.session_state.get(key_cond)
    
    if new_hp is None: return

    ref = db.collection("rooms").document(rid).collection("public_state").document("party_states")
    data = {
        trainer_name: {
            str(pid): {
                "hp": int(new_hp),
                "cond": new_cond,
                "updatedAt": str(datetime.now())
            }
        }
    }
    ref.set(data, merge=True)
    
    # Lógica de Fainted no Mapa (Visual) - Mantém igual
    if new_hp == 0:
        state_ref = db.collection("rooms").document(rid).collection("public_state").document("state")
        stt = state_ref.get().to_dict() or {}
        pieces = stt.get("pieces") or []
        dirty = False
        for p in pieces:
            if p.get("owner") == trainer_name and str(p.get("pid")) == str(pid):
                if p.get("status") != "fainted":
                    p["status"] = "fainted"
                    dirty = True
        if dirty: state_ref.update({"pieces": pieces})
            
    elif new_hp > 0:
        state_ref = db.collection("rooms").document(rid).collection("public_state").document("state")
        stt = state_ref.get().to_dict() or {}
        pieces = stt.get("pieces") or []
        dirty = False
        for p in pieces:
            if p.get("owner") == trainer_name and str(p.get("pid")) == str(pid):
                if p.get("status") == "fainted":
                    p["status"] = "active"
                    dirty = True
        if dirty: state_ref.update({"pieces": pieces})


def create_room(db, trainer_name: str, grid_size: int, theme: str, max_active: int = 5):
    my_rooms = list_my_rooms(db, trainer_name)
    if len(my_rooms) >= max_active:
        return None, f"Você já tem {len(my_rooms)} arenas ativas (limite {max_active}). Finalize/arquive uma para criar outra."

    # --- LÓGICA DE CÓDIGO ÚNICO (Tenta 5 vezes achar um livre) ---
    rid = None
    room_ref = None
    
    for _ in range(5):
        candidate = room_id_new()
        ref = db.collection("rooms").document(candidate)
        if not ref.get().exists:
            rid = candidate
            room_ref = ref
            break
    
    if not rid:
        return None, "Servidor cheio ou erro ao gerar código. Tente novamente."
    # -------------------------------------------------------------

    room_ref.set({
        "createdAt": firestore.SERVER_TIMESTAMP,
        "updatedAt": firestore.SERVER_TIMESTAMP,
        "status": "lobby",
        "gridSize": int(grid_size),
        "theme": theme,
        "owner": {"name": trainer_name},
        "ownerUid": _uid_from_trainer(trainer_name),
        "challenger": None,
        "challengers": [],
        "spectators": [],
        "turn": "owner",
        "turnNumber": 1,
    })

    # estado público inicial
    room_ref.collection("public_state").document("battle").set(
        {
            "status": "idle",
            "logs": [],
            "initiative": {},
            "createdAt": firestore.SERVER_TIMESTAMP,
        "updatedAt": firestore.SERVER_TIMESTAMP,
            "updatedAt": firestore.SERVER_TIMESTAMP,
        },
        merge=True,
    )

    add_room_to_user(db, trainer_name, rid)
    return rid, None

def get_room(db, rid: str):
    ref = db.collection("rooms").document(rid)
    doc = ref.get()
    if not doc.exists:
        return None
    data = doc.to_dict() or {}
    data["_id"] = rid
    return data

def get_perspective_color(viewer_name, player_name, room_data):
    """
    Retorna a cor baseada na perspectiva do visualizador.
    Azul: Você | Vermelho: Oponente 1 | Amarelo: Oponente 2 | Rosa: Oponente 3
    """
    if viewer_name == player_name:
        return (0, 150, 255) # Azul (Sempre você)
    
    owner = (room_data.get("owner") or {}).get("name")
    challengers = [c.get("name") for c in (room_data.get("challengers") or [])]
    
    # Criar lista de "outros" (todos menos o espectador)
    others = [owner] + challengers
    if viewer_name in others:
        others.remove(viewer_name)
    
    # Atribui cores aos oponentes na ordem em que aparecem
    opp_colors = [
        (255, 50, 50),   # Vermelho (Opp 1)
        (255, 215, 0),  # Amarelo (Opp 2)
        (255, 105, 180)  # Rosa (Opp 3)
    ]
    
    try:
        idx = others.index(player_name)
        return opp_colors[idx] if idx < len(opp_colors) else (200, 200, 200)
    except ValueError:
        return (200, 200, 200) # Cor neutra para espectadores


def join_room_as_challenger(db, rid: str, trainer_name: str, max_challengers: int = 4):
    ref = db.collection("rooms").document(rid)
    doc = ref.get()
    if not doc.exists:
        return "NOT_FOUND"

    data = doc.to_dict() or {}
    owner = (data.get("owner") or {}).get("name")
    # Agora lidamos com uma lista de desafiantes
    challengers = data.get("challengers") or []
    
    # Se já estiver na lista ou for o dono
    if owner == trainer_name:
        add_room_to_user(db, trainer_name, rid)
        return "ALREADY_OWNER"
    
    if any(c.get("name") == trainer_name for c in challengers):
        add_room_to_user(db, trainer_name, rid)
        return "ALREADY_CHALLENGER"

    # Verifica se ainda há vaga (até 4 desafiantes)
    if len(challengers) < max_challengers:
        new_challenger = {"name": trainer_name}
        ref.update({
            "challengers": firestore.ArrayUnion([new_challenger]),
            "status": "running",
        })
        add_room_to_user(db, trainer_name, rid)
        return "OK"

    return "ARENA_FULL"




def join_room_as_spectator(db, rid: str, trainer_name: str):
    ref = db.collection("rooms").document(rid)
    doc = ref.get()
    if not doc.exists:
        return "NOT_FOUND"

    data = doc.to_dict() or {}
    owner = (data.get("owner") or {}).get("name")
    challenger = (data.get("challenger") or {}).get("name") if isinstance(data.get("challenger"), dict) else data.get("challenger")

    if trainer_name in [owner, challenger]:
        add_room_to_user(db, trainer_name, rid)
        return "PLAYER"

    ref.update({
        "spectators": firestore.ArrayUnion([trainer_name]),
    })
    # evento público
    ref.collection("public_events").add({
        "type": "join_spectator",
        "by": trainer_name,
        "payload": {"room": rid},
        "ts": firestore.SERVER_TIMESTAMP,
    })
    add_room_to_user(db, trainer_name, rid)
    return "OK"

def add_public_event(db, rid: str, event_type: str, by: str, payload: dict):
    db.collection("rooms").document(rid).collection("public_events").add({
        "type": event_type,
        "by": by,
        "payload": payload or {},
        "ts": firestore.SERVER_TIMESTAMP,
    })

    # ✅ bump do estado: cria se não existir e atualiza updatedAt
    db.collection("rooms").document(rid).collection("public_state").document("state").set({
        "updatedAt": firestore.SERVER_TIMESTAMP
    }, merge=True)

    
def state_ref_for(db, rid: str):
    return (
        db.collection("rooms")
          .document(rid)
          .collection("public_state")
          .document("state")
    )

def get_state(db, rid: str) -> dict:
    doc = state_ref_for(db, rid).get()
    return doc.to_dict() if doc.exists else {}

def update_party_state(db, rid, trainer_name, pid, hp, conditions):
    ref = db.collection("rooms").document(rid).collection("public_state").document("party_states")
    
    key = f"{trainer_name}.{pid}"
    data = {
        key: {
            "hp": int(hp),
            "cond": conditions,
            "updatedAt": str(datetime.now()) # Timestamp local serve, mas muda o dado
        }
    }
    # Força um campo global de update para disparar o listener da coleção/documento
    data["last_update"] = firestore.SERVER_TIMESTAMP 
    
    ref.set(data, merge=True)
    
    # Se o HP for 0, precisamos atualizar a peça no tabuleiro para 'fainted' (se ela estiver lá)
    if hp == 0:
        # Busca peças desse treinador e desse PID
        state_doc = db.collection("rooms").document(rid).collection("public_state").document("state").get()
        if state_doc.exists:
            pieces = state_doc.to_dict().get("pieces", [])
            for p in pieces:
                if p.get("owner") == trainer_name and str(p.get("pid")) == str(pid):
                    p["status"] = "fainted"
                    upsert_piece(db, rid, p)
    # Se HP > 0 e estava fainted, revive
    elif hp > 0:
        state_doc = db.collection("rooms").document(rid).collection("public_state").document("state").get()
        if state_doc.exists:
            pieces = state_doc.to_dict().get("pieces", [])
            for p in pieces:
                if p.get("owner") == trainer_name and str(p.get("pid")) == str(pid) and p.get("status") == "fainted":
                    p["status"] = "active"
                    upsert_piece(db, rid, p)

def upsert_piece(db, rid: str, piece: dict):
    # Garante que o ID existe
    sref = state_ref_for(db, rid)
    # Precisamos ler o estado atual para não perder as outras peças
    # (Ou usar arrayUnion se fosse uma lista simples, mas é lista de dicts)
    
    # Melhor abordagem para evitar race conditions em real-time: Transação
    # Mas para simplificar no seu código atual:
    stt = get_state(db, rid)
    pieces = stt.get("pieces") or []

    # Remove a versão antiga da peça, se existir
    new_pieces = [p for p in pieces if p.get("id") != piece.get("id")]
    new_pieces.append(piece)

    sref.set({
        "pieces": new_pieces,
        "updatedAt": firestore.SERVER_TIMESTAMP,  # ✅
    }, merge=True)

def delete_piece(db, rid: str, piece_id: str):
    sref = state_ref_for(db, rid)
    stt = get_state(db, rid)
    pieces = stt.get("pieces") or []
    new_pieces = [p for p in pieces if p.get("id") != piece_id]
    sref.set({
        "pieces": new_pieces,
        "updatedAt": firestore.SERVER_TIMESTAMP,
    }, merge=True)

def find_piece_at(pieces: list[dict], row: int, col: int) -> dict | None:
    for p in pieces or []:
        if int(p.get("row", -1)) == int(row) and int(p.get("col", -1)) == int(col):
            return p
    return None


def visible_pieces_for(room: dict, viewer_name: str, pieces: list[dict]) -> list[dict]:
    # Jogador vê tudo dele; vê do oponente só o que estiver "revealed" (no campo)
    role = get_role(room, viewer_name)
    owner = (room.get("owner") or {}).get("name")
    chal = room.get("challenger") or {}
    chal_name = chal.get("name") if isinstance(chal, dict) else (chal or None)

    if role == "spectator":
        # espectador vê somente o que está no campo (revealed)
        return [p for p in pieces if p.get("revealed") is True]

    # jogador: vê os dele sempre; e do outro só se revealed
    me = owner if role == "owner" else chal_name
    out = chal_name if role == "owner" else owner

    result = []
    for p in pieces:
        if p.get("owner") == me:
            result.append(p)
        elif p.get("owner") == out and p.get("revealed") is True:
            result.append(p)
    return result

def list_public_events(db, rid: str, limit: int = 30):
    q = (db.collection("rooms").document(rid)
         .collection("public_events")
         .order_by("ts", direction=firestore.Query.DESCENDING)
         .limit(limit))
    return [d.to_dict() for d in q.stream()]
# =========================
# MAPA TÁTICO (3.1 / 3.2 / 3.3)
# =========================

TILE_SIZE = 88

THEMES = {
    "cave_water": {"base": "rock", "border": "wall"},
    "forest": {"base": "grass", "border": "tree"},
    "mountain_slopes": {"base": "stone", "border": "peak"},
    "plains": {"base": "grass", "border": "bush"},
    "dirt": {"base": "dirt", "border": "rock"},
    "river": {"base": "grass", "border": "tree"},
    "sea_coast": {"base": "sand", "border": "sea"},
    "center_lake": {"base": "grass", "border": "tree"},

# --- BIOMAS (novos) ---
"biome_grass": {"base": "grass", "border": "bush"},
"biome_forest": {"base": "grass", "border": "tree"},
"biome_meadow": {"base": "grass", "border": "flower"},
"biome_desert": {"base": "sand", "border": "rock"},
"biome_mountain": {"base": "stone", "border": "peak"},
"biome_snow": {"base": "snow", "border": "rock"},
"biome_water": {"base": "sand", "border": "sea"},
"biome_cave": {"base": "cave", "border": "wall"},
"biome_mix": {"base": "grass", "border": "tree"},
}


@st.cache_resource(show_spinner=False)
def get_biome_generator():
    base_dir = Path(__file__).resolve().parent
    assets_root = base_dir / "Novo Gerador de mapas"
    return BiomeGenerator(assets_root=assets_root)




def map_theme_to_biome(theme_key: str, no_water: bool) -> str:
    """Converte a 'theme' da sala (UI) para o 'biome' do BiomeGenerator v2.

    IMPORTANTE:
      - Esta função retorna *chaves* válidas de BIOME_CONFIG do biome_generator.py
        (ex.: grasslands, deepforest, desert, beach, snowlands, cave, mines, temple,
         seafloor, interior, lake, river).
      - O parâmetro no_water força biomas aquáticos a virarem um equivalente terrestre.
    """
    key = (theme_key or "").lower().strip()

    # Quando a sala já usa as chaves nativas do BiomeGenerator,
    # mantém o bioma exatamente como selecionado.
    if key in BIOME_CONFIG:
        if no_water and key in ["river", "lake"]:
            return "grasslands"
        if no_water and key in ["beach", "seafloor"]:
            return "desert"
        return key

    # --- Florestas / campos
    if key in ["forest", "biome_forest"]:
        return "deepforest"
    if key in ["plains", "biome_grass", "biome_meadow", "biome_mix", "biome_forest_floor"]:
        return "grasslands"

    # --- Deserto / neve / montanha
    if key in ["biome_desert"]:
        return "desert"
    if key in ["biome_snow"]:
        return "snowlands"
    if key in ["mountain_slopes", "biome_mountain", "dirt", "stone"]:
        # Não existe 'mountain' separado no BIOME_CONFIG: 'mines' dá o visual rochoso.
        return "mines"

    # --- Caverna
    if key in ["cave_water", "biome_cave", "cave"]:
        return "cave"

    # --- Água (com fallback para terrestre)
    if key in ["river"]:
        return "grasslands" if no_water else "river"
    if key in ["center_lake", "lake"]:
        return "grasslands" if no_water else "lake"
    if key in ["sea_coast", "biome_water", "sea", "coast"]:
        return "desert" if no_water else "beach"

    # Padrão
    return "grasslands"

def generate_biome_seed(seed: int | None = None) -> int:
    if seed is not None:
        return int(seed)
    return random.randint(1, 999999999)


def place_forest_objects(tiles, grid, rng, inside_fn):
    """
    Coloca objetos de floresta em duas classes:
    - árvores: bem espaçadas (blue-noise/Poisson disk simplificado)
    - moitas: mais densas perto das árvores, mas sem colar nelas
    Requer: tiles[r][c] já ter chão ('grass', 'forest_floor' etc.).
    """
    # Ajuste fino (padrão bom para grid 16..24)
    tree_min_dist = 3 if grid >= 16 else 2         # distância mínima entre árvores (em tiles)
    tree_target = int(grid * grid * 0.10)          # ~10% árvores (pode subir p/ 0.12)
    bush_target = int(grid * grid * 0.18)          # ~18% moitas (pode subir p/ 0.22)
    clearings = int(grid * grid * 0.10)            # ~10% clareiras (menos “parede”)

    # Onde pode colocar objeto?
    def is_walkable_floor(t):
        # Ajuste aqui se seus nomes forem diferentes
        return t in ("grass", "forest_floor", "sand")  # sand só se você quiser floresta costeira

    # 1) gera máscara de clareiras (respiro jogável)
    open_mask = [[False]*grid for _ in range(grid)]
    # abre alguns "blobs" de clareira
    for _ in range(max(1, grid // 4)):
        cr = rng.integers(2, grid-2)
        cc = rng.integers(2, grid-2)
        rad = 2 if grid >= 16 else 1
        for r in range(cr-rad, cr+rad+1):
            for c in range(cc-rad, cc+rad+1):
                if _in_bounds(r, c, grid) and inside_fn(r, c) and rng.random() > 0.25:
                    open_mask[r][c] = True

    # 2) Lista de candidatos (embaralhada)
    candidates = [(r, c) for r in range(1, grid-1) for c in range(1, grid-1)
                  if inside_fn(r, c) and is_walkable_floor(tiles[r][c]) and not open_mask[r][c]]
    rng.shuffle(candidates)

    # 3) coloca árvores com distância mínima (Poisson simplificado)
    trees = []
    min_d2 = tree_min_dist * tree_min_dist

    for (r, c) in candidates:
        if len(trees) >= tree_target:
            break
        ok = True
        for (tr, tc) in trees:
            if (r-tr)*(r-tr) + (c-tc)*(c-tc) < min_d2:
                ok = False
                break
        if ok:
            trees.append((r, c))

    # 4) coloca moitas:
    # - mais chance perto das árvores
    # - evita colar em árvore (1 tile de margem)
    bushes = set()
    tree_set = set(trees)

    # mapa “influência” por proximidade de árvore (barato: usa vizinhança 8 até raio 3)
    influence = [[0]*grid for _ in range(grid)]
    for tr, tc in trees:
        for rr in range(tr-3, tr+4):
            for cc in range(tc-3, tc+4):
                if _in_bounds(rr, cc, grid):
                    d2 = (rr-tr)*(rr-tr) + (cc-tc)*(cc-tc)
                    if d2 <= 9:  # raio 3
                        influence[rr][cc] += (10 - d2)  # mais perto = mais influência

    # candidatos para moitas: inclui área aberta também, mas com chance menor
    bush_candidates = [(r, c) for r in range(1, grid-1) for c in range(1, grid-1)
                       if inside_fn(r, c) and is_walkable_floor(tiles[r][c])]

    rng.shuffle(bush_candidates)

    def near_tree_margin(r, c):
        # não cola moita em árvore (margem 1 tile)
        for nr, nc in _neighbors8(r, c):
            if (nr, nc) in tree_set:
                return True
        return False

    for (r, c) in bush_candidates:
        if len(bushes) >= bush_target:
            break
        if (r, c) in tree_set:
            continue
        if near_tree_margin(r, c):
            continue

        # probabilidade:
        # - perto de árvore => maior
        # - em clareira => menor
        inf = influence[r][c]
        p = 0.10  # base
        p += min(0.35, inf / 80.0)  # boost por influência (capado)
        if open_mask[r][c]:
            p *= 0.35  # clareira: reduz muito

        if rng.random() < p:
            bushes.add((r, c))

    # 5) aplica no grid de overlay (se seu sistema usa overlay)
    # Se você NÃO tem overlay separado, você pode marcar tiles[r][c] como "tree"/"bush"
    # Aqui vou retornar listas para o seu render usar.
    return trees, list(bushes), open_mask


def _gen_tiles_legacy(grid: int, theme_key: str, seed: int | None = None, no_water: bool = False):
    # REGRA: Bloqueia água em 6x6, exceto se o tema tiver "water", "river", "lake" ou "sea" no nome
    themes_com_agua = ["water", "river", "lake", "sea", "coast"]
    if grid <= 6 and not any(word in theme_key.lower() for word in themes_com_agua):
        no_water = True
    if seed is None:
        seed = random.randint(1, 999999999)

    rng = random.Random(seed)
    theme = THEMES.get(theme_key, THEMES["cave_water"])
    base = theme["base"]
    tiles = [[base for _ in range(grid)] for _ in range(grid)]

    def inside(r, c):
        return 1 <= r <= grid - 2 and 1 <= c <= grid - 2

    # pedras leves em todos os temas (dá textura)
    for _ in range(rng.randint(grid, grid * 2)):
        rr = rng.randint(1, grid - 2)
        cc = rng.randint(1, grid - 2)
        if inside(rr, cc) and rng.random() > 0.75:
            if tiles[rr][cc] in ["grass", "dirt", "sand", "trail"]:
                tiles[rr][cc] = "rock"

    # --- features por tema ---
    if theme_key == "cave_water":
        if not no_water:
            pools = rng.randint(1, 2)
            for _ in range(pools):
                cr = rng.randint(2, grid - 3)
                cc = rng.randint(2, grid - 3)
                rad = rng.randint(1, 2)

                for rr in range(cr - rad, cr + rad + 1):
                    for cc2 in range(cc - rad, cc + rad + 1):
                        if inside(rr, cc2) and rng.random() > 0.25:
                            tiles[rr][cc2] = "water"

        # elementos sólidos continuam existindo
        spikes = rng.randint(1, max(2, grid - 3))
        for _ in range(spikes):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if inside(rr, cc) and tiles[rr][cc] == base:
                tiles[rr][cc] = "stalagmite"

    elif theme_key == "forest":
        # “carpete” de grama
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "grass" if rng.random() > 0.15 else "bush"
    
        # árvores espalhadas
        trees = rng.randint(grid, grid * 2)
        for _ in range(trees):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if inside(rr, cc) and tiles[rr][cc] in ["grass", "bush"] and rng.random() > 0.35:
                tiles[rr][cc] = "tree"

        # caminho opcional
        if rng.random() > 0.35:
            r = rng.randint(2, grid - 3)
            for c in range(1, grid - 1):
                if inside(r, c) and tiles[r][c] != "tree":
                    tiles[r][c] = "path"

        # água só se permitido
        if not no_water:
            ponds = rng.randint(0, 2)
            for _ in range(ponds):
                cr = rng.randint(2, grid - 3)
                cc = rng.randint(2, grid - 3)
                rad = rng.randint(1, 2)
                for rr in range(cr - rad, cr + rad + 1):
                    for cc2 in range(cc - rad, cc + rad + 1):
                        if inside(rr, cc2) and rng.random() > 0.35:
                            tiles[rr][cc2] = "water"
    
    elif theme_key == "mountain_slopes":
        # base rochosa
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "stone" if rng.random() > 0.25 else "rock"

        # faixas de declive (diagonais)
        bands = rng.randint(2, 4)
        for _ in range(bands):
            start_r = rng.randint(1, grid - 2)
            for c in range(1, grid - 1):
                rr = start_r + (c // 2)
                if inside(rr, c) and rng.random() > 0.25:
                    tiles[rr][c] = "slope1" if rng.random() > 0.5 else "slope2"

        # picos
        peaks = rng.randint(1, 3)
        for _ in range(peaks):
            rr = rng.randint(2, grid - 3)
            cc = rng.randint(2, grid - 3)
            tiles[rr][cc] = "peak"

    elif theme_key == "plains":
        for _ in range(rng.randint(2, grid * 2)):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if inside(rr, cc) and rng.random() > 0.5:
                tiles[rr][cc] = "flower"

        if rng.random() > 0.5:
            c = rng.randint(2, grid - 3)
            for r in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "trail"

    elif theme_key == "dirt":
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c) and rng.random() > 0.85:
                    tiles[r][c] = "stone"

        if rng.random() > 0.4:
            r = rng.randint(2, grid - 3)
            for c in range(1, grid - 1):
                if inside(r, c) and tiles[r][c] == base:
                    tiles[r][c] = "rut"

    elif theme_key == "river":
        r = rng.randint(1, grid - 2)
        width = 2 if grid >= 8 else 1
    
        for c in range(1, grid - 1):
            for w in range(width):
                rr = r + w
                if 1 <= rr <= grid - 2:
                    tiles[rr][c] = ("water" if not no_water else "trail")
    
            # margens
            if r - 1 >= 1 and rng.random() > 0.35:
                tiles[r - 1][c] = "sand" if not no_water else "stone"
            if r + width <= grid - 2 and rng.random() > 0.35:
                tiles[r + width][c] = "sand" if not no_water else "stone"
    
            step = rng.choice([-1, 0, 1])
            r = max(1, min(grid - 2 - (width - 1), r + step))

        
    elif theme_key == "sea_coast":
        for r in range(grid):
            if not no_water:
                tiles[r][0] = "sea"
                if grid > 4:
                    tiles[r][1] = "sea" if rng.random() > 0.25 else "sand"
            else:
                tiles[r][0] = "sand"
                if grid > 4:
                    tiles[r][1] = "sand" if rng.random() > 0.25 else "stone"
    
        for _ in range(rng.randint(1, 4)):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(2, grid - 2)
            if inside(rr, cc) and rng.random() > 0.45:
                tiles[rr][cc] = "rock"
    
    elif theme_key == "center_lake":
        cr = grid // 2
        cc = grid // 2
        rad = 2 if grid >= 8 else 1
        for rr in range(cr - rad, cr + rad + 1):
            for cc2 in range(cc - rad, cc + rad + 1):
                if inside(rr, cc2) and (abs(rr - cr) + abs(cc2 - cc) <= rad + 1):
                    tiles[rr][cc2] = ("water" if not no_water else "stone")

   
    
    elif theme_key == "biome_grass":
        # Campos / Rotas gramadas: várias gramas + trilha + detalhes
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "grass" if rng.random() > 0.10 else "bush"

        # trilha serpenteando
        if rng.random() > 0.25:
            r0 = rng.randint(2, grid - 3)
            for c in range(1, grid - 1):
                if inside(r0, c) and tiles[r0][c] != "tree":
                    tiles[r0][c] = "trail"
                r0 = max(1, min(grid - 2, r0 + rng.choice([-1, 0, 1])))

        # flores leves
        for _ in range(rng.randint(grid, grid * 2)):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if inside(rr, cc) and tiles[rr][cc] == "grass" and rng.random() > 0.55:
                tiles[rr][cc] = "flower"
    
    elif theme_key == "biome_forest":
        # Floresta (melhorada): manchas densas + clareiras + sub-bosque (bush) + caminho opcional
    
        # 1) Base: quase tudo grama
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "grass"
    
        def n4(rr, cc):
            for dr, dc in [(-1,0),(1,0),(0,-1),(0,1)]:
                yield rr+dr, cc+dc
    
        # 2) Clareiras (2–3): regiões que inibem árvores
        clear = set()
        n_clear = rng.randint(1, 2) if grid <= 10 else rng.randint(2, 3)
        for _ in range(n_clear):
            cr = rng.randint(2, grid - 3)
            cc = rng.randint(2, grid - 3)
            rad = rng.randint(1, 2) if grid <= 10 else rng.randint(2, 3)
            for rr in range(cr - rad, cr + rad + 1):
                for cc2 in range(cc - rad, cc + rad + 1):
                    if inside(rr, cc2) and rng.random() > 0.20:
                        clear.add((rr, cc2))
    
        # 3) Clusters densos de árvore (2–4 centros)
        forest = set()
        centers = rng.randint(3, 4) if grid <= 10 else rng.randint(5, 7)
    
        for _ in range(centers):
            sr = rng.randint(2, grid - 3)
            sc = rng.randint(2, grid - 3)
    
            steps = rng.randint(grid * 5, grid * 9)  # quanto maior, mais “mato fechado”
            rr, cc = sr, sc
            for _s in range(steps):
                if inside(rr, cc) and (rr, cc) not in clear:
                    forest.add((rr, cc))
    
                # caminhada tendendo a formar “manchas”
                dr, dc = rng.choice([(-1,0),(1,0),(0,-1),(0,1),(0,0)])
                rr = max(1, min(grid - 2, rr + dr))
                cc = max(1, min(grid - 2, cc + dc))
    
                # engrossa o cluster às vezes
                if rng.random() > 0.35:
                    for nr, nc in n4(rr, cc):
                        if inside(nr, nc) and (nr, nc) not in clear and rng.random() > 0.35:
                            forest.add((nr, nc))
    
        # 4) Aplica árvores do cluster
        for (rr, cc) in forest:
            tiles[rr][cc] = "tree"
    
        # 5) Árvores “soltas” extra (bordas + preenchimento)
        extra = rng.randint(grid * 6, grid * 9) if grid <= 10 else rng.randint(grid * 10, grid * 16)
        for _ in range(extra):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if not inside(rr, cc) or (rr, cc) in clear:
                continue
            if tiles[rr][cc] == "grass" and rng.random() > 0.25:
                tiles[rr][cc] = "tree"
    
        # 6) Arbustos (bush): bastante, principalmente perto de árvore
        bushes = rng.randint(grid * 6, grid * 10) if grid <= 10 else rng.randint(grid * 10, grid * 18)
        for _ in range(bushes):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if not inside(rr, cc) or (rr, cc) in clear:
                continue
            if tiles[rr][cc] != "grass":
                continue
    
            near_tree = False
            for nr, nc in n4(rr, cc):
                if 0 <= nr < grid and 0 <= nc < grid and tiles[nr][nc] == "tree":
                    near_tree = True
                    break
    
            p = 0.95 if near_tree else 0.50
            if rng.random() < p:
                tiles[rr][cc] = "bush"
    
        # 7) Caminho opcional: estreito, não “limpa” toda a floresta
        if rng.random() > 0.35:
            r = rng.randint(2, grid - 3)
            width = 1 if grid <= 10 else rng.choice([1, 2])
            for c in range(1, grid - 1):
                for w in range(width):
                    rr = r + w
                    if inside(rr, c) and tiles[rr][c] != "water" and tiles[rr][c] != "sea":
                        tiles[rr][c] = "path"
                # leve serpente
                r = max(1, min(grid - 2 - (width - 1), r + rng.choice([-1, 0, 1])))
    
        # 8) Poça opcional (se permitido): pequena e discreta
        if not no_water and rng.random() > 0.70:
            cr = rng.randint(2, grid - 3)
            cc = rng.randint(2, grid - 3)
            rad = 1 if grid <= 10 else 2
            for rr in range(cr - rad, cr + rad + 1):
                for cc2 in range(cc - rad, cc + rad + 1):
                    if inside(rr, cc2) and (rr, cc2) not in clear and rng.random() > 0.35:
                        tiles[rr][cc2] = "water"
    
        # 9) Mantém clareiras: limpa árvore/arbusto nelas (no fim)
        for (rr, cc) in clear:
            if inside(rr, cc) and tiles[rr][cc] in ("tree", "bush"):
                tiles[rr][cc] = "grass"
    

    elif theme_key == "biome_meadow":
        # Meadow / campo florido: MUITAS flores e plantas
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "grass"

        # grandes manchas de flores
        clusters = rng.randint(2, 4)
        for _ in range(clusters):
            cr = rng.randint(2, grid - 3)
            cc = rng.randint(2, grid - 3)
            rad = rng.randint(1, 2)
            for rr in range(cr - rad, cr + rad + 1):
                for cc2 in range(cc - rad, cc + rad + 1):
                    if inside(rr, cc2) and rng.random() > 0.30:
                        tiles[rr][cc2] = "flower"

        # ainda mais flores espalhadas
        for _ in range(rng.randint(grid * 2, grid * 4)):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if inside(rr, cc) and tiles[rr][cc] == "grass" and rng.random() > 0.45:
                tiles[rr][cc] = "flower"

        # alguns arbustos
        for _ in range(rng.randint(grid, grid * 2)):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if inside(rr, cc) and tiles[rr][cc] == "grass" and rng.random() > 0.65:
                tiles[rr][cc] = "bush"

    elif theme_key == "biome_desert":
        # Deserto / árido: areia + variações de chão seco
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "sand" if rng.random() > 0.12 else "dirt"

        # pedregulho/rochas
        for _ in range(rng.randint(grid, grid * 2)):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if inside(rr, cc) and rng.random() > 0.55:
                tiles[rr][cc] = "rock"

        # oásis rarinho
        if not no_water and rng.random() > 0.75:
            cr = rng.randint(2, grid - 3)
            cc = rng.randint(2, grid - 3)
            for rr in range(cr - 1, cr + 2):
                for cc2 in range(cc - 1, cc + 2):
                    if inside(rr, cc2) and rng.random() > 0.15:
                        tiles[rr][cc2] = "water"

    elif theme_key == "biome_mountain":
        # Montanha / rochoso: pedra/pedregulho + "cliffs" (simulados) + picos
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "stone" if rng.random() > 0.30 else "rock"

        # trilha/estrada estreita
        if rng.random() > 0.35:
            c0 = rng.randint(2, grid - 3)
            for r in range(1, grid - 1):
                if inside(r, c0) and rng.random() > 0.15:
                    tiles[r][c0] = "trail"
                c0 = max(1, min(grid - 2, c0 + rng.choice([-1, 0, 1])))

        # picos
        for _ in range(rng.randint(1, 3)):
            rr = rng.randint(2, grid - 3)
            cc = rng.randint(2, grid - 3)
            tiles[rr][cc] = "peak"

    elif theme_key == "biome_snow":
        # Neve / gelo: tiles brancos + variações
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "snow"

        # manchas de pedra/rocha pra quebrar o branco
        for _ in range(rng.randint(grid, grid * 2)):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if inside(rr, cc) and rng.random() > 0.60:
                tiles[rr][cc] = "rock"

    elif theme_key == "biome_water":
        # Água: rio/lago/mar + costa
        # Melhorado: lago "blob" + rio orgânico + margem dupla (shoreline)

        # --- base (terra) ---
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    tiles[r][c] = "sand" if rng.random() > 0.30 else "grass"

        def _neighbors4(rr, cc):
            for dr, dc in [(-1, 0), (1, 0), (0, -1), (0, 1)]:
                yield rr + dr, cc + dc

        def _neighbors8(r, c):
            for dr in (-1, 0, 1):
                for dc in (-1, 0, 1):
                    if dr == 0 and dc == 0:
                        continue
                    yield r + dr, c + dc
        
        def _dist2(a, b):
            (r1, c1), (r2, c2) = a, b
            dr = r1 - r2
            dc = c1 - c2
            return dr*dr + dc*dc
        
        def _in_bounds(r, c, grid):
            return 0 <= r < grid and 0 <= c < grid


        def _apply_double_shore():
            """
            Margem dupla:
            - ring1: vizinho direto da água -> quase sempre areia
            - ring2: vizinho do ring1 -> mistura areia/grama
            """
            ring1 = set()
            ring2 = set()

            # ring1: tudo que toca água (4-dir)
            for rr in range(1, grid - 1):
                for cc in range(1, grid - 1):
                    if tiles[rr][cc] in ("water", "sea"):
                        for nr, nc in _neighbors4(rr, cc):
                            if 1 <= nr <= grid - 2 and 1 <= nc <= grid - 2 and inside(nr, nc):
                                if tiles[nr][nc] not in ("water", "sea"):
                                    ring1.add((nr, nc))

            # ring2: tudo que toca ring1 (4-dir), excluindo ring1 e água
            for rr, cc in list(ring1):
                for nr, nc in _neighbors4(rr, cc):
                    if 1 <= nr <= grid - 2 and 1 <= nc <= grid - 2 and inside(nr, nc):
                        if (nr, nc) not in ring1 and tiles[nr][nc] not in ("water", "sea"):
                            ring2.add((nr, nc))

            # aplica ring1: areia quase sempre (costa)
            for rr, cc in ring1:
                tiles[rr][cc] = "sand" if rng.random() > 0.08 else "grass"

            # aplica ring2: mistura (transição)
            for rr, cc in ring2:
                if (rr, cc) in ring1:
                    continue
                tiles[rr][cc] = "sand" if rng.random() > 0.55 else "grass"

        def _carve_blob_lake():
            # centro levemente deslocado para parecer natural
            cr = grid // 2 + rng.choice([-1, 0, 1])
            cc = grid // 2 + rng.choice([-1, 0, 1])

            # raios do elipse (adaptam ao grid)
            rx = max(2, min(grid // 3, rng.randint(3, 5)))
            ry = max(2, min(grid // 3, rng.randint(3, 5)))

            # máscara elíptica + jitter
            water = [[False] * grid for _ in range(grid)]
            for rr in range(1, grid - 1):
                for cc2 in range(1, grid - 1):
                    dx = (cc2 - cc) / float(rx)
                    dy = (rr - cr) / float(ry)
                    d = dx * dx + dy * dy
                    jitter = (rng.random() - 0.5) * 0.25
                    if d <= 1.0 + jitter:
                        water[rr][cc2] = True

            # smooth 1-2 passadas (remove espículas)
            for _ in range(rng.randint(1, 2)):
                w2 = [[False] * grid for _ in range(grid)]
                for rr in range(1, grid - 1):
                    for cc2 in range(1, grid - 1):
                        cnt = 0
                        for nr, nc in _neighbors8(rr, cc2):
                            if 0 <= nr < grid and 0 <= nc < grid and water[nr][nc]:
                                cnt += 1
                        if water[rr][cc2]:
                            w2[rr][cc2] = cnt >= 2
                        else:
                            w2[rr][cc2] = cnt >= 6
                water = w2

            # aplica água
            for rr in range(1, grid - 1):
                for cc2 in range(1, grid - 1):
                    if water[rr][cc2] and inside(rr, cc2):
                        tiles[rr][cc2] = "water"

        def _carve_river():
            # rio serpenteando com variação de largura
            r0 = rng.randint(2, grid - 3)
            base_w = 2 if grid >= 10 else 1

            for c in range(1, grid - 1):
                # varia largura de leve (1..2 ou 2..3 dependendo do grid)
                if grid >= 12:
                    width = base_w + (1 if rng.random() > 0.80 else 0)  # às vezes engrossa
                else:
                    width = base_w

                for w in range(width):
                    rr = r0 + w
                    if 1 <= rr <= grid - 2 and inside(rr, c):
                        tiles[rr][c] = "water"

                # alargamento pequeno ocasional
                if grid >= 10 and rng.random() > 0.84:
                    rr = max(1, min(grid - 2, r0 + rng.choice([-1, 0, 1])))
                    if inside(rr, c):
                        tiles[rr][c] = "water"

                # serpenteia
                r0 = max(1, min(grid - 2 - (width - 1), r0 + rng.choice([-1, 0, 1])))

        # --- escolhe entre rio ou lago ---
        if rng.random() > 0.50:
            _carve_river()
        else:
            _carve_blob_lake()

        # --- opcional: "mar" numa borda (costa) ---
        if rng.random() > 0.65:
            for rr in range(grid):
                tiles[rr][0] = "sea"
                if grid > 4 and inside(rr, 1):
                    tiles[rr][1] = "sea" if rng.random() > 0.40 else tiles[rr][1]

        # --- aplica margem dupla por cima de tudo (única regra de margem) ---
        _apply_double_shore()



    elif theme_key == "biome_cave":
        # Caverna / Dungeon (bioma):
        # - Arena de combate em "sand" no interior
        # - "Muro" artificial com dirt_rock_edge contornando a arena (pedra para fora, terra para dentro)
        # - Fora da arena: piso rochoso mais limpo, com poucos overlays

        # base bem neutra fora da arena (menos "sopa" visual)
        for r in range(grid):
            for c in range(grid):
                tiles[r][c] = "stone" if inside(r, c) else "wall"

        # Define uma arena retangular (com alguma folga nas bordas do mapa)
        arena_h = max(6, min(grid - 4, int(grid * 0.55)))
        arena_w = max(7, min(grid - 4, int(grid * 0.60)))
        top = rng.randint(2, max(2, grid - arena_h - 2))
        left = rng.randint(2, max(2, grid - arena_w - 2))
        bottom = top + arena_h - 1
        right = left + arena_w - 1

        # Preenche arena: piso de areia + contorno de muro (tiles mascarados)
        for r in range(top, bottom + 1):
            for c in range(left, right + 1):
                if not inside(r, c):
                    continue

                on_top = (r == top)
                on_bottom = (r == bottom)
                on_left = (c == left)
                on_right = (c == right)

                if on_top or on_bottom or on_left or on_right:
                    # máscara no padrão E,N,S,W => bits 1,2,4,8
                    mask = 0
                    if on_right:
                        mask |= 1  # E
                    if on_top:
                        mask |= 2  # N
                    if on_bottom:
                        mask |= 4  # S
                    if on_left:
                        mask |= 8  # W
                    tiles[r][c] = f"arena_wall_{mask:02d}"
                else:
                    tiles[r][c] = "sand"

        # "Areia perto" do muro: reforça a faixa interna (1 tile) como areia
        for r in range(top + 1, bottom):
            for c in range(left + 1, right):
                if tiles[r][c].startswith("arena_wall_"):
                    continue
                if (r in (top + 1, bottom - 1)) or (c in (left + 1, right - 1)):
                    tiles[r][c] = "sand"

        # Poucos overlays de caverna (somente cave_overlay) fora da arena
        # (reduzido para ficar mais "limpo" e com leitura melhor)
        overlay_n = max(3, int(grid * 0.6))
        for _ in range(overlay_n):
            rr = rng.randint(1, grid - 2)
            cc = rng.randint(1, grid - 2)
            if not inside(rr, cc):
                continue
            if top <= rr <= bottom and left <= cc <= right:
                continue
            if tiles[rr][cc] in ("stone", "cave", "sand") and rng.random() > 0.55:
                tiles[rr][cc] = "cave_overlay"

        # Poça subterrânea (pequena) fora da arena
        if not no_water and rng.random() > 0.60:
            cr = rng.randint(2, grid - 3)
            cc = rng.randint(2, grid - 3)
            # evita cair dentro da arena
            if not (top - 1 <= cr <= bottom + 1 and left - 1 <= cc <= right + 1):
                rad = 1 if grid <= 10 else 2
                for rr in range(cr - rad, cr + rad + 1):
                    for cc2 in range(cc - rad, cc + rad + 1):
                        if inside(rr, cc2) and rng.random() > 0.35:
                            tiles[rr][cc2] = "water"

    elif theme_key == "biome_mix":
        # Mix: patches de biomas + 1 rio (bom pra rotas longas)
        biomes = ["grass", "forest", "meadow", "desert", "mountain", "snow", "cave"]
        centers = []
        for _ in range(rng.randint(3, 5)):
            centers.append((rng.randint(1, grid - 2), rng.randint(1, grid - 2), rng.choice(biomes)))

        def nearest_biome(rr, cc):
            best = None
            best_d = 1e9
            for r0, c0, b0 in centers:
                d = (r0 - rr) * (r0 - rr) + (c0 - cc) * (c0 - cc)
                if d < best_d:
                    best_d = d
                    best = b0
            return best or "grass"

        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if not inside(r, c):
                    continue
                b0 = nearest_biome(r, c)
                if b0 == "grass":
                    tiles[r][c] = "grass"
                elif b0 == "forest":
                    tiles[r][c] = "bush" if rng.random() > 0.30 else "grass"
                    if rng.random() > 0.65:
                        tiles[r][c] = "tree"
                elif b0 == "meadow":
                    tiles[r][c] = "flower" if rng.random() > 0.55 else "grass"
                elif b0 == "desert":
                    tiles[r][c] = "sand" if rng.random() > 0.20 else "dirt"
                elif b0 == "mountain":
                    tiles[r][c] = "stone" if rng.random() > 0.35 else "rock"
                elif b0 == "snow":
                    tiles[r][c] = "snow" if rng.random() > 0.15 else "rock"
                elif b0 == "cave":
                    tiles[r][c] = "cave" if rng.random() > 0.10 else "rock"

        # rio atravessando (se permitido)
        if not no_water:
            r0 = rng.randint(1, grid - 2)
            width = 2 if grid >= 10 else 1
            for c in range(1, grid - 1):
                for w in range(width):
                    rr = r0 + w
                    if 1 <= rr <= grid - 2:
                        tiles[rr][c] = "water"
                # margens
                if r0 - 1 >= 1 and rng.random() > 0.35 and tiles[r0 - 1][c] not in ["tree", "wall"]:
                    tiles[r0 - 1][c] = "sand"
                if r0 + width <= grid - 2 and rng.random() > 0.35 and tiles[r0 + width][c] not in ["tree", "wall"]:
                    tiles[r0 + width][c] = "sand"
                r0 = max(1, min(grid - 2 - (width - 1), r0 + rng.choice([-1, 0, 1])))

    # Chance global de flores em qualquer campo com grama
    for r in range(1, grid - 1):
        for c in range(1, grid - 1):
            if inside(r, c) and tiles[r][c] == "grass" and rng.random() > 0.97:
                tiles[r][c] = "flower"

# --- limpeza final: garante zero água se no_water=True ---
    if no_water:
        for r in range(grid):
            for c in range(grid):
                if tiles[r][c] == "water":
                    # substitui por algo coerente com o tema
                    tiles[r][c] = "path" if theme_key in ["forest", "cave_water"] else "trail"
                elif tiles[r][c] == "sea":
                    tiles[r][c] = "sand"

    return tiles, seed


def gen_tiles(grid: int, theme_key: str, seed: int | None = None, no_water: bool = False):
    """Gerador principal (com respawn). Floresta recebe geração mais coerente; demais temas usam o legado.

    - 'Respawn' aqui = reroll automático até cumprir regras mínimas (conectividade + densidade).
    - Mantém retorno (tiles, seed) e os mesmos tipos de tiles.
    """
    tkey = (theme_key or "").lower()

    # regra extra: em 6x6, evita água salvo temas explicitamente aquáticos
    if grid <= 6 and not any(w in tkey for w in ["water", "river", "lake", "sea", "coast"]):
        no_water = True

    # usa seed reprodutível
    if seed is None:
        seed = random.randint(1, 999999999)

    # só troca o algoritmo na floresta (e biome_forest)
    if tkey not in ["forest", "biome_forest"]:
        return _gen_tiles_legacy(grid, theme_key, seed=seed, no_water=no_water)

    def inside(r, c):
        return 1 <= r <= grid - 2 and 1 <= c <= grid - 2

    def is_blocking(tt: str) -> bool:
        return tt in ["tree", "wall", "peak", "stalagmite", "sea", "water"]

    def bfs_connected(start, tiles):
        from collections import deque
        q = deque([start])
        seen = {start}
        while q:
            r, c = q.popleft()
            for dr, dc in [(-1,0),(1,0),(0,-1),(0,1)]:
                nr, nc = r + dr, c + dc
                if 0 <= nr < grid and 0 <= nc < grid and (nr, nc) not in seen:
                    if not is_blocking(tiles[nr][nc]):
                        seen.add((nr, nc))
                        q.append((nr, nc))
        return seen

    def carve_disk(tiles, cr, cc, rad, tt):
        for r in range(cr - rad, cr + rad + 1):
            for c in range(cc - rad, cc + rad + 1):
                if inside(r, c) and (r - cr) * (r - cr) + (c - cc) * (c - cc) <= rad * rad + rad:
                    tiles[r][c] = tt

    def carve_path_polyline(tiles, pts, width=1, tt="path"):
        for (r0, c0), (r1, c1) in zip(pts, pts[1:]):
            steps = max(abs(r1 - r0), abs(c1 - c0), 1)
            for i in range(steps + 1):
                r = int(round(r0 + (r1 - r0) * (i / steps)))
                c = int(round(c0 + (c1 - c0) * (i / steps)))
                for dr in range(-width, width + 1):
                    for dc in range(-width, width + 1):
                        rr, cc = r + dr, c + dc
                        if inside(rr, cc) and tiles[rr][cc] not in ["water", "sea"]:
                            tiles[rr][cc] = tt

    def gen_forest_once(rng: random.Random):
        tiles = [["grass" for _ in range(grid)] for _ in range(grid)]

        # 1) noise inicial de árvores
        dense = 0.48 if tkey == "forest" else 0.58
        p_tree = max(0.20, min(0.62, dense))
        mask = [[0 for _ in range(grid)] for _ in range(grid)]
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if inside(r, c):
                    mask[r][c] = 1 if rng.random() < p_tree else 0

        # 2) cellular automata -> clusters
        def neigh(r, c):
            s = 0
            for dr in (-1, 0, 1):
                for dc in (-1, 0, 1):
                    if dr == 0 and dc == 0:
                        continue
                    rr, cc = r + dr, c + dc
                    if 0 <= rr < grid and 0 <= cc < grid:
                        s += mask[rr][cc]
            return s

        iters = 2 if grid <= 8 else 3
        for _ in range(iters):
            newm = [[0 for _ in range(grid)] for _ in range(grid)]
            for r in range(1, grid - 1):
                for c in range(1, grid - 1):
                    if not inside(r, c):
                        continue
                    n = neigh(r, c)
                    if mask[r][c] == 1:
                        newm[r][c] = 1 if n >= 3 else 0
                    else:
                        newm[r][c] = 1 if n >= 5 else 0
            mask = newm

        # aplica
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                tiles[r][c] = "tree" if mask[r][c] else "grass"

        # 3) clareiras (spawn) + caminho conectando
        left = (grid // 2, 2)
        right = (grid // 2, grid - 3)
        rad_clear = 1 if grid <= 6 else 2
        carve_disk(tiles, left[0], left[1], rad_clear, "grass")
        carve_disk(tiles, right[0], right[1], rad_clear, "grass")

        width = 1 if grid <= 8 else 2
        mid = (rng.randint(2, grid - 3), rng.randint(2, grid - 3))
        carve_path_polyline(tiles, [left, mid, right], width=width, tt="path")

        # 4) suaviza entorno do caminho (evita gargalos)
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if tiles[r][c] == "tree":
                    near_path = any(
                        0 <= r + dr < grid and 0 <= c + dc < grid and tiles[r + dr][c + dc] == "path"
                        for dr, dc in [(-1,0),(1,0),(0,-1),(0,1)]
                    )
                    if near_path and rng.random() > 0.35:
                        tiles[r][c] = "bush"

        # 5) detalhes leves (não poluir)
        for r in range(1, grid - 1):
            for c in range(1, grid - 1):
                if tiles[r][c] == "grass":
                    if rng.random() < 0.10:
                        tiles[r][c] = "bush"
                    elif rng.random() < 0.03:
                        tiles[r][c] = "flower"

        # 6) água opcional bem contida
        if (not no_water) and grid >= 8 and rng.random() > 0.82:
            cr = rng.randint(2, grid - 3)
            cc = rng.randint(2, grid - 3)
            rad = 1 if grid <= 10 else 2
            carve_disk(tiles, cr, cc, rad, "water")
            for r in range(cr - rad - 1, cr + rad + 2):
                for c in range(cc - rad - 1, cc + rad + 2):
                    if inside(r, c) and tiles[r][c] != "water" and (abs(r - cr) + abs(c - cc)) <= rad + 2:
                        if rng.random() > 0.35:
                            tiles[r][c] = "sand"

        # limpeza final se no_water
        if no_water:
            for r in range(grid):
                for c in range(grid):
                    if tiles[r][c] == "water":
                        tiles[r][c] = "path"
                    if tiles[r][c] == "sea":
                        tiles[r][c] = "sand"

        return tiles, left, right

    def validate_forest(tiles, left, right) -> bool:
        conn = bfs_connected(left, tiles)
        if right not in conn:
            return False

        total = (grid - 2) * (grid - 2)
        trees = sum(1 for r in range(1, grid - 1) for c in range(1, grid - 1) if tiles[r][c] == "tree")
        walk = sum(1 for r in range(1, grid - 1) for c in range(1, grid - 1) if not is_blocking(tiles[r][c]))

        tree_ratio = trees / max(1, total)
        walk_ratio = walk / max(1, total)

        # floresta precisa "parecer floresta" e ainda ser jogável
        if tree_ratio < 0.18:
            return False
        if walk_ratio < 0.35:
            return False
        return True

    # respawn/retry
    MAX_TRIES = 35
    for k in range(MAX_TRIES):
        rng = random.Random(int(seed) + k * 9973)
        tiles, left, right = gen_forest_once(rng)
        if validate_forest(tiles, left, right):
            return tiles, int(seed) + k * 9973

    # fallback: devolve o último mesmo se não passou (evita travar)
    return tiles, int(seed)


@st.cache_data(show_spinner=False)
def fetch_image_pil(url: str) -> Image.Image | None:
    try:
        r = requests.get(url, timeout=5)
        r.raise_for_status()
        img = Image.open(BytesIO(r.content)).convert("RGBA")
        return img
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def fetch_image_pil(url: str) -> Image.Image | None:
    try:
        r = requests.get(url, timeout=5)
        r.raise_for_status()
        img = Image.open(BytesIO(r.content)).convert("RGBA")
        return img
    except Exception:
        return None

def draw_tile_asset(img, r, c, tiles, assets, rng):
    """Desenha 1 tile (utilitário). Mantido por compatibilidade; usa a mesma lógica do render_map_png."""
    grid = len(tiles)
    t_type = tiles[r][c]
    x, y = c * TILE_SIZE, r * TILE_SIZE

    # pick variante rápida (se existir)
    def pick(base_key: str) -> str | None:
        if base_key in assets:
            # tenta variantes __v
            variants = [k for k in assets.keys() if k.startswith(base_key + "__v")]
            if variants:
                return rng.choice(variants)
            return base_key
        # tenta variantes mesmo sem base
        variants = [k for k in assets.keys() if k.startswith(base_key + "__v")]
        return rng.choice(variants) if variants else None

    # chão base
    under = {
        "water": "grama_1",
        "sea": "areia_1",
        "grass": "grama_1",
        "tree": "grama_1",
        "bush": "grama_1",
        "flower": "grama_1",
        "path": "terra_1",
        "trail": "terra_1",
        "rut": "terra_1",
        "dirt": "terra_1",
        "sand": "areia_1",
        "stone": "pedra_1",
        "rock": "pedra_1",
        "snow": "neve_1",
        "cave": "cave_1",
        "wall": "cave_1",
        "stalagmite": "cave_1",
        "cave_overlay": "pedra_1",
        "peak": "pedra_1",
    }.get(t_type, "grama_1")

    k_under = pick(under)
    if k_under:
        img.alpha_composite(assets[k_under], (x, y))

    # camada 2: piso específico / água shore
    asset_to_draw = None
    if t_type in ("water", "sea"):
                # 1) base embaixo (lago em grama / mar em areia)
        base_under = "grama_1" if t_type == "water" else "areia_1"
        base_choices = floor_variants.get(base_under, [base_under])
        chosen_base = rng.choice(base_choices)
        if chosen_base in assets:
            img.paste(assets[chosen_base], (x, y))

        # 2) calcula land_mask compatível com water_*_mXX
        # índice do pack: E,N,S,W => bits 1,2,4,8
        land_mask = 0
        for bit, (dr, dc) in enumerate([(0, 1), (-1, 0), (1, 0), (0, -1)]):  # E,N,S,W
            nr, nc = r + dr, c + dc
            if 0 <= nr < grid and 0 <= nc < grid:
                if tiles[nr][nc] not in ("water", "sea"):
                    land_mask |= (1 << bit)
            else:
                # fora do mapa conta como terra -> fecha margens
                land_mask |= (1 << bit)

        # 3) decide família da borda (não misturar grass/sand)
        # regra simples: se qualquer vizinho for "sand", usa sand; caso contrário grass
        shore_family = "sand" if t_type == "sea" else "grass"
        if shore_family == "grass":
            for dr, dc in [(-1, 0), (0, 1), (1, 0), (0, -1)]:
                nr, nc = r + dr, c + dc
                if 0 <= nr < grid and 0 <= nc < grid and tiles[nr][nc] == "sand":
                    shore_family = "sand"
                    break

        # 4) escolhe tile (miolo vs borda)
        if land_mask == 0:
            if "agua_shallow" in assets or any(k.startswith("agua_shallow__v") for k in assets):
                asset_to_draw = "agua_shallow"
            elif "agua_deep" in assets or any(k.startswith("agua_deep__v") for k in assets):
                asset_to_draw = "agua_deep"
            else:
                asset_to_draw = rng.choice([k for k in ["agua_1", "agua_3", "agua_2"] if k in assets] or [None])
        else:
            asset_to_draw = f"agua_shore_{shore_family}_{land_mask:02d}"


    elif t_type.startswith("arena_wall_"):
        # Paredes artificiais (dirt_rock_edge) ao redor da arena:
        # o sufixo é a máscara (00..15) no padrão E,N,S,W => bits 1,2,4,8.
        try:
            mask = int(t_type.rsplit("_", 1)[-1])
        except Exception:
            mask = 0
        asset_to_draw = f"dirt_rock_edge_{mask:02d}"
    elif t_type in ("sand", "stone", "dirt", "grass", "snow", "cave"):
        prefix_map = {
            "sand": "areia",
            "stone": "pedra",
            "dirt": "terra",
            "grass": "grama",
            "snow": "neve",
            "cave": "cave",
        }
        pref = prefix_map[t_type]
        # tenta 1..3; se não existir, cai no _1
        candidate = f"{pref}_{rng.randint(1,3)}"
        if candidate in assets:
            asset_to_draw = candidate
        elif f"{pref}_1" in assets:
            asset_to_draw = f"{pref}_1"

    # aplica variação se existir
    if asset_to_draw and asset_to_draw in assets:
        choices = floor_variants.get(asset_to_draw)
        if not choices:
            # suporte para famílias mascaradas (ex.: dirt_rock_edge_02__v01)
            choices = [k for k in assets.keys() if k == asset_to_draw or k.startswith(asset_to_draw + "__v")]
        asset_choice = rng.choice(choices)
        img.alpha_composite(assets[asset_choice], (x, y))

    if t_type == "tree":
        # pega qualquer asset que pareça árvore
        pool = [k for k in assets.keys()
                if k.lower().startswith("tree") and not k.lower().startswith("treetop")]
        # fallback se seu projeto usa nomes tipo "arvore_*"
        if not pool:
            pool = [k for k in assets.keys() if "tree" in k.lower() or "arvore" in k.lower()]
        obj_asset = rng.choice(pool) if pool else None
    
    elif t_type == "bush":
        # pega qualquer asset que pareça arbusto/moita/brush
        pool = [k for k in assets.keys()
                if ("brush" in k.lower()) or ("bush" in k.lower()) or ("moita" in k.lower())]
        obj_asset = rng.choice(pool) if pool else None
    elif t_type == "cave_overlay":
        pool = [k for k in assets.keys() if k == "cave_overlay" or k.startswith("cave_overlay__v")]
        obj_asset = rng.choice(pool) if pool else None
    elif t_type == "stalagmite":
        obj_asset = "estalagmite_1" if "estalagmite_1" in assets else None
    elif t_type == "peak":
        obj_asset = "pico_1" if "pico_1" in assets else None
    elif t_type == "flower":
        obj_asset = "flower" if "flower" in assets else None

    # rochas aleatórias em pisos (reduzido na caverna/areia pra não poluir)
    rock_p = 0.03 if t_type in ("cave", "sand") else 0.10
    if t_type in ("grass", "stone", "dirt", "sand", "snow", "cave") and rng.random() < rock_p:
        pool = [k for k in ["rochas", "rochas_2"] if k in assets]
        obj_asset = rng.choice(pool) if pool else obj_asset
    elif t_type == "rock":
        pool = [k for k in ["rochas", "rochas_2"] if k in assets]
        obj_asset = rng.choice(pool) if pool else obj_asset

    if obj_asset and obj_asset in assets:
        img.alpha_composite(assets[obj_asset], (x, y))

    _apply_water_shading_and_foam(img, tiles, grid, TILE_SIZE)


    # --- CAMADA 4: GRID TÁTICO FINO ---
    if show_grid:
        _draw_tactical_grid(img, grid, TILE_SIZE)

    return img.convert("RGB")


@st.cache_data(show_spinner=False)
def _generate_map_cached(grid: int, theme_key: str, seed: int, no_water: bool, show_grid: bool = True):
    """Gera o mapa e retorna (png_bytes, map_data).

    Armazenado em cache por parâmetros, então png_bytes e map_data sempre
    correspondem ao mesmo generate() – sem race condition com o singleton.
    """
    import io as _io
    biome = map_theme_to_biome(theme_key, no_water)
    generator = get_biome_generator()
    img = None
    map_data: dict = {}
    try:
        generated = generator.generate(
            biome=biome,
            grid_w=grid,
            grid_h=grid,
            tile_px=TILE_SIZE,
            seed=int(seed or 0),
        )
        map_data = generator.export_map_data() or {}
        if generated is not None:
            img = generated.convert("RGBA")
    except Exception as e:
        print(f"[_generate_map_cached] generate failed: biome={biome} grid={grid} seed={seed} error={e}")

    if img is None:
        img = Image.new("RGBA", (max(1, grid) * TILE_SIZE, max(1, grid) * TILE_SIZE), (16, 24, 40, 255))

    if show_grid:
        _draw_tactical_grid(img, grid, TILE_SIZE)

    buf = _io.BytesIO()
    img.convert("RGB").save(buf, format="PNG")
    return buf.getvalue(), map_data


def render_biome_map_png(grid: int, theme_key: str, seed: int, no_water: bool, show_grid: bool = True):
    import io as _io
    png_bytes, _ = _generate_map_cached(grid, theme_key, seed, no_water, show_grid)
    return Image.open(_io.BytesIO(png_bytes)).convert("RGB")


def battle_site_url_for_room(rid: str, trainer_name: str | None = None) -> str:
    base = "https://peucaptured.github.io/battle-site/"
    params = {"rid": str(rid)}
    if trainer_name:
        params["trainer"] = str(trainer_name)
    return base + "?" + urllib.parse.urlencode(params)

def redirect_to_battle_site(rid: str, trainer_name: str | None = None) -> None:
    url = battle_site_url_for_room(rid, trainer_name)
    # Tenta redirecionar a janela "top" (fora do iframe do Streamlit). Se falhar, cai no location local.
    components.html(
        "<script>(function(){const u=" + json.dumps(url) + "; try{window.top.location.href=u;}catch(e){window.location.href=u;}})();</script>",
        height=0, width=0
    )
    st.stop()

def render_map_with_pieces(grid, theme_key, seed, no_water, pieces, viewer_name, room, effects=None, show_grid: bool = True):
    
    # 1. Base do Mapa (Cacheada)␊
    img = render_biome_map_png(grid, theme_key, seed, no_water, show_grid=show_grid).convert("RGBA")
    draw = ImageDraw.Draw(img)
    
    # 2. CAMADA DE EFEITOS (Agora usando Imagens Reais)
    if effects:
        # Mapeamento: Emoji -> Caminho do Arquivo (Deve coincidir com a variável usada abaixo)
        EMOJI_TO_PATH = {
            "🔥": "Assets/fogo.png",
            "🧊": "Assets/gelo.png",
            "💧": "Assets/agua.png",
            "🪨": "Assets/rocha.png",
            "☁️": "Assets/nuvem.png",
            "☀️": "Assets/sol.png",
            "🍃": "Assets/terrenograma.png",
            "⚡": "Assets/terrenoeletrico.png",
        }

        for eff in effects:
            try:
                r, c = int(eff.get("row")), int(eff.get("col"))
                icon_char = eff.get("icon", "?")
                x, y = c * TILE_SIZE, r * TILE_SIZE
                
                path = EMOJI_TO_PATH.get(icon_char)
                
                if path and os.path.exists(path):
                    base_icon = load_effect_icon_rgba(path)
                    if base_icon is not None:
                        icon_img = base_icon.copy()  # importante: não mutar o cache
                        icon_img.thumbnail((int(TILE_SIZE * 0.7), int(TILE_SIZE * 0.7)))
                        ix = x + (TILE_SIZE - icon_img.size[0]) // 2
                        iy = y + (TILE_SIZE - icon_img.size[1]) // 2
                        img.alpha_composite(icon_img, (ix, iy))
                    else:
                        draw.ellipse([x+16, y+16, x+TILE_SIZE-16, y+TILE_SIZE-16], fill=(255, 255, 255, 150))

            except Exception as e:
                # Opcional: imprimir o erro no console para debug
                print(f"Erro ao renderizar efeito {icon_char}: {e}")
                continue

    # 3. CAMADA DE POKÉMONS
    local_cache = {}
    
    for p in pieces or []:
        r = int(p.get("row", -1))
        c = int(p.get("col", -1))
        if r < 0 or c < 0: continue

        owner = p.get("owner")
        border_color = get_perspective_color(viewer_name, owner, room)

        x = c * TILE_SIZE
        y = r * TILE_SIZE
        
        # Borda
        draw.rectangle([x, y, x + TILE_SIZE - 1, y + TILE_SIZE - 1], outline=border_color, width=4)

        sprite = None
        if p.get("kind") == "trainer":
            avatar_name = p.get("avatar")
            avatar_lookup = build_trainer_avatar_lookup()
            avatar_info = avatar_lookup.get(avatar_name) if avatar_name else None
            avatar_path = avatar_info.get("path") if avatar_info else None

            if avatar_path:
                if avatar_path not in local_cache:
                    local_cache[avatar_path] = load_trainer_avatar_image(avatar_path)
                sprite = local_cache[avatar_path]
        else:
            pid = str(p.get("pid", ""))
            is_p_shiny = p.get("shiny", False)
            url = pokemon_pid_to_image(pid, mode="sprite", shiny=is_p_shiny)

            if url not in local_cache:
                local_cache[url] = fetch_image_pil(url)
            sprite = local_cache[url]

        if sprite is None:
            continue

        sp = sprite.copy()
        max_size = int(TILE_SIZE * 0.75) if p.get("kind") == "trainer" else TILE_SIZE
        sp.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
        
        x0 = x + (TILE_SIZE - sp.size[0]) // 2
        y0 = y + (TILE_SIZE - sp.size[1]) // 2
        img.alpha_composite(sp, (x0, y0))

    return img.convert("RGB")
    
def normalize_text(text):
    if not isinstance(text, str): return str(text)
    return unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode('utf-8').lower().strip()


_LYCANROC_FORM_ALIASES = {
    "midday": "lycanroc-midday",
    "day": "lycanroc-midday",
    "dia": "lycanroc-midday",
    "midnight": "lycanroc-midnight",
    "night": "lycanroc-midnight",
    "noite": "lycanroc-midnight",
    "dusk": "lycanroc-dusk",
    "twilight": "lycanroc-dusk",
    "crepusculo": "lycanroc-dusk",
}


def _pokemon_lookup_candidates(pname: str) -> list[tuple[str, str]]:
    """Gera candidatos de busca para nomes com forma e fallback para a espécie base."""
    raw = str(pname or "").strip()
    if not raw:
        return []

    candidates: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()

    def add(kind: str, value: str):
        value = str(value or "").strip()
        if not value:
            return
        token = (kind, value.lower())
        if token in seen:
            return
        seen.add(token)
        candidates.append((kind, value))

    add("raw", raw)
    add("normalized", normalize_text(raw))

    api_name = to_pokeapi_name(raw)
    add("api", api_name)

    base_raw = ""
    parsed = re.match(r"^(.*?)\s*\((.*?)\)\s*$", raw)
    if parsed:
        base_raw = parsed.group(1).strip()
        form_raw = parsed.group(2).strip()
        form_key = normalize_text(form_raw).replace(" ", "-")
        base_api = to_pokeapi_name(base_raw)

        region = REGION_ALIASES.get(form_key)
        if region:
            add("api", f"{base_api}-{region}")
        elif normalize_text(base_raw) == "lycanroc":
            lycanroc_alias = _LYCANROC_FORM_ALIASES.get(form_key)
            if lycanroc_alias:
                add("api", lycanroc_alias)

    if api_name.startswith("lycanroc-"):
        add("api", "lycanroc")
        if not base_raw:
            base_raw = "Lycanroc"

    for region_suffix in ("-alola", "-galar", "-hisui", "-paldea"):
        if api_name.endswith(region_suffix):
            add("api", api_name[: -len(region_suffix)])
            break

    if base_raw:
        add("raw", base_raw)
        add("normalized", normalize_text(base_raw))
        add("api", to_pokeapi_name(base_raw))

    return candidates



def resolve_pokemon_pid(df_pokedex: pd.DataFrame, pname: str) -> str:
    """Resolve o PID numérico da Pokédex por nome, com fallback normalizado."""
    name = str(pname or "").strip()
    if not name or df_pokedex is None or df_pokedex.empty:
        return "0"

    if "Nome" not in df_pokedex.columns or "Nº" not in df_pokedex.columns:
        return "0"

    try:
        nomes = df_pokedex["Nome"].astype(str)

        # 1) match exato sem espaços extras
        hit = df_pokedex[nomes.str.strip().str.lower() == name.lower()]
        if not hit.empty:
            return str(int(float(hit.iloc[0]["Nº"])))

        # 2) match por normalização (acentos/símbolos)
        nkey = normalize_text(name)
        hit = df_pokedex[nomes.apply(normalize_text) == nkey]
        if not hit.empty:
            return str(int(float(hit.iloc[0]["Nº"])))

        # 3) match por nome canônico da API (útil para variações de escrita)
        api_key = to_pokeapi_name(name)
        if api_key:
            hit = df_pokedex[nomes.apply(lambda x: to_pokeapi_name(str(x))) == api_key]
            if not hit.empty:
                return str(int(float(hit.iloc[0]["Nº"])))
    except Exception:
        return "0"

    return "0"


def resolve_pokemon_pid(df_pokedex: pd.DataFrame, pname: str) -> str:
    """Resolve o PID numérico com suporte a formas salvas e fallback para a espécie base."""
    name = str(pname or "").strip()
    if not name or df_pokedex is None or df_pokedex.empty:
        return "0"

    if "Nome" in df_pokedex.columns:
        for candidate in ("N\u00ba", "N\u00c2\u00ba", "N\u00c3\u0082\u00c2\u00ba"):
            if candidate in df_pokedex.columns:
                df_pokedex = df_pokedex.rename(columns={candidate: "N\u00c3\u0082\u00c2\u00ba"})
                break

    if "Nome" not in df_pokedex.columns or "NÂº" not in df_pokedex.columns:
        return "0"

    try:
        nomes = df_pokedex["Nome"].astype(str)
        nomes_lower = nomes.str.strip().str.lower()
        nomes_norm = nomes.apply(normalize_text)
        nomes_api = nomes.apply(lambda x: to_pokeapi_name(str(x)))

        def _pid_from_hit(hit: pd.DataFrame) -> str:
            return str(int(float(hit.iloc[0]["NÂº"])))

        for kind, candidate in _pokemon_lookup_candidates(name):
            if kind == "raw":
                hit = df_pokedex[nomes_lower == candidate.lower()]
            elif kind == "normalized":
                hit = df_pokedex[nomes_norm == candidate]
            elif kind == "api":
                hit = df_pokedex[nomes_api == candidate]
            else:
                continue

            if not hit.empty:
                return _pid_from_hit(hit)
    except Exception:
        return "0"

    return "0"


def get_pid_from_name(user_name: str, name_map: dict) -> str | None:
    if not isinstance(user_name, str):
        return None

    pre_clean = user_name.replace('♀', '-f').replace('♂', '-m')
    # --- Normalização de formas regionais no formato "Nome (Galar)" / "Nome (Alola)" etc ---
    # Isso garante que sprites usem o ID correto do PokeAPI (ex: ponyta-galar -> 10162)
    try:
        _pc = str(pre_clean)
        _pc = re.sub(r"\s*\(\s*galar\s*\)\s*", "-galar", _pc, flags=re.IGNORECASE)
        _pc = re.sub(r"\s*\(\s*alola\s*\)\s*", "-alola", _pc, flags=re.IGNORECASE)
        _pc = re.sub(r"\s*\(\s*hisui\s*\)\s*", "-hisui", _pc, flags=re.IGNORECASE)
        _pc = re.sub(r"\s*\(\s*paldea\s*\)\s*", "-paldea", _pc, flags=re.IGNORECASE)

        if re.search(r"\bgalarian\b", _pc, flags=re.IGNORECASE):
            _pc = re.sub(r"\bgalarian\b", "", _pc, flags=re.IGNORECASE).strip() + "-galar"
        if re.search(r"\balolan\b", _pc, flags=re.IGNORECASE):
            _pc = re.sub(r"\balolan\b", "", _pc, flags=re.IGNORECASE).strip() + "-alola"
        if re.search(r"\bhisuian\b", _pc, flags=re.IGNORECASE):
            _pc = re.sub(r"\bhisuian\b", "", _pc, flags=re.IGNORECASE).strip() + "-hisui"
        if re.search(r"\bpaldean\b", _pc, flags=re.IGNORECASE):
            _pc = re.sub(r"\bpaldean\b", "", _pc, flags=re.IGNORECASE).strip() + "-paldea"

        pre_clean = _pc
    except Exception:
        pass

    clean = normalize_text(pre_clean).replace('.', '').replace("'", '').replace(' ', '-')

    # exceções / formas (as mesmas que você já usa)
    if clean == 'mimikyu': clean = 'mimikyu-disguised'
    if clean == 'aegislash': clean = 'aegislash-blade'
    if clean == 'giratina': clean = 'giratina-origin'
    if clean == 'wishiwashi': clean = 'wishiwashi-solo'
    if clean == 'pumpkaboo': clean = 'pumpkaboo-average'
    if clean == 'gourgeist': clean = 'gourgeist-average'
    if clean == 'lycanroc': clean = 'lycanroc-midday'
    if clean == 'deoxys': clean = 'deoxys-normal'
    if clean == 'wormadam': clean = 'wormadam-plant'
    if clean == 'shaymin': clean = 'shaymin-land'

    if clean == 'toxtricity': clean = 'toxtricity-amped'
    if clean == 'eiscue': clean = 'eiscue-ice'
    if clean == 'indeedee': clean = 'indeedee-male'
    if clean == 'morpeko': clean = 'morpeko-full-belly'
    if clean == 'urshifu': clean = 'urshifu-single-strike'

    if clean == 'basculegion': clean = 'basculegion-male'
    if clean == 'enamorus': clean = 'enamorus-incarnate'
    if clean == 'keldeo': clean = 'keldeo-ordinary'
    if clean == 'meloetta': clean = 'meloetta-aria'
    if clean == 'darmanitan': clean = 'darmanitan-standard'
    if clean == 'minior': clean = 'minior-red-meteor'

    # formas especiais que costumam ser digitadas invertidas (forma + espécie)
    if clean in ('eternal-floette', 'floette-eternal-forme', 'floette-eternal-form'):
        clean = 'floette-eternal'
    if clean in ('bloodmoon-ursaluna', 'blood-moon-ursaluna', 'ursaluna-blood-moon'):
        clean = 'ursaluna-bloodmoon'

    # regionais (as mesmas)
    if clean.endswith('-a'): clean = clean[:-2] + '-alola'
    if clean.endswith('-g'): clean = clean[:-2] + '-galar'
    if clean.endswith('-h'): clean = clean[:-2] + '-hisui'
    if clean.endswith('-p'): clean = clean[:-2] + '-paldea'
    if clean.startswith('g-'): clean = clean[2:] + '-galar'
    if clean.startswith('a-'): clean = clean[2:] + '-alola'
    if clean.startswith('h-'): clean = clean[2:] + '-hisui'
    if clean.startswith('p-'): clean = clean[2:] + '-paldea'

    p_id = name_map.get(clean)
    if not p_id:
        base_name = clean.split('-')[0]
        p_id = name_map.get(base_name)

    return p_id
    
def get_pokemon_image_url(user_name: str, name_map: dict, mode: str = "artwork", shiny: bool = False) -> str:
    # Tenta achar o ID da National Dex pelo nome
    p_id = get_pid_from_name(user_name, name_map)
    
    if not p_id:
        return "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"

    if mode == "sprite":
        return get_pokemon_sprite_url(p_id, shiny=shiny)
    return get_pokemon_artwork_url(p_id, shiny=shiny)


def get_image_from_name(user_name, name_map):
    if not isinstance(user_name, str): return "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"
    pre_clean = user_name.replace('♀', '-f').replace('♂', '-m')
    clean = normalize_text(pre_clean).replace('.', '').replace("'", '').replace(' ', '-')
    
    # --- LISTA DE EXCEÇÕES E FORMAS (ATUALIZADA) ---
    if clean == 'mimikyu': clean = 'mimikyu-disguised'
    if clean == 'aegislash': clean = 'aegislash-blade'
    if clean == 'giratina': clean = 'giratina-origin'
    if clean == 'wishiwashi': clean = 'wishiwashi-solo'
    if clean == 'pumpkaboo': clean = 'pumpkaboo-average'
    if clean == 'gourgeist': clean = 'gourgeist-average'
    if clean == 'lycanroc': clean = 'lycanroc-midday'
    if clean == 'deoxys': clean = 'deoxys-normal'
    if clean == 'wormadam': clean = 'wormadam-plant'
    if clean == 'shaymin': clean = 'shaymin-land'
    
    # Correções da 8ª Geração (Seus pedidos)
    if clean == 'toxtricity': clean = 'toxtricity-amped'
    if clean == 'eiscue': clean = 'eiscue-ice'
    if clean == 'indeedee': clean = 'indeedee-male'
    if clean == 'morpeko': clean = 'morpeko-full-belly'
    if clean == 'urshifu': clean = 'urshifu-single-strike'
    
    # Outras correções úteis
    if clean == 'basculegion': clean = 'basculegion-male'
    if clean == 'enamorus': clean = 'enamorus-incarnate'
    if clean == 'keldeo': clean = 'keldeo-ordinary'
    if clean == 'meloetta': clean = 'meloetta-aria'

    # Sufixos Regionais
    if clean.endswith('-a'): clean = clean[:-2] + '-alola'
    if clean.endswith('-g'): clean = clean[:-2] + '-galar'
    if clean.endswith('-h'): clean = clean[:-2] + '-hisui'
    if clean.endswith('-p'): clean = clean[:-2] + '-paldea' # Paldea
    if clean.startswith('g-'): clean = clean[2:] + '-galar'
    if clean.startswith('a-'): clean = clean[2:] + '-alola'
    if clean.startswith('h-'): clean = clean[2:] + '-hisui'
    if clean.startswith('p-'): clean = clean[2:] + '-paldea'
    p_id = name_map.get(clean)
    if not p_id:
        base_name = clean.split('-')[0]
        p_id = name_map.get(base_name)

    if p_id:
        return f"https://raw.githubusercontent.com/PokeAPI/sprites/master/sprites/pokemon/other/official-artwork/{p_id}.png"
    else:
        return "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"


_DEFAULT_POKEMON_API_FORMS = {
    "mimikyu": "mimikyu-disguised",
    "aegislash": "aegislash-blade",
    "giratina": "giratina-origin",
    "wishiwashi": "wishiwashi-solo",
    "pumpkaboo": "pumpkaboo-average",
    "gourgeist": "gourgeist-average",
    "lycanroc": "lycanroc-midday",
    "deoxys": "deoxys-normal",
    "wormadam": "wormadam-plant",
    "shaymin": "shaymin-land",
    "toxtricity": "toxtricity-amped",
    "eiscue": "eiscue-ice",
    "indeedee": "indeedee-male",
    "morpeko": "morpeko-full-belly",
    "urshifu": "urshifu-single-strike",
    "basculegion": "basculegion-male",
    "enamorus": "enamorus-incarnate",
    "keldeo": "keldeo-ordinary",
    "meloetta": "meloetta-aria",
    "darmanitan": "darmanitan-standard",
    "minior": "minior-red-meteor",
}

_GENERIC_FORM_SUFFIX_ALIASES = {
    "attack-forme": "attack",
    "attack-form": "attack",
    "defense-forme": "defense",
    "defense-form": "defense",
    "defence-forme": "defense",
    "defence-form": "defense",
    "speed-forme": "speed",
    "speed-form": "speed",
    "altered-forme": "altered",
    "origin-forme": "origin",
    "sky-forme": "sky",
    "land-forme": "land",
    "zen-mode": "zen",
    "standard-mode": "standard",
    "hero-form": "hero",
    "school-form": "school",
    "solo-form": "solo",
    "lowkey": "low-key",
    "full": "full-belly",
}

_POKEMON_FORM_SUFFIX_ALIASES = {
    "lycanroc": {
        "midday": "midday",
        "day": "midday",
        "dia": "midday",
        "midnight": "midnight",
        "night": "midnight",
        "noite": "midnight",
        "dusk": "dusk",
        "twilight": "dusk",
        "crepusculo": "dusk",
    },
    "deoxys": {
        "normal": "normal",
        "attack": "attack",
        "defense": "defense",
        "defence": "defense",
        "speed": "speed",
    },
    "giratina": {
        "altered": "altered",
        "origin": "origin",
    },
    "shaymin": {
        "land": "land",
        "sky": "sky",
    },
    "wormadam": {
        "plant": "plant",
        "sandy": "sandy",
        "trash": "trash",
    },
    "meloetta": {
        "aria": "aria",
        "pirouette": "pirouette",
    },
    "keldeo": {
        "ordinary": "ordinary",
        "resolute": "resolute",
    },
    "wishiwashi": {
        "solo": "solo",
        "school": "school",
    },
    "aegislash": {
        "blade": "blade",
        "shield": "shield",
    },
    "mimikyu": {
        "disguised": "disguised",
        "busted": "busted",
    },
    "toxtricity": {
        "amped": "amped",
        "low-key": "low-key",
        "lowkey": "low-key",
    },
    "eiscue": {
        "ice": "ice",
        "noice": "noice",
    },
    "indeedee": {
        "male": "male",
        "female": "female",
    },
    "basculegion": {
        "male": "male",
        "female": "female",
    },
    "enamorus": {
        "incarnate": "incarnate",
        "therian": "therian",
    },
    "morpeko": {
        "full-belly": "full-belly",
        "hangry": "hangry",
    },
    "urshifu": {
        "single-strike": "single-strike",
        "rapid-strike": "rapid-strike",
    },
    "darmanitan": {
        "standard": "standard",
        "zen": "zen",
        "galar": "galar-standard",
        "galar-standard": "galar-standard",
        "galar-zen": "galar-zen",
    },
    "minior": {
        "red": "red-meteor",
        "orange": "orange-meteor",
        "yellow": "yellow-meteor",
        "green": "green-meteor",
        "blue": "blue-meteor",
        "indigo": "indigo-meteor",
        "violet": "violet-meteor",
        "meteor": "red-meteor",
        "core": "red",
    },
    "hoopa": {
        "confined": "confined",
        "unbound": "unbound",
    },
    "palafin": {
        "zero": "zero",
        "hero": "hero",
    },
    "oricorio": {
        "baile": "baile",
        "pom-pom": "pom-pom",
        "pau": "pau",
        "sensu": "sensu",
    },
    "rotom": {
        "heat": "heat",
        "wash": "wash",
        "frost": "frost",
        "fan": "fan",
        "mow": "mow",
    },
    "zygarde": {
        "10": "10",
        "10-percent": "10",
        "50": "50",
        "complete": "complete",
    },
    "necrozma": {
        "dusk-mane": "dusk",
        "dawn-wings": "dawn",
        "ultra": "ultra",
    },
    "floette": {
        "eternal": "eternal",
    },
}

_SPECIAL_FORM_CANONICAL_NAMES = {
    "eternal-floette": "floette-eternal",
    "floette-eternal-forme": "floette-eternal",
    "floette-eternal-form": "floette-eternal",
    "bloodmoon-ursaluna": "ursaluna-bloodmoon",
    "blood-moon-ursaluna": "ursaluna-bloodmoon",
    "ursaluna-blood-moon": "ursaluna-bloodmoon",
}


def _basic_pokeapi_slug(user_text: str) -> str:
    s = (user_text or "").strip().lower()
    s = (
        s.replace("♀", " f")
        .replace("♂", " m")
        .replace("â™€", " f")
        .replace("â™‚", " m")
    )
    s = re.sub(r"[\s_]+", "-", s)
    s = re.sub(r"-{2,}", "-", s).strip("-")

    if s in ("nidoran", "nidoran-"):
        return "nidoran"
    if s in ("nidoran-f", "nidoranf", "nidoran-female", "nidoran-fem", "nidoran-f."):
        return "nidoran-f"
    if s in ("nidoran-m", "nidoranm", "nidoran-male", "nidoran-masc", "nidoran-m."):
        return "nidoran-m"

    if re.match(r"^[aghp]-", s):
        tag, base = s.split("-", 1)
        region = REGION_ALIASES.get(tag)
        if region and base:
            return f"{base}-{region}"

    m = re.match(r"^(.+)-([aghp])$", s)
    if m:
        base, tag = m.group(1), m.group(2)
        region = REGION_ALIASES.get(tag)
        if region:
            return f"{base}-{region}"

    parts = s.split("-")
    if len(parts) >= 2:
        last = parts[-1]
        region = REGION_ALIASES.get(last)
        if region:
            base = "-".join(parts[:-1])
            return f"{base}-{region}"

    return s


def _normalize_form_slug(form_text: str) -> str:
    form = normalize_text(form_text).replace("%", " percent ")
    form = form.replace("'", "").replace(".", " ").replace("/", " ").replace("_", " ")
    form = re.sub(r"\bforme\b", "", form)
    form = re.sub(r"\bform\b", "", form)
    form = re.sub(r"\bmode\b", "", form)
    form = re.sub(r"\s+", " ", form).strip()
    form = form.replace(" percent", "")
    return form.replace(" ", "-")


def _canonical_pokemon_api_name(user_text: str, apply_default_form: bool = True) -> str:
    raw = str(user_text or "").strip()
    if not raw:
        return ""

    parsed = re.match(r"^(.*?)\s*\((.*?)\)\s*$", raw)
    if parsed:
        base_raw = parsed.group(1).strip()
        form_raw = parsed.group(2).strip()
        base_slug = _basic_pokeapi_slug(base_raw)
        form_slug = _normalize_form_slug(form_raw)

        if base_slug == "nidoran":
            if form_slug in ("f", "female"):
                return "nidoran-f"
            if form_slug in ("m", "male"):
                return "nidoran-m"

        region = REGION_ALIASES.get(form_slug)
        if region:
            return f"{base_slug}-{region}"

        species_forms = _POKEMON_FORM_SUFFIX_ALIASES.get(base_slug, {})
        suffix = species_forms.get(form_slug)
        if suffix is None:
            suffix = _GENERIC_FORM_SUFFIX_ALIASES.get(form_slug, form_slug)

        if suffix:
            return suffix if suffix.startswith(f"{base_slug}-") else f"{base_slug}-{suffix}"
        return base_slug

    slug = _basic_pokeapi_slug(raw)
    slug = _SPECIAL_FORM_CANONICAL_NAMES.get(slug, slug)

    if apply_default_form:
        slug = _DEFAULT_POKEMON_API_FORMS.get(slug, slug)

    return slug


def to_pokeapi_name(user_text: str) -> str:
    return _canonical_pokemon_api_name(user_text, apply_default_form=True)


def _find_pokedex_number_column(df_pokedex: pd.DataFrame) -> str | None:
    for col in df_pokedex.columns:
        compact = re.sub(r"[^a-z0-9]+", "", normalize_text(str(col)))
        if compact in ("n", "no", "num", "numero", "nmero"):
            return col
    return None


def resolve_pokemon_pid(df_pokedex: pd.DataFrame, pname: str) -> str:
    """Resolve o PID numérico suportando nomes salvos com forma."""
    name = str(pname or "").strip()
    if not name or df_pokedex is None or df_pokedex.empty:
        return "0"

    num_col = _find_pokedex_number_column(df_pokedex)
    if "Nome" not in df_pokedex.columns or not num_col:
        return "0"

    try:
        nomes = df_pokedex["Nome"].astype(str)
        nomes_lower = nomes.str.strip().str.lower()
        nomes_norm = nomes.apply(normalize_text)
        nomes_api = nomes.apply(lambda x: _canonical_pokemon_api_name(str(x), apply_default_form=False))

        candidates = []
        seen = set()

        def add_candidate(value: str):
            value = str(value or "").strip()
            if not value:
                return
            token = value.lower()
            if token in seen:
                return
            seen.add(token)
            candidates.append(value)

        add_candidate(name)
        add_candidate(normalize_text(name))
        add_candidate(_canonical_pokemon_api_name(name, apply_default_form=False))

        parsed = re.match(r"^(.*?)\s*\((.*?)\)\s*$", name)
        if parsed:
            base_name = parsed.group(1).strip()
            add_candidate(base_name)
            add_candidate(normalize_text(base_name))
            add_candidate(_canonical_pokemon_api_name(base_name, apply_default_form=False))

        for candidate in candidates:
            hit = df_pokedex[nomes_lower == candidate.lower()]
            if hit.empty:
                hit = df_pokedex[nomes_norm == candidate]
            if hit.empty:
                hit = df_pokedex[nomes_api == candidate]
            if not hit.empty:
                return str(int(float(hit.iloc[0][num_col])))
    except Exception:
        return "0"

    return "0"


def get_pid_from_name(user_name: str, name_map: dict) -> str | None:
    if not isinstance(user_name, str):
        return None

    canonical_name = to_pokeapi_name(user_name)
    p_id = name_map.get(canonical_name)
    if not p_id:
        base_name = canonical_name.split('-')[0]
        p_id = name_map.get(base_name)
    return p_id


def get_image_from_name(user_name, name_map):
    if not isinstance(user_name, str):
        return "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"

    p_id = get_pid_from_name(user_name, name_map)
    if p_id:
        return f"https://raw.githubusercontent.com/PokeAPI/sprites/master/sprites/pokemon/other/official-artwork/{p_id}.png"
    return "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"


@st.cache_data
def get_official_pokemon_map():
    try:
        url = "https://pokeapi.co/api/v2/pokemon?limit=10000"
        response = requests.get(url, timeout=5)
        data = response.json()
        name_map = {}
        for p in data['results']:
            p_id = p['url'].split('/')[-2]
            name_map[p['name']] = p_id
        return name_map
    except:
        return {}
        
def get_pokemon_artwork_url(p_id: str, shiny: bool = False) -> str:
    # Garante que é numérico para API
    try:
        n = int(str(p_id).lstrip("0") or "0")
    except ValueError:
        return "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"
        
    if shiny:
        return f"https://raw.githubusercontent.com/PokeAPI/sprites/master/sprites/pokemon/other/official-artwork/shiny/{n}.png"
    return f"https://raw.githubusercontent.com/PokeAPI/sprites/master/sprites/pokemon/other/official-artwork/{n}.png"


def get_pokemon_sprite_url(p_id: str, shiny: bool = False) -> str:
    try:
        n = int(str(p_id).lstrip("0") or "0")
    except ValueError:
        return "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"

    if shiny:
        return f"https://raw.githubusercontent.com/PokeAPI/sprites/master/sprites/pokemon/shiny/{n}.png"
    return f"https://raw.githubusercontent.com/PokeAPI/sprites/master/sprites/pokemon/{n}.png"

TRAINER_AVATAR_DIR = Path("trainer")

def _trainer_avatar_base(name: str) -> str:
    return (name or "").split("_")[0].strip().lower()

@st.cache_data(show_spinner=False)
def list_trainer_avatar_paths() -> list[Path]:
    if not TRAINER_AVATAR_DIR.exists():
        return []
    return sorted([p for p in TRAINER_AVATAR_DIR.glob("*.png") if p.is_file()])

@st.cache_data(show_spinner=False)
def build_trainer_avatar_catalog() -> dict[str, list[dict]]:
    catalog: dict[str, list[dict]] = {}
    for path in list_trainer_avatar_paths():
        name = path.stem
        base = _trainer_avatar_base(name)
        catalog.setdefault(base, []).append({
            "name": name,
            "path": str(path),
            "base": base,
        })
    return catalog

import io
import numpy as np
from PIL import Image

def _crop_center_square(img: Image.Image) -> Image.Image:
    w, h = img.size
    s = min(w, h)
    left = (w - s) // 2
    top = (h - s) // 2
    return img.crop((left, top, left + s, top + s))

def _feature_from_image(img: Image.Image, size: int = 96) -> np.ndarray:
    """
    Feature mais robusta para fotos e sprites:
      - Cor: HSV hist + stats (global e centro)
      - Forma: HOG multi-célula + hist de magnitude
      - Textura/mascara: silhueta suavizada
      - Iluminação: stats em LAB
    Para fotos (sem alpha), reduz fundo por máscara radial + distância da borda.
    Retorna vetor L2-normalizado.
    """
    # 1) normaliza e redimensiona
    img_rgba = _crop_center_square(img.convert("RGBA")).resize((size, size), Image.Resampling.LANCZOS)
    arr = np.asarray(img_rgba).astype(np.float32)  # (H,W,4)
    alpha = arr[..., 3] / 255.0
    rgb = arr[..., :3]

    # 2) máscara: sprite (tem transparência) vs foto (sem transparência)
    if float(alpha.min()) < 0.98:
        mask = alpha > 0.10
    else:
        yy, xx = np.mgrid[0:size, 0:size].astype(np.float32)
        cx = (size - 1) * 0.5
        cy = (size - 1) * 0.5
        xx = (xx - cx) / cx
        yy = (yy - cy) / cy
        r = np.sqrt(xx * xx + yy * yy)
        radial = r <= 0.88

        border = np.concatenate([
            rgb[0, :, :],
            rgb[-1, :, :],
            rgb[:, 0, :],
            rgb[:, -1, :],
        ], axis=0)
        bg = np.median(border, axis=0)
        dist = np.sqrt(((rgb - bg) ** 2).sum(axis=2))
        color_fg = dist > 18.0
        mask = radial & color_fg

    if mask.sum() < (size * size) * 0.08:
        mask = np.ones((size, size), dtype=bool)

    # 3) HSV hist + stats (global e centro)
    rgb_u8 = rgb.clip(0, 255).astype(np.uint8)
    hsv = np.asarray(Image.fromarray(rgb_u8, mode="RGB").convert("HSV")).astype(np.float32)
    h = (hsv[..., 0] / 255.0)
    s = (hsv[..., 1] / 255.0)
    v = (hsv[..., 2] / 255.0)

    def _masked_hist(hh, ss, vv, m, bins=(12, 4, 4)) -> np.ndarray:
        vals = np.stack([hh[m], ss[m], vv[m]], axis=1)
        hist, _ = np.histogramdd(
            vals,
            bins=bins,
            range=((0, 1), (0, 1), (0, 1)),
        )
        hist = hist.flatten().astype(np.float32)
        return hist / (hist.sum() + 1e-8)

    hist_hsv = _masked_hist(h, s, v, mask, bins=(12, 4, 4))

    mid = size // 2
    span = int(size * 0.30)
    y0, y1 = max(0, mid - span), min(size, mid + span)
    x0, x1 = max(0, mid - span), min(size, mid + span)
    center_mask = mask[y0:y1, x0:x1]
    if center_mask.sum() < 25:
        center_mask = np.ones_like(center_mask, dtype=bool)
    center_hist = _masked_hist(h[y0:y1, x0:x1], s[y0:y1, x0:x1], v[y0:y1, x0:x1], center_mask, bins=(8, 3, 3))

    hsv_stats = np.array(
        [
            h[mask].mean(),
            s[mask].mean(),
            v[mask].mean(),
            h[mask].std(),
            s[mask].std(),
            v[mask].std(),
        ],
        dtype=np.float32,
    )

    # 4) LAB stats (iluminação e contraste)
    lab = np.asarray(Image.fromarray(rgb_u8, mode="RGB").convert("LAB")).astype(np.float32)
    lab_scaled = lab / np.array([255.0, 255.0, 255.0], dtype=np.float32)
    lab_stats = np.array(
        [
            lab_scaled[..., 0][mask].mean(),
            lab_scaled[..., 1][mask].mean(),
            lab_scaled[..., 2][mask].mean(),
            lab_scaled[..., 0][mask].std(),
            lab_scaled[..., 1][mask].std(),
            lab_scaled[..., 2][mask].std(),
        ],
        dtype=np.float32,
    )

    # 5) gradientes (bordas)
    gray = (0.299 * rgb[..., 0] + 0.587 * rgb[..., 1] + 0.114 * rgb[..., 2]) / 255.0
    gray = gray.astype(np.float32)
    gx = np.zeros_like(gray)
    gy = np.zeros_like(gray)
    gx[:, 1:-1] = gray[:, 2:] - gray[:, :-2]
    gy[1:-1, :] = gray[2:, :] - gray[:-2, :]
    mag = np.sqrt(gx * gx + gy * gy)
    ang = (np.arctan2(gy, gx) + np.pi) / (2 * np.pi)

    cells = 5
    bins = 8
    cell = size // cells
    hog = np.zeros((cells, cells, bins), dtype=np.float32)

    for cy in range(cells):
        for cx in range(cells):
            y0, y1 = cy * cell, min(size, (cy + 1) * cell)
            x0, x1 = cx * cell, min(size, (cx + 1) * cell)

            m = mask[y0:y1, x0:x1]
            if m.sum() <= 0:
                continue

            mag_c = mag[y0:y1, x0:x1][m]
            ang_c = ang[y0:y1, x0:x1][m]
            idx = np.clip((ang_c * bins).astype(np.int32), 0, bins - 1)

            hist = np.zeros((bins,), dtype=np.float32)
            np.add.at(hist, idx, mag_c)
            hist /= (hist.sum() + 1e-8)
            hog[cy, cx, :] = hist

    hog = hog.flatten().astype(np.float32)

    mag_hist, _ = np.histogram(
        mag[mask],
        bins=6,
        range=(0, mag.max() + 1e-6),
        density=False,
    )
    mag_hist = mag_hist.astype(np.float32)
    mag_hist /= (mag_hist.sum() + 1e-8)

    # 6) silhueta (mask downsample)
    mask_u8 = (mask.astype(np.uint8) * 255)
    mask_img = Image.fromarray(mask_u8, mode="L").resize((16, 16), Image.Resampling.BILINEAR)
    mask_feat = (np.asarray(mask_img).astype(np.float32) / 255.0).flatten()

    # 7) junta com pesos
    feat = np.concatenate([
        0.55 * hist_hsv,
        0.20 * center_hist,
        0.20 * hog,
        0.10 * mag_hist,
        0.10 * hsv_stats,
        0.10 * lab_stats,
        0.20 * mask_feat,
    ], axis=0).astype(np.float32)

    feat /= (np.linalg.norm(feat) + 1e-8)
    return feat

@st.cache_data(show_spinner=False)
def build_trainer_avatar_index() -> list[dict]:
    """
    Usa seu catalog (base/name/path) e acrescenta feat para cada skin.
    Cacheado => custo pago uma vez (mesmo com 1200 imagens).
    """
    catalog = build_trainer_avatar_catalog()
    entries: list[dict] = []

    for base, items in catalog.items():
        for it in items:
            path = it["path"]
            name = it["name"]
            try:
                img = Image.open(path)
                img.load()
                feat = _feature_from_image(img, size=96)
            except Exception:
                continue

            entries.append({
                "base": base,
                "name": name,
                "path": path,
                "feat": feat,
            })

    return entries

def suggest_bases_and_best_skins(
    photo_img: Image.Image,
    index_entries: list[dict],
    top_bases: int = 5,
    per_base_limit: int = 1,
) -> tuple[list[str], dict[str, str]]:
    if not index_entries:
        return [], {}

    q = _feature_from_image(photo_img, size=96)

    feats = np.stack([e["feat"] for e in index_entries], axis=0)
    sims = feats @ q  # cosine similarity

    # ✅ desempate determinístico (muda por foto, mas não fica “aleatório”)
    seed = int.from_bytes(hashlib.blake2b(q.tobytes(), digest_size=8).digest(), "big")
    rng = np.random.default_rng(seed)
    sims = sims + rng.normal(0.0, 1e-6, size=sims.shape).astype(np.float32)

    base_feats: dict[str, list[np.ndarray]] = {}
    for entry in index_entries:
        base_feats.setdefault(entry["base"], []).append(entry["feat"])

    base_scores: dict[str, float] = {}
    best_skin_by_base: dict[str, str] = {}
    best_sim_by_base: dict[str, float] = {}

    for base, feats_list in base_feats.items():
        base_stack = np.stack(feats_list, axis=0)
        centroid = base_stack.mean(axis=0)
        centroid /= (np.linalg.norm(centroid) + 1e-8)
        centroid_sim = float(centroid @ q)

        idxs = [i for i, e in enumerate(index_entries) if e["base"] == base]
        if idxs:
            base_sims = sims[idxs]
            top_idx = idxs[int(np.argmax(base_sims))]
            top_sim = float(sims[top_idx])
            best_sim_by_base[base] = top_sim
            best_skin_by_base[base] = index_entries[top_idx]["name"]

            k = max(1, min(per_base_limit, base_sims.size))
            topk = np.sort(base_sims)[-k:]
            top_sim_agg = float(topk.mean())
        else:
            top_sim_agg = centroid_sim

        base_scores[base] = 0.60 * centroid_sim + 0.40 * top_sim_agg

    ranked_bases = sorted(base_scores.items(), key=lambda item: -item[1])
    bases_sugeridas = [b for b, _ in ranked_bases[:top_bases]]

    # preenche melhor skin para bases sugeridas (por segurança)
    for base in bases_sugeridas:
        if base in best_skin_by_base:
            continue
        base_idxs = [i for i, e in enumerate(index_entries) if e["base"] == base]
        if not base_idxs:
            continue
        top_idx = base_idxs[int(np.argmax(sims[base_idxs]))]
        best_skin_by_base[base] = index_entries[top_idx]["name"]

    return bases_sugeridas, best_skin_by_base


@st.cache_data(show_spinner=False)
def build_trainer_avatar_lookup() -> dict[str, dict]:
    catalog = build_trainer_avatar_catalog()
    lookup = {}
    for items in catalog.values():
        for item in items:
            lookup[item["name"]] = item
    return lookup

def crop_center_square(image: Image.Image) -> Image.Image:
    w, h = image.size
    size = min(w, h)
    left = (w - size) // 2
    top = (h - size) // 2
    return image.crop((left, top, left + size, top + size))

def image_fingerprint(image: Image.Image, size: int = 16) -> list[float]:
    img = image.convert("RGB").resize((size, size), Image.Resampling.LANCZOS)
    pixels = list(img.getdata())
    return [channel / 255.0 for pixel in pixels for channel in pixel]

@st.cache_data(show_spinner=False)
def avatar_fingerprint_for_path(path: str) -> list[float] | None:
    try:
        img = Image.open(path).convert("RGB")
    except Exception:
        return None
    return image_fingerprint(img)

@st.cache_data(show_spinner=False)
def build_avatar_base_vectors() -> dict[str, list[float]]:
    catalog = build_trainer_avatar_catalog()
    base_vectors: dict[str, list[float]] = {}
    for base, items in catalog.items():
        vectors = []
        for item in items:
            vec = avatar_fingerprint_for_path(item["path"])
            if vec:
                vectors.append(vec)
        if vectors:
            base_vectors[base] = [sum(vals) / len(vals) for vals in zip(*vectors)]
    return base_vectors

def pick_similar_avatar_bases(image: Image.Image, limit: int = 5) -> list[str]:
    base_vectors = build_avatar_base_vectors()
    if not base_vectors:
        return []
    target_vec = image_fingerprint(image)
    scored = []
    for base, vec in base_vectors.items():
        dist = sum((a - b) ** 2 for a, b in zip(target_vec, vec))
        scored.append((dist, base))
    scored.sort(key=lambda item: item[0])
    return [base for _, base in scored[:limit]]

@st.cache_data(show_spinner=False)
def load_trainer_avatar_image(path: str) -> Image.Image | None:
    try:
        return Image.open(path).convert("RGBA")
    except Exception:
        return None

def get_selected_trainer_avatar(user_data: dict) -> tuple[str | None, str | None]:
    profile = user_data.get("trainer_profile", {})
    choice = profile.get("avatar_choice")
    if not choice:
        return None, None
    lookup = build_trainer_avatar_lookup()
    avatar = lookup.get(choice)
    if not avatar:
        return choice, None
    return choice, avatar.get("path")

def get_trainer_photo_thumb(user_data: dict) -> str | None:
    profile = user_data.get("trainer_profile", {})
    thumb_b64 = profile.get("photo_thumb_b64")
    if not thumb_b64:
        return None
    return f"data:image/png;base64,{thumb_b64}"


    
def extract_strategies(text):
    if not isinstance(text, str): return []
    pattern = r'(?:^|\n)\s*(?:\*\*|[\-\>])?\s*([CFS][ODFIC][RL])\b'
    matches = re.findall(pattern, text)
    return matches


# -----------------------------
# Viabilidade -> Arquétipos (presets)
# -----------------------------

def _viab_split_blocks(viab_text: str) -> list[str]:
    """Divide um texto de Viabilidade em blocos por código (ex: CDL/FIR/...)."""
    if not isinstance(viab_text, str):
        return []
    t = viab_text.strip()
    if not t:
        return []
    # split por início de linha com código de 3 letras (tolerando bullets)
    pattern = r"(?=(?:^|\n)\s*(?:\*\*|[\-•>])?\s*[CFS][ODFIC][RL]\b)"
    parts = re.split(pattern, t, flags=re.IGNORECASE)
    blocks = [p.strip() for p in parts if p.strip()]
    return blocks if blocks else [t]

def _viab_extract_code(block: str) -> str | None:
    if not isinstance(block, str):
        return None
    m = re.match(r"^\s*(?:\*\*|[\-•>])?\s*([CFS][ODFIC][RL])\b", block.strip(), flags=re.IGNORECASE)
    return m.group(1).upper() if m else None

def _viab_extract_ability(block: str) -> str | None:
    if not isinstance(block, str):
        return None
    # "Habilidade: chlorophyll" / "habilidade: Torrent" / "a habilidade Rain Dish"
    m = re.search(r"\bhabilidade\b\s*:\s*([^\.\n;]+)", block, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip()
    m2 = re.search(r"\ba\s+habilidade\b\s*[:\-]?\s*([A-Za-z0-9' _\-]+)", block, flags=re.IGNORECASE)
    if m2:
        return m2.group(1).strip()
    return None

def _viab_extract_partners(block: str) -> list[str]:
    if not isinstance(block, str):
        return []
    m = re.search(r"\bparceiros\b\s*:\s*(.+?)(?:\.|\n|\bexplica[cç][aã]o\b|$)", block, flags=re.IGNORECASE)
    if not m:
        return []
    seg = m.group(1)
    # remove parênteses e descrições
    seg = re.sub(r"\(.*?\)", "", seg)
    seg = seg.replace(" e ", ", ")
    parts = [p.strip(" .;:-").upper() for p in seg.split(",") if p.strip(" .;:-")]
    # remove vazios/duplicados mantendo ordem
    seen = set()
    out = []
    for p in parts:
        if p and p not in seen:
            out.append(p)
            seen.add(p)
    return out

def _viab_extract_moves_tokens(block: str) -> list[str]:
    """Extrai uma lista 'crua' de tokens de golpes a partir do texto de Viabilidade."""
    if not isinstance(block, str):
        return []
    b = block.strip()
    if not b:
        return []

    seg = None
    m = re.search(
        r"\bcom\s+os\s+golpes\b\s*:\s*(.+?)(?:\.|\n|\be\s+a\s+habilidade\b|\bhabilidade\b\s*:|\bparceiros\b\s*:|\bexplica[cç][aã]o\b|$)",
        b,
        flags=re.IGNORECASE,
    )
    if m:
        seg = m.group(1)
    else:
        # fallback: tudo após o primeiro ":" (até habilidade/parceiros/explicação)
        if ":" in b:
            seg = b.split(":", 1)[1]
            seg = re.split(r"\b(?:e\s+a\s+habilidade|habilidade\s*:|parceiros\s*:|explica[cç][aã]o)\b", seg, flags=re.IGNORECASE)[0]

    if not seg:
        return []

    seg = re.sub(r"\(.*?\)", "", seg)
    seg = seg.replace(" e ", ", ")
    seg = seg.replace(";", ",")
    tokens = [t.strip(" .;:-") for t in seg.split(",") if t.strip(" .;:-")]
    # remove duplicados mantendo ordem
    seen = set()
    out = []
    for t in tokens:
        key = _norm(t)
        if key and key not in seen:
            out.append(t)
            seen.add(key)
    return out

def _viab_parse_archetypes(viab_text: str) -> list[dict]:
    """Retorna lista de arquétipos extraídos da coluna Viabilidade."""
    blocks = _viab_split_blocks(viab_text)
    out = []
    for b in blocks:
        code = _viab_extract_code(b)
        moves_tokens = _viab_extract_moves_tokens(b)
        ability = _viab_extract_ability(b)
        partners = _viab_extract_partners(b)
        # label amigável
        label = (b.split(":", 1)[0] if ":" in b else (code or "Arquétipo")).strip()
        if code and not label.upper().startswith(code):
            label = f"{code} — {label}"
        out.append({
            "code": code,
            "label": label[:80] + ("…" if len(label) > 80 else ""),
            "raw": b,
            "moves_tokens": moves_tokens,
            "ability": ability,
            "partners": partners,
        })
    return out

def _viab_find_moves_in_text(db: "MoveDB", text_blob: str) -> list["Move"]:
    """Fallback: encontra golpes do banco que aparecem como substring em um trecho de texto."""
    if db is None or not isinstance(text_blob, str) or not text_blob.strip():
        return []
    t = text_blob.lower()
    # normaliza pontuação -> espaço
    t = re.sub(r"[^a-z0-9'\-\s]", " ", t)
    t = t.replace("-", " ")
    t = re.sub(r"\s+", " ", t).strip()

    found: list[Move] = []
    seen = set()

    # varre nomes do banco (875 aprox. -> ok)
    try:
        names = db.df["Nome"].dropna().astype(str).tolist()
    except Exception:
        names = []

    # prioriza nomes maiores para reduzir falsos positivos
    names.sort(key=lambda x: len(str(x)), reverse=True)

    for nm in names:
        nml = str(nm).lower()
        nml = re.sub(r"[^a-z0-9'\-\s]", " ", nml).replace("-", " ")
        nml = re.sub(r"\s+", " ", nml).strip()
        if not nml:
            continue
        # boundary simples
        if re.search(r"(?<![a-z0-9])" + re.escape(nml) + r"(?![a-z0-9])", t):
            mv = db.get_by_name(nm)
            if mv:
                k = _norm(mv.name)
                if k and k not in seen:
                    found.append(mv)
                    seen.add(k)
        if len(found) >= 20:
            break

    return found

def _viab_suggest_moves(db: "MoveDB", archetype: dict, include_fulltext: bool = True) -> tuple[list["Move"], list[str]]:
    """Dado um arquétipo, sugere Moves (encontrados no banco) e lista tokens não encontrados.

    - Quando include_fulltext=True, faz um *fallback* varrendo o texto completo do arquétipo para achar nomes do banco.
    - Para o modo "núcleo", use include_fulltext=False para ficar fiel apenas aos golpes explicitamente listados.
    """
    if db is None or not archetype:
        return [], []

    tokens = archetype.get("moves_tokens", []) or []
    found: list[Move] = []
    missing: list[str] = []
    seen = set()

    for tok in tokens:
        mv = _try_match_move_in_db(db, tok)
        if mv:
            k = _norm(mv.name)
            if k and k not in seen:
                found.append(mv)
                seen.add(k)
            continue

        # token pode vir colado (ex: "Withdraw Dark Pulse")
        sub = _viab_find_moves_in_text(db, tok)
        if sub:
            for mv2 in sub:
                k = _norm(mv2.name)
                if k and k not in seen:
                    found.append(mv2)
                    seen.add(k)
        else:
            missing.append(tok)

    # fallback final: varrer no texto completo (apenas no modo completo)
    if include_fulltext:
        full_hits = _viab_find_moves_in_text(db, archetype.get("raw", ""))
        for mv in full_hits:
            k = _norm(mv.name)
            if k and k not in seen:
                found.append(mv)
                seen.add(k)

    return found, missing


def _viab_apply_to_session(
    archetype: dict,
    rank_default: int,
    excel_path: str,
    mode: str = "full",
    core_n: int = 3,
) -> tuple[list[str], list[str], str | None]:
    """Aplica um arquétipo: adiciona golpes encontrados e tenta definir habilidade sugerida.

    mode:
      - "full": adiciona o kit completo (lista inteira + fallback no texto).
      - "core": adiciona só o núcleo (primeiros core_n golpes explicitamente listados; sem fallback).
    """
    try:
        db = load_move_db(excel_path)
    except Exception:
        db = None

    if db is None:
        return [], ["(Não consegui carregar o banco de golpes)"], None

    mode = (mode or "full").strip().lower()
    core_n = int(core_n or 0)

    if mode == "core":
        arch2 = dict(archetype or {})
        toks = (arch2.get("moves_tokens") or [])
        if core_n > 0:
            toks = toks[:core_n]
        arch2["moves_tokens"] = toks
        moves_found, missing = _viab_suggest_moves(db, arch2, include_fulltext=False)
    else:
        moves_found, missing = _viab_suggest_moves(db, archetype, include_fulltext=True)

    existing = {_norm(m.get("name")) for m in (st.session_state.get("cg_moves") or []) if isinstance(m, dict)}
    added_names: list[str] = []

    for mv in moves_found:
        k = _norm(mv.name)
        if not k or k in existing:
            continue
        # Usa o motor MM3e para calcular PP quando disponível
        move_entry = _cg_confirm_move_with_engine(mv, int(rank_default))
        st.session_state["cg_moves"].append(move_entry)
        existing.add(k)
        added_names.append(mv.name)

    # habilidade sugerida
    ability_suggested = archetype.get("ability")
    if isinstance(ability_suggested, str):
        ability_suggested = ability_suggested.strip()
    else:
        ability_suggested = None

    return added_names, missing, ability_suggested

def calculate_power_level(row, cols_map):
    score = 0
    rarity = normalize_text(row.get('Raridade', ''))
    if 'trio' in rarity: score += 10
    elif 'fóssil' in rarity or 'fossil' in rarity: score += 7
    elif 'ultra' in rarity or 'super' in rarity: score += 5
    elif 'raro' in rarity: score += 3
    else: score += 1
    
    types = str(row.get('Tipo', '')).split('/')
    for t in types:
        t = normalize_text(t)
        if t in ['psychic', 'ghost', 'dragon', 'psiquico', 'fantasma', 'dragao']: score += 3
        elif t in ['steel', 'ice', 'dark', 'fighting', 'fairy', 'poison', 'metal', 'gelo', 'noturno', 'lutador', 'fada', 'veneno']: score += 2
        else: score += 1
        
    try:
        if cols_map.get('estagio'):
            val = pd.to_numeric(row.get(cols_map.get('estagio'), 0), errors='coerce')
            score += 0 if pd.isna(val) else val
    except: pass

    try:
        if cols_map.get('evolucao'):
            val = pd.to_numeric(row.get(cols_map.get('evolucao'), 0), errors='coerce')
            score += 0 if pd.isna(val) else val
    except: pass
    
    final_score = int(score)
    if final_score > 15: final_score = 15
    return final_score



# ----------------------------
# Pokédex Card UI helpers (Tipos, Viabilidade)
# ----------------------------

TYPE_COLORS = {
    # padrão de cores (franquia / TCG-like)
    "normal":  "#A8A77A",
    "fogo":    "#EE8130",
    "fire":    "#EE8130",
    "agua":    "#6390F0",
    "água":    "#6390F0",
    "water":   "#6390F0",
    "planta":  "#7AC74C",
    "grass":   "#7AC74C",
    "eletrico":"#F7D02C",
    "elétrico":"#F7D02C",
    "electric":"#F7D02C",
    "gelo":    "#96D9D6",
    "ice":     "#96D9D6",
    "lutador": "#C22E28",
    "fighting":"#C22E28",
    "veneno":  "#A33EA1",
    "poison":  "#A33EA1",
    "terra":   "#E2BF65",
    "ground":  "#E2BF65",
    "voador":  "#A98FF3",
    "flying":  "#A98FF3",
    "psiquico":"#F95587",
    "psíquico":"#F95587",
    "psychic": "#F95587",
    "inseto":  "#A6B91A",
    "bug":     "#A6B91A",
    "pedra":   "#B6A136",
    "rock":    "#B6A136",
    "fantasma":"#735797",
    "ghost":   "#735797",
    "dragao":  "#6F35FC",
    "dragão":  "#6F35FC",
    "dragon":  "#6F35FC",
    "noturno": "#705746",
    "dark":    "#705746",
    "metal":   "#B7B7CE",
    "steel":   "#B7B7CE",
    "fada":    "#D685AD",
    "fairy":   "#D685AD",
}

_VIAB_RE = re.compile(r"\b([CFS])\s*([ODFIC])\s*([RL])\b", re.IGNORECASE)

def _extract_viab_code_from_text(txt: str) -> str:
    if not txt:
        return "---"
    m = _VIAB_RE.search(str(txt))
    if not m:
        return "---"
    return (m.group(1) + m.group(2) + m.group(3)).upper()

def _split_types(raw: str) -> list[str]:
    if not raw:
        return []
    parts = [p.strip() for p in str(raw).split("/") if str(p).strip()]
    return parts[:2]

def _type_color(t: str) -> str:
    if not t:
        return "#64748b"
    key = normalize_text(str(t))
    return TYPE_COLORS.get(key, "#64748b")

# Status icons (SVG inline - sem assets externos)
_SVG_POKEBALL = """<svg viewBox='0 0 24 24' width='16' height='16' fill='none' xmlns='http://www.w3.org/2000/svg'><circle cx='12' cy='12' r='10' stroke='white' stroke-width='2'/><path d='M2 12h20' stroke='white' stroke-width='2'/><circle cx='12' cy='12' r='3' fill='white'/></svg>"""
_SVG_STAR = """<svg viewBox='0 0 24 24' width='16' height='16' fill='white' xmlns='http://www.w3.org/2000/svg'><path d='M12 17.3l-6.18 3.25 1.18-6.88L1.99 8.95l6.91-1 3.1-6.26 3.1 6.26 6.91 1-5 4.72 1.18 6.88z'/></svg>"""
_SVG_EYE = """<svg viewBox='0 0 24 24' width='16' height='16' fill='none' xmlns='http://www.w3.org/2000/svg'><path d='M2 12s3.5-7 10-7 10 7 10 7-3.5 7-10 7S2 12 2 12z' stroke='white' stroke-width='2'/><circle cx='12' cy='12' r='3' fill='white'/></svg>"""

@st.cache_data(ttl=7200, show_spinner=False)
def load_excel_data():
    # Preferência: usar a pokédex nova se estiver no projeto com esse nome.
    # Mantém compatibilidade com o nome antigo (pokedex.xlsx).
    candidates = ["pokedex Nova.xlsx", "pokedex.xlsx"]
    file_name = next((n for n in candidates if os.path.exists(n)), None)
    if not file_name:
        return None, None
    try:
        df = pd.read_excel(file_name)
        df.columns = [c.strip() for c in df.columns]
        cols_map = {}
        for col in df.columns:
            norm_col = normalize_text(col)
            if 'estagio' in norm_col: cols_map['estagio'] = col
            if 'evolucao' in norm_col or 'evolution' in norm_col: cols_map['evolucao'] = col

        df['Região'] = df['Região'].fillna('Desconhecida').astype(str)
        df['Biomas'] = df['Biomas'].fillna('Desconhecido').astype(str)
        df['Nome'] = df['Nome'].fillna('Desconhecido')
        df['Viabilidade'] = df['Viabilidade'].fillna('Sem dados.')
        if 'Nº' in df.columns:
            df['Nº'] = df['Nº'].astype(str).str.replace('#', '')
            
        df['Codigos_Estrategia'] = df['Viabilidade'].apply(extract_strategies)
        df['Nivel_Poder'] = df.apply(lambda row: calculate_power_level(row, cols_map), axis=1)
        return df, cols_map
    except Exception as e:
        st.error(f"Erro ao ler Excel: {e}")
        return None, None

api_name_map = get_official_pokemon_map()

def pokemon_pid_to_image(pid: str, mode: str = "artwork", shiny: bool = False) -> str:
    if not pid:
        return "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"
    pid_str = str(pid).strip()
    pid_norm = _norm_pid(pid_str)
    
    # Caso 1: Visitante (EXT)
    if pid_str.startswith("EXT:"):
        name = pid_str.replace("EXT:", "")
        return get_pokemon_image_url(name, api_name_map, mode=mode, shiny=shiny)
        
    # Caso 2: Busca no EXCEL (Correção do ID Regional)
    # Procura o ID no Excel para pegar o NOME correto
    if 'df' in globals() or 'df' in st.session_state:
        # Tenta pegar o df de onde estiver disponível
        local_df = st.session_state.get('df_data') if 'df_data' in st.session_state else df
        
        # Compara usando PID normalizado para evitar falhas em ids como "18" vs "18.0"
        row = local_df[local_df["Nº"].apply(_norm_pid) == pid_norm]
        if not row.empty:
            # Pega o nome (ex: "MyStarter")
            real_name = row.iloc[0]["Nome"]
            # Busca a imagem pelo NOME, não pelo número
            return get_pokemon_image_url(real_name, api_name_map, mode=mode, shiny=shiny)

    # Fallback: Se não achou no Excel, retorna erro ou tenta direto (mas evita erro de imagem quebrada)
    return "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"


# load_excel_data já tem @st.cache_data – basta chamar diretamente
# (result é cacheado em memória entre reruns e sessões; session_state serve só de alias)
if 'df_data' not in st.session_state:
    st.session_state['df_data'], st.session_state['cols_map'] = load_excel_data()

df = st.session_state['df_data']
cols_map = st.session_state.get('cols_map', {})

# =========================
# Limpeza única: remove resquícios do sistema __dex_uid
# =========================
try:
    _dex_cleanup_once(user_data, df)
except Exception:
    pass

if st.session_state.get("show_login_menu"):
    render_login_menu(trainer_name, user_data)
    st.stop()


@st.cache_data(show_spinner=False)
def _enc_load_pokedex_cached():
    if not _ENC_AVAILABLE or _enc_carregar_pokedex is None:
        return [], True
    return _enc_carregar_pokedex()


def _enc_region_options() -> list[str]:
    if isinstance(_ENC_CLIMAS_POR_REGIAO, dict) and _ENC_CLIMAS_POR_REGIAO:
        return sorted(_ENC_CLIMAS_POR_REGIAO.keys())
    return []


def _enc_biome_options(pokedex: list, regiao: str) -> list[str]:
    biomas = set()
    for p in pokedex or []:
        regs = getattr(p, "regioes", []) or []
        if any(str(regiao or "").strip().lower() in str(r).strip().lower() for r in regs):
            for b in (getattr(p, "biomas", []) or []):
                b = str(b).strip()
                if b:
                    biomas.add(b)
    return sorted(biomas)


def render_encounter_generator_page() -> None:
    st.title("🎲 Gerador de Encontros")
    st.caption("Ferramenta para criar encontros narrativos de Ga'Al usando a Pokédex.")

    if trainer_name != "Ezenek":
        st.warning("Esta aba é exclusiva do perfil Ezenek.")
        return

    if not _ENC_AVAILABLE:
        st.error("Não foi possível carregar 'encounter_generator.py'.")
        return

    pokedex, is_mock = _enc_load_pokedex_cached()
    if not pokedex:
        st.error("Pokédex indisponível para gerar encontros.")
        return

    if is_mock:
        st.info("Pokédex em modo mock (arquivo principal não encontrado).")

    col1, col2, col3 = st.columns(3)
    with col1:
        region_opts = _enc_region_options() or ["Deserto Irrigado"]
        regiao = st.selectbox("Região", options=region_opts, index=0)

    with col2:
        biome_opts = _enc_biome_options(pokedex, regiao)
        if not biome_opts:
            biome_opts = ["Dunas"]
        bioma = st.selectbox("Bioma", options=biome_opts, index=0)

    with col3:
        horario = st.selectbox("Horário", ["Manhã", "Tarde", "Anoitecer", "Noite", "Madrugada"], index=1)

    quantidade = st.slider("Quantidade de encontros", min_value=1, max_value=5, value=1)

    if st.button("🎲 Gerar encontros", type="primary"):
        for i in range(quantidade):
            txt = _enc_gerar_encontro(pokedex, regiao, bioma, horario)
            if quantidade > 1:
                st.markdown(f"### Encontro {i + 1}")
            st.code(txt)

# --- INTERFACE PRINCIPAL ---

st.sidebar.title("📱 Menu")
st.sidebar.markdown(f"**Treinador:** {trainer_name}")

if st.sidebar.button("💾 Salvar na Nuvem"):
    if save_data_cloud(trainer_name, user_data):
        st.sidebar.success("Salvo com sucesso!")

if st.sidebar.button("🚪 Sair (Logout)"):
    del st.session_state['trainer_name']
    st.rerun()


if st.sidebar.button("🔄 Recarregar Excel"):
    load_excel_data.clear()  # limpa cache do @st.cache_data
    st.session_state.pop('df_data', None)
    st.session_state.pop('cols_map', None)
    st.rerun()

# --- navegação programática (antes do radio key="page") ---
if "nav_to" in st.session_state:
    st.session_state.update({"page": st.session_state.pop("nav_to")})

st.sidebar.markdown("---")
menu_pages = [
    "Pokédex (Busca)",
    "Trainer Hub (Meus Pokémons)",
    "Criação Guiada de Fichas",
    "Minhas Fichas",
    "PvP – Arena Tática",
    "Mochila",
    "Compendium de Ga'Al",
]
if trainer_name == "Ezenek":
    menu_pages.append("Gerador de Encontros")

page = st.sidebar.radio(
    "Ir para:",
    menu_pages,
    key="page",
)
# 👇 COLE AQUI (uma linha abaixo do radio)
track_path = "music/geral.mp3"

if page == "Compendium de Ga'Al":
    track_path = "music/compendium.mp3"

elif page == "PvP – Arena Tática":
    if st.session_state.get("pvp_view", "lobby") == "battle":
        track_path = "music/pvp.mp3"
    else:
        track_path = "music/geral.mp3"

render_bgm(track_path, volume=0.25)

if page != "PvP – Arena Tática":
    stop_pvp_sync_listener()



# ==============================================================================
# COMPENDIUM DE GA'AL — Aba completa (UPGRADED)
# - Split view (lista à esquerda / dossiê à direita)
# - Breadcrumbs + navegação Próximo/Anterior
# - Favoritos + Recentes (Locais e NPCs)
# - Busca global full-text (Cidades, Sublocais, NPCs, Regiões)
# - Cross-links (cidades/NPCs mencionados)
# - Sprites de Pokémon (cache local offline)
# - Tags automáticas + overrides em JSON (exportável)
# ==============================================================================

from pathlib import Path

try:
    from docx import Document
except Exception:
    Document = None


# ----------------------------
# CONFIG
# ----------------------------
# (Preferência atual) Compendium em JSON
COMP_JSON_LOCAIS = "gaal_locais.json"
COMP_JSON_GINASIOS = "gaal_ginasios.json"          # opcional (detalhes/staff via NPCs)
COMP_JSON_NPCS_VIVOS = "gaal_npcs_vivos.json"
COMP_JSON_NPCS_MORTOS = "gaal_npcs_mortos.json"    # opcional

# (Fallback legado) Compendium em DOCX
COMP_DOC_LOCAIS = "GAAL_Banco_de_Dados_Locais_Unificado_v2.docx"
COMP_DOC_NPCS_VIVOS = "Npcs_Pokemon_vivos_profundo_v2.docx"
COMP_DOC_NPCS_MORTOS = "Npcs_Pokemon_mortos.docx"  # opcional
COMP_DOC_GINASIOS = "ginasios.docx"                # opcional (detalhes extras de ginásios)
COMP_DEFAULT_MAP = "GaAl_2.jpg"  # mapa geral (fallback)


COMP_REGIOES_PRINCIPAIS = [
    "Baía Morta",
    "Baixo Deserto",
    "Campos Longos",
    "Chifre de Ga'Al",
    "Deserto Irrigado",
    "Terras Fluviais",
]

# Overrides opcionais via arquivo (sem mexer no código):
# Estrutura:
# {
#   "cities": {"Obsidian": ["Industrial","Portuária"]},
#   "npcs": {"Dra. Tulsi": ["Cientista","Ranger"]}
# }
COMP_TAGS_JSON = "compendium_tags.json"
COMP_SESSIONS_JSON = "compendium_sessions.json"


# ----------------------------
# THEME
# ----------------------------
def apply_compendium_theme() -> None:
    st.markdown(
        """
        <style>
        [data-testid="stAppViewContainer"]{
            background:
                radial-gradient(900px 500px at 12% 0%, rgba(250, 204, 21, 0.10), transparent 60%),
                radial-gradient(1200px 700px at 92% 10%, rgba(244, 114, 182, 0.10), transparent 60%),
                linear-gradient(180deg, #0c0a0a 0%, #0f172a 55%, #09080b 100%);
            color: #f8fafc;
        }
        [data-testid="stHeader"]{ background: transparent !important; }
        h1,h2,h3,h4{ color:#f8fafc !important; letter-spacing:-0.02em; }
        .stMarkdown, .stMarkdown p, label { color: #e2e8f0 !important; }

        .comp-shell{
            border: 1px solid rgba(148,163,184,0.18);
            background: rgba(15,23,42,0.55);
            border-radius: 22px;
            padding: 18px 18px;
            box-shadow: 0 24px 70px rgba(0,0,0,0.45);
            backdrop-filter: blur(12px);
        }
        .comp-hero{
            border: 1px solid rgba(203,213,225,0.18);
            background: linear-gradient(135deg, rgba(17,24,39,0.72), rgba(2,6,23,0.55));
            border-radius: 20px;
            padding: 18px;
            box-shadow: 0 18px 50px rgba(0,0,0,0.35);
        }
        .comp-hero-title{
            font-size: 32px;
            font-weight: 800;
            letter-spacing: -0.02em;
            margin-bottom: 6px;
        }
        .comp-hero-kicker{
            color: rgba(248,250,252,0.70);
            font-size: 14px;
        }
        .comp-muted{ color: rgba(226,232,240,0.72); }
        .comp-cta{
            display: inline-flex;
            align-items: center;
            gap: 8px;
            padding: 8px 14px;
            border-radius: 999px;
            border: 1px solid rgba(250,204,21,0.35);
            background: rgba(250,204,21,0.12);
            color: #fde68a;
            font-weight: 600;
        }
        .comp-panel{
            border: 1px solid rgba(148,163,184,0.16);
            background: rgba(15,23,42,0.45);
            border-radius: 18px;
            padding: 14px;
        }
        .comp-card{
            border: 1px solid rgba(148,163,184,0.18);
            background: rgba(2,6,23,0.55);
            border-radius: 16px;
            padding: 12px;
            box-shadow: 0 18px 45px rgba(0,0,0,0.30);
        }
        .comp-divider{
            height: 1px;
            background: rgba(148,163,184,0.18);
            margin: 12px 0;
        }
        .comp-block{
            border: 1px solid rgba(148,163,184,0.12);
            background: rgba(15,23,42,0.45);
            border-radius: 16px;
            padding: 12px;
            margin: 12px 0;
        }
        .comp-row{
            border: 1px solid rgba(148,163,184,0.12);
            background: rgba(2,6,23,0.35);
            border-radius: 16px;
            padding: 10px 12px;
            margin: 10px 0;
        }
        .comp-mini{
            font-size: 12px;
            color: rgba(226,232,240,0.70);
        }
        .comp-map{
            border-radius: 16px;
            border: 1px solid rgba(148,163,184,0.2);
            box-shadow: 0 12px 30px rgba(0,0,0,0.35);
        }
        .comp-stat{
            border: 1px solid rgba(148,163,184,0.2);
            background: rgba(15,23,42,0.65);
            border-radius: 14px;
            padding: 10px 12px;
            text-align: center;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

def render_top_menu_compendium(selected: str):
    st.markdown("<div class='ds-home'>", unsafe_allow_html=True)

    opts = [("menu", "Menu"),
            ("npcs", "NPCs"),
            ("ginasios", "Ginásios"),
            ("locais", "Locais"),
            ("sair", "Sair")]

    for key, label in opts:
        if st.button(label, key=f"top_{key}"):
            st.session_state["comp_view"] = key
            st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)


# ----------------------------
# ASSETS (index rápido)
# ----------------------------
def _comp_base_dirs() -> list[str]:
    try:
        base = os.path.dirname(__file__)
    except Exception:
        base = os.getcwd()

    roots = [
        os.getcwd(),
        base,
        os.path.join(base, "assets"),
        os.path.join(base, "Assets"),
        os.path.join(os.getcwd(), "assets"),
        os.path.join(os.getcwd(), "Assets"),
        os.path.join(base, "data"),
        os.path.join(os.getcwd(), "data"),
        os.path.join(base, "assets", "compendium"),
        os.path.join(base, "assets", "locais"),
        os.path.join(base, "assets", "cities"),
        os.path.join(base, "assets", "npcs"),
        os.path.join(base, "assets", "portraits"),
        os.path.join(base, "cidades"),
        os.path.join(base, "treinadores"),
        os.path.join(os.getcwd(), "assets", "compendium"),
        os.path.join(os.getcwd(), "assets", "locais"),
        os.path.join(os.getcwd(), "assets", "cities"),
        os.path.join(os.getcwd(), "assets", "npcs"),
        os.path.join(os.getcwd(), "assets", "portraits"),
        os.path.join(os.getcwd(), "cidades"),
        os.path.join(os.getcwd(), "treinadores"),
        os.path.join(base, "Assets", "insignias"),
        os.path.join(os.getcwd(), "Assets", "insignias"),
    ]
    uniq = []
    for r in roots:
        if r and os.path.exists(r) and r not in uniq:
            uniq.append(r)
    return uniq


def _stem_key(s: str) -> str:
    x = _norm(s)
    x = x.replace("vila de ", "")
    x = x.replace("—", "-")
    x = re.sub(r"[^\w\s-]+", "", x)
    x = x.replace("-", " ")
    x = re.sub(r"\s+", " ", x).strip()
    x = x.replace(" ", "_")
    return x


@st.cache_data(show_spinner=False)
def _build_image_index(roots: tuple[str, ...], exclude_dirs: tuple[str, ...] | None = None) -> dict[str, dict]:
    exts = {".png", ".jpg", ".jpeg", ".webp"}
    by_key: dict[str, str] = {}
    exclude = {d.lower() for d in (exclude_dirs or ())}
    for root in roots:
        for dirpath, dirnames, files in os.walk(root):
            if exclude:
                dirnames[:] = [d for d in dirnames if d.lower() not in exclude]
            for fn in files:
                ext = os.path.splitext(fn)[1].lower()
                if ext not in exts:
                    continue
                stem = os.path.splitext(fn)[0]
                k = _stem_key(stem)
                p = os.path.join(dirpath, fn)
                if k and (k not in by_key or len(p) < len(by_key[k])):
                    by_key[k] = p

                # index adicional sem sufixos numéricos
                base = re.sub(r"[_-]?\d+$", "", stem).strip()
                if base:
                    kb = _stem_key(base)
                    if kb and kb not in by_key:
                        by_key[kb] = p
    return {"by_key": by_key, "keys": sorted(by_key.keys())}


def _filter_roots(roots: list[str], exclude_dirs: set[str] | None = None) -> list[str]:
    if not exclude_dirs:
        return roots
    trimmed = []
    for root in roots:
        parts = {part.lower() for part in Path(root).parts}
        if parts.isdisjoint(exclude_dirs):
            trimmed.append(root)
    return trimmed


def comp_find_image(
    name: str,
    exclude_dirs: set[str] | None = None,
    roots: list[str] | None = None,
) -> str | None:
    roots = roots or _comp_base_dirs()
    roots = _filter_roots(roots, exclude_dirs)
    roots = tuple(roots)
    exclude = tuple(sorted(d.lower() for d in (exclude_dirs or set())))
    idx = _build_image_index(roots, exclude or None)
    by_key = idx.get("by_key", {})
    all_keys = idx.get("keys", [])

    variants = [
        _stem_key(name),
        _stem_key(name).replace("_", ""),
        _stem_key(name).replace("_", "-"),
    ]
    for k in variants:
        if k in by_key and os.path.exists(by_key[k]):
            return by_key[k]

    # tenta aproximação por similaridade
    for target in variants:
        if not target:
            continue
        close = difflib.get_close_matches(target, all_keys if all_keys else [], n=1, cutoff=0.78)
        if close:
            p = by_key.get(close[0])
            if p and os.path.exists(p):
                return p
        if all_keys:
            for k in all_keys:
                if target in k or k in target:
                    p = by_key.get(k)
                    if p and os.path.exists(p):
                        return p

    exts = [".png", ".jpg", ".jpeg", ".webp"]
    for ext in exts:
        fname = f"{_stem_key(name)}{ext}"
        try:
            p = _resolve_asset_path(fname)
        except Exception:
            p = fname
        if p and os.path.exists(p):
            return p

    return None


def comp_find_npc_image(name: str) -> str | None:
    try:
        base = os.path.dirname(__file__)
    except Exception:
        base = os.getcwd()
    npc_roots = [
        os.path.join(base, "treinadores"),
        os.path.join(os.getcwd(), "treinadores"),
    ]
    return comp_find_image(name, roots=npc_roots)


# ----------------------------
# TAG OVERRIDES (JSON + sessão)
# ----------------------------
def _load_comp_tags_overrides() -> dict:
    roots = _comp_base_dirs()
    # tenta localizar o json nos roots
    candidates = []
    for r in roots:
        candidates.append(os.path.join(r, COMP_TAGS_JSON))
    # também tenta via resolve
    try:
        candidates.insert(0, _resolve_asset_path(COMP_TAGS_JSON))
    except Exception:
        pass

    for p in candidates:
        if p and os.path.exists(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    obj = json.load(f)
                if isinstance(obj, dict):
                    obj.setdefault("cities", {})
                    obj.setdefault("npcs", {})
                    return obj
            except Exception:
                return {"cities": {}, "npcs": {}}
    return {"cities": {}, "npcs": {}}


def _get_live_overrides() -> dict:
    if "comp_live_overrides" not in st.session_state:
        st.session_state["comp_live_overrides"] = {"cities": {}, "npcs": {}}
    return st.session_state["comp_live_overrides"]


def _merge_tags(entity_type: str, name: str, inferred: list[str]) -> list[str]:
    base = _load_comp_tags_overrides()
    live = _get_live_overrides()
    manual = (base.get(entity_type, {}) or {}).get(name, [])
    manual2 = (live.get(entity_type, {}) or {}).get(name, [])
    tags = []
    for t in (manual + manual2 + inferred):
        if t and t not in tags:
            tags.append(t)
    return tags[:6]


# ----------------------------
# DOCX PARSERS
# ----------------------------
def _heading_level(p) -> int:
    try:
        style = (p.style.name or "")
    except Exception:
        style = ""
    m = re.match(r"Heading\s+(\d+)", style)
    return int(m.group(1)) if m else 0


def _clean_title(s: str) -> str:
    s = (s or "").strip()
    s = s.replace("—", "-")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _split_name_desc(text: str) -> tuple[str, str]:
    parts = [p.strip() for p in (text or "").split("\n") if p.strip()]
    if not parts:
        return "", ""
    if len(parts) == 1:
        return parts[0], ""
    return parts[0], "\n\n".join(parts[1:])


def _is_fake_city_h2(h2: str) -> bool:
    t = _norm(h2)
    if t.startswith("visao geral"):
        return True
    if t.startswith("anexo"):
        return True
    if t.startswith("como ler"):
        return True
    if t.startswith("ginasios "):
        return True
    if t.startswith("locais em "):
        return True
    if t.startswith("ginasio de "):
        return True
    return False


def _h2_city_from_pattern(h2: str) -> str | None:
    s = _clean_title(h2)
    m = re.match(r"^(.+?)\s*-\s*vis[aã]o geral$", _norm(s))
    if m:
        left = re.split(r"[—-]", s, maxsplit=1)[0].strip()
        return left
    return None

def _is_city_line(line: str, known_cities: set[str]) -> bool:
    t = (line or "").strip()
    if not t:
        return False
    if ":" in t:
        return False
    if len(t) > 40:
        return False
    if _norm(t) in {_norm(x) for x in known_cities}:
        return True
    if re.fullmatch(r"[A-Za-zÀ-ÿ'’\- ]+", t) and t[0].isupper():
        return True
    return False


@st.cache_data(show_spinner=False)
def parse_ginasios_docx(doc_path: str, known_cities: list[str]) -> dict:
    gyms: dict[str, dict] = {}
    if Document is None or not os.path.exists(doc_path):
        return gyms

    known = set([c for c in known_cities if c])
    doc = Document(doc_path)

    cur_city: str | None = None
    cur_person: dict | None = None
    mode = "normal"  # "history"
    history_item: dict | None = None

    def ensure_city(city: str):
        gyms.setdefault(city, {"leaders": [], "vices": [], "people": [], "history": []})

    def flush_person():
        nonlocal cur_person
        if not cur_city or not cur_person:
            cur_person = None
            return
        ensure_city(cur_city)

        occ = _norm(cur_person.get("fields", {}).get("Ocupação atual", ""))

        if "líder" in occ or "lider" in occ:
            gyms[cur_city]["leaders"].append(cur_person)
        elif "vice" in occ:
            gyms[cur_city]["vices"].append(cur_person)
        else:
            gyms[cur_city]["people"].append(cur_person)

        cur_person = None

    def flush_history():
        nonlocal history_item
        if not cur_city or not history_item:
            history_item = None
            return
        ensure_city(cur_city)
        gyms[cur_city]["history"].append(history_item)
        history_item = None

    def new_person(name: str):
        return {"name": name.strip(), "fields": {}, "pokemons": [], "text": ""}

    def add_text(obj: dict, txt: str):
        txt = (txt or "").strip()
        if not txt:
            return
        obj["text"] = (obj["text"] + "\n\n" + txt).strip()

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue

        if "registro histórico" in _norm(t):
            mode = "history"
            flush_person()
            continue

        if mode == "normal" and _is_city_line(t, known):
            flush_person()
            cur_city = _clean_title(t)
            ensure_city(cur_city)
            continue

        if mode == "history" and cur_city is None:
            cur_city = "__REGISTRO_HISTORICO__"
            ensure_city(cur_city)

        if mode == "history" and re.match(r"^\d+\.\s", t):
            flush_history()
            history_item = {"title": t.strip(), "fields": {}, "pokemons": [], "text": ""}
            continue

        if mode == "normal":
            if cur_city and cur_person is None and ":" not in t and len(t) <= 40:
                cur_person = new_person(t)
                continue

        if ":" in t:
            k, v = t.split(":", 1)
            kN = _norm(k)
            v = v.strip()

            if "pokemon" in kN:
                pokes = [x.strip() for x in re.split(r",|;|/|\|", v) if x.strip()]
                if mode == "history" and history_item is not None:
                    history_item["pokemons"].extend(pokes)
                elif cur_person is not None:
                    cur_person["pokemons"].extend(pokes)
                continue

            if mode == "history" and history_item is not None:
                history_item["fields"][k.strip()] = v
            elif cur_person is not None:
                cur_person["fields"][k.strip()] = v
            else:
                if cur_city:
                    ensure_city(cur_city)
                    gyms[cur_city].setdefault("notes", [])
                    gyms[cur_city]["notes"].append({k.strip(): v})
            continue

        if mode == "history":
            if history_item is not None:
                add_text(history_item, t)
            else:
                ensure_city(cur_city or "__REGISTRO_HISTORICO__")
                gyms[cur_city or "__REGISTRO_HISTORICO__"].setdefault("history_text", "")
                gyms[cur_city or "__REGISTRO_HISTORICO__"]["history_text"] = (
                    gyms[cur_city or "__REGISTRO_HISTORICO__"]["history_text"] + "\n\n" + t
                ).strip()
        else:
            if cur_person is not None:
                add_text(cur_person, t)
            else:
                ensure_city(cur_city or "__SEM_CIDADE__")
                gyms[cur_city or "__SEM_CIDADE__"].setdefault("city_text", "")
                gyms[cur_city or "__SEM_CIDADE__"]["city_text"] = (
                    gyms[cur_city or "__SEM_CIDADE__"]["city_text"] + "\n\n" + t
                ).strip()

    flush_person()
    flush_history()
    return gyms



@st.cache_data(show_spinner=False)
def load_compendium_data(
    locais_path: str,
    locais_mtime: float,
    npcs_vivos_path: str,
    npcs_vivos_mtime: float,
    npcs_mortos_path: str,
    npcs_mortos_mtime: float,
) -> dict:
    if Document is None:
        raise RuntimeError("python-docx não está disponível (instale python-docx).")

    data = {
        "regions": {},   # reg -> {"intro": str, "sections": {title: text}, "cities": [city]}
        "cities": {},    # city -> {"region": reg, "sections": {title: text}, "sublocais": [{"name","text"}]}
        "npcs": {},      # name -> {...}
    }

    # ---------- LOCAIS ----------
    if os.path.exists(locais_path):
        doc = Document(locais_path)

        cur_reg: str | None = None
        cur_city: str | None = None
        cur_h3: str | None = None
        inside_sublocais = False
        chifre_city_mode: str | None = None

        def ensure_region(r: str):
            data["regions"].setdefault(r, {"intro": "", "sections": {}, "cities": []})

        def ensure_city(c: str, r: str):
            data["cities"].setdefault(c, {"region": r, "sections": {}, "sublocais": []})
            if c not in data["regions"][r]["cities"]:
                data["regions"][r]["cities"].append(c)

        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if not t:
                continue

            lvl = _heading_level(p)
            t_clean = _clean_title(t)

            if lvl == 1:
                cur_reg = t_clean
                ensure_region(cur_reg)
                cur_city = None
                cur_h3 = None
                inside_sublocais = False
                chifre_city_mode = None
                continue

            if cur_reg is None:
                continue

            if lvl == 2:
                maybe_city = _h2_city_from_pattern(t_clean)
                if maybe_city:
                    cur_city = maybe_city
                    chifre_city_mode = maybe_city
                    ensure_city(cur_city, cur_reg)
                    cur_h3 = "Visão geral"
                    inside_sublocais = False
                    continue

                tn = _norm(t_clean)
                m_gym = re.match(r"^ginasio de (.+)$", tn)
                m_loc = re.match(r"^locais em (.+)$", tn)
                if m_gym and chifre_city_mode and _norm(m_gym.group(1)) in {_norm(chifre_city_mode), _norm(chifre_city_mode).replace("vila de ", "")}:
                    cur_city = chifre_city_mode
                    cur_h3 = f"Ginásio de {chifre_city_mode}"
                    inside_sublocais = False
                    continue
                if m_loc and chifre_city_mode and _norm(m_loc.group(1)) in {_norm(chifre_city_mode), _norm(chifre_city_mode).replace("vila de ", "")}:
                    cur_city = chifre_city_mode
                    cur_h3 = "Sublocais e pontos de interesse"
                    inside_sublocais = True
                    continue

                if _is_fake_city_h2(t_clean):
                    cur_city = None
                    cur_h3 = t_clean
                    inside_sublocais = False
                    chifre_city_mode = None
                    data["regions"][cur_reg]["sections"].setdefault(cur_h3, "")
                    continue

                cur_city = t_clean
                if _norm(cur_city).startswith("vila de "):
                    cur_city = cur_city.split(" ", 2)[-1].strip()

                ensure_city(cur_city, cur_reg)
                cur_h3 = None
                inside_sublocais = False
                chifre_city_mode = cur_city
                continue

            if lvl == 3:
                if cur_city:
                    cur_h3 = t_clean
                    inside_sublocais = ("sublocais" in _norm(cur_h3) or "pontos de interesse" in _norm(cur_h3))
                    data["cities"][cur_city]["sections"].setdefault(cur_h3, "")
                else:
                    cur_h3 = t_clean
                    inside_sublocais = False
                    data["regions"][cur_reg]["sections"].setdefault(cur_h3, "")
                continue

            if lvl == 4:
                if cur_city and inside_sublocais:
                    nm, desc = _split_name_desc(t)
                    data["cities"][cur_city]["sublocais"].append({"name": nm, "text": desc})
                else:
                    if cur_city and cur_h3:
                        k = f"{cur_h3} • {t_clean}"
                        data["cities"][cur_city]["sections"].setdefault(k, "")
                    elif (not cur_city) and cur_h3:
                        k = f"{cur_h3} • {t_clean}"
                        data["regions"][cur_reg]["sections"].setdefault(k, "")
                continue

            # texto normal
            if cur_city:
                ensure_city(cur_city, cur_reg)

                if inside_sublocais:
                    nm, desc = _split_name_desc(t)
                    if nm and desc:
                        data["cities"][cur_city]["sublocais"].append({"name": nm, "text": desc})
                    else:
                        if data["cities"][cur_city]["sublocais"]:
                            last = data["cities"][cur_city]["sublocais"][-1]
                            last["text"] = (last["text"] + "\n\n" + t).strip()
                        else:
                            data["cities"][cur_city]["sections"].setdefault("Notas", "")
                            data["cities"][cur_city]["sections"]["Notas"] = (data["cities"][cur_city]["sections"]["Notas"] + "\n\n" + t).strip()
                else:
                    if cur_h3:
                        data["cities"][cur_city]["sections"].setdefault(cur_h3, "")
                        data["cities"][cur_city]["sections"][cur_h3] = (data["cities"][cur_city]["sections"][cur_h3] + "\n\n" + t).strip()
                    else:
                        data["cities"][cur_city]["sections"].setdefault("Visão geral", "")
                        data["cities"][cur_city]["sections"]["Visão geral"] = (data["cities"][cur_city]["sections"]["Visão geral"] + "\n\n" + t).strip()
            else:
                ensure_region(cur_reg)
                if cur_h3:
                    data["regions"][cur_reg]["sections"].setdefault(cur_h3, "")
                    data["regions"][cur_reg]["sections"][cur_h3] = (data["regions"][cur_reg]["sections"][cur_h3] + "\n\n" + t).strip()
                else:
                    data["regions"][cur_reg]["intro"] = (data["regions"][cur_reg]["intro"] + "\n\n" + t).strip()

        for r, obj in data["regions"].items():
            if not obj["intro"]:
                for k in list(obj["sections"].keys()):
                    if "visão geral" in _norm(k) and obj["sections"].get(k):
                        obj["intro"] = obj["sections"][k]
                        break

    # ---------- NPCs ----------
    def parse_npcs(doc_path: str, fallback_status: str | None = None) -> dict[str, dict]:
        if not os.path.exists(doc_path):
            return {}
        d = Document(doc_path)
        npcs: dict[str, dict] = {}
        cur: str | None = None
        cur_section: str | None = None

        ID_KEYS = {
            "status": "status",
            "idade": "idade",
            "local de origem": "origem",
            "ocupação atual": "ocupacao",
            "ocupacao atual": "ocupacao",
            "pokémons conhecidos": "pokemons",
            "pokemons conhecidos": "pokemons",
        }

        def ensure(nm: str):
            npcs.setdefault(nm, {"name": nm, "status": fallback_status or "", "idade": "", "origem": "", "ocupacao": "", "pokemons": [], "sections": {}})

        for p in d.paragraphs:
            t = (p.text or "").strip()
            if not t:
                continue
            lvl = _heading_level(p)

            if lvl == 2:
                cur = _clean_title(t)
                ensure(cur)
                cur_section = None
                continue

            if cur is None:
                continue

            if re.fullmatch(r"[—\-]{5,}", t):
                continue

            if t.endswith(":") and len(t) <= 60 and ":" not in t[:-1]:
                cur_section = t[:-1].strip()
                npcs[cur]["sections"].setdefault(cur_section, "")
                continue

            if ":" in t:
                k, v = t.split(":", 1)
                k2 = _norm(k)
                v = v.strip()
                if k2 in ID_KEYS:
                    field = ID_KEYS[k2]
                    if field == "pokemons":
                        pokes = [x.strip() for x in re.split(r",|;|\|", v) if x.strip()]
                        npcs[cur]["pokemons"] = pokes
                    else:
                        npcs[cur][field] = v
                    if not npcs[cur].get("status") and fallback_status:
                        npcs[cur]["status"] = fallback_status
                    continue

            sec = cur_section or "Lore"
            npcs[cur]["sections"].setdefault(sec, "")
            npcs[cur]["sections"][sec] = (npcs[cur]["sections"][sec] + "\n\n" + t).strip()

        for nm, obj in npcs.items():
            if not obj.get("status"):
                obj["status"] = fallback_status or "Vivo(a)"
        return npcs

    vivos = parse_npcs(npcs_vivos_path, fallback_status="Vivo(a)")
    mortos = parse_npcs(npcs_mortos_path, fallback_status="Morto(a)") if os.path.exists(npcs_mortos_path) else {}
    data["npcs"] = {**vivos, **mortos}

    
    return data


# ----------------------------
# JSON loaders (NOVO)
# ----------------------------
def _json_read(path: str) -> dict:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f) or {}
    except Exception:
        return {}


@st.cache_data(show_spinner=False)
def load_compendium_json_data(
    locais_path: str,
    locais_mtime: float,
    npcs_vivos_path: str,
    npcs_vivos_mtime: float,
    npcs_mortos_path: str,
    npcs_mortos_mtime: float,
) -> dict:
    """
    Carrega Compendium em JSON mantendo o MESMO contrato do loader legado (DOCX):
    {
      "regions": {...},
      "cities": {...},
      "npcs": {...},
    }
    """
    data = {"regions": {}, "cities": {}, "npcs": {}}

    # LOCAIS (obrigatório para a aba de Locais)
    if locais_path and os.path.exists(locais_path):
        j = _json_read(locais_path)
        data["regions"] = (j.get("regions") or {})
        data["cities"] = (j.get("cities") or {})



    # NPCs (vivos + mortos)
    npcs_all: dict = {}
    if npcs_vivos_path and os.path.exists(npcs_vivos_path):
        jv = _json_read(npcs_vivos_path)
        npcs_all.update(jv.get("npcs") or {})
    if npcs_mortos_path and os.path.exists(npcs_mortos_path):
        jm = _json_read(npcs_mortos_path)
        # se tiver colisão de nome, mortos não sobrescrevem vivos
        for k, v in (jm.get("npcs") or {}).items():
            npcs_all.setdefault(k, v)
    data["npcs"] = npcs_all
    return data

def _extract_gym_meta_from_text(t: str) -> dict:
    """Extrai metadados do bloco 'Ginásio de X' (locais.json).

    Suporta formatos como:
      - 'Líder: X | Tipo: Y | Localização: Z'
      - 'Tipo: Y | Líder: X | Vice-líder: W | Localização: Z'
      - 'Liderança: A (registro antigo) • B (controle atual)'
    """
    out: dict = {}
    if not t:
        return out

    def _clean(v: str) -> str:
        v = (v or "").strip()
        # remove bullets e espaços extras nas pontas
        v = re.sub(r"^[•\-–—\s]+", "", v).strip()
        v = re.sub(r"[•\-–—\s]+$", "", v).strip()
        return v

    def _as_list(x) -> list[str]:
        if not x:
            return []
        if isinstance(x, list):
            return [str(i).strip() for i in x if str(i).strip()]
        if isinstance(x, str):
            return [s.strip() for s in re.split(r"[;,•/]+", x) if s.strip()]
        return []

    # tenta interpretar a PRIMEIRA linha (normalmente contém os pipes)
    head = ""
    for ln in str(t).splitlines():
        if ln.strip():
            head = ln.strip()
            break

    if head and "|" in head:
        parts = [p.strip() for p in head.split("|") if p.strip()]
        for part in parts:
            # Tipo
            m = re.match(r"(?i)^tipo\s*:\s*(.+)$", part)
            if m:
                out["tipo"] = _clean(m.group(1))
                continue

            # Vice-líder
            m = re.match(r"(?i)^vice\s*[-\s]*l[íi]der\s*:\s*(.+)$", part)
            if m:
                out["vice_lider"] = _clean(m.group(1))
                continue

            # Líder / Liderança
            m = re.match(r"(?i)^(l[íi]der|lideran[cç]a)\s*:\s*(.+)$", part)
            if m:
                k = _norm(m.group(1))
                val = _clean(m.group(2))
                if "lideranca" in k:
                    out["lideranca_raw"] = val
                    # tenta escolher o atual: '(controle atual)' ou último item
                    pieces = [x.strip() for x in re.split(r"[•;/]+", val) if x.strip()]
                    cur = next((x for x in pieces if "controle atual" in _norm(x)), (pieces[-1] if pieces else val))
                    cur_name = re.sub(r"\s*\(.*?\)\s*", "", cur).strip()
                    if cur_name:
                        out["lider"] = cur_name
                    # ex-líderes = demais itens
                    ex = []
                    for x in pieces:
                        if x == cur:
                            continue
                        nm = re.sub(r"\s*\(.*?\)\s*", "", x).strip()
                        if nm:
                            ex.append(nm)
                    if ex:
                        out["ex_lideres"] = ex
                else:
                    out["lider"] = val
                continue

            # Localização
            m = re.match(r"(?i)^localiza[cç][aã]o\s*:\s*(.+)$", part)
            if m:
                out["localizacao"] = _clean(m.group(1))
                continue

            # Status
            m = re.match(r"(?i)^status\s*:\s*(.+)$", part)
            if m:
                out["status"] = _clean(m.group(1))
                continue

            # Ex-líder(es) explícito
            m = re.match(r"(?i)^ex\s*[-\s]*l[íi]der(?:es)?\s*:\s*(.+)$", part)
            if m:
                out["ex_lideres"] = _as_list(_clean(m.group(1)))
                continue

    # fallback: procura em qualquer parte (captura até '|' ou quebra de linha)
    def _grab(key: str, pat: str):
        if key in out and out.get(key):
            return
        mm = re.search(pat, str(t), flags=re.IGNORECASE)
        if mm:
            out[key] = _clean(mm.group(1))

    _grab("tipo", r"tipo\s*:\s*([^|\n]+)")
    _grab("lider", r"l[íi]der\s*:\s*([^|\n]+)")
    _grab("vice_lider", r"vice\s*[-\s]*l[íi]der\s*:\s*([^|\n]+)")
    _grab("localizacao", r"localiza[cç][aã]o\s*:\s*([^|\n]+)")
    _grab("status", r"status\s*:\s*([^|\n]+)")
    # ex-líder(es)
    if "ex_lideres" not in out:
        mm = re.search(r"ex\s*[-\s]*l[íi]der(?:es)?\s*:\s*([^|\n]+)", str(t), flags=re.IGNORECASE)
        if mm:
            out["ex_lideres"] = _as_list(_clean(mm.group(1)))

    # saneia campos (evita pegar o resto do texto após pipes)
    for k in ("tipo", "lider", "vice_lider", "localizacao", "status"):
        if k in out and isinstance(out[k], str) and "|" in out[k]:
            out[k] = _clean(out[k].split("|")[0])

    return out

    # Alternativo: linhas soltas
    m1 = re.search(r"L[íi]der\s*:\s*(.+)", t, flags=re.IGNORECASE)
    if m1:
        out["lider"] = (m1.group(1) or "").strip()
    m2 = re.search(r"Tipo\s*:\s*(.+)", t, flags=re.IGNORECASE)
    if m2:
        out["tipo"] = (m2.group(1) or "").strip()
    m3 = re.search(r"Localiza[cç][aã]o\s*:\s*(.+)", t, flags=re.IGNORECASE)
    if m3:
        out["localizacao"] = (m3.group(1) or "").strip()
    return out


@st.cache_data(show_spinner=False)
def load_compendium_gym_data_json(
    locais_path: str,
    locais_mtime: float,
    ginasios_path: str,
    ginasios_mtime: float,
) -> dict:
    """
    Constrói bundle do ginásio no mesmo formato do legado:
      {"gyms": {city: {"meta":{}, "staff":{}, "staff_npcs":{}, "narrative": ""}},
       "npcs_extra": {name: npc_obj}}
    - 'meta' vem do texto 'Ginásio de <Cidade>' no JSON de locais.
    - 'staff_npcs' e 'npcs_extra' vêm do JSON de ginasios (que é uma coleção de NPCs).
    """
    gyms: dict[str, dict] = {}
    npcs_extra: dict[str, dict] = {}

    # 1) Meta + narrativa via locais.json
    if locais_path and os.path.exists(locais_path):
        jl = _json_read(locais_path)
        cities = (jl.get("cities") or {})
        for city, cobj in cities.items():
            sec = (cobj.get("sections") or {})
            # tenta chave exata "Ginásio de {city}" (normal)
            key = None
            for k in sec.keys():
                if _norm(k).startswith(_norm("Ginásio de " + str(city))):
                    key = k
                    break
            txt_gym = (sec.get(key) if key else "") or ""
            meta = _extract_gym_meta_from_text(txt_gym)
            narrative = txt_gym.strip()
            if meta or narrative:
                gyms[city] = {"meta": meta, "staff": {}, "staff_npcs": {}, "narrative": narrative}

    # 2) Staff via ginasios.json (NPCs)

    if ginasios_path and os.path.exists(ginasios_path):
        jg = _json_read(ginasios_path) or {}
    
        npcs_src = jg.get("npcs") or {}
    
        # aceita tanto {"npcs": { "Nome": {...} }} quanto {"npcs": [ {...}, {...} ]}
        it = []
        if isinstance(npcs_src, list):
            for i, npc in enumerate(npcs_src):
                if not isinstance(npc, dict):
                    continue
                nm = (npc.get("name") or npc.get("nome") or "").strip()
                if not nm:
                    nm = f"npc_{i}"
                it.append((nm, npc))
        elif isinstance(npcs_src, dict):
            for nm, npc in npcs_src.items():
                if not isinstance(npc, dict):
                    continue
                nm2 = str(nm).strip()
                if not nm2:
                    continue
                it.append((nm2, npc))
    
        # candidatos de cidade (reusa cities do bloco 1 se já existir; sem reler JSON)
        cand = set((gyms or {}).keys())
        try:
            cand.update((cities or {}).keys())  # 'cities' já carregado no bloco 1
        except Exception:
            pass
        cand_list = list(cand)
    
        # pré-normaliza para match rápido e consistente
        cand_norm = [(c, _norm(c)) for c in cand_list]
        cand_norm = [(c, cn) for (c, cn) in cand_norm if cn]
    
        for nm, npc in it:
            if not isinstance(npc, dict):
                continue
    
            npcs_extra[nm] = npc
    
            # tenta vincular a um ginásio/cidade pelo texto de ocupação
            occ_raw = npc.get("ocupacao") or ""
            occ = _norm(occ_raw)
    
            # papel (líder / vice / ex)
            role = "staff"
            if ("ex" in occ and "lider" in occ):
                role = "ex_lider"
            elif "vice" in occ:
                role = "vice"
            elif "lider" in occ:
                role = "lider"
    
            # tenta achar a cidade pelo padrão "ginasio_de_<cidade_norm>" dentro do texto normalizado
            city_hit = None
            for c, c_norm in cand_norm:
                if (
                    f"ginasio de {c_norm}" in occ
                    or f"ginasio do {c_norm}" in occ
                    or f"ginasio da {c_norm}" in occ
                    or f"ginasio_de_{c_norm}" in occ
                    or f"ginasio_{c_norm}" in occ
                ):
                    city_hit = c
                    break
    
            # heurística: ex-líder sem cidade explícita -> tenta inferir pelas seções (histórico/ginásio)
            if not city_hit and role == "ex_lider":
                secs = npc.get("sections") or {}
                blob_parts = []
    
                for kk in ("Ginásio", "Ginasio", "Histórico", "Historico", "História", "Historia"):
                    v = secs.get(kk)
                    if v:
                        blob_parts.append(str(v))
    
                if not blob_parts and isinstance(secs, dict):
                    blob_parts = [str(v) for v in secs.values() if v]
    
                blob_n = _norm(" ".join(blob_parts))
                for c, c_norm in cand_norm:
                    if c_norm and c_norm in blob_n:
                        city_hit = c
                        break
    
            if not city_hit:
                continue
    
            g = gyms.setdefault(city_hit, {"meta": {}, "staff": {}, "staff_npcs": {}, "narrative": ""})
            staff = g.setdefault("staff", {})
    
            if role == "ex_lider":
                staff.setdefault("ex_lideres", [])
                if nm not in staff["ex_lideres"]:
                    staff["ex_lideres"].append(nm)
            elif role == "vice":
                staff["vice"] = nm
            elif role == "lider":
                staff["lider"] = nm
    
            g.setdefault("staff_npcs", {})
            g["staff_npcs"][nm] = npc
    
    return {"gyms": gyms, "npcs_extra": npcs_extra}
    
    




# ----------------------------
# GINÁSIOS (integração opcional via ginasios.docx + fichas no DOCX de locais)
# ----------------------------
def _parse_ficha_block(lines: list[str]) -> dict:
    out: dict[str, str] = {}
    for ln in lines:
        ln = (ln or "").strip()
        if not ln or ":" not in ln:
            continue
        k, v = ln.split(":", 1)
        k = _norm(k)
        v = v.strip()
        if not k:
            continue
        out[k] = v
    return out


def _canon_city_key(s: str) -> str:
    # mantém acentos no display, mas cria chave normalizada para match
    return _norm(s or "")


@st.cache_data(show_spinner=False)
def load_compendium_gym_data(
    locais_path: str,
    locais_mtime: float,
    ginasios_path: str,
    ginasios_mtime: float,
) -> dict:
    """Carrega dados de ginásios e (se necessário) cria/enriquece NPCs a partir do arquivo ginasios.docx.
    - Meta estruturada (tipo/status/localização/arena) vem do DOCX de locais via blocos [FICHA]...[/FICHA].
    - Detalhes de líder/vice e Pokémons podem vir de ginasios.docx.
    Retorna: {"gyms": {city: {...}}, "npcs_extra": {name: npc_obj}}
    """
    if Document is None:
        return {"gyms": {}, "npcs_extra": {}}

    gyms: dict[str, dict] = {}
    npcs_extra: dict[str, dict] = {}

    # -------- 1) Meta via [FICHA] no DOCX de locais --------
    if locais_path and os.path.exists(locais_path):
        try:
            d = Document(locais_path)
            cur_city: str | None = None
            in_ficha = False
            ficha_lines: list[str] = []
            narrative_lines: list[str] = []

            def flush():
                nonlocal cur_city, ficha_lines, narrative_lines
                if not cur_city:
                    return
                meta = _parse_ficha_block(ficha_lines)
                # normaliza campos esperados
                city = meta.get("cidade") or cur_city
                city_disp = (city or cur_city).strip()
                g = gyms.setdefault(city_disp, {"meta": {}, "staff": {}, "staff_npcs": {}, "narrative": ""})
                g["meta"] = {**g.get("meta", {}), **meta}
                if narrative_lines:
                    txt = "\n\n".join([x.strip() for x in narrative_lines if x.strip()]).strip()
                    if txt:
                        g["narrative"] = (g.get("narrative","") + "\n\n" + txt).strip()
                ficha_lines = []
                narrative_lines = []

            for p in d.paragraphs:
                t = (p.text or "").strip()
                if not t:
                    continue

                # detecta nome do ginásio "Ginásio de X" (muito comum no seu doc unificado)
                m = re.match(r"^Gin[aá]sio\s+de\s+(.+)$", t.strip(), flags=re.IGNORECASE)
                if m and not in_ficha:
                    # muda de ginásio
                    flush()
                    cur_city = m.group(1).strip()
                    continue

                if t.strip() == "[FICHA]":
                    in_ficha = True
                    ficha_lines = []
                    continue
                if t.strip() == "[/FICHA]":
                    in_ficha = False
                    # não flush aqui; vem narrativa depois
                    continue

                if in_ficha:
                    ficha_lines.append(t)
                else:
                    # narrativa do ginásio (logo após a ficha, até o próximo "Ginásio de ...")
                    if cur_city:
                        narrative_lines.append(t)

            flush()
        except Exception:
            pass

    # -------- 2) Staff via ginasios.docx (opcional) --------
    # Formato do ginasios.docx: blocos por cidade, com nomes e linhas "Campo: valor" + seção "Pokémons"
    if ginasios_path and os.path.exists(ginasios_path):
        try:
            d2 = Document(ginasios_path)
            cur_city: str | None = None
            cur_person: str | None = None
            cur_person_lines: list[str] = []
            in_pokes = False

            def ensure_person(nm: str) -> dict:
                npcs_extra.setdefault(nm, {"name": nm, "status": "Vivo(a)", "idade": "", "origem": "", "ocupacao": "", "pokemons": [], "sections": {}})
                return npcs_extra[nm]

            def commit_person():
                nonlocal cur_person, cur_person_lines, cur_city, in_pokes
                if not (cur_city and cur_person):
                    cur_person = None
                    cur_person_lines = []
                    in_pokes = False
                    return

                npc = ensure_person(cur_person)
                # parse linhas "Campo: valor" e narrativa
                section = "Ginásio (importado)"
                npc["sections"].setdefault(section, "")

                # tenta inferir papel
                role = None
                blob_role = _norm(" ".join(cur_person_lines))
                if "vice" in blob_role:
                    role = "vice_lider"
                if "lider" in blob_role or "líder" in blob_role:
                    # se tiver vice e líder, o primeiro costuma ser líder
                    role = role or "lider"

                # campos
                pokes: list[str] = []
                narrative: list[str] = []
                for ln in cur_person_lines:
                    ln = (ln or "").strip()
                    if not ln:
                        continue
                    if ":" in ln:
                        k, v = ln.split(":", 1)
                        k2 = _norm(k)
                        v = v.strip()
                        if k2 == "idade":
                            npc["idade"] = npc.get("idade") or v
                            continue
                        if k2 in {"local de nascimento", "local de origem", "origem"}:
                            npc["origem"] = npc.get("origem") or v
                            continue
                        if k2 in {"ocupacao atual", "ocupação atual", "ocupacao", "ocupação"}:
                            npc["ocupacao"] = npc.get("ocupacao") or v
                            continue
                        if k2 in {"pokemons", "pokémons", "pokemons conhecidos", "pokémons conhecidos"}:
                            pokes = [x.strip() for x in re.split(r",|;|\|", v) if x.strip()]
                            continue
                        # outros campos viram narrativa curta
                        narrative.append(ln)
                    else:
                        narrative.append(ln)

                if pokes:
                    # não sobrescreve se já existir
                    if not npc.get("pokemons"):
                        npc["pokemons"] = pokes

                if narrative:
                    txt = "\n\n".join(narrative).strip()
                    if txt:
                        npc["sections"][section] = (npc["sections"][section] + "\n\n" + txt).strip()

                # pluga no ginásio da cidade
                g = gyms.setdefault(cur_city, {"meta": {}, "staff": {}, "staff_npcs": {}, "narrative": ""})
                g["staff_npcs"][cur_person] = npc

                # se meta já indicar líder/vice, respeita. senão usa heurística.
                meta = g.get("meta") or {}
                lider = (meta.get("lider") or "").strip()
                vice = (meta.get("vice_lider") or meta.get("vice-lider") or "").strip()

                if role == "vice_lider" and not vice:
                    g["staff"]["vice_lider"] = cur_person
                elif role == "lider" and not lider:
                    g["staff"]["lider"] = cur_person
                else:
                    # fallback: se ainda não tem líder, assume líder; senão vice
                    if not (g["staff"].get("lider") or lider):
                        g["staff"]["lider"] = cur_person
                    elif not (g["staff"].get("vice_lider") or vice):
                        g["staff"]["vice_lider"] = cur_person

                # reset
                cur_person = None
                cur_person_lines = []
                in_pokes = False

            # cidades candidatas: nomes que aparecem também no compendium (heurística: palavra só / curto)
            for p in d2.paragraphs:
                t = (p.text or "").strip()
                if not t:
                    continue

                # detecta cidade: linha curta sem ":" e sem muitos espaços
                if ":" not in t and len(t) <= 40 and len(t.split()) <= 3 and t.lower() not in {"pokémons", "pokemons"}:
                    # se já estávamos em uma pessoa, fecha
                    if cur_person:
                        commit_person()
                    cur_city = t.strip()
                    gyms.setdefault(cur_city, {"meta": {}, "staff": {}, "staff_npcs": {}, "narrative": ""})
                    continue

                # detecta nome de pessoa: linha sem ":" e com 1-4 palavras (evita títulos)
                if ":" not in t and len(t) <= 60 and 1 <= len(t.split()) <= 5 and t.lower() not in {"pokémons", "pokemons"}:
                    # evita capturar descrições como "Líder:" na mesma linha
                    if t.lower().startswith("líder") or t.lower().startswith("lider") or t.lower().startswith("vice"):
                        # trata como narrativa
                        if cur_person:
                            cur_person_lines.append(t)
                        else:
                            if cur_city:
                                gyms[cur_city]["narrative"] = (gyms[cur_city].get("narrative","") + "\n\n" + t).strip()
                        continue
                    # troca pessoa
                    if cur_person:
                        commit_person()
                    cur_person = t.strip()
                    cur_person_lines = []
                    continue

                # conteúdo da pessoa/cidade
                if cur_person:
                    cur_person_lines.append(t)
                elif cur_city:
                    gyms[cur_city]["narrative"] = (gyms[cur_city].get("narrative","") + "\n\n" + t).strip()

            if cur_person:
                commit_person()

        except Exception:
            pass

    # -------- 3) Normalização final / merge de leader/vice --------
    for city, g in list(gyms.items()):
        meta = g.get("meta") or {}
        # se meta define nomes, usa para staff
        lider = (meta.get("lider") or "").strip()
        vice = (meta.get("vice_lider") or meta.get("vice-lider") or "").strip()
        if lider and "lider" not in (g.get("staff") or {}):
            g.setdefault("staff", {})["lider"] = lider
        if vice and "vice_lider" not in (g.get("staff") or {}):
            g.setdefault("staff", {})["vice_lider"] = vice

        # garante objetos para leader/vice no staff_npcs se existirem em npcs_extra
        for role_key in ["lider", "vice_lider"]:
            nm = (g.get("staff") or {}).get(role_key)
            if nm and nm in npcs_extra and nm not in (g.get("staff_npcs") or {}):
                g.setdefault("staff_npcs", {})[nm] = npcs_extra[nm]

    return {"gyms": gyms, "npcs_extra": npcs_extra}




def comp_load() -> dict:
    # ----------------------------
    # Preferência: JSON
    # ----------------------------
    loc_j = _resolve_asset_path(COMP_JSON_LOCAIS)
    npc_v_j = _resolve_asset_path(COMP_JSON_NPCS_VIVOS)
    npc_m_j = _resolve_asset_path(COMP_JSON_NPCS_MORTOS)
    gin_j = _resolve_asset_path(COMP_JSON_GINASIOS)

    use_json = bool(loc_j and os.path.exists(loc_j))

    if use_json:
        data = load_compendium_json_data(
            loc_j, _comp_mtime(loc_j),
            npc_v_j, _comp_mtime(npc_v_j),
            npc_m_j, _comp_mtime(npc_m_j),
        )
    
        bundle = load_compendium_gym_data_json(
            loc_j, _comp_mtime(loc_j),
            gin_j, _comp_mtime(gin_j),
        ) if (loc_j or gin_j) else {"gyms": {}, "npcs_extra": {}}

        # =========================
        # OVERRIDES FIXOS (STAFF)
        # =========================
        vice_overrides = {
            "Dainise": "Grassa",
            "Obsidian": "Clay",
        }

        gyms_bundle = (bundle.get("gyms") or {})
        for city, vice_name in vice_overrides.items():
            g = gyms_bundle.get(city)
            if not isinstance(g, dict):
                continue

            g.setdefault("staff", {})

            # só força se estiver vazio (pra não sobrescrever JSON/meta quando você corrigir depois)
            meta = (g.get("meta") or {})
            if not (g["staff"].get("vice_lider") or meta.get("vice_lider") or meta.get("vice-lider")):
                g["staff"]["vice_lider"] = vice_name
    
        # =========================
        # NPCs gerais (Vivos + Mortos + Extras)
        # =========================
        import re as _re

        def _keycanon(x: str) -> str:
            # normaliza e remove pontuação/aspas para dedupe
            x = _norm(x)
            x = x.replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
            x = _re.sub(r"[^a-z0-9 ]+", "", x)
            x = _re.sub(r"\s+", " ", x).strip()
            return x
        
        npcs_gerais = {}
        npcs_gerais.update(data.get("npcs", {}) or {})
        
        extras = (bundle.get("npcs_extra", {}) or {})
        for nm, obj in extras.items():
            k = _keycanon(nm)
            hit = None
            for existing in npcs_gerais.keys():
                if _keycanon(existing) == k:
                    hit = existing
                    break
            if hit:
                npcs_gerais[hit] = obj
            else:
                npcs_gerais[nm] = obj
        

    else:
        # ----------------------------
        # Fallback legado: DOCX
        # ----------------------------
        loc_p = _resolve_asset_path(COMP_DOC_LOCAIS)
        npc_v_p = _resolve_asset_path(COMP_DOC_NPCS_VIVOS)
        npc_m_p = _resolve_asset_path(COMP_DOC_NPCS_MORTOS)
        gin_p = _resolve_asset_path(COMP_DOC_GINASIOS)

        data = load_compendium_data(
            loc_p, _comp_mtime(loc_p),
            npc_v_p, _comp_mtime(npc_v_p),
            npc_m_p, _comp_mtime(npc_m_p),
        )

        # integra ginásios (meta + staff), sem quebrar se o arquivo não existir
        bundle = load_compendium_gym_data(
            loc_p, _comp_mtime(loc_p),
            gin_p, _comp_mtime(gin_p),
        ) if (loc_p or gin_p) else {"gyms": {}, "npcs_extra": {}}

    # ----------------------------
    # Merge: gyms + NPCs extras
    # ----------------------------
    gyms = (bundle.get("gyms") or {})
    if gyms:
        data["gyms"] = gyms
    else:
        data.setdefault("gyms", {})

    # adiciona NPCs extras (líder/vice que não existam no JSON/DOCX de NPCs)
    extra = bundle.get("npcs_extra") or {}
    for nm, obj in extra.items():
        if nm not in data.get("npcs", {}):
            data.setdefault("npcs", {})
            data["npcs"][nm] = obj
        else:
            # preenche só o que está vazio
            dst = data["npcs"][nm]
            for k in ["idade", "origem", "ocupacao"]:
                if not (dst.get(k) or "").strip() and (obj.get(k) or "").strip():
                    dst[k] = obj.get(k)
            if (not dst.get("pokemons")) and obj.get("pokemons"):
                dst["pokemons"] = obj.get("pokemons")
            # injeta seção importada sem sobrescrever
            sec = "Ginásio (importado)"
            if obj.get("sections", {}).get(sec):
                dst.setdefault("sections", {})
                if not dst["sections"].get(sec):
                    dst["sections"][sec] = obj["sections"][sec]

    # -------------------------------------------------
    # Sync persistente: NPC Users (Trainer Hub) -> NPCs
    # -------------------------------------------------
    try:
        data.setdefault("npcs", {})
        data["npcs"] = _sync_npc_users_and_overrides(data.get("npcs") or {})
    except Exception:
        pass

    # -------------------------------------------------
    # Overrides por sessão (UI): sempre por último
    # -------------------------------------------------
    if "npc_sync_overrides" in st.session_state:
        data["npcs"] = apply_npc_overrides(
            data.get("npcs") or {},
            st.session_state["npc_sync_overrides"]
        )

    return data

if trainer_name == "Ezenek":
    if st.sidebar.button("🔄 Atualizar NPCs"):
        try:
            st.session_state.pop("comp_data", None)  # se você cacheia em comp_data
            st.sidebar.success("Compendium será recarregado com NPCs do Trainer Hub.")
        except Exception as e:
            st.sidebar.error(f"Falha ao atualizar NPCs: {e}")

# ----------------------------
# SESSION TRACKER (JSON)
# ----------------------------
def _session_now_iso() -> str:
    return datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def _sessions_file_path() -> str:
    """
    Sempre usa ./data/compendium_sessions.json como fonte principal (gravável).
    Se não existir ainda, tenta ler de um arquivo existente em assets/raiz (fallback).
    """
    # 1) caminho gravável principal
    data_dir = os.path.join(os.getcwd(), "data")
    try:
        os.makedirs(data_dir, exist_ok=True)
    except Exception:
        # se não conseguir criar, cai pro cwd mesmo
        data_dir = os.getcwd()

    primary = os.path.join(data_dir, COMP_SESSIONS_JSON)

    # Se já existe no local gravável, usa ele
    if os.path.exists(primary):
        return primary

    # 2) fallback: se existir algum "resolved" em assets/raiz, usa só para leitura inicial
    resolved = _resolve_asset_path(COMP_SESSIONS_JSON)
    if resolved and os.path.exists(resolved):
        return resolved

    # 3) se não existir em lugar nenhum ainda, já aponta pro local gravável
    return primary

def _sessions_file_path() -> str:
    resolved = _resolve_asset_path(COMP_SESSIONS_JSON)
    if resolved and os.path.exists(resolved):
        return resolved
    return os.path.join(os.getcwd(), COMP_SESSIONS_JSON)


def load_sessions_data(db=None, trainer_name: str | None = None) -> dict:
    # 1) tenta Firestore
    try:
        if db is None:
            db = globals().get("db")
        if trainer_name is None:
            trainer_name = st.session_state.get("trainer_name") or st.session_state.get("player_name")
        if db is not None and trainer_name:
            trainer_id = safe_doc_id(str(trainer_name))
            ref = (
                db.collection("trainers")
                  .document(trainer_id)
                  .collection("compendium")
                  .document("sessions")
            )
            snap = ref.get()
            if snap.exists:
                data = snap.to_dict() or {}
                if isinstance(data, dict):
                    data.setdefault("meta", {})
                    data.setdefault("sessions", {})
                    data["meta"].setdefault("schema", 1)
                    data["meta"].setdefault("updated_at", _session_now_iso())
                    if not isinstance(data["sessions"], dict):
                        data["sessions"] = {}
                    return data
    except Exception:
        pass

    # 2) fallback: arquivo local (seu comportamento atual)
    path = _sessions_file_path()
    if not os.path.exists(path):
        return _sessions_default_payload()
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return _sessions_default_payload()

    if not isinstance(data, dict):
        return _sessions_default_payload()
    data.setdefault("meta", {})
    data.setdefault("sessions", {})
    data["meta"].setdefault("schema", 1)
    data["meta"].setdefault("updated_at", _session_now_iso())
    if not isinstance(data["sessions"], dict):
        data["sessions"] = {}
    return data


def save_sessions_data(data: dict, db=None, trainer_name: str | None = None) -> None:
    if not isinstance(data, dict):
        return
    data.setdefault("meta", {})
    data.setdefault("sessions", {})
    data["meta"]["updated_at"] = _session_now_iso()

    # 1) tenta Firestore
    try:
        if db is None:
            db = globals().get("db")
        if trainer_name is None:
            trainer_name = st.session_state.get("trainer_name") or st.session_state.get("player_name")
        if db is not None and trainer_name:
            trainer_id = safe_doc_id(str(trainer_name))
            ref = (
                db.collection("trainers")
                  .document(trainer_id)
                  .collection("compendium")
                  .document("sessions")
            )
            ref.set(data, merge=True)
            # ✅ avisa a sala toda que “algo mudou”
            db.collection("rooms").document(rid).collection("public_state").document("state").set({
                "updatedAt": firestore.SERVER_TIMESTAMP
            }, merge=True)
            return  # sucesso no Firestore -> não precisa arquivo
    except Exception:
        pass

    # 2) fallback: arquivo local (seu comportamento atual)
    path = _sessions_file_path()
    folder = os.path.dirname(path) or os.getcwd()
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M")
    backup_path = os.path.join(folder, f"compendium_sessions_backup_{ts}.json")

    payload = json.dumps(data, ensure_ascii=False, indent=2)
    try:
        with open(backup_path, "w", encoding="utf-8") as backup:
            backup.write(payload)
    except Exception:
        pass

    with open(path, "w", encoding="utf-8") as f:
        f.write(payload)



def _session_id_from_number(number: int) -> str:
    return f"S{int(number):04d}"


def _session_label(session_id: str, session: dict) -> str:
    num = session.get("number")
    title = (session.get("title") or "").strip() or "Sem título"
    date = (session.get("date") or "").strip()
    base = f"{session_id}"
    if isinstance(num, int) or (isinstance(num, str) and str(num).isdigit()):
        base = f"S{int(num):04d}"
    suffix = f" — {title}"
    if date:
        suffix += f" ({date})"
    return base + suffix


def _ensure_session_links(session: dict) -> dict:
    session.setdefault("links", {})
    session["links"].setdefault("npcs", [])
    session["links"].setdefault("places", [])
    session["links"].setdefault("gyms", [])
    session["links"].setdefault("items", [])
    return session


def _add_entity_to_session(session: dict, link_key: str, entity_id: str) -> None:
    _ensure_session_links(session)
    if entity_id and entity_id not in session["links"].get(link_key, []):
        session["links"][link_key].append(entity_id)


def _append_session_event(session: dict, event: dict) -> None:
    session.setdefault("events", [])
    if isinstance(event, dict):
        session["events"].append(event)


def add_entity_to_active_session(link_key: str, entity_id: str, event_type: str, text: str, refs: dict) -> bool:
    active_sid = st.session_state.get("comp_session_active_id")
    if not active_sid or not entity_id:
        return False
    db = _sessions_db(globals().get("db"))
    trainer_name = _sessions_owner_name()
    sessions_data = load_sessions_data_cloud_first(db, trainer_name)


    sessions = sessions_data.get("sessions", {}) or {}
    session = sessions.get(active_sid)
    if not isinstance(session, dict):
        return False
    _add_entity_to_session(session, link_key, entity_id)
    _append_session_event(
        session,
        {
            "t": datetime.utcnow().strftime("%H:%M"),
            "type": event_type,
            "text": text,
            "refs": refs,
        },
    )
    save_sessions_data_cloud_first(db, trainer_name, sessions_data)
    return True

def _sessions_firestore_ref(db, trainer_name: str):
    trainer_id = safe_doc_id(trainer_name)
    return (
        db.collection("trainers")
          .document(trainer_id)
          .collection("compendium")
          .document("sessions")
    )


def _sessions_owner_name(explicit_name: str | None = None) -> str:
    """
    Resolve o dono dos dados de sessões no Compendium.
    Prioriza o nome explicitamente informado e, em seguida,
    os nomes já presentes na sessão de login.
    """
    return (
        (explicit_name or "").strip()
        or (st.session_state.get("trainer_name") or "").strip()
        or (st.session_state.get("player_name") or "").strip()
    )


def _sessions_db(db=None):
    """Retorna cliente Firestore ativo, tentando inicializar quando necessário."""
    if db is not None:
        return db
    db = globals().get("db")
    if db is not None:
        return db
    try:
        db, _ = init_firebase()
        return db
    except Exception:
        return None

def load_sessions_data_cloud_first(db, trainer_name: str) -> dict:
    db = _sessions_db(db)
    trainer_name = _sessions_owner_name(trainer_name)
    # 1) tenta Firestore
    try:
        if db is not None and trainer_name:
            snap = _sessions_firestore_ref(db, trainer_name).get()
            if snap.exists:
                data = snap.to_dict() or {}
                if isinstance(data, dict):
                    data.setdefault("meta", {})
                    data.setdefault("sessions", {})
                    data["meta"].setdefault("schema", 1)
                    data["meta"].setdefault("updated_at", _session_now_iso())
                    if not isinstance(data["sessions"], dict):
                        data["sessions"] = {}
                    return data
    except Exception as e:
        st.warning(f"Falha ao carregar sessões do Firestore: {e}")

    # 2) fallback: JSON local (o seu atual)
    return load_sessions_data()

def save_sessions_data_cloud_first(db, trainer_name: str, data: dict) -> None:
    if not isinstance(data, dict):
        return
    db = _sessions_db(db)
    trainer_name = _sessions_owner_name(trainer_name)
    data.setdefault("meta", {})
    data.setdefault("sessions", {})
    data["meta"]["updated_at"] = _session_now_iso()

    # 1) tenta Firestore
    try:
        if db is not None and trainer_name:
            _sessions_firestore_ref(db, trainer_name).set(data, merge=True)
            return
    except Exception as e:
        st.warning(f"Falha ao salvar sessões no Firestore: {e}")

    # 2) fallback: JSON local
    save_sessions_data(data)

# ----------------------------
# INFERÊNCIA DE TAGS + MENCÕES
# ----------------------------
def infer_city_tags(city_obj: dict) -> list[str]:
    blob = _norm(" ".join([city_obj.get("region","")] + list((city_obj.get("sections") or {}).values())))
    tags = []
    def add(tag, cond):
        nonlocal tags
        if cond and tag not in tags:
            tags.append(tag)

    add("Industrial", any(w in blob for w in ["industr", "usina", "refin", "fabrica", "minas", "petroleo", "petróleo"]))
    add("Arruinada", any(w in blob for w in ["ruina", "ruína", "abandon", "decad", "colaps"]))
    add("Portuária", any(w in blob for w in ["porto", "baia", "baía", "mar", "oceano", "cais"]))
    add("Desértica", any(w in blob for w in ["deserto", "areia", "duna", "arido", "árido"]))
    add("Florestal", any(w in blob for w in ["floresta", "bosque", "selva", "ranger", "copas"]))
    add("Montanhosa", any(w in blob for w in ["monte", "montanha", "serra", "pedra", "vulcao", "vulcão"]))
    add("Urbana", any(w in blob for w in ["avenida", "estadio", "estádio", "shopping", "metro", "aeroporto", "centro"]))
    return tags[:4]


def infer_npc_tags(npc_obj: dict) -> list[str]:
    blob = _norm(" ".join([npc_obj.get("ocupacao","")] + list(((npc_obj.get("sections") or {}) or {}).values())))
    tags = []
    def add(tag, cond):
        nonlocal tags
        if cond and tag not in tags:
            tags.append(tag)

    add("Líder de Ginásio", ("lider" in blob or "líder" in blob) and "ginas" in blob)
    add("Cientista", any(w in blob for w in ["cient", "laborat", "pesquis", "dra.", "dr."]))
    add("Ranger", "ranger" in blob)
    add("Treinador", "treinador" in blob)
    add("Mercenário", any(w in blob for w in ["mercen", "caca", "caça", "cacador", "caçador"]))
    add("Político", any(w in blob for w in ["prefeit", "govern", "senador", "movimento civil", "liga"]))
    return tags[:4]


def city_mentions_npcs(city_obj: dict, npc_names: list[str]) -> list[str]:
    text = _norm(" ".join(list((city_obj.get("sections") or {}).values())))
    hits = []
    for nm in npc_names:
        if len(nm) < 3:
            continue
        if _norm(nm) in text:
            hits.append(nm)
    return hits[:16]


def npc_mentions_cities(npc_obj: dict, city_names: list[str]) -> list[str]:
    text = _norm(" ".join(list((npc_obj.get("sections") or {}).values())))
    hits = []
    for cn in city_names:
        if len(cn) < 3:
            continue
        if _norm(cn) in text:
            hits.append(cn)
    return hits[:16]


def npc_mentions_npcs(npc_obj: dict, npc_names: list[str], self_name: str) -> list[str]:
    text = _norm(" ".join(list((npc_obj.get("sections") or {}).values())))
    hits = []
    for nm in npc_names:
        if nm == self_name:
            continue
        if len(nm) < 3:
            continue
        if _norm(nm) in text:
            hits.append(nm)
    return hits[:16]


# ----------------------------
# SPRITES (cache local offline)
# ----------------------------
def _sprite_cache_dir() -> Path:
    p = Path(os.getcwd()) / "assets" / "sprites_cache"
    try:
        p.mkdir(parents=True, exist_ok=True)
    except Exception:
        # fallback
        p = Path(os.getcwd()) / "sprites_cache"
        p.mkdir(parents=True, exist_ok=True)
    return p


def _to_pokeapi_simple(name: str) -> str:
    x = _norm(name)
    x = x.replace("♀", "-f").replace("♂", "-m")
    x = x.replace(".", "").replace("'", "").replace("’", "")
    x = x.replace(" ", "-")
    x = re.sub(r"[^a-z0-9\-]+", "", x)
    return x


def _extract_pokemon_name(raw: str) -> str:
    text = (raw or "").strip()
    if not text:
        return ""
    text = re.split(r"[|/;,]", text)[0]
    text = re.sub(r"\(.*?\)", "", text)
    text = re.sub(r"\[.*?\]", "", text)
    text = re.sub(r"\b(?:lvl|lv|level)\s*\d+\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"[^A-Za-zÀ-ÿ0-9'\-\. ]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _closest_pokeapi_name(raw: str) -> str:
    base = to_pokeapi_name(raw)
    if not api_name_map:
        return base
    if base in api_name_map:
        return base
    base_simple = base.split("-")[0]
    if base_simple in api_name_map:
        return base_simple
    keys = list(api_name_map.keys())
    match = difflib.get_close_matches(base, keys, n=1, cutoff=0.78)
    if not match and base_simple != base:
        match = difflib.get_close_matches(base_simple, keys, n=1, cutoff=0.76)
    return match[0] if match else base


@st.cache_data(ttl=60*60*24, show_spinner=False)
def _poke_sprite_cached(poke_name: str) -> str | None:
    """
    Retorna um caminho local (cache) se possível; fallback para URL.
    """
    raw = _extract_pokemon_name(poke_name)
    if not raw:
        return None

    # resolve query
    try:
        q = _closest_pokeapi_name(raw)
    except Exception:
        q = _to_pokeapi_simple(raw)

    # nome de arquivo robusto
    key = hashlib.sha1(q.encode("utf-8")).hexdigest()[:12]
    out = _sprite_cache_dir() / f"{key}_{q}.png"

    if out.exists():
        return str(out)

    # tenta baixar da PokeAPI
    try:
        url = f"{POKEAPI_BASE}/pokemon/{q}"
        r = requests.get(url, timeout=12)
        if r.status_code != 200:
            # sem internet? retorna None (não quebra)
            return str(out) if out.exists() else None
        pjson = r.json()
        sprites = (pjson.get("sprites") or {})
        sprite_url = sprites.get("front_default") or sprites.get("front_shiny")
        if not sprite_url:
            # tenta gen8 icons
            vers = (sprites.get("versions") or {})
            gen8 = (vers.get("generation-viii") or {})
            icons = (gen8.get("icons") or {})
            sprite_url = icons.get("front_default")

        if not sprite_url:
            return None

        img = requests.get(sprite_url, timeout=12)
        if img.status_code != 200:
            return None

        out.write_bytes(img.content)
        return str(out)

    except Exception:
        return str(out) if out.exists() else None


# ----------------------------
# BUSCA GLOBAL (full-text)
# ----------------------------
def _norm_tokens(q: str) -> list[str]:
    qn = _norm_loc(q)
    toks = [t for t in re.split(r"\s+", qn) if t]
    # remove tokens muito curtos
    toks = [t for t in toks if len(t) >= 2]
    return toks


def _snippet(text: str, toks: list[str], n: int = 220) -> str:
    if not text:
        return ""
    tnorm = _norm(text)
    idx = None
    for tk in toks:
        j = tnorm.find(tk)
        if j != -1:
            idx = j
            break
    if idx is None:
        return (text[:n] + "…") if len(text) > n else text
    # tenta aproximar no texto original (sem normalizar)
    start = max(0, idx - 80)
    end = min(len(text), idx + (n - 80))
    chunk = text[start:end].strip()
    if start > 0:
        chunk = "…" + chunk
    if end < len(text):
        chunk = chunk + "…"
    return chunk


@st.cache_data(show_spinner=False)

def build_comp_search_index(data: dict) -> list[dict]:
    idx: list[dict] = []
    regions = data.get("regions", {}) or {}
    cities = data.get("cities", {}) or {}
    npcs = data.get("npcs", {}) or {}
    gyms = data.get("gyms", {}) or {}

    # regiões
    for r, robj in regions.items():
        blob = " ".join([robj.get("intro","")] + list((robj.get("sections") or {}).values()))
        idx.append({"type":"region", "key": r, "title": r, "text": blob, "region": r})

    # cidades e sublocais
    for c, cobj in cities.items():
        blob = " ".join(list((cobj.get("sections") or {}).values()))
        idx.append({"type":"city", "key": c, "title": c, "text": blob, "region": cobj.get("region","")})
        for sl in (cobj.get("sublocais") or []):
            nm = (sl.get("name") or "").strip()
            tx = (sl.get("text") or "").strip()
            if not nm and not tx:
                continue
            idx.append({"type":"sublocal", "key": f"{c}::{nm}", "title": nm or "(Sublocal)", "text": tx, "city": c, "region": cobj.get("region","")})

    # ginásios
    for city, g in gyms.items():
        meta = g.get("meta") or {}
        staff = g.get("staff") or {}
        # tenta região via cidades conhecidas
        reg = (cities.get(city, {}) or {}).get("region", "")
        title = f"Ginásio — {city}"
        blob = " ".join([
            city,
            str(reg),
            str(meta.get("status","")),
            str(meta.get("tipo","")),
            str(meta.get("localizacao","")),
            str(meta.get("arena_extra","")),
            str(meta.get("observacao","")),
            str(meta.get("lider","")),
            str(meta.get("vice_lider","")),
            str(staff.get("lider","")),
            str(staff.get("vice_lider","")),
            str(g.get("narrative","")),
        ])
        idx.append({"type":"gym", "key": city, "title": title, "text": blob, "city": city, "region": reg})

    # npcs
    for n, nobj in npcs.items():
        blob = " ".join([nobj.get("ocupacao",""), nobj.get("origem","")] + list(((nobj.get("sections") or {}) or {}).values()))
        idx.append({"type":"npc", "key": n, "title": n, "text": blob, "status": nobj.get("status","")})

    return idx

def comp_search(index: list[dict], query: str, types: set[str] | None = None, limit: int = 40) -> list[dict]:
    toks = _norm_tokens(query)
    if not toks:
        return []
    res = []
    for d in index:
        if types and d.get("type") not in types:
            continue
        t = _norm(d.get("text","") + " " + d.get("title",""))
        score = 0
        for tk in toks:
            if tk in t:
                score += 3
                score += t.count(tk)
        if score > 0:
            res.append((score, d))
    res.sort(key=lambda x: x[0], reverse=True)
    out = []
    for score, d in res[:limit]:
        dd = dict(d)
        dd["score"] = score
        dd["snippet"] = _snippet(d.get("text",""), toks)
        out.append(dd)
    return out


# ----------------------------
# Favoritos/Recentes
# ----------------------------
def _init_comp_state():
    st.session_state.setdefault("comp_axis", "🌍 Locais")
    st.session_state.setdefault("comp_fav_cities", [])
    st.session_state.setdefault("comp_fav_npcs", [])
    st.session_state.setdefault("comp_recent_cities", [])
    st.session_state.setdefault("comp_recent_npcs", [])
    st.session_state.setdefault("comp_region", None)
    st.session_state.setdefault("comp_city", None)
    st.session_state.setdefault("comp_npc_selected", None)


def _touch_recent(key: str, value: str, limit: int = 12):
    # key = "comp_recent_cities" ou "comp_recent_npcs"
    arr = st.session_state.get(key, []) or []
    arr = [x for x in arr if x != value]
    arr.insert(0, value)
    st.session_state[key] = arr[:limit]


def _toggle_fav(key: str, value: str):
    # key = "comp_fav_cities" ou "comp_fav_npcs"
    arr = st.session_state.get(key, []) or []
    if value in arr:
        arr = [x for x in arr if x != value]
    else:
        arr = [value] + arr
    st.session_state[key] = arr


def _is_fav(key: str, value: str) -> bool:
    return value in (st.session_state.get(key, []) or [])


# ----------------------------
# UI helpers
# ----------------------------
def _chip_row(tags: list[str]) -> None:
    if not tags:
        return
    chips = "".join([f'<span class="comp-chip">{t}</span>' for t in tags])
    st.markdown(chips, unsafe_allow_html=True)


def _breadcrumb(items: list[str]) -> None:
    st.markdown("**" + " › ".join(items) + "**")


def _nav_prev_next(current: str, ordered: list[str], prefix: str):
    if not current or current not in ordered or len(ordered) <= 1:
        return
    i = ordered.index(current)
    c1, c2, c3 = st.columns([0.55, 0.55, 1.3], gap="small")
    with c1:
        if st.button("⬅ Anterior", key=f"{prefix}_prev_{_stem_key(current)}"):
            st.session_state[prefix] = ordered[(i - 1) % len(ordered)]
            st.rerun()
    with c2:
        if st.button("Próximo ➡", key=f"{prefix}_next_{_stem_key(current)}"):
            st.session_state[prefix] = ordered[(i + 1) % len(ordered)]
            st.rerun()
    with c3:
        st.caption(f"{i+1}/{len(ordered)}")


def _render_region_panel(region_name: str, region_obj: dict) -> None:
    img_region = comp_find_image(region_name)
    img_map = comp_find_image(COMP_DEFAULT_MAP)
    st.markdown('<div class="comp-hero">', unsafe_allow_html=True)
    cA, cB = st.columns([1, 1.7], gap="large")
    with cA:
        if img_map:
            st.image(img_map, use_container_width=True)
            st.caption("🗺️ Mapa geral")
        elif img_region:
            st.image(img_region, use_container_width=True)
        else:
            st.caption("🖼️ (sem imagem encontrada para esta região)")
    with cB:
        st.markdown(f"## 🌍 {region_name}")
        if img_region and img_map != img_region:
            st.image(img_region, use_container_width=True)
        intro = (region_obj.get("intro") or "").strip()
        if intro:
            paras = [p.strip() for p in intro.split("\n\n") if p.strip()]
            st.markdown("\n\n".join(paras[:2]))
            if len(paras) > 2:
                with st.expander("Ler mais"):
                    st.markdown("\n\n".join(paras[2:]))
        else:
            st.caption("(sem texto de identidade detectado)")

        cities = region_obj.get("cities") or []
        if cities:
            st.markdown("**Cidades para começar a leitura:**")
            for cname in cities[:4]:
                if st.button(f"➡ {cname}", key=f"region_city_pick_{_stem_key(region_name)}_{_stem_key(cname)}"):
                    st.session_state["comp_city"] = cname
                    _touch_recent("comp_recent_cities", cname)
                    st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)


def _render_city_dossier(city_name: str, city_obj: dict, npcs: dict[str, dict], cities: dict[str, dict], gyms: dict[str, dict] | None = None, region_obj: dict | None = None) -> None:
    img_city = comp_find_image(city_name)
    img_map = comp_find_image(city_obj.get("region","")) or comp_find_image(COMP_DEFAULT_MAP)

    # header + ações
    st.markdown('<div class="comp-hero">', unsafe_allow_html=True)
    hA, hB, hC = st.columns([1, 1, 1.2], gap="large")
    with hA:
        if img_city:
            st.image(img_city, use_container_width=True)
            st.caption("📸 Paisagem da cidade")
        else:
            st.caption("🖼️ (sem imagem da cidade)")
    with hB:
        if img_map:
            st.image(img_map, use_container_width=True)
            st.caption("🗺️ Mapa da região")
        else:
            st.caption("🗺️ (sem mapa)")
    with hC:
        st.markdown(f"## 🏙️ {city_name}")
        st.markdown(f'<div class="comp-muted">Região: <b>{city_obj.get("region","")}</b></div>', unsafe_allow_html=True)

        fav = _is_fav("comp_fav_cities", city_name)
        if st.button("⭐ Remover favorito" if fav else "⭐ Favoritar", key=f"fav_city_{_stem_key(city_name)}"):
            _toggle_fav("comp_fav_cities", city_name)
            st.rerun()

        npc_names = list(npcs.keys())
        mentions = city_mentions_npcs(city_obj, npc_names)
        if mentions:
            st.markdown("**NPCs citados:**")
            cols = st.columns(3)
            for i, nm in enumerate(mentions[:9]):
                with cols[i % 3]:
                    if st.button(f"🧑 {nm}", key=f"comp_city_npc_{_stem_key(city_name)}_{_stem_key(nm)}"):
                        st.session_state["comp_axis"] = "🧑‍🤝‍🧑 NPCs"
                        st.session_state["comp_npc_selected"] = nm
                        _touch_recent("comp_recent_npcs", nm)
                        st.rerun()

        # cidades relacionadas (mesma região)
        same_region = [c for c, obj in cities.items() if obj.get("region") == city_obj.get("region") and c != city_name]
        if same_region:
            st.markdown("**Continue lendo nesta região:**")
            for cname in same_region[:4]:
                if st.button(f"➡ {cname}", key=f"rel_city_{_stem_key(city_name)}_{_stem_key(cname)}"):
                    st.session_state["comp_city"] = cname
                    _touch_recent("comp_recent_cities", cname)
                    st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)

    # trilha de leitura
    if region_obj:
        st.markdown('<div class="comp-panel">', unsafe_allow_html=True)
        st.markdown("### 📖 Próximo capítulo")
        region_cities = [c for c in (region_obj.get("cities") or []) if c in cities and c != city_name]
        if region_cities:
            next_city = region_cities[0]
            if st.button(f"🚪 Ir para {next_city}", key=f"next_city_{_stem_key(city_name)}"):
                st.session_state["comp_city"] = next_city
                _touch_recent("comp_recent_cities", next_city)
                st.rerun()
        mentions = city_mentions_npcs(city_obj, list(npcs.keys()))
        if mentions:
            st.markdown("**NPCs para conhecer:**")
            cols = st.columns(min(3, len(mentions)))
            for i, nm in enumerate(mentions[:3]):
                with cols[i % len(cols)]:
                    if st.button(f"🧑 {nm}", key=f"city_npc_next_{_stem_key(city_name)}_{_stem_key(nm)}"):
                        st.session_state["comp_axis"] = "🧑‍🤝‍🧑 NPCs"
                        st.session_state["comp_npc_selected"] = nm
                        _touch_recent("comp_recent_npcs", nm)
                        st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    
    # ----------------------------
    # 🏅 GINÁSIO (estruturado)
    # ----------------------------
    gyms = gyms or {}
    g = gyms.get(city_name)
    if g:
        meta = g.get("meta") or {}
        staff = g.get("staff") or {}
        staff_npcs = g.get("staff_npcs") or {}

        st.markdown("## 🏅 Ginásio")
        chips = []
        if meta.get("tipo"): chips.append(f"Tipo: {meta.get('tipo')}")
        if meta.get("status"): chips.append(f"Status: {meta.get('status')}")
        if chips:
            _chip_row(chips)

        # bloco meta
        st.markdown('<div class="comp-block">', unsafe_allow_html=True)
        if meta.get("localizacao"):
            st.markdown(f"**Localização:** {meta.get('localizacao')}")
        if meta.get("arena_extra"):
            st.markdown(f"**Arena/Extra:** {meta.get('arena_extra')}")
        if meta.get("observacao"):
            st.markdown(f"**Observação:** {meta.get('observacao')}")
        st.markdown("</div>", unsafe_allow_html=True)


        lider = (staff.get("lider") or meta.get("lider") or "").strip()
        vice  = (staff.get("vice") or staff.get("vice_lider") or meta.get("vice_lider") or meta.get("vice-lider") or "").strip()
        
        def _canon_name(x: str) -> str:
            if not x:
                return ""
            return (x.replace("“", '"').replace("”", '"')
                     .replace("’", "'").replace("‘", "'")
                     .strip())
        
        def _norm_key(x: str) -> str:
            try:
                return _norm(x)
            except Exception:
                return str(x).strip().lower()
        
        def _norm_light(x: str) -> str:
            x = (x or "").strip().lower()
            x = x.replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
            x = _re.sub(r"\s+", " ", x)
            return x
        
        lider = _canon_name(lider)
        vice  = _canon_name(vice)
        
        def _find_npc(name: str):
            if not name:
                return None
            name2 = _canon_name(name)
            if name2 in npcs:
                return name2, npcs[name2]
            if name2 in staff_npcs:
                return name2, staff_npcs[name2]
        
            k = _norm_key(name2)
            for d in (npcs, staff_npcs):
                for nm, obj in d.items():
                    if _norm_key(nm) == k:
                        return nm, obj
            return None
        
        # -------- fallback: inferir líder/vice pelos NPCs (ocupacao) --------
        city_hit_name = (g.get("name") or g.get("city") or "").strip() if isinstance(g, dict) else ""
        if not city_hit_name:
            city_hit_name = city_name  # usa a cidade atual do compendium
        
        c_light = _norm_light(city_hit_name)
        
        if (not lider or not vice) and c_light:
            for nm, obj in npcs.items():
                occ = _norm_light(obj.get("ocupacao") or "")
                if not lider and ("lider de ginasio de" in occ and f"ginasio de {c_light}" in occ):
                    lider = nm
                if not vice and ("vice lider" in occ and f"ginasio de {c_light}" in occ):
                    vice = nm
        
        # -------- cards (fora do for/if) --------
        def person_card(role_label: str, name: str):
            if not name:
                return
            hit = _find_npc(name)
            npc_obj = (hit[1] if hit else {})
            shown_name = (hit[0] if hit else name)
        
            st.markdown('<div class="comp-card">', unsafe_allow_html=True)
            cA, cB = st.columns([0.65, 1.35], gap="large")
            with cA:
                imgp = comp_find_image(shown_name)  # usa a chave real encontrada
                if imgp:
                    st.image(imgp, use_container_width=True)
                else:
                    st.caption("🖼️ (sem retrato)")
            with cB:
                st.markdown(f"### {role_label}: {shown_name}")
                if npc_obj.get("ocupacao"):
                    st.caption(npc_obj.get("ocupacao"))
        
                pokes = npc_obj.get("pokemons") or []
                if pokes:
                    st.markdown("**Pokémons (sprites):**")
                    cols = st.columns(min(6, max(1, len(pokes))))
                    for i, pnm in enumerate(pokes[:12]):
                        with cols[i % len(cols)]:
                            spr = _poke_sprite_cached(pnm)
                            if spr:
                                st.image(spr, width=56)
                            st.caption(pnm)
        
                if shown_name in npcs:
                    if st.button("Abrir NPC", key=f"open_gym_npc_{_stem_key(city_name)}_{_stem_key(shown_name)}"):
                        st.session_state["comp_axis"] = "🧑‍🤝‍🧑 NPCs"
                        st.session_state["comp_npc_selected"] = shown_name
                        _touch_recent("comp_recent_npcs", shown_name)
                        st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
        
        if lider or vice:
            colL, colV = st.columns(2, gap="large")
            with colL:
                person_card("🏅 Líder", lider)
            with colV:
                person_card("🥈 Vice", vice)
        
        if (g.get("narrative") or "").strip():
            with st.expander("📜 Detalhes do Ginásio (texto)", expanded=False):
                st.markdown((g.get("narrative") or "").strip())



    # Seções (modo operação: seletor)
    sections = city_obj.get("sections", {}) or {}
    sublocs = city_obj.get("sublocais", []) or []

    ordered_keys = list(sections.keys())

    # ordem preferencial
    pref = ["Visão geral", "Como é viver", "Ginásio", "Sublocais", "Treinadores", "Economia", "Riscos", "Rumores", "Missões"]
    def kscore(k: str):
        nk = _norm(k)
        for i, p in enumerate(pref):
            if _norm(p) in nk:
                return (i, k)
        return (999, k)

    ordered_keys = [k for k in sorted(ordered_keys, key=kscore) if (sections.get(k) or "").strip()]

    opts = ["Tudo"] + ordered_keys + (["📍 Locais Importantes"] if sublocs else [])
    sel = st.selectbox("Seção do dossiê", opts, index=0, key=f"city_section_{_stem_key(city_name)}")

    if sel == "Tudo":
        for title in ordered_keys:
            body = (sections.get(title) or "").strip()
            if not body:
                continue
            st.markdown(f"### {title}")
            st.markdown(body)
    elif sel == "📍 Locais Importantes":
        st.markdown("### 📍 Locais Importantes")
    else:
        st.markdown(f"### {sel}")
        st.markdown((sections.get(sel) or "").strip())

    if sublocs and (sel in ["Tudo", "📍 Locais Importantes"]):
        st.markdown("### 📍 Locais Importantes")
        for it in sublocs:
            nm = (it.get("name") or "").strip()
            tx = (it.get("text") or "").strip()
            if not nm and not tx:
                continue
            with st.expander(f"📍 {nm}" if nm else "📍 Local"):
                if tx:
                    st.markdown(tx)
                else:
                    st.caption("(sem descrição)")



def _render_npc_dossier(nm: str, npc: dict, cities: dict[str, dict], npcs: dict[str, dict], gyms: dict[str, dict] | None = None) -> None:
    img = comp_find_npc_image(nm)

    st.markdown('<div class="comp-hero">', unsafe_allow_html=True)
    hA, hB = st.columns([1, 1.6], gap="large")
    with hA:
        if img:
            st.image(img, use_container_width=True)
        else:
            st.caption("🖼️ (sem retrato encontrado)")
        fav = _is_fav("comp_fav_npcs", nm)
        if st.button("⭐ Remover favorito" if fav else "⭐ Favoritar", key=f"fav_npc_{_stem_key(nm)}"):
            _toggle_fav("comp_fav_npcs", nm)
            st.rerun()

    with hB:
        st.markdown(f"## 🧑 {nm}")

        status = npc.get("status","")
        idade = npc.get("idade","")
        origem = npc.get("origem","")
        ocup = npc.get("ocupacao","")

        st.markdown('<div class="comp-divider"></div>', unsafe_allow_html=True)
        st.markdown("### 1️⃣ Identidade")
        c1, c2 = st.columns(2)
        with c1:
            if status: st.markdown(f"**Status:** {status}")
            if idade: st.markdown(f"**Idade:** {idade}")
            if origem: st.markdown(f"**Origem:** {origem}")
        with c2:
            if ocup: st.markdown(f"**Ocupação:** {ocup}")

        pokes = npc.get("pokemons") or []
        if pokes:
            st.markdown("**Pokémons conhecidos:**")
            spr_cols = st.columns(min(6, max(1, len(pokes))))
            for i, pnm in enumerate(pokes[:12]):
                with spr_cols[i % len(spr_cols)]:
                    spr = _poke_sprite_cached(pnm)
                    if spr:
                        st.image(spr, width=56)
                    st.caption(pnm)

        st.markdown('<div class="comp-panel">', unsafe_allow_html=True)
        st.markdown("### 🔮 Continue lendo")
        origem = (npc.get("origem") or "").strip()
        status = (npc.get("status") or "").strip()
        related = []
        for other, obj in npcs.items():
            if other == nm:
                continue
            if origem and _norm(obj.get("origem","")) == _norm(origem):
                related.append(other)
            elif status and _norm(obj.get("status","")) == _norm(status):
                related.append(other)
        if related:
            for other in related[:3]:
                if st.button(f"🧑 Conhecer {other}", key=f"npc_related_{_stem_key(nm)}_{_stem_key(other)}"):
                    st.session_state["comp_npc_selected"] = other
                    _touch_recent("comp_recent_npcs", other)
                    st.rerun()
        else:
            st.caption("Abra outro NPC da lista ao lado para continuar a história.")
        st.markdown("</div>", unsafe_allow_html=True)

    # ----------------------------
    # 🏅 Referências de Ginásio
    # ----------------------------
    gyms = gyms or {}
    gym_refs = []
    for city, g in gyms.items():
        meta = g.get("meta") or {}
        lider = (meta.get("lider") or (g.get("staff") or {}).get("lider") or "").strip()
        vice = (meta.get("vice_lider") or meta.get("vice-lider") or (g.get("staff") or {}).get("vice_lider") or "").strip()
        if _norm(lider) == _norm(nm):
            gym_refs.append(("🏅 Líder", city, meta.get("tipo",""), meta.get("status","")))
        if _norm(vice) == _norm(nm):
            gym_refs.append(("🥈 Vice", city, meta.get("tipo",""), meta.get("status","")))

    if gym_refs:
        st.markdown("### 🏅 Ginásio")
        for role, city, tipo, status in gym_refs:
            line = f"**{role}** em **{city}**"
            meta_bits = []
            if tipo:
                meta_bits.append(f"Tipo: {tipo}")
            if status:
                meta_bits.append(f"Status: {status}")
            if meta_bits:
                line += " — " + " • ".join(meta_bits)
            st.markdown(line)
            if city in cities:
                if st.button(f"Abrir cidade: {city}", key=f"npc_open_city_{_stem_key(nm)}_{_stem_key(city)}"):
                    st.session_state["comp_axis"] = "🌍 Locais"
                    st.session_state["comp_region"] = cities[city].get("region")
                    st.session_state["comp_city"] = city
                    _touch_recent("comp_recent_cities", city)
                    st.rerun()

    # Cross-links: cidades e NPCs citados
    city_names = list(cities.keys())
    npc_names = list(npcs.keys())
    cm = npc_mentions_cities(npc, city_names)
    nm_hits = npc_mentions_npcs(npc, npc_names, nm)
    if cm:
        st.markdown("**Cidades citadas:**")
        cols = st.columns(3)
        for i, cn in enumerate(cm[:9]):
            with cols[i % 3]:
                if st.button(f"🏙️ {cn}", key=f"npc_city_{_stem_key(nm)}_{_stem_key(cn)}"):
                    st.session_state["comp_axis"] = "🌍 Locais"
                    # tenta inferir região
                    st.session_state["comp_region"] = cities.get(cn, {}).get("region")
                    st.session_state["comp_city"] = cn
                    _touch_recent("comp_recent_cities", cn)
                    st.rerun()

    if nm_hits:
        st.markdown("**Outros NPCs citados:**")
        cols = st.columns(3)
        for i, other in enumerate(nm_hits[:9]):
            with cols[i % 3]:
                if st.button(f"🧑 {other}", key=f"npc_npc_{_stem_key(nm)}_{_stem_key(other)}"):
                    st.session_state["comp_npc_selected"] = other
                    _touch_recent("comp_recent_npcs", other)
                    st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("### 2️⃣ Lore completa")
    sections = npc.get("sections") or {}
    pref = ["Histórico", "Rumores", "Objetivos", "Estado Atual", "Lore"]
    def score(k):
        nk = _norm(k)
        for i, p in enumerate(pref):
            if _norm(p) in nk:
                return (i, k)
        return (999, k)

    for k in sorted(sections.keys(), key=score):
        v = (sections.get(k) or "").strip()
        if not v:
            continue
        if k != "Lore":
            st.markdown(f"#### {k}")
        st.markdown(v)


def render_compendium_sessions(comp_data: dict) -> None:
    st.markdown("## 📒 Tracker de Sessões")
    st.markdown(
    """
    <style>
      .sess-card {
        border: 1px solid rgba(255,255,255,0.10);
        border-radius: 14px;
        padding: 14px 14px 10px 14px;
        background: rgba(255,255,255,0.03);
      }
      .sess-card h4 { margin: 0 0 8px 0; }
      .sess-muted { opacity: 0.75; font-size: 0.92rem; }
      .sess-danger { border-color: rgba(255,0,0,0.20); background: rgba(255,0,0,0.06); }
    </style>
    """,
    unsafe_allow_html=True
)

    db = _sessions_db(globals().get("db"))
    trainer_name = _sessions_owner_name()
    sessions_data = load_sessions_data_cloud_first(db, trainer_name)

    sessions = sessions_data.setdefault("sessions", {})
    # --- RESTAURA UI (persistente em Firestore) ---
    sessions_data.setdefault("meta", {})
    sessions_data["meta"].setdefault("ui", {})
    ui = sessions_data["meta"]["ui"]
    
    # restaura seleção/ativa salvas no banco
    if "comp_session_selected" not in st.session_state:
        st.session_state["comp_session_selected"] = ui.get("selected_sid")
    
    if "comp_session_active_id" not in st.session_state:
        st.session_state["comp_session_active_id"] = ui.get("active_sid")
    
    # fallback seguro: se não tiver nada, seleciona a primeira existente
    if not st.session_state.get("comp_session_selected") and sessions:
        st.session_state["comp_session_selected"] = sorted(sessions.keys())[0]


    npcs = sorted((comp_data.get("npcs") or {}).keys())
    cities = sorted((comp_data.get("cities") or {}).keys())
    gyms = sorted((comp_data.get("gyms") or {}).keys())

    def _sorted_session_items() -> list[tuple[str, dict]]:
        def _sort_key(item: tuple[str, dict]):
            sid, sess = item
            date = sess.get("date") or ""
            num = sess.get("number") or 0
            return (date, int(num) if str(num).isdigit() else 0, sid)
        return sorted(sessions.items(), key=_sort_key)

    items_sorted = _sorted_session_items()
    default_number = 1
    if items_sorted:
        nums = [s.get("number") for _, s in items_sorted if str(s.get("number")).isdigit()]
        if nums:
            default_number = max(int(n) for n in nums) + 1

    with st.expander("➕ Nova Sessão", expanded=not bool(items_sorted)):
        st.markdown(
            """
            <div class="sess-card">
              <h4>Adicionar sessão</h4>
              <div class="sess-muted">Crie uma sessão e já deixe pronta para virar “ativa”.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.write("")
    
        with st.form("session_create_form", clear_on_submit=False):
            c1, c2, c3 = st.columns([1, 1, 2])
    
            with c1:
                number = st.number_input(
                    "Número",
                    min_value=1,
                    value=default_number,
                    step=1,
                    help="Ex.: 129"
                )
            with c2:
                date_val = st.date_input(
                    "Data",
                    value=datetime.utcnow().date(),
                    help="Data da sessão"
                )
            with c3:
                title = st.text_input(
                    "Título",
                    placeholder="Ex.: Hospital pós Bronzong",
                )
    
            tags_raw = st.text_input(
                "Tags (separadas por vírgula)",
                placeholder="Ex.: amber, shintaro, joice",
            )
    
            submitted = st.form_submit_button("✅ Criar sessão", use_container_width=True)
    
        if submitted:
            sid = _session_id_from_number(int(number))
            if not title.strip():
                st.error("Dê um título para a sessão antes de criar.")
            elif sid in sessions:
                st.error(f"Já existe a sessão {sid}.")
            else:
                sessions[sid] = {
                    "number": int(number),
                    "date": str(date_val),
                    "title": title.strip(),
                    "summary": "",
                    "tags": [t.strip() for t in tags_raw.split(",") if t.strip()],
                    "links": {"npcs": [], "places": [], "gyms": [], "items": []},
                    "events": [],
                    "flags": {},
                }
    
                # importante: garante que o payload carregado recebe o dict atualizado
                sessions_data["sessions"] = sessions
                save_sessions_data_cloud_first(db, trainer_name, sessions_data)
    
                st.session_state["comp_session_selected"] = sid
                st.success(f"Sessão {sid} criada.")
                st.rerun()

    if not sessions:
        st.info("Nenhuma sessão registrada ainda.")
        return

    options = [sid for sid, _ in items_sorted]
    labels = {sid: _session_label(sid, sess) for sid, sess in items_sorted}
    default_sid = st.session_state.get("comp_session_selected") or options[0]
    if default_sid not in sessions:
        default_sid = options[0]

    sel_label = labels.get(default_sid, default_sid)
    chosen_label = st.selectbox("Sessão atual", [labels[sid] for sid in options], index=[labels[sid] for sid in options].index(sel_label))
    selected_sid = next(sid for sid, label in labels.items() if label == chosen_label)
    st.session_state["comp_session_selected"] = selected_sid

    active_sid = st.session_state.get("comp_session_active_id")

    col_set, col_del, col_info = st.columns([1.1, 1.1, 2])
    
    with col_set:
        if st.button("✅ Definir como sessão ativa", key="set_active_session", use_container_width=True):
            st.session_state["comp_session_active_id"] = selected_sid
            st.success(f"Sessão ativa: {selected_sid}")
            st.rerun()
    
    with col_del:
        with st.popover("🗑️ Excluir sessão", use_container_width=True):
            st.markdown(
                f"""
                <div class="sess-card sess-danger">
                  <b>Excluir {selected_sid}</b><br/>
                  <span class="sess-muted">Esta ação não pode ser desfeita.</span>
                </div>
                """,
                unsafe_allow_html=True,
            )
            st.write("")
            confirm = st.checkbox("Confirmo que quero excluir esta sessão", key="sess_del_confirm")
            typed = st.text_input("Digite o ID da sessão para confirmar", value="", key="sess_del_typed", placeholder=selected_sid)
    
            if st.button("❌ Excluir definitivamente", key="sess_del_btn", use_container_width=True, disabled=not (confirm and typed.strip() == selected_sid)):
                # remove
                if selected_sid in sessions:
                    sessions.pop(selected_sid, None)
    
                # se era ativa, limpa
                if st.session_state.get("comp_session_active_id") == selected_sid:
                    st.session_state["comp_session_active_id"] = None
    
                # ajusta selecionada
                items_left = sorted(
                    sessions.items(),
                    key=lambda it: ((it[1].get("date") or ""), int(it[1].get("number") or 0), it[0])
                )
                if items_left:
                    st.session_state["comp_session_selected"] = items_left[0][0]
                else:
                    st.session_state["comp_session_selected"] = None
    
                sessions_data["sessions"] = sessions
                # --- PERSISTE UI NO BANCO ---
                sessions_data.setdefault("meta", {})
                sessions_data["meta"].setdefault("ui", {})
                sessions_data["meta"]["ui"]["active_sid"] = st.session_state.get("comp_session_active_id")
                sessions_data["meta"]["ui"]["selected_sid"] = st.session_state.get("comp_session_selected")
                save_sessions_data_cloud_first(db, trainer_name, sessions_data)

                save_sessions_data_cloud_first(db, trainer_name, sessions_data)
    
                st.success(f"Sessão {selected_sid} excluída.")
                st.rerun()
    
    with col_info:
        if active_sid:
            st.caption(f"Sessão ativa atual: {active_sid}")
        else:
            st.caption("Nenhuma sessão ativa definida.")

    session = sessions.get(selected_sid, {})
    _ensure_session_links(session)

    tabs = st.tabs(["Resumo", "Eventos", "Entidades", "Flags", "Timeline", "Sessões"])

    with tabs[0]:
        with st.form("session_summary_form"):
            summary = st.text_area("Resumo da sessão", value=session.get("summary") or "", height=140)
            tags_raw = st.text_input("Tags", value=", ".join(session.get("tags") or []))
            st.markdown("**Links rápidos (IDs do compendium):**")
            col_n, col_p = st.columns(2)
            with col_n:
                link_npcs = st.multiselect("NPCs", npcs, default=session["links"].get("npcs", []))
                link_items = st.multiselect("Itens", [], default=session["links"].get("items", []))
            with col_p:
                link_places = st.multiselect("Locais", cities, default=session["links"].get("places", []))
                link_gyms = st.multiselect("Ginásios", gyms, default=session["links"].get("gyms", []))
            save_summary = st.form_submit_button("Salvar resumo")

        if save_summary:
            session["summary"] = summary.strip()
            session["tags"] = [t.strip() for t in tags_raw.split(",") if t.strip()]
            session["links"]["npcs"] = link_npcs
            session["links"]["places"] = link_places
            session["links"]["gyms"] = link_gyms
            session["links"]["items"] = link_items
            save_sessions_data_cloud_first(db, trainer_name, sessions_data)
            st.success("Resumo atualizado.")

    with tabs[1]:
        st.markdown("### ⚡ Adicionar evento rápido")
        event_types = {
            "battle": "Batalha",
            "reveal": "Revelação",
            "travel": "Viagem",
            "loot": "Item",
            "npc": "NPC apareceu",
            "quest": "Quest",
        }
        st.session_state.setdefault("sess_event_type", "battle")
        st.session_state.setdefault("sess_event_text", "")

        col_type, col_time = st.columns([1.6, 1])
        with col_type:
            ev_type = st.selectbox("Tipo", list(event_types.keys()), index=list(event_types.keys()).index(st.session_state["sess_event_type"]))
        with col_time:
            ev_time = st.text_input("Hora (HH:MM)", value=datetime.utcnow().strftime("%H:%M"))

        ev_text = st.text_input("Descrição do evento", value=st.session_state["sess_event_text"])
        st.session_state["sess_event_type"] = ev_type
        st.session_state["sess_event_text"] = ev_text
        col_r1, col_r2, col_r3, col_r4 = st.columns(4)
        with col_r1:
            ref_npc = st.selectbox("Ref NPC", [""] + npcs, index=0)
        with col_r2:
            ref_place = st.selectbox("Ref Local", [""] + cities, index=0)
        with col_r3:
            ref_gym = st.selectbox("Ref Ginásio", [""] + gyms, index=0)
        with col_r4:
            ref_item = st.text_input("Ref Item (ID)")

        if st.button("Adicionar evento", key="add_event_button"):
            refs = {}
            if ref_npc:
                refs["npc"] = ref_npc
                _add_entity_to_session(session, "npcs", ref_npc)
            if ref_place:
                refs["place"] = ref_place
                _add_entity_to_session(session, "places", ref_place)
            if ref_gym:
                refs["gym"] = ref_gym
                _add_entity_to_session(session, "gyms", ref_gym)
            if ref_item:
                refs["item"] = ref_item
                _add_entity_to_session(session, "items", ref_item)

            _append_session_event(
                session,
                {
                    "t": ev_time.strip(),
                    "type": ev_type,
                    "text": ev_text.strip(),
                    "refs": refs,
                },
            )
            save_sessions_data_cloud_first(db, trainer_name, sessions_data)
            st.success("Evento adicionado.")
            st.rerun()

        st.markdown("### 📜 Eventos registrados")
        icon_map = {
            "battle": "⚔️",
            "reveal": "🔎",
            "travel": "🗺️",
            "loot": "🎁",
            "npc": "🧑",
            "quest": "📜",
        }
        for ev in session.get("events", []):
            icon = icon_map.get(ev.get("type"), "•")
            ref_txt = ""
            if ev.get("refs"):
                ref_txt = " — refs: " + ", ".join(f"{k}:{v}" for k, v in ev["refs"].items())
            st.markdown(f"{icon} **{ev.get('t','')}** {ev.get('text','')}{ref_txt}")

    with tabs[2]:
        st.markdown("### 🔗 Entidades vinculadas")
        st.markdown("**NPCs:** " + (", ".join(session["links"].get("npcs", [])) or "—"))
        st.markdown("**Locais:** " + (", ".join(session["links"].get("places", [])) or "—"))
        st.markdown("**Ginásios:** " + (", ".join(session["links"].get("gyms", [])) or "—"))
        st.markdown("**Itens:** " + (", ".join(session["links"].get("items", [])) or "—"))

        st.markdown("### 🔎 Buscar por entidade")
        entity_type = st.selectbox("Tipo", ["NPC", "Local", "Ginásio", "Item"])
        if entity_type == "NPC":
            entity_id = st.selectbox("NPC", npcs)
            link_key = "npcs"
        elif entity_type == "Local":
            entity_id = st.selectbox("Local", cities)
            link_key = "places"
        elif entity_type == "Ginásio":
            entity_id = st.selectbox("Ginásio", gyms)
            link_key = "gyms"
        else:
            entity_id = st.text_input("Item (ID)")
            link_key = "items"

        if entity_id:
            matched = []
            for sid, sess in sessions.items():
                links = (sess.get("links") or {}).get(link_key, [])
                if entity_id in links:
                    matched.append(_session_label(sid, sess))
            if matched:
                st.markdown("**Aparece nas sessões:**")
                for line in matched:
                    st.markdown(f"- {line}")
            else:
                st.caption("Nenhuma sessão encontrada para essa entidade.")

    with tabs[3]:
        st.markdown("### 🚩 Flags")
        flags = session.get("flags") or {}
        with st.form("session_flags_form"):
            new_flags = {}
            for key, val in flags.items():
                new_flags[key] = st.checkbox(key, value=bool(val), key=f"flag_{selected_sid}_{key}")
            new_key = st.text_input("Novo flag")
            new_val = st.checkbox("Ativo", value=True, key=f"flag_new_{selected_sid}")
            save_flags = st.form_submit_button("Salvar flags")

        if save_flags:
            flags.update(new_flags)
            if new_key.strip():
                flags[new_key.strip()] = bool(new_val)
            session["flags"] = flags
            save_sessions_data_cloud_first(db, trainer_name, sessions_data)
            st.success("Flags atualizadas.")

    with tabs[4]:
        st.markdown("### 🕒 Timeline global")
        for sid, sess in items_sorted:
            line = _session_label(sid, sess)
            st.markdown(f"- {line}")

    with tabs[5]:
        st.markdown("### 📋 Lista de sessões")
        tag_filter = st.text_input("Filtrar por tag")
        for sid, sess in items_sorted:
            tags = sess.get("tags") or []
            if tag_filter and tag_filter not in tags:
                continue
            st.markdown(f"- {_session_label(sid, sess)}")

@st.cache_data
def get_font_base64(font_path):
    """Lê o arquivo de fonte e converte para base64 para uso no CSS."""
    try:
        with open(font_path, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except Exception:
        return None
# ==============================================================================
# 📚 COMPENDIUM NOVO (JSON + DARK SOULS) - CORRIGIDO
# ==============================================================================

def render_compendium_page() -> None:
    # garante default
    if "comp_view" not in st.session_state:
        st.session_state["comp_view"] = "home"
    
    # --- INÍCIO DA INSERÇÃO ---
    font_b64 = get_font_base64("fonts/DarkSouls.ttf")
    font_css = f"@font-face {{ font-family: 'DarkSouls'; src: url('data:font/ttf;base64,{font_b64}') format('truetype'); }}" if font_b64 else ""

    st.markdown(f"""
    <style>
        {font_css}
        :root {{
            --ds-font: 'DarkSouls', serif;
            --ds-gold-dim: rgba(255,215,0,0.55);
        }}
        /* Fundo preto APENAS quando o Compendium (ds-home) estiver presente */
        html:has(.ds-home),
        body:has(.ds-home),
        .stApp:has(.ds-home),
        [data-testid="stAppViewContainer"]:has(.ds-home) {{
          background: #000 !important;
          color: #f8fafc;
        }}
        
        /* Fonte do Compendium só dentro do Compendium */
        [data-testid="stAppViewContainer"]:has(.ds-home),
        [data-testid="stAppViewContainer"]:has(.ds-home) * {{
          font-family: var(--ds-font) !important;
        }}
        /* Aplica a fonte em tudo na aba compendium */
        [data-testid="stAppViewContainer"], [data-testid="stAppViewContainer"] * {{
            font-family: var(--ds-font) !important;
        }}
        /* ESTILO DOS BOTÕES (O que você gostou) */
        /* Remove a caixa/borda padrão e deixa só o texto */
        .stButton > button {{
            background: transparent !important;
            border: none !important;
            color: #888 !important; /* Cor normal (cinza escuro) */
            font-size: 24px !important;
            text-shadow: 0px 0px 5px rgba(0,0,0,0.8);
            font-family: 'DarkSouls', serif !important;
            text-transform: uppercase;
            transition: transform 0.2s, color 0.2s;
        }}
        .stButton > button:hover {{
            color: #FFD700 !important; text-shadow: 0 0 10px #FFD700; transform: scale(1.1);
        }}
        .stButton > button:active, .stButton > button:focus {{
            color: #FFD700 !important; outline: none !important; border: none !important; box-shadow: none !important;
        }}
        
    /* Esconde elementos padrão do Streamlit para imersão */
    [data-testid="stHeader"] {{ visibility: hidden; }}
    [data-testid="stSidebar"] {{ display: none !important; }}
    
    /* Esconde o resto da UI padrão do Streamlit (menu, toolbar, rodapé) */
    #MainMenu {{ visibility: hidden; }}
    footer {{ visibility: hidden; }}
    header {{ visibility: hidden; }}
    [data-testid="stToolbar"] {{ visibility: hidden !important; height: 0px !important; }}
    [data-testid="stStatusWidget"] {{ visibility: hidden !important; }}
    [data-testid="stDeployButton"] {{ display: none !important; }}

    
/* HOME (igual app 35) */
.ds-home {{
    height: 100vh;
    box-sizing: border-box;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: flex-start;
    gap: 12px;
    padding: 10vh 0 220px 0; /* topo (logo/press) + reserva para o menu fixo */
    overflow: hidden;
}}

.ds-logo{{
  width: min(78vw, 760px);
  max-height: 34vh;  /* impede a logo de descer e colidir com o menu */
  height: auto;
  display: block;
  margin: 0;
  filter: drop-shadow(0 0 18px rgba(0,0,0,0.75));
}}
.ds-title {{
    text-align: center;
    color: var(--ds-white);
    font-size: 56px;
    letter-spacing: 0.28em;
    text-transform: uppercase;
    margin: 0;
}}
.ds-press {{
    position: static;
    width: 100%;
    pointer-events: none;
    text-align: center;
    color: var(--ds-faint);
    font-size: 14px;
    letter-spacing: 0.34em;
    text-transform: uppercase;
    margin: 0 0 12px 0;
}}
@keyframes dsBlink {{
    0%, 48% {{ opacity: 0.10; }}
    60%, 100% {{ opacity: 0.88; }}
}}
.ds-blink {{ animation: dsBlink 1.05s ease-in-out infinite; }}

/* Tabs (radio) no rodapé — SOMENTE NA HOME */
[data-testid="stAppViewContainer"]:has(.ds-home) div[data-testid="stRadio"] {{
    position: fixed !important;
    left: 50% !important;
    transform: translateX(-50%) !important;
    bottom: 48px !important;
    z-index: 10000 !important;
    padding: 10px 18px !important;
    background: rgba(0,0,0,0.0) !important;
}}
[data-testid="stAppViewContainer"]:has(.ds-home) div[data-testid="stRadio"] > label {{ display: none !important; }}
/* Linha dourada fina acima dos tabs */
[data-testid="stAppViewContainer"]:has(.ds-home) div[data-testid="stRadio"]::before {{
    content: "";
    display: block;
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--ds-gold-dim), transparent);
    margin-bottom: 10px;
}}

/* Estilo das opções */
[data-testid="stAppViewContainer"]:has(.ds-home) div[role="radiogroup"] {{
    display: flex !important;
    flex-direction: column !important;
    gap: 18px !important;
    justify-content: center !important;
    align-items: center !important;
}}
[data-testid="stAppViewContainer"]:has(.ds-home) div[role="radiogroup"] > label {{
    position: relative !important;
    padding: 6px 34px !important;
    cursor: pointer !important;
    color: rgba(255,255,255,0.70) !important;
    text-transform: uppercase !important;
    letter-spacing: 0.24em !important;
    font-size: 12px !important;
    user-select: none !important;
    transition: all 120ms ease !important;
}}
[data-testid="stAppViewContainer"]:has(.ds-home) div[role="radiogroup"] > label::before {{
    content: "" !important;
    position: absolute !important;
    inset: 0 !important;
    border-radius: 999px !important;
    background: radial-gradient(ellipse at center, rgba(255, 200, 64, 0.55), rgba(255, 200, 64, 0.05) 70%, transparent 100%) !important;
    opacity: 0 !important;
    filter: blur(0.4px) !important;
    transition: opacity 140ms ease, transform 140ms ease !important;
    transform: scaleX(0.85) !important;
    z-index: -1 !important;
}}
[data-testid="stAppViewContainer"]:has(.ds-home) div[role="radiogroup"] > label:hover {{
    color: #FFD700 !important;
    text-shadow: 0 0 10px rgba(255,215,0,0.65) !important;
    transform: translateY(-1px) !important;
}}
.ds-home div[role="radiogroup"] > label:hover {{
    color: #FFD700 !important;
    text-shadow: 0 0 10px rgba(255,215,0,0.65) !important;
    transform: translateY(-1px) !important;
}}
[data-testid="stAppViewContainer"]:has(.ds-home) div[role="radiogroup"] > label[data-checked="true"] {{
    color: #FFD700 !important;
    text-shadow: 0 0 10px rgba(255,215,0,0.65) !important;
}}
[data-testid="stAppViewContainer"]:has(.ds-home) div[role="radiogroup"] > label:hover::before,
[data-testid="stAppViewContainer"]:has(.ds-home) div[role="radiogroup"] > label[data-checked="true"]::before {{
    opacity: 1 !important;
    transform: scaleX(1) !important;
}}
/* Trava scroll apenas na HOME do compendium */
html:has(.ds-home),
body:has(.ds-home),
[data-testid="stAppViewContainer"]:has(.ds-home),
section.main:has(.ds-home),
div[data-testid="stMain"]:has(.ds-home),
div[data-testid="stMainBlockContainer"]:has(.ds-home) {{
    overflow: hidden !important;
    height: 100vh !important;
    overscroll-behavior: none !important;
}}
/* Esconde o bolinha padrão do radio */
[data-testid="stAppViewContainer"]:has(.ds-home) div[role="radiogroup"] input {{ display: none !important; }}
[data-testid="stAppViewContainer"]:has(.ds-home) div[role="radiogroup"] > label > div:first-child {{
    display: none !important;
}}


.ds-frame {{
            background: rgba(0,0,0,0.55);
            border: 2px solid rgba(176,143,60,0.55);
            box-shadow: 0 0 45px rgba(0,0,0,0.9);
            border-radius: 12px;
            padding: 26px 26px 18px 26px;
            position: relative;
        }}
        .ds-frame::after {{
            content: "";
            position: absolute;
            top: 10px; left: 10px; right: 10px; bottom: 10px;
            border: 1px solid rgba(255,255,255,0.10);
            border-radius: 10px;
            pointer-events: none;
        }}
        .ds-name {{
            font-size: 52px;
            text-transform: uppercase;
            letter-spacing: 0.22em;
            text-align: center;
            margin: 0 0 10px 0;
            padding-bottom: 14px;
            border-bottom: 1px solid rgba(255,255,255,0.10);
        }}
        .ds-meta {{
            text-align: center;
            color: var(--ds-faint);
            letter-spacing: 0.20em;
            text-transform: uppercase;
            font-size: 13px;
            margin-bottom: 18px;
        }}
        .ds-portrait {{
            display:flex; justify-content:center; margin: 12px 0 10px 0;
        }}
        .ds-portrait img {{
            max-width: 320px; width:100%;
            border-radius: 10px;
            border: 1px solid rgba(255,255,255,0.12);
            box-shadow: 0 0 26px rgba(0,0,0,0.75);
        }}
        .ds-section-title {{
            margin-top: 14px; margin-bottom: 10px;
            color: var(--ds-faint);
            letter-spacing: 0.18em;
            text-transform: uppercase;
            font-size: 14px;
            border-bottom: 1px solid rgba(255,255,255,0.08);
            padding-bottom: 8px;
            text-align:center;
        }}
        .ds-poke-container {{
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            background: rgba(0,0,0,0.4);
            border: 1px solid #443311;
            border-radius: 8px;
            padding: 8px;
            transition: all 0.3s ease;
        }}
        .ds-poke-container:hover {{
            border-color: #FFD700;
            background: rgba(20,20,10,0.8);
            transform: translateY(-3px);
        }}
        .ds-poke-img {{
            width: 60px;
            height: 60px;
            object-fit: contain;
            filter: sepia(1) brightness(0.8) contrast(1.2);
            opacity: 0.85;
            transition: all 0.4s ease;
        }}
        .ds-poke-container:hover .ds-poke-img {{
            filter: sepia(0) brightness(1) contrast(1);
            opacity: 1;
        }}
        .ds-poke-name {{
            font-size: 10px;
            text-transform: uppercase;
            color: #887766;
            margin-top: 5px;
            text-align: center;
            letter-spacing: 1px;
        }}
        .ds-poke-container:hover .ds-poke-name {{
            color: #FFD700;
        }}
        .ds-history p {{
            color: rgba(255,255,255,0.88);
            font-size: 18px;
            line-height: 1.68;
            text-align: justify;
            margin: 0 0 14px 0;
        }}
        /* TOP NAV — texto puro */
        .ds-tab div[data-testid="stButton"] > button {{
          background: transparent !important;
          border: none !important;
          box-shadow: none !important;
          color: rgba(255,255,255,0.72) !important;
        }}
        
        /* Hover dourado */
        .ds-tab div[data-testid="stButton"] > button:hover {{
          color: rgba(255,215,0,0.95) !important;
          text-shadow: 0 0 14px rgba(255,215,0,0.35) !important;
        }}
        
        /* Selecionado = dourado sempre */
        .ds-tab.selected div[data-testid="stButton"] > button {{
          color: rgba(255,215,0,0.98) !important;
          text-shadow: 0 0 18px rgba(255,215,0,0.45) !important;
        }}
        .ds-gold-top{{
          height: 1px;
          width: 100%;
          margin: 12px 0 18px 0;
          background: linear-gradient(90deg, transparent, var(--ds-gold-dim), transparent);
        }}
        .ds-nav-item,
        .ds-nav-item:visited {{
          color: rgba(255,255,255,0.72) !important;
          text-decoration: none !important;
        }}
        
        .ds-nav-item:hover {{
          color: rgba(255,215,0,0.95) !important;
          text-shadow: 0 0 14px rgba(255,215,0,0.35) !important;
          text-decoration: none !important;
        }}
        
        .ds-nav-item.selected {{
          color: rgba(255,215,0,0.98) !important;
          text-shadow: 0 0 16px rgba(255,215,0,0.45) !important;
          text-decoration: none !important;
        }}
        
        /* Remove outline padrão */
        .ds-tab div[data-testid="stButton"] > button:focus {{
          outline: none !important;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )
        
    components.html("<div class='ds-gold-top'></div>", height=10)

    # 4️⃣ Menu
   
    # ----------------------------
    # Estado do Compendium — NÃO cai direto em NPCs
    # ----------------------------
    st.session_state.setdefault("comp_view", "home")
    st.session_state.setdefault("comp_selected_npc", None)
    try:
        comp_data = comp_load()  # retorna {"regions","cities","npcs","gyms",...}
    except Exception as e:
        st.error(f"Falha ao carregar Compendium: {e}")
        return
    
    npcs_gerais = comp_data.get("npcs", {}) or {}
    # ----------------------------
    # Navegação via query params (?cv=...) — evita click-detector na HOME e no TOP NAV
    # ----------------------------
    def _qp_get(key: str):
        qp = st.query_params
        v = qp.get(key)
        if v is None:
            return None
        if isinstance(v, list):
            return v[0] if v else None
        return str(v)
    
    def _clear_qp():
        st.query_params.clear()
    
    
    

    
    def _consume_comp_qp():
        cv = _qp_get("cv")  # home / npcs / ginasios / locais / sair
        if not cv:
            return
    
        if cv == "sair":
            st.session_state["nav_to"] = "Pokédex (Busca)"
        else:
            st.session_state["comp_view"] = cv
    
            # se voltou pro HOME, reseta o radio do rodapé
            if cv == "home":
                st.session_state["ds_home_tabs"] = "__home__"
                st.session_state["ds_home_tabs_prev"] = "__home__"
    
            if cv != "npcs":
                st.session_state["comp_selected_npc"] = None
    
        _clear_qp()
        st.rerun()
    



            
    # ----------------------------
    # Navegação (sempre no topo)
    # ----------------------------
    # ----------------------------
    # Navegação (HOME embaixo / outras páginas em cima e sticky)
    # ----------------------------
    def _go(view: str):
        st.session_state["comp_view"] = view
        if view != "npcs":
            st.session_state["comp_selected_npc"] = None
        st.rerun()
    
    
        st.markdown("</div>", unsafe_allow_html=True)



    _consume_comp_qp()
    

    # =====================================================================

    # NPCs (VERSÃO CORRIGIDA - SAFE IDs)
    # =====================================================================
    if st.session_state["comp_view"] == "npcs":
        try:
            from st_click_detector import click_detector
        except ImportError:
            st.error("Biblioteca não instalada. Adicione 'st-click-detector' ao requirements.txt e reinicie o app.")
            return

        render_ds_tools_nav(st.session_state["comp_view"])
        st.markdown("""
        <style>
        /* Fundo global do app (mata rgb(14,17,23)) */
        [data-testid="stAppViewContainer"],
        [data-testid="stAppViewContainer"] > .main,
        [data-testid="stMain"],
        [data-testid="stMainBlockContainer"],
        section.main {
          background: #000 !important;
        }
        
        /* Wrappers internos */
        [data-testid="stVerticalBlock"],
        [data-testid="stHorizontalBlock"],
        [data-testid="stBlock"]{
          background: transparent !important;
        }
        
        /* pega exatamente o div com style inline rgb(14, 17, 23) */
        section.main div[style*="background: rgb(14, 17, 23)"]{
          background: #000 !important;
        }
        </style>
        """, unsafe_allow_html=True)
    
        
        # CSS das molduras (não interfere no click)
        css = """
        <style>
          .ds-npc-banner{
              position: fixed;
              top: 0;
              left: 0;
              right: 0;
              z-index: 12000;
              text-align: center;
              background: rgba(0,0,0,0.85);
              border-bottom: 1px solid rgba(176,143,60,0.6);
              box-shadow: 0 6px 18px rgba(0,0,0,0.6);
              padding: 14px 12px 12px 12px;
              text-transform: uppercase;
              letter-spacing: 0.25em;
            }
        
          .ds-npc-banner-title{
            font-size: 20px;
            color: rgba(255,255,255,0.95);
            margin: 0 0 6px 0;
          }
          .ds-npc-banner-sub{
            font-size: 11px;
            color: rgba(255,255,255,0.65);
            margin: 0;
          }
          .ds-npc-banner-spacer{
            height: 14px;
          }
        
          .ds-npc-panel{
            background-repeat:no-repeat;
            background-position:center;
            background-size:100% 100%;
            padding: 28px 28px 26px 28px;
            min-height: 0px; /* deixa crescer pelo conteúdo */
          }
        
          .ds-npc-panel.left{
            background: transparent !important;
            background-image: none !important;
            border: none !important;
            box-shadow: none !important;
            padding-top: 6px !important; /* bem pequeno */
          }
        
          /* (opcional) tira o padding padrão do Streamlit só nessa tela */
          [data-testid="stMainBlockContainer"]{
            padding-top: 0 !important;
          }
        
          /* ✅ FIX: remove o “contorno preto” do st_click_detector
             (APENAS dentro do painel de NPCs; não mexe no iframe do Compendium) */
          .ds-npc-panel div[data-testid="stComponentFrame"],
          .ds-npc-panel div[data-testid="stCustomComponentV1"],
          .ds-npc-panel div[data-testid="stCustomComponent"]{
            background: transparent !important;
            border: 0 !important;
            box-shadow: none !important;
            padding: 0 !important;
            margin: 0 !important;
          }
        
          .ds-npc-panel div[data-testid="stComponentFrame"] > iframe,
          .ds-npc-panel div[data-testid="stCustomComponentV1"] > iframe,
          .ds-npc-panel div[data-testid="stCustomComponent"] > iframe,
          .ds-npc-panel iframe[title^="st_click_detector"],
          .ds-npc-panel iframe[title*="click_detector"]{
            background: transparent !important;
            border: 0 !important;
            box-shadow: none !important;
            outline: none !important;
          }
        
          /* 3) remove padding extra do wrapper do Streamlit */
          .ds-npc-panel div[data-testid="stElementContainer"]{
            padding: 0 !important;
            margin: 0 !important;
            background: transparent !important;
          }
        
          /* remove padding/margem que às vezes vira “caixa” */
          .ds-npc-panel.left div[data-testid="stElementContainer"]{
            padding: 0 !important;
            margin: 0 !important;
          }
        
          /* grid automático */
          .ds-grid{
            display:grid;
            grid-template-columns:repeat(auto-fill, minmax(140px, 1fr));
            gap: 10px;
            width:100%;
          }
        
          /* evita qualquer camada bloquear cliques */
          .ds-npc-panel, .ds-npc-panel * { pointer-events:auto; }
        </style>
        """

        
        st.markdown(css, unsafe_allow_html=True)
        st.markdown(
            "<div class='ds-npc-banner'>"
            "<div class='ds-npc-banner-title'>Selecione um NPC</div>"
            "<div class='ds-npc-banner-sub'>Clique em um retrato à esquerda</div>"
            "</div>"
            "<div class='ds-npc-banner-spacer'></div>",
            unsafe_allow_html=True,
        )

        st.markdown(
            "<div class='ds-npc-banner'>"
            "<div class='ds-npc-banner-title'>Selecione um NPC</div>"
            "<div class='ds-npc-banner-sub'>Clique em um retrato à esquerda</div>"
            "</div>",
            unsafe_allow_html=True,
        )
        


        # --- LAYOUT PRINCIPAL (SEM INDENTAÇÃO ERRADA) ---
        left, right = st.columns([1.25, 2.15], gap="large")
    
        # --- COLUNA ESQUERDA ---
        with left:
            st.markdown("<div class='ds-npc-panel left'>", unsafe_allow_html=True)
            search = st.text_input(
                "Buscar personagem",
                key="ds_npc_search",
                placeholder="Nome ou história...",
            ).strip()
    
            def _norm(s: str) -> str:
                if not isinstance(s, str):
                    return ""
                return re.sub(r"\s+", " ", s).strip().lower()
    
            q = _norm(search)
    
            # Preparar lista
            items = []
            for nome, obj in (npcs_gerais or {}).items():
                if not isinstance(obj, dict):
                    continue
    
                historia = ""
                secs = obj.get("sections") or {}
                if isinstance(secs, dict):
                    historia = secs.get("História") or secs.get("Historia") or ""
                historia = _strip_html_if_any(historia)
                

    
                hay = _norm(nome) + " " + _norm(historia)
                if not q or q in hay:
                    items.append((nome, obj))
    
            items.sort(key=lambda x: x[0])
    
            # Cache de imagens
            @st.cache_data(show_spinner=False)
            def _thumb_data_uri(path: str, max_w: int = 360, max_h: int = 520) -> str:
                try:
                    from PIL import Image
                    import io, base64, os
    
                    if not path or not os.path.exists(path):
                        return ""
                    img = Image.open(path).convert("RGB")
                    img.thumbnail((max_w, max_h))
                    buf = io.BytesIO()
                    img.save(buf, format="JPEG", quality=70)
                    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
                    return f"data:image/jpeg;base64,{b64}"
                except:
                    return ""
    
            if not items:
                st.info("Nenhum NPC encontrado.")
            else:
                content_html = """
                <style>
                  html, body{
                    background: #000 !important;
                    margin: 0 !important;
                    padding: 6px !important;  /* pinta também a “folga” */

                  }
                :root{ --background-color:#000 !important; --text-color:#fff !important; }

                  body, div { background:#000 !important; }
                
                  div[style*="background: rgb(14, 17, 23)"]{
                    background:#000 !important;
                  }
                                
                  .ds-npc-grid{
                    display:grid;
                    grid-template-columns:repeat(4, 1fr);
                    gap: 10px;
                    width:100%;
                    row-gap:12px;

                    background: #000 !important;
                  }
                                
                  .ds-card {
                    position: relative;
                    aspect-ratio: 3/4;
                    border: 2px solid #554422;
                    background: #000 !important;   /* <- ISSO remove o “cinza” */

                    border-radius: 8px;
                    overflow: hidden;
                    cursor: pointer;
                    transition: transform 0.1s;
                  }
                
                  .ds-card:hover { border-color: #FFD700; transform: scale(1.02); }
                
                  .ds-card img {
                    width: 100%;
                    height:100%;
                    object-fit: cover;
                    filter: brightness(0.8);
                    display: block;
                    background: #000 !important;   /* reforço (se o browser/iframe insistir) */

                  }
                
                  .ds-name-tag {
                    position: absolute;
                    bottom: 0; left: 0; right: 0;
                    background: rgba(0,0,0,0.85);
                    color: #ddd;
                    font-size: 10px;
                    text-align: center;
                    padding: 4px;
                    font-weight: bold;
                    text-transform: uppercase;
                  }
                
                  a { text-decoration: none; display: block; }
                </style>
                
                <div class="ds-npc-grid">

                """
   
                id_map = {}
    
                for idx, (nome, obj) in enumerate(items):
                    safe_id = str(idx)
                    id_map[safe_id] = nome
    
                    img_path = None
                    try:
                        img_path = comp_find_npc_image(nome)
                    except:
                        pass
    
                    src = _thumb_data_uri(img_path) if img_path else ""
                    img_html = f"<img src='{src}' />" if src else "<div style='width:100%;height:100%;background:#222;'></div>"
    
                    safe_name = html.escape(str(nome), quote=True)

                    content_html += (
                        f'<a href="javascript:void(0)" id="{safe_id}">'
                        f'  <div class="ds-card">'
                        f'    {img_html}'  # img_html já é seu <img ...> ou <div ...>
                        f'    <div class="ds-name-tag">{safe_name}</div>'
                        f'  </div>'
                        f'</a>'
                    )
                        
                content_html += "</div>"
                clicked_safe_id = click_detector(content_html)
    
                if clicked_safe_id is not None:
                    nome_selecionado = id_map.get(str(clicked_safe_id))
                    if nome_selecionado and nome_selecionado != st.session_state.get("comp_selected_npc"):
                        st.session_state["comp_selected_npc"] = nome_selecionado
                        st.rerun()
                
                st.markdown("</div>", unsafe_allow_html=True)  # NÃO remove: fecha <div class='ds-npc-panel left'>

            
        # --- COLUNA DIREITA ---
        with right:
            sel = st.session_state.get("comp_selected_npc")
            if not sel:
                panel_html = (
                    "<div class='ds-npc-panel right'>"
                    "<div class='ds-frame'>"
                    "<div class='ds-name' style='font-size:30px;'>NENHUM NPC SELECIONADO</div>"
                    "</div>"
                    "</div>"
                )
                st.markdown(panel_html, unsafe_allow_html=True)
            else:
                active_sid = st.session_state.get("comp_session_active_id")
                if active_sid:
                    if st.button("➕ Adicionar à sessão atual", key=f"add_npc_session_{_stem_key(sel)}"):
                        if add_entity_to_active_session(
                            "npcs",
                            sel,
                            "npc",
                            f"NPC apareceu: {sel}",
                            {"npc": sel},
                        ):
                            st.success(f"{sel} adicionado à sessão {active_sid}.")
                npc = npcs_gerais.get(sel, {}) or {}
                ocupacao = npc.get("ocupacao", "")
                idade = npc.get("idade", "")
                status = npc.get("status", "")

                # retrato grande (base64 p/ garantir)
                portrait_b64 = ""
                portrait_path = None
                try:
                    portrait_path = comp_find_npc_image(sel)
                except Exception:
                    portrait_path = None

                ext = "png"
                if portrait_path and os.path.exists(portrait_path):
                    try:
                        with open(portrait_path, "rb") as f:
                            portrait_b64 = base64.b64encode(f.read()).decode("utf-8")
                        ext = os.path.splitext(portrait_path)[1].lower().replace(".", "")
                        if ext not in ("png", "jpg", "jpeg", "webp"):
                            ext = "png"
                    except Exception:
                        portrait_b64 = ""
                        ext = "png"

                # sprites dos pokemons (usa mapa oficial)
                pokemons = npc.get("pokemons") or npc.get("pokemons_conhecidos") or []
                if not isinstance(pokemons, list):
                    pokemons = []
                seen = set()
                clean = []
                for x in pokemons:
                    s = str(x).strip()
                    if not s:
                        continue
                    k = s.lower()
                    if s.isdigit():
                        try:
                            k = _get_pokemon_name(s).strip().lower()
                        except Exception:
                            k = s.lower()
                    if k in seen:
                        continue
                    seen.add(k)
                    clean.append(s)
                pokemons = clean

                try:
                    name_map = get_official_pokemon_map() or {}
                except Exception:
                    name_map = {}

                sprite_imgs = []
                for pkm in pokemons:
                    try:
                        s = str(pkm).strip()

                        # se vier número (ex: "224") ou EXT:forma, usa o resolvedor por PID
                        if s.isdigit() or s.startswith("EXT:"):
                            url = pokemon_pid_to_image(s, mode="sprite", shiny=False)
                        else:
                            url = get_pokemon_image_url(s, name_map, mode="sprite", shiny=False)

                    except Exception:
                        url = ""
                    if url:
                        sprite_imgs.append(url)

                # história
                historia = ""
                secs = npc.get("sections") or {}
                if isinstance(secs, dict):
                    historia = secs.get("História") or secs.get("Historia") or ""

                historia = _strip_html_if_any(historia)


                # highlight busca na história
                q2 = st.session_state.get("ds_npc_search", "").strip()
                h_html = ""
                if isinstance(historia, str) and historia.strip():
                    paragraphs = [p.strip() for p in historia.split("\n\n") if p.strip()]
                    for para in paragraphs:
                        safe = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        if q2:
                            pattern = re.compile(re.escape(q2), re.IGNORECASE)
                            safe = pattern.sub(lambda m: f"<span class='ds-mark'>{m.group(0)}</span>", safe)
                        h_html += f"<p>{safe}</p>"
                else:
                    h_html = "<p>(Sem história cadastrada)</p>"

                meta_parts = []
                if ocupacao:
                    meta_parts.append(str(ocupacao))
                if status:
                    meta_parts.append(f"STATUS: {status}")
                if idade:
                    meta_parts.append(f"IDADE: {idade}")
                meta_line = " | ".join(meta_parts) if meta_parts else ""

                portrait_html = ""
                if portrait_b64:
                    portrait_html = f"<div class='ds-portrait'><img src='data:image/{ext};base64,{portrait_b64}' /></div>"
                
                sprites_html = ""
                if sprite_imgs:
                    sprites_html = "<div class='ds-sprites'>" + "".join(
                        f"<img src='{u}' alt='sprite' style='width:72px;height:72px;image-rendering:pixelated;'/>"
                        for u in sprite_imgs
                    ) + "</div>"

                
                parts = [
                    '<div class="ds-npc-panel right">',
                    '  <div class="ds-frame">',
                    f'    <div class="ds-name">{sel}</div>',
                    f'    <div class="ds-meta">{meta_line}</div>',
                ]

                if portrait_html:
                    parts.append(portrait_html)

                if sprites_html:
                    parts.append(sprites_html)

                parts += [
                    '    <div class="ds-section-title">História</div>',
                    f'    <div class="ds-history">{h_html}</div>',
                    '  </div>',
                    '</div>',
                ]

                st.markdown("\n".join(parts), unsafe_allow_html=True)
        
        return

    
    # =====================================================================
    # Ginásios / Locais (placeholder)
    # =====================================================================
    elif st.session_state["comp_view"] == "ginasios":
        render_ds_tools_nav(st.session_state["comp_view"])
        render_compendium_ginasios()
        return

    if st.session_state["comp_view"] == "sessoes":
        render_ds_tools_nav(st.session_state["comp_view"])
        render_compendium_sessions(comp_data)
        return
    

    
     # ===================LOCAIS==================================================
    if st.session_state["comp_view"] == "locais":
        render_ds_tools_nav(st.session_state["comp_view"])

        cities: dict = (comp_data.get("cities") or {})
        
        # "regions" pode vir como dict (antigo) OU list[str] (novo)
        regions_raw = (comp_data.get("regions") or {})
        if isinstance(regions_raw, dict):
            regions_meta = regions_raw
        elif isinstance(regions_raw, list):
            # lista de nomes -> vira dict só pra compatibilizar com o resto do código
            regions_meta = {str(r).strip(): {} for r in regions_raw if str(r).strip()}
        else:
            regions_meta = {}
        
        if not isinstance(cities, dict) or not cities:
            st.error("Não encontrei cidades em Ga'Al.")
            return
    

        # ----------------------------
        # Helpers locais
        # ----------------------------

        def _norm_loc(s: str) -> str:
            if not isinstance(s, str):
                return ""
            return re.sub(r"\s+", " ", s).strip().lower()

        def _region_order(regs: list[str]) -> list[str]:
            ordered = []
            for r in (COMP_REGIOES_PRINCIPAIS or []):
                if r in regs and r not in ordered:
                    ordered.append(r)
            for r in sorted(regs):
                if r and r not in ordered:
                    ordered.append(r)
            return ordered

        def _cities_by_region() -> dict[str, list[str]]:
            by: dict[str, list[str]] = {}
            for cname, cobj in (cities or {}).items():
                if not isinstance(cobj, dict):
                    continue
                reg = (cobj.get("region") or "").strip() or "Sem região"
                by.setdefault(reg, []).append(cname)
            for r in by:
                by[r] = sorted(by[r])
            return by

        def _pick(d: dict, *keys: str) -> str:
            for k in keys:
                v = d.get(k)
                if isinstance(v, str) and v.strip():
                    return v.strip()
            return ""

        @st.cache_data(show_spinner=False)
        def _img_data_uri(path: str, max_w: int = 1280) -> str:
            try:
                import base64, io, os
                from PIL import Image
                if not path or not os.path.exists(path):
                    return ""
                img = Image.open(path).convert("RGB")
                img.thumbnail((max_w, 2000))
                buf = io.BytesIO()
                img.save(buf, format="JPEG", quality=80)
                b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
                return f"data:image/jpeg;base64,{b64}"
            except Exception:
                return ""

        def _map_path() -> str | None:
            # prioridade: GaAl_2.jpg na raiz/pastas -> resolve -> find_image
            p = None
            try:
                p = _resolve_asset_path(COMP_DEFAULT_MAP)
            except Exception:
                p = COMP_DEFAULT_MAP
            if p and os.path.exists(p):
                return p

            # fallback: tenta achar por index
            p2 = comp_find_image("GaAl_2") or comp_find_image("GaAl_2.jpg") or comp_find_image("GaAl")
            if p2 and os.path.exists(p2):
                return p2
            return None

        def _city_image_path(city_name: str) -> str | None:
            p = comp_find_image(city_name)
            if p and os.path.exists(p):
                return p
            # tenta variantes
            for v in [f"{city_name}.png", f"{city_name}.jpg", f"{city_name}.jpeg", f"{city_name}.webp"]:
                pv = comp_find_image(v)
                if pv and os.path.exists(pv):
                    return pv
            return None

        by_region = _cities_by_region()
        all_regions = _region_order(list(set(list(by_region.keys()) + list((regions_meta or {}).keys()))))

        st.session_state.setdefault("comp_loc_region", all_regions[0] if all_regions else "Sem região")
        st.session_state.setdefault("comp_loc_city", None)
        st.session_state.setdefault("comp_loc_sublocal", "__visao__")  # "__visao__" = visão geral da cidade
        st.session_state.setdefault("comp_loc_search", "")

        
            
        # CSS (escopado) — menu DS para Locais
        st.markdown(
            """
            <style>
              .ds-loc-shell{
                margin-top: 6px;
                padding: 0 24px 12px 24px;
                box-sizing: border-box;
              }
                html:has(.ds-loc-shell),
                body:has(.ds-loc-shell),
                div[data-testid="stAppViewContainer"]:has(.ds-loc-shell),
                section.main:has(.ds-loc-shell),
                div[data-testid="stMain"]:has(.ds-loc-shell){
                  overflow: hidden !important;
                }
              .ds-frame-marker{
                display: none;
              }
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left),
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-center),
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right),
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left),
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-center),
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right){
                background: rgba(0,0,0,0.55);
                border: 2px solid rgba(176,143,60,0.55);
                box-shadow: 0 0 45px rgba(0,0,0,0.9);
                border-radius: 12px;
                padding: 18px 18px 14px 18px !important;
                position: relative;
                box-sizing: border-box;
              }
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left)::after,
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-center)::after,
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right)::after,
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left)::after,
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-center)::after,
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right)::after{
                content: "";
                position: absolute;
                top: 10px; left: 10px; right: 10px; bottom: 10px;
                border: 1px solid rgba(255,255,255,0.10);
                border-radius: 10px;
                pointer-events: none;
              }
              /* Dá espaçamento entre blocos */
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left) .comp-divider,
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left) .comp-divider{
                margin: 14px 0 14px 0 !important;
              }
              /* Selectbox com cara de "encaixado" dentro do frame */
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left) div[data-testid="stSelectbox"],
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left) div[data-testid="stSelectbox"]{
                width: 100% !important;
              }
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left) div[data-testid="stSelectbox"] > div,
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left) div[data-testid="stSelectbox"] > div{
                background: rgba(0,0,0,0.25) !important;
                border: 1px solid rgba(176,143,60,0.45) !important;
                border-radius: 12px !important;
                padding: 6px 10px !important;
                box-shadow: 0 0 18px rgba(255,215,0,0.06) !important;
              }
              /* Texto do select com fonte DS */
              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left) div[data-testid="stSelectbox"] *,
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left) div[data-testid="stSelectbox"] *{
                font-family: "DarkSouls", serif !important;
                letter-spacing: 0.18em !important;
                text-transform: uppercase !important;
                color: rgba(255,255,255,0.82) !important;
              }

              div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left) .stTextInput input,
              div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left) .stTextInput input{
                background: rgba(0,0,0,0.35) !important;
                border: 1px solid rgba(176,143,60,0.35) !important;
                color: rgba(255,255,255,0.88) !important;
              }


              /* Radio vertical DS */
              .ds-loc-menu div[data-testid="stRadio"] [role="radiogroup"]{
                gap: 8px !important;
              }
              .ds-loc-menu div[data-testid="stRadio"] label{
                border: 1px solid rgba(255,255,255,0.10) !important;
                background: rgba(0,0,0,0.25) !important;
                border-radius: 10px !important;
                padding: 10px 12px !important;
                margin: 0 !important;
                transition: all 120ms ease !important;
              }
              .ds-loc-menu div[data-testid="stRadio"] label:hover{
                border-color: rgba(255,215,0,0.35) !important;
                box-shadow: 0 0 18px rgba(255,215,0,0.10) !important;
              }
              .ds-loc-menu div[data-testid="stRadio"] label p{
                font-family: "DarkSouls", serif !important;
                letter-spacing: 0.20em !important;
                text-transform: uppercase !important;
                font-size: 12px !important;
                margin: 0 !important;
                color: rgba(255,255,255,0.72) !important;
              }
              .ds-loc-menu div[data-testid="stRadio"] label[data-checked="true"]{
                border-color: rgba(255,215,0,0.55) !important;
                box-shadow: 0 0 26px rgba(255,215,0,0.14) !important;
                background: rgba(25,18,6,0.35) !important;
              }
              .ds-loc-menu div[data-testid="stRadio"] label[data-checked="true"] p{
                color: #FFD700 !important;
                text-shadow: 0 0 12px rgba(255,215,0,0.55) !important;
              }
              .ds-loc-menu div[data-testid="stRadio"] input{ display:none !important; }

            .ds-loc-hint{
              font-family: "DarkSouls", serif;
              letter-spacing: 0.18em;
              text-transform: uppercase;
              font-size: 11px;
              color: rgba(255,255,255,0.55);
              }
            /* remove margens estranhas dos blocos de widget dentro do frame */
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-center) img.comp-map,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-center) img.comp-map{
              width: 100% !important;
              height: auto !important;
              display: block !important;
              border-radius: 12px !important;
            }
            
            /* espaçamento controlado entre widgets */
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-center) .comp-divider,
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right) .comp-divider,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-center) .comp-divider,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right) .comp-divider{
              margin: 14px 0 12px 0 !important;
            }
            /* frame da lore */
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right),
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right){
              display: flex;
              flex-direction: column;
              height: 78vh;
              min-height: 0;
              overflow: hidden;
            }
            /* ESQUERDA com altura fixa + scroll interno */
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left),
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left){
              height: 78vh;
              overflow-y: auto;
              overflow-x: hidden;
              padding-right: 8px;
              scrollbar-width: none;
              -ms-overflow-style: none;
            }
            
            /* scrollbar discreto (esquerda) */
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left)::-webkit-scrollbar,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left)::-webkit-scrollbar{ width: 0; height: 0; }
            
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left)::-webkit-scrollbar-thumb,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left)::-webkit-scrollbar-thumb{
              background: transparent;
            }
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-left)::-webkit-scrollbar-track,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-left)::-webkit-scrollbar-track{
              background: transparent;
            }
            
             /* 🔧 stMarkdownContainer que contém a lore precisa virar o item flex “principal” */
             div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right) div[data-testid="stMarkdownContainer"]:has(.ds-loc-right-content),
             div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right) div[data-testid="stMarkdownContainer"]:has(.ds-loc-right-content){
               flex: 1 1 auto;
               min-height: 0;
               overflow: hidden;
             }

             /* wrapper interno do markdown precisa ter altura */
             div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right) div[data-testid="stMarkdownContainer"]:has(.ds-loc-right-content) > div,
             div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right) div[data-testid="stMarkdownContainer"]:has(.ds-loc-right-content) > div{
               height: 100%;
               min-height: 0;
             }

             /* conteúdo da direita vira uma coluna com altura real */
             div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right) .ds-loc-right-content,
             div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right) .ds-loc-right-content{
               display: flex;
               flex-direction: column;
               min-height: 0;
             }
            
            /* área rolável da lore */
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right) .ds-lore-scroll,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right) .ds-lore-scroll{
              flex: 1 1 auto;
              min-height: 0;
              overflow-y: auto;
              padding-right: 8px;
            
              /* impede o scroll de “vazar” pro scroll principal */
              overscroll-behavior: contain;
            
              /* opcional: mantém espaço do scrollbar estável */
              scrollbar-gutter: stable;
            }

            
            /* scrollbar discreto */
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right) .ds-lore-scroll::-webkit-scrollbar,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right) .ds-lore-scroll::-webkit-scrollbar{ width: 8px; }
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right) .ds-lore-scroll::-webkit-scrollbar-thumb,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right) .ds-lore-scroll::-webkit-scrollbar-thumb{
              background: rgba(255,215,0,0.18);
              border-radius: 10px;
            }
            /* Fallback robusto (independente dos wrappers do Streamlit) */
            .ds-loc-right-content{
              height: 78vh;
              min-height: 0;
              display: flex;
              flex-direction: column;
            }
            
            .ds-loc-right-content .ds-lore-scroll{
              /* não depende do flex/wrapper do Streamlit */
              max-height: calc(78vh - 210px);
              overflow-y: auto !important;
            
              /* evita "vazar" pro scroll da página */
              overscroll-behavior: contain;
            
              padding-right: 8px;
            }
            div[data-testid="column"]:has(.ds-frame-marker.ds-loc-right) .ds-lore-scroll::-webkit-scrollbar-track,
            div[data-testid="stColumn"]:has(.ds-frame-marker.ds-loc-right) .ds-lore-scroll::-webkit-scrollbar-track{
              background: rgba(255,255,255,0.06);
              
            }

            </style>
            """,
            unsafe_allow_html=True,
        )

        # ---------------------------
        # Layout (3 colunas): esquerda (menus) | centro (imagem) | direita (lore)
        # ----------------------------
        st.markdown("<div class='ds-loc-shell'>", unsafe_allow_html=True)
        col_left, col_center, col_right = st.columns([1.05, 1.70, 1.70], gap="large")

        # ============================
        # ESQUERDA: Região -> Cidade -> Sublocal (sem busca)
        # ============================
        with col_left:
            st.markdown("<div class='ds-frame-marker ds-loc-left'></div>", unsafe_allow_html=True)
            # --- Região ---
            st.markdown("<div class='ds-meta'>REGIÃO</div>", unsafe_allow_html=True)

            prev_reg = st.session_state.get("comp_loc_region")
            region = st.selectbox(
                "Região",
                options=all_regions,
                index=all_regions.index(prev_reg) if prev_reg in all_regions else 0,
                key="comp_loc_region_sel",
                label_visibility="collapsed",
            )

            if region != st.session_state.get("comp_loc_region"):
                st.session_state["comp_loc_region"] = region
                st.session_state["comp_loc_city"] = None
                st.session_state["comp_loc_sublocal"] = "__visao__"
                st.rerun()

            # cidades disponíveis na região
            region_cities = by_region.get(region, []) or []
            if not region_cities:
                region_cities = sorted(list(cities.keys()))

            st.markdown("<div class='comp-divider'></div>", unsafe_allow_html=True)

            # --- Cidade ---
            st.markdown("<div class='ds-meta'>CIDADES</div>", unsafe_allow_html=True)
            if not region_cities:
                st.markdown("<div class='ds-loc-hint'>Sem cidades nesta região.</div>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)  # fecha ds-loc-shell
                return

            cur_city = st.session_state.get("comp_loc_city")
            if cur_city not in region_cities:
                cur_city = region_cities[0]
                st.session_state["comp_loc_city"] = cur_city

            city = st.selectbox(
                "Cidades",
                options=region_cities,
                index=region_cities.index(cur_city),
                key="comp_loc_city_sel",
                label_visibility="collapsed",
            )

            if city != st.session_state.get("comp_loc_city"):
                st.session_state["comp_loc_city"] = city
                st.session_state["comp_loc_sublocal"] = "__visao__"
                st.rerun()

            st.markdown("<div class='comp-divider'></div>", unsafe_allow_html=True)

            # --- Sublocais ---
            st.markdown("<div class='ds-meta'>SUBLOCAIS</div>", unsafe_allow_html=True)

            cobj = cities.get(city) or {}
            sublocais = (cobj.get("sublocais") or [])
            sub_names = ["__visao__"] + [
                (sl.get("name") or "").strip() for sl in sublocais
                if isinstance(sl, dict) and (sl.get("name") or "").strip()
            ]

            # remove duplicados
            seen = set()
            sub_names = [x for x in sub_names if (x not in seen and not seen.add(x))]

            def _fmt_sublocal(v: str) -> str:
                return "Visão geral" if v == "__visao__" else v

            if len(sub_names) <= 1:
                st.markdown("<div class='ds-loc-hint'>Sem sublocais.</div>", unsafe_allow_html=True)
                st.session_state["comp_loc_sublocal"] = "__visao__"
            else:
                cur_sl = st.session_state.get("comp_loc_sublocal", "__visao__")
                if cur_sl not in sub_names:
                    cur_sl = "__visao__"
                    st.session_state["comp_loc_sublocal"] = "__visao__"
                sl_choice = st.selectbox(
                    "Sublocal",
                    options=sub_names,
                    index=sub_names.index(cur_sl),
                    key="comp_loc_sublocal_sel",
                    label_visibility="collapsed",
                    format_func=_fmt_sublocal,
                )

                if sl_choice != st.session_state.get("comp_loc_sublocal"):
                    st.session_state["comp_loc_sublocal"] = sl_choice
                    st.rerun()
    

        # ============================
       
        # CENTRO: imagem da CIDADE (sempre)
      
        with col_center:
            city_now = st.session_state.get("comp_loc_city") or (region_cities[0] if region_cities else "")
            st.markdown("<div class='ds-frame-marker ds-loc-center'></div>", unsafe_allow_html=True)

            # sempre imagem da cidade, mesmo em sublocal
            city_img = _city_image_path(city_now)
            img_uri = _img_data_uri(city_img, max_w=1600) if city_img else ""

            # título compacto
            reg_now = st.session_state.get("comp_loc_region") or ""
            st.markdown(
                f"<div class='ds-meta' style='text-align:center'>{reg_now} — {city_now}</div>",
                unsafe_allow_html=True
            )

            if img_uri:
                st.markdown(f"<img class='comp-map' src='{img_uri}'/>", unsafe_allow_html=True)
            else:
                st.markdown(
                    "<div class='ds-loc-hint' style='text-align:center'>Sem imagem para esta cidade.</div>",
                    unsafe_allow_html=True
                )
        


        # ============================
        # DIREITA: Lore (cidade ou sublocal)
        # ============================
        with col_right:
            reg_now = st.session_state.get("comp_loc_region") or ""
            city_now = st.session_state.get("comp_loc_city") or ""
            sl_now = st.session_state.get("comp_loc_sublocal") or "__visao__"

            st.markdown("<div class='ds-frame-marker ds-loc-right'></div>", unsafe_allow_html=True)
            crumb = f"{reg_now} — {city_now}"
            if sl_now and sl_now != "__visao__":
                crumb += f" — {sl_now}"

            active_sid = st.session_state.get("comp_session_active_id")
            if active_sid and city_now:
                place_id = sl_now if sl_now != "__visao__" else city_now
                label = f"{place_id}" if sl_now == "__visao__" else f"{sl_now} ({city_now})"
                if st.button("➕ Adicionar à sessão atual", key=f"add_loc_session_{_stem_key(place_id)}"):
                    if add_entity_to_active_session(
                        "places",
                        place_id,
                        "travel",
                        f"Local visitado: {label}",
                        {"place": place_id},
                    ):
                        st.success(f"{label} adicionado à sessão {active_sid}.")

            cobj = cities.get(city_now) or {}
            secs = (cobj.get("sections") or {}) if isinstance(cobj.get("sections"), dict) else {}

            # Tags
            try:
                inferred = infer_city_tags(cobj)
                tags = _merge_tags("cities", city_now, inferred)
            except Exception:
                tags = []

            chips_html = ""
            if tags:
                chips = " ".join([f"<span class='comp-chip'>{t}</span>" for t in tags])
                chips_html = f"<div style='text-align:center'>{chips}</div><div class='comp-divider'></div>"


            # Conteúdo rolável (abre)
            # Conteúdo rolável (render em 1 markdown só)
            lore_html_parts = []

            def _push_loc_section(title_txt: str, body_txt: str):
                if not isinstance(body_txt, str) or not body_txt.strip():
                    return
                ttl = html.escape(str(title_txt))
                body = html.escape(body_txt.strip()).replace("\n", "<br>")
                lore_html_parts.append(f"<div class='ds-subtitle'>{ttl}</div>")
                lore_html_parts.append(f"<div class='ds-history'>{body}</div>")
                lore_html_parts.append("<div class='comp-divider'></div>")

            # Lore
            if sl_now == "__visao__":
                prefer = [
                    "Visão geral",
                    "Como é viver em " + city_now,
                    "Como é viver em " + _clean_title(city_now),
                    "Treinadores, estrangeiros e controle",
                    "Ginásio de " + city_now,
                    "Sublocais e pontos de interesse",
                ]
                used = set()

                for k in prefer:
                    if k in secs and k not in used and isinstance(secs.get(k), str) and (secs.get(k) or "").strip():
                        used.add(k)
                        _push_loc_section(k, secs.get(k) or "")

                for k, v in secs.items():
                    if k in used:
                        continue
                    if not isinstance(v, str) or not v.strip():
                        continue
                    _push_loc_section(k, v)

            else:
                sublocs = cobj.get("sublocais") or []
                if isinstance(sublocs, dict):
                    sl_obj = sublocs.get(sl_now) or {}
                elif isinstance(sublocs, list):
                    sl_obj = next(
                        (x for x in sublocs
                         if isinstance(x, dict) and _norm(x.get("name","")) == _norm(sl_now)),
                        {}
                    )
                else:
                    sl_obj = {}
            
                txt = (sl_obj.get("text") or "").strip()
                if txt:
                    _push_loc_section(sl_now, txt)
            
            lore_html = "".join(lore_html_parts) or "<div class='ds-history'>(Sem lore cadastrada)</div>"


            right_html = f"""
            <div class='ds-loc-right-content'>
              <div class='ds-name'>LOCAIS</div>
              <div class='ds-meta'>{crumb}</div>
              <div class='comp-divider'></div>
              {chips_html}
              <div class='ds-lore-scroll'>{lore_html}</div>
            </div>
            """
            st.markdown(right_html, unsafe_allow_html=True)



        st.markdown("</div>", unsafe_allow_html=True)  # ds-loc-shell


    # =========================================================
    # VIEW: HOME (estilo do app 35 — sem clicker e sem ENTER)
    # =========================================================
    if st.session_state["comp_view"] == "home":
        logo_src = comp_img_data_uri("logo.png")

        html_home = (
            "<div class='ds-home'>"
            f"<img class='ds-logo' src='{logo_src}' alt=\"Ga'Al\" />"
            "<div class='ds-press ds-blink'>Press Any Button</div>"
            "</div>"
        )
    
        st.markdown(html_home, unsafe_allow_html=True)

        tab_key = st.radio(
            "Compendium Tabs",
            ["__home__", "npcs", "ginasios", "locais", "sessoes", "sair"],  # <-- placeholder
            index=0,
            horizontal=True,
            label_visibility="collapsed",
            key="ds_home_tabs",
            format_func=lambda v: {
                "__home__": "",      # não mostra texto
                "npcs": "NPCs",
                "ginasios": "Ginásios",
                "locais": "Locais",
                "sessoes": "Sessões",
                "sair": "Sair",
            }[v],
        )

        # esconde visualmente o primeiro item (placeholder)
        st.markdown(
            """
            <style>
              /* some o 1º item do radio (placeholder) */
              div[data-testid="stRadio"] label:first-child { display:none !important; }
            </style>
            """,
            unsafe_allow_html=True,
        )
        
        # Evita rerun em loop na primeira renderização
        if "ds_home_tabs_prev" not in st.session_state:
            st.session_state["ds_home_tabs_prev"] = tab_key
            return
        
        if st.session_state["ds_home_tabs_prev"] != tab_key:
            st.session_state["ds_home_tabs_prev"] = tab_key
        
            if tab_key == "__home__":
                return
        
            if tab_key == "sair":
                st.session_state["nav_to"] = "Pokédex (Busca)"
            else:
                st.session_state["comp_view"] = tab_key  # <-- já vem sem acento
        
            st.rerun()
        
        return
           
    
    def _tentar_achar_imagem_compendium(nome):
        if not nome:
            return None
        for tentativa in [nome, nome.replace(" ", "_"), nome.replace(" ", ""), nome.lower()]:
            for ext in [".png", ".jpg", ".jpeg"]:
                if os.path.exists(tentativa + ext):
                    return tentativa + ext
                if os.path.exists("assets/" + tentativa + ext):
                    return "assets/" + tentativa + ext
        return None
if page == "Pokédex (Busca)":
    # WRAPPER ÚNICO da página (pra escopar blindagens sem afetar Compendium)
    st.markdown('<div class="pokedex-root">', unsafe_allow_html=True)

    # =========================
    # CSS PRINCIPAL (mantido)
    # =========================
    st.markdown("""
    <style>
    /* ============================================================
       1. ESTILO GLOBAL E GERAL (MANTIDO)
       ============================================================ */
    [data-testid="stAppViewContainer"] {
        animation: pageFade 0.35s ease-in;
    }
    [data-testid="stMainBlockContainer"] {
        animation: contentSlide 0.35s ease-in;
    }
    h1, h2, h3 {
        color: #f8fafc;
        text-shadow: 0 1px 2px rgba(15, 23, 42, 0.6);
    }
    .pokedex-shell {
        border-radius: 18px;
        padding: 18px 18px 8px 18px;
        border: 2px solid rgba(148, 163, 184, 0.55);
        box-shadow: inset 0 0 18px rgba(15, 23, 42, 0.55);
        background: rgba(15, 23, 42, 0.65);
    }

    /* Fontes P2 */
    .pokedex-shell, .pokedex-card, .pokedex-info-value, .pokedex-info-title, .pokedex-header {
        font-family: "Press Start 2P", cursive !important;
    }

    .pokedex-header {
        display: flex;
        justify-content: space-between;
        gap: 12px;
        background: rgba(226, 232, 240, 0.9);
        padding: 6px 16px;
        border-radius: 16px;
        font-size: 12px;
        color: #0f172a;
    }
    .pokedex-grid-note {
        font-size: 12px;
        color: #e2e8f0;
        text-align: center;
        margin: 6px 0 10px 0;
    }
    .pokedex-card {
        background: rgba(15, 23, 42, 0.82);
        color: #f8fafc;
        padding: 18px;
        border-radius: 16px;
        border: 2px solid rgba(255,255,255,0.35);
        margin-top: 18px;
    }
    .pokedex-detail-grid { display: grid; gap: 12px; }

    .pokedex-info-card {
        background: rgba(15, 23, 42, 0.9) !important;
        border: 1px solid rgba(56, 189, 248, 0.4) !important;
        color: #f8fafc !important;
        padding: 10px;
        border-radius: 8px;
    }
    .pokedex-info-card--dark { background: #e2e8f0; }

    .pokedex-info-title {
        font-size: 12px;
        color: #38bdf8 !important;
        margin-bottom: 4px;
    }
    .pokedex-info-value {
        font-size: 14px;
        color: #0f172a;
        line-height: 1.6;
    }
    .pokedex-info-card--wide { padding: 12px 14px; }
    .pokedex-info-card--wide .pokedex-info-value { font-size: 12px; }
    .pokedex-info-card--wide .section-title { margin-top: 0; }

    .pokedex-tags span {
        display: inline-block;
        padding: 2px 8px;
        border-radius: 999px;
        font-size: 10px;
        margin-right: 6px;
        margin-bottom: 4px;
        background: rgba(0,0,0,0.35);
        color: #ffffff;
    }
    .pokedex-carousel {
        display: flex; gap: 12px; overflow-x: auto; padding: 10px 4px;
    }
    .pokedex-carousel img {
        width: 72px; height: 72px; image-rendering: pixelated;
        background: rgba(255,255,255,0.25); border-radius: 10px; padding: 6px;
    }

    /* CARROSSEL INFERIOR */
    .pokedex-footer-carousel {
        display: flex; flex-wrap: nowrap; overflow-x: auto; gap: 12px; padding: 14px;
        background: rgba(0, 0, 0, 0.30); border-radius: 15px;
        border: 1px solid rgba(255,255,255,0.18); scroll-behavior: smooth;
    }
    .pokedex-footer-carousel::-webkit-scrollbar { height: 8px; }
    .pokedex-footer-carousel::-webkit-scrollbar-thumb { background: #FFCC00; border-radius: 10px; }

    .carousel-item {
        flex: 0 0 auto; width: 70px; height: 70px; border-radius: 12px;
        display: grid; place-items: center; cursor: pointer;
        background: rgba(255,255,255,0.08); border: 1px solid rgba(255,255,255,0.18);
        transition: transform 0.15s;
    }
    .carousel-item:hover { transform: scale(1.12); }
    .carousel-item img { width: 54px; height: 54px; image-rendering: pixelated; }
    .carousel-item-active { border: 2px solid #FFCC00; background: rgba(255, 204, 0, 0.10); }

    .info-label { color: #ffd166; font-weight: 800; }
    .section-title { color: #80ed99; font-weight: 900; margin-top: 10px; }
    .hi-red { color: #ff5c5c; font-weight: 900; }
    .hi-cyan { color: #4dd6ff; font-weight: 900; }
    .hi-purple { color: #b197ff; font-weight: 900; }
    .power-badge {
        display: block; width: fit-content; margin: 10px auto 0 auto;
        padding: 6px 12px; border-radius: 999px;
        background: rgba(255,255,255,0.10); border: 1px solid rgba(255,255,255,0.25);
        color: #ffd166; font-weight: 900; text-align: center;
    }
    @keyframes pageFade { from { opacity: 0.92; } to { opacity: 1; } }
    @keyframes contentSlide { from { transform: translateY(8px); opacity: 0.92; } to { transform: translateY(0); opacity: 1; } }

    /* ============================================================
       2. O NOVO CSS DOS TILES (CORRIGIDO PARA BATER COM O PYTHON)
       ============================================================ */
    .dex-card-frame {
        display: flex;
        justify-content: center;
        align-items: center;
        height: 110px;
        width: 100%;
        border-radius: 12px;
        margin-bottom: 8px;
        transition: transform 0.2s ease;
        background: rgba(15, 23, 42, 0.6);
        position: relative;
    }
    .dex-card-frame:hover { transform: scale(1.02); }

    .dex-sprite-img {
        max-width: 80px;
        max-height: 80px;
        width: auto;
        height: auto;
        object-fit: contain;
        image-rendering: pixelated;
        filter: drop-shadow(0 4px 6px rgba(0,0,0,0.5));
        z-index: 1;
    }

    .dex-frame--caught {
        border: 2px solid #00ff41;
        box-shadow: 0 0 12px rgba(0, 255, 65, 0.25), inset 0 0 15px rgba(0, 60, 20, 0.6);
        background: rgba(0, 60, 20, 0.4);
    }
    .dex-frame--seen {
        border: 2px solid #00d0ff;
        box-shadow: 0 0 12px rgba(0, 208, 255, 0.25), inset 0 0 15px rgba(0, 40, 60, 0.6);
        background: rgba(0, 40, 60, 0.4);
    }
    .dex-frame--wish {
        border: 2px solid #ffd700;
        box-shadow: 0 0 12px rgba(255, 215, 0, 0.25), inset 0 0 15px rgba(60, 50, 0, 0.6);
        background: rgba(60, 50, 0, 0.4);
    }
    .dex-frame--default {
        border: 2px solid rgba(255, 255, 255, 0.15);
        background: rgba(255, 255, 255, 0.03);
    }
    .dex-card-link {
        display: block;
        text-decoration: none;
        color: inherit;
        cursor: pointer;
    }

    /* ============================================================
       3. CARDS "TCG" DA POKEDEX (NOVO)
       ============================================================ */
    .dex-tcg-card{
        border-radius: 14px;
        overflow: hidden;
        position: relative;
        height: 170px;
        width: 100%;
        cursor: pointer;
        box-shadow: 0 10px 22px rgba(0,0,0,0.35);
        border: 4px solid rgba(255,255,255,0.12);
        transition: transform .15s ease, filter .15s ease;
    }
    .dex-tcg-card:hover{
        transform: translateY(-2px) scale(1.01);
        filter: brightness(1.03);
    }

    .dex-tcg-header{
        min-height: 30px;
        display: flex;
        align-items: flex-start;
        gap: 8px;
        padding: 4px 8px 2px;
        background: rgba(15,23,42,0.62);
        border-bottom: 1px solid rgba(255,255,255,0.12);
    }
    .dex-tcg-statusicon{
        width: 22px; height: 22px;
        display:flex; align-items:center; justify-content:center;
        opacity: .95;
        filter: drop-shadow(0 2px 3px rgba(0,0,0,.45));
    }
    
        .dex-tcg-left{
          display: flex;
          align-items: flex-start;
          gap: 6px;
          flex-shrink: 0;
        }
        
        .dex-tcg-num{
          font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", monospace;
          font-size: 10px;
          font-weight: 800;
          letter-spacing: .6px;
        
          padding: 1px 6px;
          border-radius: 999px;
          background: rgba(15,23,42,0.55);
          border: 1px solid rgba(255,255,255,0.18);
          color: rgba(226,232,240,0.95);
          text-shadow: 0 1px 2px rgba(0,0,0,.55);
        }
        .dex-tcg-name{
          font-family: 'Trebuchet MS', 'Segoe UI', system-ui, sans-serif;
          font-size: 11px;
          font-weight: 900;
          letter-spacing: .4px;
          color: #f8fafc;
          text-shadow: 0 2px 2px rgba(0,0,0,.55);
          text-transform: uppercase;
          line-height: 1.1;
          white-space: nowrap;
          overflow: hidden;
          text-overflow: ellipsis;
          text-align: left;
          margin-top: 2px;
          flex: 1;
        }
        
        .dex-tcg-id{
            font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", monospace;
            font-size: 9px;
            font-weight: 600;
            opacity: .9;
        }

    }
    .pokedex-grid{
      display: grid;
      grid-template-columns: repeat(6, minmax(0, 1fr));
      gap: 14px;
      overflow-x: auto;
      padding-bottom: 8px;
      -webkit-overflow-scrolling: touch;
    }
    .dex-tcg-card{ min-width: 165px; }

    .dex-tcg-np{
        font-size: 8px;
        color: #e2e8f0;
        background: rgba(15,23,42,0.7);
        border: 1px solid rgba(255,255,255,0.14);
        border-radius: 999px;
        padding: 2px 7px;
        text-shadow: 0 1px 2px rgba(0,0,0,.55);
        margin-top: 1px;
        margin-left: auto;
        flex-shrink: 0;
    }

    .dex-tcg-body{
        height: 110px;
        display:flex;
        align-items:center;
        justify-content:center;
        background: rgba(2,6,23,0.18);
    }
    .dex-tcg-sprite{
        width: 86px;
        height: 86px;
        object-fit: contain;
        image-rendering: pixelated;
        filter: drop-shadow(0 7px 10px rgba(0,0,0,0.55));
    }

    .dex-tcg-footer{
        height: 30px;
        display:flex;
        align-items:center;
        justify-content:center;
        background: rgba(15,23,42,0.62);
        border-top: 1px solid rgba(255,255,255,0.12);
    }
    .dex-tcg-viab{
        font-size: 14px;
        letter-spacing: 2px;
        color: #ffd166;
        text-shadow: 0 2px 2px rgba(0,0,0,.55);
    }

    /* ===== Blindagem: só dentro da Pokédex ===== */
    .pokedex-root div[data-testid="stCustomComponentV1"],
    .pokedex-root div[data-testid="stCustomComponent"],
    .pokedex-root div[data-testid="stComponentFrame"]{
      background: transparent !important;
      border: 0 !important;
      box-shadow: none !important;
      padding: 0 !important;
    }
    .pokedex-root div[data-testid="stCustomComponentV1"] > iframe,
    .pokedex-root div[data-testid="stCustomComponent"] > iframe,
    .pokedex-root div[data-testid="stComponentFrame"] > iframe{
      background: transparent !important;
      border: 0 !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # =========================
    # Captura query param (?dex=XXX) uma vez (mantido)
    # =========================
    dex_param = st.query_params.get("dex", None)
    if dex_param:
        st.session_state["pokedex_selected"] = str(dex_param)
        st.query_params.clear()
        st.rerun()

    # =========================
    # Sidebar filtros (mantido)
    # =========================
    st.sidebar.header("🔍 Filtros")
    search_query = st.sidebar.text_input("Buscar (Nome ou Nº)", "")

    # --- FIX REAL: remove a “moldura preta” do click_detector sem afetar o iframe do Compendium ---
    st.markdown("""
    <style>
    .pokedex-clickwrap,
    .pokedex-clickwrap *{
      background: transparent !important;
    }
    .pokedex-clickwrap div[data-testid="stComponentFrame"],
    .pokedex-clickwrap div[data-testid="stComponentFrame"] > div,
    .pokedex-clickwrap div[data-testid="stElementContainer"],
    .pokedex-clickwrap div[data-testid="stElementContainer"] > div{
      background: transparent !important;
      border: none !important;
      box-shadow: none !important;
      outline: none !important;
      padding: 0 !important;
      margin: 0 !important;
    }
    .pokedex-clickwrap iframe[title^="st_click_detector"],
    .pokedex-clickwrap iframe[title*="click_detector"]{
      background: transparent !important;
      border: none !important;
      box-shadow: none !important;
      outline: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # 1) FILTRO DE REGIÃO
    all_regions = sorted(list(set([r.strip() for region in df["Região"].unique() for r in str(region).split("/")])) )
    selected_regions = st.sidebar.multiselect("Região", all_regions)

    # 2) FILTRO DE BIOMA (CASCATA)
    if selected_regions:
        df_for_biomes = df[df["Região"].apply(lambda x: any(reg in str(x) for reg in selected_regions))]
        raw_biomes = df_for_biomes["Biomas"].unique()
    else:
        raw_biomes = df["Biomas"].unique()

    all_biomes = sorted(list(set([b.strip() for biome in raw_biomes for b in str(biome).split("/")])) )
    biomes_clean = [b for b in all_biomes if "toda" not in b.lower() and "ga" not in b.lower()]
    selected_biomes = st.sidebar.multiselect("Bioma", biomes_clean)

    # 3) FILTRO DE TIPO (COMBINAÇÃO)
    all_types = sorted(list(set([t.strip() for t_str in df["Tipo"].unique() for t in str(t_str).split("/")])) )
    selected_types = st.sidebar.multiselect("Tipo Elementar (Combinação)", all_types)

    # 4) NÍVEL DE PODER
    min_p, max_p = int(df["Nivel_Poder"].min()), int(df["Nivel_Poder"].max())
    power_range = st.sidebar.slider("⚡ Nível de Poder", min_p, max_p, (min_p, max_p))

    # 5) ESTRATÉGIA
    st.sidebar.subheader("⚔️ Estratégia")
    sel_func = st.sidebar.selectbox("Função", ["Todos", "C - Controlador", "F - Finalizador", "S - Suporte"])
    sel_style = st.sidebar.selectbox("Estilo", ["Todos", "O - Ofensivo", "D - Defensivo", "F - Furtivo", "I - Incompleto", "C - Completo"])
    sel_speed = st.sidebar.selectbox("Velocidade", ["Todos", "R - Rápido", "L - Lento"])

    l1 = sel_func[0] if sel_func != "Todos" else ""
    l2 = sel_style[0] if sel_style != "Todos" else ""
    l3 = sel_speed[0] if sel_speed != "Todos" else ""

    # -----------------------------
    # APLICAÇÃO DOS FILTROS (mantido)
    # -----------------------------
    filtered_df = df.copy()

    if search_query:
        filtered_df = filtered_df[
            filtered_df["Nome"].str.contains(search_query, case=False, na=False)
            | filtered_df["Nº"].astype(str).str.contains(search_query, case=False, na=False)
        ]

    if selected_regions:
        filtered_df = filtered_df[filtered_df["Região"].apply(lambda x: any(region in str(x) for region in selected_regions))]

    if selected_biomes:
        filtered_df = filtered_df[
            filtered_df["Biomas"].apply(
                lambda x: ("toda" in str(x).lower() and "ga" in str(x).lower())
                or any(b in str(x) for b in selected_biomes)
            )
        ]

    if selected_types:
        filtered_df = filtered_df[filtered_df["Tipo"].apply(lambda x: all(t in str(x) for t in selected_types))]

    filtered_df = filtered_df[
        (filtered_df["Nivel_Poder"] >= power_range[0]) & (filtered_df["Nivel_Poder"] <= power_range[1])
    ]

    if l1 or l2 or l3:
        def _match_codes(codes):
            if not isinstance(codes, list):
                codes = [c.strip() for c in str(codes).split(",") if c.strip()]
            for c in codes:
                if len(c) >= 3:
                    ok = (not l1 or c[0] == l1) and (not l2 or c[1] == l2) and (not l3 or c[2] == l3)
                    if ok:
                        return True
            return False
        filtered_df = filtered_df[filtered_df["Codigos_Estrategia"].apply(_match_codes)]

    # -----------------------------
    # CSS DO CARROSSEL INFERIOR (mantido)
    # -----------------------------
    st.markdown(
        """
        <style>
        .footer-carousel {
            display: flex;
            flex-wrap: nowrap;
            overflow-x: auto;
            gap: 15px;
            padding: 15px;
            background: rgba(0, 0, 0, 0.4);
            border-radius: 15px;
            border: 1px solid rgba(255,255,255,0.2);
        }
        .footer-carousel::-webkit-scrollbar { height: 8px; }
        .footer-carousel::-webkit-scrollbar-thumb { background: #FFCC00; border-radius: 10px; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # -----------------------------
    # SESSION STATE (mantido)
    # -----------------------------
    if "pokedex_selected" not in st.session_state:
        st.session_state["pokedex_selected"] = None

    def select_pokedex_entry(pid: str) -> None:
        st.session_state["pokedex_selected"] = str(pid)

    # Se veio por query param (?dex=XXX), seleciona o Pokémon na MESMA ABA (mantido)
    try:
        _dex_param = st.query_params.get("dex")
        if isinstance(_dex_param, list):
            _dex_param = _dex_param[0] if _dex_param else None
        if _dex_param:
            select_pokedex_entry(str(_dex_param))
            try:
                del st.query_params["dex"]
            except Exception:
                st.query_params.clear()
    except Exception:
        pass

    selected_id = st.session_state.get("pokedex_selected")


    # ==============================================================================
    # VISÃO DE FOCO (selecionado)
    # ==============================================================================
    if selected_id:
        

        selected_df = df[df["Nº"].astype(str) == str(selected_id)]
        if selected_df.empty:
            st.session_state["pokedex_selected"] = None
            st.rerun()

        row = selected_df.iloc[0]
        dex_num = str(row["Nº"])
        p_name = row["Nome"]
        codes = row.get("Codigos_Estrategia", [])
        if not isinstance(codes, list):
            codes = [c.strip() for c in str(codes).split(",") if c.strip()]

        if "wishlist" not in user_data:
            user_data["wishlist"] = []

        # Botão sair
        if st.button("⬅️ Sair da Visão de Foco"):
            st.session_state["pokedex_selected"] = None
            st.rerun()

        # Helpers locais (não depende do resto do arquivo)
        def build_info_entries():
            hidden = {
                "Estágio",
                "Tipo de Evolução",
                "Nivel_Poder",
                "Nível de Poder",
            }
        
            entries = []
            for col in row.index:
                if col in {"Nome", "Nº", "Codigos_Estrategia"}:
                    continue
                if col in hidden:
                    continue
        
                value = row[col]
                if pd.isna(value):
                    continue
                value_str = str(value).strip()
                if not value_str or value_str.lower() == "nan":
                    continue
                entries.append((col, value_str))
            return entries


        def render_info_columns(entries):
            st.markdown("<div class='pokedex-detail-grid'>", unsafe_allow_html=True)
            for label, value in entries:
                # Título: Descrição
                if label == "Descrição da Pokedex":
                    value_html = str(value).replace("\n", "<br>")
                    st.markdown(
                        f"""
                        <div class='pokedex-info-card pokedex-info-card--wide'>
                            <div class='section-title'>📘 Descrição da Pokédex</div>
                            <div class='pokedex-info-value'>{value_html}</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                    continue

                # Título: Viabilidade (texto grande)
                if label == "Viabilidade":
                    viab = (
                        str(value)
                        .replace("PARCEIROS:", "<br><br><strong>👥 PARCEIROS:</strong>")
                        .replace("Explicação:", "<br><br><strong>💡 EXPLICAÇÃO:</strong>")
                        .replace("Habilidade:", "<strong>✨ Habilidade:</strong>")
                    )

                    viab = viab.replace("\n", "<br>")

                    # pinta o FIR (e outros códigos se quiser)
                    viab = viab.replace("FIR", "<span class='hi-red'>FIR</span>")

                    # destaca os códigos de estratégia no texto (mantém o que você já tinha)
                    for code in codes:
                        viab = re.sub(rf"\b{re.escape(code)}\b", f"<span class='hi-purple'>{code}</span>", viab)

                    st.markdown(
                        f"""
                        <div class='pokedex-info-card pokedex-info-card--wide pokedex-info-card--dark'>
                            <div class='section-title'>🧠 Viabilidade</div>
                            <div class='pokedex-info-value'>{viab}</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                    continue

                # Campos normais (Tipo, Raridade, Biomas, Região etc.)
                value_html = str(value).replace("\n", "<br>")
                st.markdown(
                    f"""
                    <div class='pokedex-info-card'>
                        <div class='pokedex-info-title'>{label}</div>
                        <div class='pokedex-info-value'>{value_html}</div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            st.markdown("</div>", unsafe_allow_html=True)


        def render_info_tags():
            tags_html = "".join([f"<span>{c}</span>" for c in codes])
            st.markdown(f"<div class='pokedex-tags'>{tags_html}</div>", unsafe_allow_html=True)

        def render_status_controls():
            try:
                _dex_cleanup_once(user_data, df)
            except Exception:
                pass
            # precisa existir no save
            if "seen" not in user_data:
                user_data["seen"] = []
            if "caught" not in user_data:
                user_data["caught"] = []

            is_seen = dex_num in user_data["seen"]
            is_caught = dex_num in user_data["caught"]
            is_wished = dex_num in user_data["wishlist"]

            s1, s2, s3 = st.columns(3)

            with s1:
                label = "👁️ Visto" if not is_seen else "✅ Visto"
                if st.button(label, key=f"seen_{dex_num}"):
                    if dex_num not in user_data["seen"]:
                        user_data["seen"].append(dex_num)
                    save_data_cloud(trainer_name, user_data)
                    st.rerun()

            with s2:
                label = "🔴 Capturar" if not is_caught else "✅ Capturado"
                if st.button(label, key=f"caught_{dex_num}"):
                    if dex_num in user_data["caught"]:
                        user_data["caught"].remove(dex_num)
                    else:
                        user_data["caught"].append(dex_num)
                        if dex_num not in user_data["seen"]:
                            user_data["seen"].append(dex_num)
                    save_data_cloud(trainer_name, user_data)
                    st.rerun()

            with s3:
                label = "⭐ Desejar" if not is_wished else "✅ Na Lista"
                if st.button(label, key=f"wish_{dex_num}"):
                    if dex_num in user_data["wishlist"]:
                        user_data["wishlist"].remove(dex_num)
                    else:
                        user_data["wishlist"].append(dex_num)
                    save_data_cloud(trainer_name, user_data)
                    st.rerun()

        # --- LAYOUT DO FOCO (seu estilo) ---
        st.markdown("<div class='pokedex-card'>", unsafe_allow_html=True)
        st.markdown(f"### #{dex_num} • {p_name}")

        info_entries = build_info_entries()
        midpoint = (len(info_entries) + 1) // 2
        top_left, top_center, top_right = st.columns([1.3, 1.7, 1.3])

        with top_left:
            render_info_columns(info_entries[:midpoint])

        with top_center:
            # --- LÓGICA DE FORMAS VISUAIS (Dex) ---
            # Por padrão, usa o nome do Excel
            target_visual_name = p_name 
            
            # Se for Lycanroc, exibe o seletor
            if "lycanroc" in p_name.lower().strip():
                lyc_dex_form = st.radio(
                    "Visualizar Forma:",
                    ["Midday", "Midnight", "Dusk"],
                    horizontal=True,
                    key="dex_lyc_visual_selector"
                )
                
                # Mapeia a escolha para o nome da API
                if lyc_dex_form == "Midnight":
                    target_visual_name = "lycanroc-midnight"
                elif lyc_dex_form == "Dusk":
                    target_visual_name = "lycanroc-dusk"
                else:
                    target_visual_name = "lycanroc-midday"

            # --- RENDERIZA A IMAGEM PRINCIPAL ---
            # Usa o nome específico da forma para buscar a arte oficial
            final_img_url = get_pokemon_image_url(target_visual_name, api_name_map, mode="artwork", shiny=False)
            st.image(final_img_url, use_container_width=True)
        
            # Nível de Poder (mantido)
            np = row.get("Nivel_Poder", row.get("Nível de Poder", ""))
            if str(np).strip() != "" and str(np).lower() != "nan":
                st.markdown(
                    f"<div class='power-badge'>⚡ Nível de Poder: {np}</div>",
                    unsafe_allow_html=True
                )

        with top_right:
            render_info_columns(info_entries[midpoint:])

        render_status_controls()
        render_info_tags()

        # --- CARROSSEL DE SPRITES (Atualizado para seguir a forma) ---
        st.markdown("#### 🎞️ Variações (Sprites)")
        
        # Gera URLs baseadas no nome da forma (target_visual_name) em vez do número da Dex
        sprite_normal = get_pokemon_image_url(target_visual_name, api_name_map, mode="sprite", shiny=False)
        sprite_shiny  = get_pokemon_image_url(target_visual_name, api_name_map, mode="sprite", shiny=True)
        
        sprites_html = f"""
            <div class='pokedex-carousel'>
                <div style="text-align:center; font-size:10px;">
                    <img src='{sprite_normal}' style='width:70px; image-rendering: pixelated;'><br>Normal
                </div>
                <div style="text-align:center; font-size:10px;">
                    <img src='{sprite_shiny}' style='width:70px; image-rendering: pixelated;'><br>Shiny
                </div>
            </div>
        """
        st.markdown(sprites_html, unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        st.divider()

        # --- CARROSSEL INFERIOR (navegação) ---
        from st_click_detector import click_detector
        items_html_list = []
        carousel_id_map = {}
        for idx, (_, r_car) in enumerate(filtered_df.iterrows()):
            pid = str(r_car["Nº"])
            sprite = pokemon_pid_to_image(pid, mode="sprite", shiny=False)
            is_active = "carousel-item-active" if pid == str(st.session_state.get("pokedex_selected")) else ""
            
            safe_id = f"carousel_{idx}"
            carousel_id_map[safe_id] = pid

            item_node = (
                f'<a href="javascript:void(0)" id="{safe_id}" class="dex-card-link">'
                f'  <div class="carousel-item {is_active}">'
                f'    <img src="{sprite}" alt="{pid}">'
                "  </div>"
                "</a>"
            )
            
            items_html_list.append(item_node)
        
        all_items_string = "".join(items_html_list)
        
        html_template = '''
        <style>
            .pokedex-footer-carousel {
                display: flex; flex-wrap: nowrap; overflow-x: auto; gap: 12px; padding: 14px;
                background: rgba(0, 0, 0, 0.30); border-radius: 15px;
                border: 1px solid rgba(255,255,255,0.18); scroll-behavior: smooth;
            }
            .pokedex-footer-carousel::-webkit-scrollbar { height: 8px; }
            .pokedex-footer-carousel::-webkit-scrollbar-thumb { background: #FFCC00; border-radius: 10px; }
            
            .carousel-item {
                flex: 0 0 auto; width: 70px; height: 70px; border-radius: 12px;
                display: grid; place-items: center; cursor: pointer;
                background: rgba(255,255,255,0.08); border: 1px solid rgba(255,255,255,0.18);
                transition: transform 0.15s;
            }
            .carousel-item:hover { transform: scale(1.12); border-color: #FFCC00; }
            .carousel-item-active { border: 2px solid #FFCC00; background: rgba(255, 204, 0, 0.10); }
            .carousel-item img { width: 54px; height: 54px; image-rendering: pixelated; }
        </style>
        
        <div class="pokedex-footer-carousel" id="pokedex-footer-carousel" tabindex="0">
            REPLACE_ME
        </div>
        <script>
            const carousel = document.getElementById("pokedex-footer-carousel");
            if (carousel) {
                carousel.addEventListener("wheel", (event) => {
                    event.preventDefault();
                    const delta = Math.abs(event.deltaY) > Math.abs(event.deltaX) ? event.deltaY : event.deltaX;
                    carousel.scrollLeft += delta;
                }, { passive: false });

                const centerActiveItem = () => {
                    const activeItem = carousel.querySelector(".carousel-item-active");
                    if (!activeItem) return;
                    const targetLeft =
                        activeItem.offsetLeft - (carousel.clientWidth - activeItem.offsetWidth) / 2;
                    carousel.scrollLeft = Math.max(0, targetLeft);
                };

                requestAnimationFrame(centerActiveItem);
                window.addEventListener("resize", centerActiveItem);
            }
        </script>
        '''
        
        final_carousel_html = html_template.replace("REPLACE_ME", all_items_string)
        st.markdown("<div class='pokedex-clickwrap'>", unsafe_allow_html=True)
        clicked_carousel_id = click_detector(final_carousel_html, key="pokedex_carousel")
        st.markdown("</div>", unsafe_allow_html=True)
        
        if clicked_carousel_id is not None:
            selected_pid = carousel_id_map.get(str(clicked_carousel_id))
            if selected_pid and selected_pid != st.session_state.get("pokedex_selected"):
                st.session_state["pokedex_selected"] = selected_pid
                st.rerun()                
               
        

    # ==============================================================================
    # GRID (visão geral)
    # ==============================================================================
    else:
        if filtered_df.empty:
            st.warning("Nenhum Pokémon encontrado.")
        else:
            st.title("📕 Pokédex Universal")
            st.markdown(f"**Resultados:** {len(filtered_df)}")

            # contadores (igual seu arquivo)
            obtained_count = len(user_data.get("caught", []))
            seen_count = len(user_data.get("seen", []))

            st.markdown("<div class='pokedex-shell'>", unsafe_allow_html=True)
            st.markdown(
                f"<div class='pokedex-header'><span>Pokémon Obtidos {obtained_count}</span><span>Pokémon Vistos {seen_count}</span></div>",
                unsafe_allow_html=True,
            )
            st.markdown("<div class='pokedex-grid-note'>Clique em um Pokémon para ver os detalhes.</div>", unsafe_allow_html=True)

            grid_cols = 6 # Reduzi para 6 para as bordas não ficarem espremidas
            rows = list(filtered_df.iterrows())

            
            # Renderiza o grid com o visual original + clique (click_detector roda em iframe, então precisa CSS inline)
            from st_click_detector import click_detector

            POKEDEX_IFRAME_CSS = r"""@import url('https://fonts.googleapis.com/css2?family=Press+Start+2P&display=swap');
           
            html, body {
              background: transparent !important;
              margin: 0 !important;
              padding: 0 !important;
            }
            /* o st_click_detector costuma ter um root/wrapper */
            #root, .root, .container, .main, .app {
              background: transparent !important;
            }
            /* --- BLINDAGEM DO FUNDO PRETO (INLINE) DO st_click_detector --- */
            /* Ele injeta um <div style="background:rgb(14, 17, 23)"> dentro do iframe */
            div[style*="background: rgb(14, 17, 23)"],
            div[style*="background:rgb(14, 17, 23)"],
            div[style*="background-color: rgb(14, 17, 23)"],
            div[style*="background-color:rgb(14,17,23)"],
            div[style*="background: rgb(14,17,23)"],
            div[style*="background:rgb(14,17,23)"]{
              background: transparent !important;
              background-color: transparent !important;
            }
            
            /* Segurança extra: alguns builds usam #app/#stApp dentro do iframe */
            #app, #stApp, .stApp, [data-testid="stAppViewContainer"]{
              background: transparent !important;
              background-color: transparent !important;
            }


            .pokedex-grid-note {
                    font-size: 12px;
                    color: #e2e8f0;
                    text-align: center;
                    margin: 6px 0 10px 0;
                }

                /* GRID (sempre 6 colunas; em telas pequenas vira scroll horizontal) */
                .pokedex-grid-wrap { overflow-x: auto; padding-bottom: 6px; }
                .pokedex-grid {
                    display: grid;
                    grid-template-columns: repeat(var(--cols, 6), minmax(140px, 1fr));
                    gap: 14px;
                    align-items: stretch;
                    min-width: calc(var(--cols, 6) * 150px);
                }
                .dex-tcg-card { height: 160px; } /* ajuda a caber mais na viewport */
                .dex-tcg-body { display:flex; justify-content:center; align-items:center; height: 96px; }
                .dex-tcg-sprite { width: 64px; height: 64px; image-rendering: pixelated; }


                .pokedex-card {
                    background: rgba(15, 23, 42, 0.82);
                    color: #f8fafc;
                    padding: 18px;
                    border-radius: 16px;
                    border: 2px solid rgba(255,255,255,0.35);
                    margin-top: 18px;
                }
                .pokedex-detail-grid { display: grid; gap: 12px; }
    
                .pokedex-info-card {
                    background: rgba(15, 23, 42, 0.9) !important;
                    border: 1px solid rgba(56, 189, 248, 0.4) !important;
                    color: #f8fafc !important;
                    padding: 10px;
                    border-radius: 8px;
                }
                .pokedex-info-card--dark { background: #e2e8f0; }
    
                .pokedex-info-title {
                    font-size: 12px;
                    color: #38bdf8 !important;
                    margin-bottom: 4px;
                }
                .pokedex-info-value {
                    font-size: 14px;
                    color: #0f172a;
                    line-height: 1.6;
                }
                .pokedex-info-card--wide { padding: 12px 14px; }
                .pokedex-info-card--wide .pokedex-info-value { font-size: 12px; }
                .pokedex-info-card--wide .section-title { margin-top: 0; }
    
                .pokedex-tags span {
                    display: inline-block;
                    padding: 2px 8px;
                    border-radius: 999px;
                    font-size: 10px;
                    margin-right: 6px;
                    margin-bottom: 4px;
                    background: rgba(0,0,0,0.35);
                    color: #ffffff;
                }
                .pokedex-carousel {
                    display: flex; gap: 12px; overflow-x: auto; padding: 10px 4px;
                }
                .pokedex-carousel img {
                    width: 72px; height: 72px; image-rendering: pixelated;
                    background: rgba(255,255,255,0.25); border-radius: 10px; padding: 6px;
                }

                /* CARROSSEL INFERIOR */
                .pokedex-footer-carousel {
                    display: flex; flex-wrap: nowrap; overflow-x: auto; gap: 12px; padding: 14px;
                    background: rgba(0, 0, 0, 0.30); border-radius: 15px;
                    border: 1px solid rgba(255,255,255,0.18); scroll-behavior: smooth;
                }
                .pokedex-footer-carousel::-webkit-scrollbar { height: 8px; }
                .pokedex-footer-carousel::-webkit-scrollbar-thumb { background: #FFCC00; border-radius: 10px; }

                .carousel-item {
                    flex: 0 0 auto; width: 70px; height: 70px; border-radius: 12px;
                    display: grid; place-items: center; cursor: pointer;
                    background: rgba(255,255,255,0.08); border: 1px solid rgba(255,255,255,0.18);
                    transition: transform 0.15s;
                }
                .carousel-item:hover { transform: scale(1.12); }
                .carousel-item img { width: 54px; height: 54px; image-rendering: pixelated; }
                .carousel-item-active { border: 2px solid #FFCC00; background: rgba(255, 204, 0, 0.10); }

                .info-label { color: #ffd166; font-weight: 800; }
                .section-title { color: #80ed99; font-weight: 900; margin-top: 10px; }
                .hi-red { color: #ff5c5c; font-weight: 900; }
                .hi-cyan { color: #4dd6ff; font-weight: 900; }
                .hi-purple { color: #b197ff; font-weight: 900; }
                .power-badge {
                    display: block; width: fit-content; margin: 10px auto 0 auto;
                    padding: 6px 12px; border-radius: 999px;
                    background: rgba(255,255,255,0.10); border: 1px solid rgba(255,255,255,0.25);
                    color: #ffd166; font-weight: 900; text-align: center;
                }
                @keyframes pageFade { from { opacity: 0.92; } to { opacity: 1; } }
                @keyframes contentSlide { from { transform: translateY(8px); opacity: 0.92; } to { transform: translateY(0); opacity: 1; } }


                /* ============================================================
                   2. O NOVO CSS DOS TILES (CORRIGIDO PARA BATER COM O PYTHON)
                   ============================================================ */

                /* MOLDURA DO CARD */
                .dex-card-frame {
                    display: flex;
                    justify-content: center;
                    align-items: center;
                    height: 110px;        
                    width: 100%;
                    border-radius: 12px;
                    margin-bottom: 8px;
                    transition: transform 0.2s ease;
                    background: rgba(15, 23, 42, 0.6);
                    position: relative;
                }
    
                .dex-card-frame:hover {
                    transform: scale(1.02);
                }

                /* A IMAGEM */
                .dex-sprite-img {
                    max-width: 80px;
                    max-height: 80px;
                    width: auto;
                    height: auto;
                    object-fit: contain;
                    image-rendering: pixelated;
                    filter: drop-shadow(0 4px 6px rgba(0,0,0,0.5));
                    z-index: 1;
                }

                /* --- AQUI ESTAVA O ERRO: NOMES DAS CLASSES --- */
                /* Antes estava .frame-caught, agora é .dex-frame--caught */

                /* 🟢 CAPTURADO (Verde) */
                .dex-frame--caught {
                    border: 2px solid #00ff41;
                    box-shadow: 0 0 12px rgba(0, 255, 65, 0.25), inset 0 0 15px rgba(0, 60, 20, 0.6);
                    background: rgba(0, 60, 20, 0.4);
                }

                /* 🔵 VISTO (Azul) */
                .dex-frame--seen {
                    border: 2px solid #00d0ff;
                    box-shadow: 0 0 12px rgba(0, 208, 255, 0.25), inset 0 0 15px rgba(0, 40, 60, 0.6);
                    background: rgba(0, 40, 60, 0.4);
                }

                /* ⭐ WISHLIST (Dourado) */
                .dex-frame--wish {
                    border: 2px solid #ffd700;
                    box-shadow: 0 0 12px rgba(255, 215, 0, 0.25), inset 0 0 15px rgba(60, 50, 0, 0.6);
                    background: rgba(60, 50, 0, 0.4);
                }

                /* ⚪ PADRÃO (Cinza) */
                .dex-frame--default {
                    border: 2px solid rgba(255, 255, 255, 0.15);
                    background: rgba(255, 255, 255, 0.03);
                }

                .dex-card-link {
                    display: block;
                    text-decoration: none;
                    color: inherit;
                    cursor: pointer;
                }
    

                /* ============================================================
                   3. CARDS "TCG" DA POKEDEX (NOVO)
                   ============================================================ */
                .dex-tcg-card{
                    border-radius: 14px;
                    overflow: hidden;
                    position: relative;
                    height: 170px;
                    width: 100%;
                    cursor: pointer;
                    box-shadow: 0 10px 22px rgba(0,0,0,0.35);
                    border: 4px solid rgba(255,255,255,0.12); /* será "dominada" pela classe de status */
                    transition: transform .15s ease, filter .15s ease;
                }
                .dex-tcg-card:hover{
                    transform: translateY(-2px) scale(1.01);
                    filter: brightness(1.03);
                }

                .dex-tcg-header{
                    min-height: 30px;
                    display: flex;
                    align-items: flex-start;
                    gap: 8px;
                    padding: 4px 8px 2px;
                    background: rgba(15,23,42,0.62);
                    border-bottom: 1px solid rgba(255,255,255,0.12);
                    position: relative;
                }
                .dex-tcg-statusicon{
                    width: 22px; height: 22px;
                    display:flex; align-items:center; justify-content:center;
                    opacity: .95;
                    filter: drop-shadow(0 2px 3px rgba(0,0,0,.45));
                }
                .dex-tcg-left{
                    display: flex;
                    align-items: flex-start;
                    gap: 6px;
                    flex-shrink: 0;
                }
                .dex-tcg-num{
                    font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", monospace;
                    font-size: 10px;
                    font-weight: 800;
                    letter-spacing: .6px;
                    padding: 1px 6px;
                    border-radius: 999px;
                    background: rgba(15,23,42,0.55);
                    border: 1px solid rgba(255,255,255,0.18);
                    color: rgba(226,232,240,0.95);
                    text-shadow: 0 1px 2px rgba(0,0,0,.55);
                }
                .dex-tcg-name{
                    font-family: 'Trebuchet MS', 'Segoe UI', system-ui, sans-serif;
                    font-size: 11px;
                    font-weight: 900;
                    letter-spacing: .4px;
                    color: #f8fafc;
                    text-shadow: 0 2px 2px rgba(0,0,0,.55);
                    text-transform: uppercase;
                    line-height: 1.1;
                    white-space: nowrap;
                    overflow: hidden;
                    text-overflow: ellipsis;
                    text-align: center;
                    margin-top: 2px;
                    position: absolute;
                    left: 50%;
                    transform: translateX(-50%);
                    max-width: calc(100% - 145px);
                    pointer-events: none;
                }
                .dex-tcg-np{
                    font-size: 8px;
                    color: #e2e8f0;
                    background: rgba(15,23,42,0.7);
                    border: 1px solid rgba(255,255,255,0.14);
                    border-radius: 999px;
                    padding: 2px 7px;
                    text-shadow: 0 1px 2px rgba(0,0,0,.55);
                    margin-top: 1px;
                    margin-left: auto;
                    flex-shrink: 0;
                }

                .dex-tcg-body{
                    height: 110px;
                    display:flex;
                    align-items:center;
                    justify-content:center;
                    background: rgba(2,6,23,0.18);
                }
                .dex-tcg-sprite{
                    width: 86px;
                    height: 86px;
                    object-fit: contain;
                    image-rendering: pixelated;
                    filter: drop-shadow(0 7px 10px rgba(0,0,0,0.55));
                }

                .dex-tcg-footer{
                    height: 30px;
                    display:flex;
                    align-items:center;
                    justify-content:center;
                    background: rgba(15,23,42,0.62);
                    border-top: 1px solid rgba(255,255,255,0.12);
                }
                .dex-tcg-viab{
                    font-size: 14px;
                    letter-spacing: 2px;
                    color: #ffd166;
                    text-shadow: 0 2px 2px rgba(0,0,0,.55);
                }

            """

            card_nodes = []
            grid_id_map = {}

            for _, row_g in filtered_df.iterrows():
                dex_num = str(row_g["Nº"])
                p_name = row_g["Nome"]

                sprite_url = pokemon_pid_to_image(dex_num, mode="sprite", shiny=False)

                is_caught = dex_num in user_data.get("caught", [])
                is_seen = dex_num in user_data.get("seen", [])
                is_wished = dex_num in user_data.get("wishlist", [])

                if is_caught:
                    status_class = "dex-frame--caught"
                elif is_wished:
                    status_class = "dex-frame--wish"
                elif is_seen:
                    status_class = "dex-frame--seen"
                else:
                    status_class = "dex-frame--default"

                types = _split_types(row_g.get("Tipo", ""))
                t1 = types[0] if len(types) > 0 else ""
                t2 = types[1] if len(types) > 1 else ""
                c1 = _type_color(t1)
                c2 = _type_color(t2) if t2 else ""

                if c2:
                    bg_style = f"background: linear-gradient(135deg, {c1} 0%, {c1} 50%, {c2} 50%, {c2} 100%);"
                else:
                    bg_style = f"background: {c1};"

                viab_code = _extract_viab_code_from_text(row_g.get("Viabilidade", ""))
                try:
                    np_val = int(row_g.get("Nivel_Poder", 0) or 0)
                except Exception:
                    np_val = 0

                if is_caught:
                    status_svg = _SVG_POKEBALL
                elif is_wished:
                    status_svg = _SVG_STAR
                elif is_seen:
                    status_svg = _SVG_EYE
                else:
                    status_svg = ""

                safe_id = f"dex_{dex_num}"
                safe_id_card = f"{safe_id}__card"
                safe_id_img = f"{safe_id}__img"

                grid_id_map[safe_id] = dex_num
                grid_id_map[safe_id_card] = dex_num
                grid_id_map[safe_id_img] = dex_num
                try:
                    dex_id_txt = f"#{int(dex_num):03d}"
                except Exception:
                    dex_id_txt = f"#{dex_num}"


                card_nodes.append("\n".join([
                    f'<a href="javascript:void(0)" id="{safe_id}" '
                    f'style="text-decoration:none;color:inherit;display:block;">',
                
                    f'  <div id="{safe_id_card}" class="dex-tcg-card {status_class}" '
                    f'style="{bg_style}; cursor:pointer;" role="button" tabindex="0">',
                                    
                    '    <div class="dex-tcg-header" style="pointer-events:none;">',
                    f'      <div class="dex-tcg-left">'
                    f'        <div class="dex-tcg-statusicon">{status_svg}</div>'
                    f'        <div class="dex-tcg-num">{dex_id_txt}</div>'
                    f'      </div>',
                    f'      <div class="dex-tcg-name">{p_name}</div>',
                    f'      <div class="dex-tcg-np">NP {np_val}</div>',
                    '    </div>',
                
                    '    <div class="dex-tcg-body" style="pointer-events:none;">',
                    f'      <img id="{safe_id_img}" src="{sprite_url}" class="dex-tcg-sprite" alt="{p_name}" />',
                    '    </div>',
                
                    '    <div class="dex-tcg-footer" style="pointer-events:none;">',
                    f'      <div class="dex-tcg-viab" title="{viab_code}">{viab_code}</div>',
                    '    </div>',
                
                    '  </div>',
                    '</a>',
                ]))

            grid_html = f"<style>{POKEDEX_IFRAME_CSS}</style><div class='pokedex-grid-wrap'><div class='pokedex-grid' style='--cols:{grid_cols};'>" + "".join(card_nodes) + "</div></div>"
            st.markdown("<div class='pokedex-clickwrap'>", unsafe_allow_html=True)
            clicked_id = click_detector(grid_html, key="pokedex_grid")
            st.markdown("</div>", unsafe_allow_html=True)

            
            if clicked_id is not None:
                clicked_id = str(clicked_id)
                selected_pid = grid_id_map.get(clicked_id)

                if selected_pid is None and "__" in clicked_id:
                    selected_pid = grid_id_map.get(clicked_id.split("__", 1)[0])

                if selected_pid:
                    st.session_state["pokedex_selected"] = str(selected_pid)
                    st.rerun()


    # FECHAMENTO ÚNICO do wrapper da página
    st.markdown("</div>", unsafe_allow_html=True)


# ==============================================================================
# PÁGINA 2: TRAINER HUB
# ==============================================================================
if page == "Trainer Hub (Meus Pokémons)":
    st.title("🏕️ Central do Treinador")

    # ----------------------------
    # GBA HUB CSS (visual)
    # ----------------------------

    # ----------------------------
    # Init de dados persistidos
    # ----------------------------
    user_data.setdefault("stats", {})          # agora guarda também stgr/int/notes
    user_data.setdefault("wishlist", [])
    user_data.setdefault("shinies", [])
    user_data.setdefault("favorite_moves", {}) # {pid: [move_name,...]}
    user_data.setdefault("forms", {})
    try:
        _dex_cleanup_once(user_data, df)
    except Exception:
        pass


    # estados de UI
    st.session_state.setdefault("hub_selected_pid", None)      # abre ficha
    st.session_state.setdefault("hub_context_pid", None)       # menu contexto BOX
    st.session_state.setdefault("hub_context_action", None)    # "move" ou "view"
    st.session_state.setdefault("hub_box_page", 1)             # paginação box

    # ----------------------------
    # Helpers
    # ----------------------------
    def _pid_key(pid: str) -> str:
        return str(pid).strip()

    def _get_pokemon_row(pid: str):
        pidn = _norm_pid(pid)
        # df é seu dataframe da dex
        hit = df[df["Nº"].apply(_norm_pid) == pidn]
        if len(hit) == 0:
            return None
        return hit.iloc[0].to_dict()

    def _get_pokemon_name(pid: str) -> str:
        pid_str = str(pid)
        if pid_str.startswith("EXT:"):
            return pid_str.replace("EXT:", "").strip() or "Visitante"
    
        pidn = _norm_pid(pid_str)
        row = _get_pokemon_row(pidn)
        return str(row["Nome"]) if row is not None and "Nome" in row else f"ID {pidn}"
    
    
    def _get_sprite(pid: str) -> str:
        pid_str = str(pid)
        if pid_str.startswith("EXT:"):
            return get_pokemon_image_url(_get_pokemon_name(pid_str), api_name_map, mode="sprite", shiny=False)
    
        pidn = _norm_pid(pid_str)
    
        # shinies pode estar salvo como "16" ou "16.0" dependendo da versão antiga
        shinies_norm = set(_norm_pid(x) for x in user_data.get("shinies", []))
        is_shiny = pidn in shinies_norm
    
        # forms idem
        forms_map = user_data.get("forms", {}) or {}
        saved_form = forms_map.get(pidn) or forms_map.get(pid_str)
        if saved_form:
            return pokemon_pid_to_image(f"EXT:{saved_form}", mode="sprite", shiny=is_shiny)
    
        return pokemon_pid_to_image(pidn, mode="sprite", shiny=is_shiny)
    def _get_pokemon_type(pid: str) -> str:
        if str(pid).startswith("EXT:"):
            return "Visitante"
        row = _get_pokemon_row(pid)
        return str(row["Tipo"]) if row is not None and "Tipo" in row else "—"

    def _get_pokemon_np(pid: str):
        if str(pid).startswith("EXT:"):
            return "—"
        row = _get_pokemon_row(pid)
        return row.get("Nivel_Poder") if row is not None else "—"



    def _get_artwork(pid: str) -> str:
        if str(pid).startswith("EXT:"):
            return get_pokemon_image_url(_get_pokemon_name(pid), api_name_map, mode="artwork", shiny=False)
            
        is_shiny = str(pid) in user_data.get("shinies", [])
        
        # ✅ NOVO: Verifica se tem forma salva na Artwork também
        saved_form = user_data.get("forms", {}).get(str(pid))
        if saved_form:
            return pokemon_pid_to_image(f"EXT:{saved_form}", mode="artwork", shiny=is_shiny)
            
        return pokemon_pid_to_image(str(pid), mode="artwork", shiny=is_shiny)

    def _ensure_stats_slot(pid: str) -> dict:
        pid = _pid_key(pid)
        user_data["stats"].setdefault(pid, {})
        # campos mínimos (se não existir, cria)
        for k in ["dodge", "parry", "thg", "stgr", "int", "will", "fortitude", "notes"]:
            user_data["stats"][pid].setdefault(k, 0 if k != "notes" else "")
        return user_data["stats"][pid]

    def _select_pid(pid: str):
        st.session_state["hub_selected_pid"] = _pid_key(pid)
        st.session_state["hub_context_pid"] = None
        st.session_state["hub_context_action"] = None

    def _open_box_context(pid: str):
        st.session_state["hub_context_pid"] = _pid_key(pid)
        st.session_state["hub_context_action"] = None

    def _move_box_to_party(pid: str):
        pid = _pid_key(pid)
        if pid in [str(x) for x in user_data.get("party", [])]:
            return
        # limite 8 (visitantes também contam)
        if len(user_data.get("party", [])) >= 8:
            st.warning("Sua equipe ativa já está cheia (máximo 8). Remova alguém antes.")
            return
        user_data.setdefault("party", [])
        user_data["party"].append(pid)
        save_data_cloud(trainer_name, user_data)
        st.success("Movido para a equipe!")
        st.rerun()

    def _remove_from_party(pid: str):
        pid = _pid_key(pid)
        user_data["party"] = [p for p in user_data.get("party", []) if str(p) != pid]
        save_data_cloud(trainer_name, user_data)
        st.rerun()

    def _delete_pokemon_from_hub(pid: str):
        pid = _pid_key(pid)

        user_data["caught"] = [p for p in user_data.get("caught", []) if str(p) != pid]
        user_data["party"] = [p for p in user_data.get("party", []) if str(p) != pid]
        user_data["seen"] = [p for p in user_data.get("seen", []) if str(p) != pid]
        user_data["wishlist"] = [p for p in user_data.get("wishlist", []) if str(p) != pid]
        user_data["shinies"] = [p for p in user_data.get("shinies", []) if str(p) != pid]

        if isinstance(user_data.get("stats"), dict):
            user_data["stats"].pop(pid, None)
        if isinstance(user_data.get("forms"), dict):
            user_data["forms"].pop(pid, None)
        if isinstance(user_data.get("favorite_moves"), dict):
            user_data["favorite_moves"].pop(pid, None)

        if st.session_state.get("hub_selected_pid") == pid:
            st.session_state["hub_selected_pid"] = None
        st.session_state["hub_context_pid"] = None
        st.session_state["hub_context_action"] = None

        save_data_cloud(trainer_name, user_data)
        st.success(f"{_get_pokemon_name(pid)} removido da sua conta.")
        st.rerun()

    # ----------------------------
    # Carrega fichas do Firebase (mapa pid -> sheet)
    # ----------------------------
    sheets_map = {}



    def _register_sheet(sheet_payload: dict, pid_value) -> None:
        pid_norm = _normalize_hub_pid(pid_value)
        if not pid_norm:
            return

        prev_sheet = sheets_map.get(pid_norm)
        if prev_sheet is None:
            sheets_map[pid_norm] = sheet_payload
            return

        # Se houver conflito, mantém a ficha mais recente.
        prev_updated = str(prev_sheet.get("updated_at") or "")
        curr_updated = str(sheet_payload.get("updated_at") or "")
        if curr_updated >= prev_updated:
            sheets_map[pid_norm] = sheet_payload

    try:
        db, bucket = init_firebase()
        for sh in list_sheets(db, trainer_name) or []:
            p = (sh.get("pokemon") or {})
            _register_sheet(sh, p.get("id"))
            _register_sheet(sh, sh.get("linked_pid"))
    except Exception:
        sheets_map = {}

    # ----------------------------
    # Tabs do Hub
    # ----------------------------
    t_main, t_wish, t_seen, t_trainer = st.tabs(
        ["🎮 Equipe & BOX", "🌟 Lista de Interesses", "👁️ Pokédex (Vistos)", "🧑‍🎤 Meu Treinador"]
    )

    # ==========================
    # TAB PRINCIPAL: BOX + PARTY
    # ==========================
    with t_main:

        # ==========================
        # FICHA (painel superior)
        # ==========================
        sel = st.session_state.get("hub_selected_pid")
        if sel:
            pid = str(sel)
            is_ext = pid.startswith("EXT:")
            pname = _get_pokemon_name(pid)

            st.markdown('<div class="gba-window summary">', unsafe_allow_html=True)
            st.markdown(f"""
            <div class="gba-header">
              <div class="left">
                <span class="gba-chip">SUMMARY</span>
                <span class="gba-title" style="font-size:0.75rem;">{pname}</span>
              </div>
            </div>
            """, unsafe_allow_html=True)

            cA, cB = st.columns([1, 1.6], gap="large")

            # painel esquerdo: imagem + notas + ações rápidas
            with cA:
                # --- LÓGICA DE TROCA DE FORMA/SPRITE NO HUB ---
                final_hub_image = _get_artwork(pid) # Imagem padrão (ou shiny se já estiver marcado)

                normalized_pname = normalize_text(pname).replace(" ", "-")
                form_selector_options = None

                if "lycanroc" in normalized_pname:
                    form_selector_options = [
                        ("Midday", "lycanroc-midday"),
                        ("Midnight", "lycanroc-midnight"),
                        ("Dusk", "lycanroc-dusk"),
                    ]
                elif "darmanitan" in normalized_pname:
                    form_selector_options = [
                        ("Standard", "darmanitan-standard"),
                        ("Zen", "darmanitan-zen"),
                        ("Galar Standard", "darmanitan-galar-standard"),
                        ("Galar Zen", "darmanitan-galar-zen"),
                    ]
                elif "minior" in normalized_pname:
                    form_selector_options = [
                        ("Red Meteor", "minior-red-meteor"),
                        ("Orange Meteor", "minior-orange-meteor"),
                        ("Yellow Meteor", "minior-yellow-meteor"),
                        ("Green Meteor", "minior-green-meteor"),
                        ("Blue Meteor", "minior-blue-meteor"),
                        ("Indigo Meteor", "minior-indigo-meteor"),
                        ("Violet Meteor", "minior-violet-meteor"),
                        ("Red Core", "minior-red"),
                        ("Orange Core", "minior-orange"),
                        ("Yellow Core", "minior-yellow"),
                        ("Green Core", "minior-green"),
                        ("Blue Core", "minior-blue"),
                        ("Indigo Core", "minior-indigo"),
                        ("Violet Core", "minior-violet"),
                    ]
                elif "floette" in normalized_pname:
                    form_selector_options = [
                        ("Padrão", "floette"),
                        ("Eternal", "floette-eternal"),
                    ]
                elif "ursaluna" in normalized_pname:
                    form_selector_options = [
                        ("Ursaluna", "ursaluna"),
                        ("Bloodmoon", "ursaluna-bloodmoon"),
                    ]

                if form_selector_options:
                    st.caption("Visualizar Forma:")

                    label_to_api = {label: api_name for label, api_name in form_selector_options}
                    api_to_label = {api_name: label for label, api_name in form_selector_options}
                    default_api = form_selector_options[0][1]
                    current_saved = user_data.get("forms", {}).get(pid, default_api)

                    options = [label for label, _ in form_selector_options]
                    default_idx = options.index(api_to_label.get(current_saved, api_to_label[default_api]))

                    selected_label = st.radio(
                        "Forma",
                        options,
                        index=default_idx,
                        horizontal=True,
                        label_visibility="collapsed",
                        key=f"hub_form_selector_{pid}"
                    )

                    selected_api_name = label_to_api[selected_label]
                    if selected_api_name != current_saved:
                        user_data["forms"][pid] = selected_api_name
                        save_data_cloud(trainer_name, user_data)
                        st.rerun()

                # Renderiza a imagem final (Os helpers _get_artwork agora leem user_data['forms'] automaticamente)
                final_hub_image = _get_artwork(pid) 
                st.image(final_hub_image, use_container_width=True)
                
                # --- Resto do código original da coluna esquerda ---
                stats_slot = _ensure_stats_slot(pid)

                st.markdown('<div class="gba-divider"></div>', unsafe_allow_html=True)
                st.markdown("#### Ações rápidas")
                if st.button("Fechar ficha", key="hub_close_sheet", use_container_width=True):
                    st.session_state["hub_selected_pid"] = None
                    st.rerun()

                st.markdown('<div class="gba-divider"></div>', unsafe_allow_html=True)
                st.markdown('<div class="gba-notes"><div class="gba-caption">NOTAS</div></div>', unsafe_allow_html=True)

                notes = st.text_area(
                    "Notas",
                    value=str(stats_slot.get("notes", "")),
                    height=140,
                    key=f"hub_notes_{pid}",
                    label_visibility="collapsed",
                )
                if notes != stats_slot.get("notes", ""):
                    stats_slot["notes"] = notes
                    user_data["stats"][pid] = stats_slot
                    save_data_cloud(trainer_name, user_data)

            # painel direito: stats + resumo + golpes
            with cB:
                sheet = sheets_map.get(_normalize_hub_pid(pid)) if (not is_ext) else None

                if sheet is None:
                    st.warning("Este Pokémon não tem ficha salva. Preencha os atributos mínimos para usar no Hub.")
                    stats_slot = _ensure_stats_slot(pid)

                    c1, c2, c3 = st.columns(3)
                    with c1:
                        dodge = st.number_input("Dodge", min_value=0, max_value=99, value=int(stats_slot.get("dodge", 0)), key=f"hub_dodge_{pid}")
                        parry = st.number_input("Parry", min_value=0, max_value=99, value=int(stats_slot.get("parry", 0)), key=f"hub_parry_{pid}")
                        thg = st.number_input("Thg", min_value=0, max_value=99, value=int(stats_slot.get("thg", 0)), key=f"hub_thg_{pid}")
                    with c2:
                        stgr = st.number_input("Stgr", min_value=0, max_value=99, value=int(stats_slot.get("stgr", 0)), key=f"hub_stgr_{pid}")
                        intel = st.number_input("Int", min_value=0, max_value=99, value=int(stats_slot.get("int", 0)), key=f"hub_int_{pid}")
                        will = st.number_input("Will", min_value=0, max_value=99, value=int(stats_slot.get("will", 0)), key=f"hub_will_{pid}")
                    with c3:
                        fort = st.number_input("Fortitude", min_value=0, max_value=99, value=int(stats_slot.get("fortitude", 0)), key=f"hub_fort_{pid}")

                    new_stats = {
                        "dodge": int(dodge),
                        "parry": int(parry),
                        "thg": int(thg),
                        "stgr": int(stgr),
                        "int": int(intel),
                        "will": int(will),
                        "fortitude": int(fort),
                        "notes": str(stats_slot.get("notes", "")),
                    }
                    if new_stats != {k: stats_slot.get(k) for k in new_stats.keys()}:
                        user_data["stats"][pid] = new_stats
                        save_data_cloud(trainer_name, user_data)

                    st.info("Sem ficha salva: golpes, advantages e skills não disponíveis aqui ainda.")
                else:
                    stats = sheet.get("stats") or {}
                    moves = sheet.get("moves") or []
                    advantages = sheet.get("advantages") or []
                    skills_raw = sheet.get("skills") or []

                    stats_slot = _ensure_stats_slot(pid)

                    st.markdown(
                        f"""
                        <div class="gba-chips">
                            <span class="gba-chip">Dodge {int(stats.get("dodge", 0))}</span>
                            <span class="gba-chip">Parry {int(stats.get("parry", 0))}</span>
                            <span class="gba-chip">Thg {int(stats.get("thg", 0))}</span>
                            <span class="gba-chip">Fort {int(stats.get("fortitude", 0))}</span>
                            <span class="gba-chip">Will {int(stats.get("will", 0))}</span>
                            <span class="gba-chip">Stgr {int(stats.get("stgr", 0))}</span>
                            <span class="gba-chip">Int {int(stats.get("int", 0))}</span>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )


                    def _skills_list(raw) -> list[str]:
                        items = []
                        if isinstance(raw, list):
                            for row in raw:
                                if isinstance(row, dict):
                                    name = str(row.get("name", "")).strip()
                                    ranks = int(row.get("ranks", 0) or 0)
                                    if name and ranks > 0:
                                        items.append(f"{name} ({ranks})")
                                elif isinstance(row, str):
                                    name = row.strip()
                                    if name:
                                        items.append(name)
                        elif isinstance(raw, dict):
                            for name, ranks in raw.items():
                                try:
                                    ranks = int(ranks)
                                except Exception:
                                    ranks = 0
                                if name and ranks > 0:
                                    items.append(f"{name} ({ranks})")
                        elif isinstance(raw, str):
                            for line in raw.splitlines():
                                line = line.strip().lstrip("-•").strip()
                                if line:
                                    items.append(line)
                        return items

                    skills_list = _skills_list(skills_raw)

                    adv_html = "<div class='hub-summary-block'><strong>⭐ Advantages</strong>"
                    if advantages:
                        adv_html += "<ul>" + "".join(f"<li>{a}</li>" for a in advantages) + "</ul>"
                    else:
                        adv_html += "<div class='hub-muted'>Sem advantages registradas.</div>"
                    adv_html += "</div>"

                    skills_html = "<div class='hub-summary-block'><strong>🎯 Skills</strong>"
                    if skills_list:
                        skills_html += "<ul>" + "".join(f"<li>{s}</li>" for s in skills_list) + "</ul>"
                    else:
                        skills_html += "<div class='hub-muted'>Sem skills registradas.</div>"
                    skills_html += "</div>"

                    summary_html = f"""
                    <div class="hub-summary-grid">
                        {adv_html}
                        {skills_html}
                    </div>
                    """
                    st.markdown(summary_html, unsafe_allow_html=True)

                    try:
                        mvdb = load_move_db(file_name)
                    except Exception:
                        mvdb = None

                    def _based_stat(move_name: str, move_meta: dict | None = None, move_obj: dict | None = None) -> str:
                        move_obj = move_obj or {"name": move_name, "meta": (move_meta or {})}
                        b = _move_based_stat_from_meta(move_obj)
                        if b != "Stgr" or mvdb is None:
                            return b

                        mv = mvdb.get_by_name(move_name)
                        if mv is None:
                            return b

                        cat = (mv.categoria or "").strip().lower()
                        if "status" in cat:
                            return "—"
                        if "especial" in cat or "special" in cat:
                            return "Int"
                        if "físico" in cat or "fisico" in cat or "physical" in cat:
                            return "Stgr"

                        return b


                    def _final_rank(m: dict) -> tuple[int, str]:
                        base = int(m.get("rank", 0) or 0)
                        bstat = _based_stat(m.get("name", ""), m.get("meta") or {}, m)
                    
                        if bstat == "Stgr":
                            bonus = int(stats.get("stgr", 0) or 0)
                        elif bstat == "Int":
                            bonus = int(stats.get("int", 0) or 0)
                        else:  # "—" (Status)
                            bonus = 0
                    
                        return base + bonus, bstat


                    fav = user_data.get("favorite_moves", {}).get(pid, [])
                    if not isinstance(fav, list):
                        fav = []

                    st.markdown("### ⭐ Golpes Favoritos (até 4)")
                    all_names = [m.get("name", "Golpe") for m in moves]
                    fav = [n for n in fav if n in all_names]

                    shown = 0
                    for m in moves:
                        name = m.get("name", "Golpe")
                        if name in fav and shown < 4:
                            base = int(m.get("rank", 0) or 0)
                            fr, bstat = _final_rank(m)
    
                            if bstat == "—":  # Status
                                st.write(f"**{name}** — Rank base {base} → **{fr}** (Status)")
                            else:             # Physical/Special
                                st.write(f"**{name}** — Rank base {base} + {bstat} → **{fr}**")
                            build_txt = (m.get("build") or "").strip()
                            if build_txt:
                                st.code(build_txt, language="text")
                            shown += 1
                    if shown == 0:
                        st.caption("Nenhum favorito definido.")

                    st.markdown('<div class="gba-divider"></div>', unsafe_allow_html=True)

                    st.markdown("### 📜 Lista completa de golpes")
                    for idx, m in enumerate(moves):
                        name = m.get("name", "Golpe")
                        fr, bstat = _final_rank(m)
                        checked = name in fav
                        c1, c2 = st.columns([0.15, 0.85])
                        with c1:
                            star = st.checkbox("⭐", value=checked, key=f"hub_star_{pid}_{idx}")
                        with c2:
                            base = int(m.get("rank", 0) or 0)

                            if bstat == "—":  # Golpe de Status
                                st.write(f"**{name}** — base {base} → **{fr}** (Status)")
                            else:             # Physical / Special
                                st.write(f"**{name}** — base {base} + {bstat} → **{fr}**")

                            build_txt = (m.get("build") or "").strip()
                            if build_txt:
                                with st.expander("Ingredientes do golpe"):
                                    st.code(build_txt, language="text")
                        if star and name not in fav:
                            fav.append(name)
                        if (not star) and name in fav:
                            fav.remove(name)

                    if len(fav) > 4:
                        fav = fav[:4]
                        st.warning("Favoritos limitados a 4. Mantive os 4 primeiros que você marcou.")

                    if user_data.get("favorite_moves", {}).get(pid, []) != fav:
                        user_data.setdefault("favorite_moves", {})
                        user_data["favorite_moves"][pid] = fav
                        save_data_cloud(trainer_name, user_data)

        

        
        st.markdown('<div class="hub-split">', unsafe_allow_html=True)
        
        col_left, col_right = st.columns([1, 1])

        # ---------
        # BOX (capturados)
        # ---------
        with col_left:
            # Container principal da BOX
            st.markdown('<div class="grass-box">', unsafe_allow_html=True)
        
            st.markdown("""
            <div class="gba-header">
              <div class="left">
                <span class="gba-chip">BOX</span>
                <span class="gba-title" style="font-size:0.75rem;">Capturados</span>
              </div>
              <div class="right">Pokémons</div>
            </div>
            """, unsafe_allow_html=True)
        
            # ✅ NOVO FILTRO: Identifica quem está na party ativa
            party_ids = [str(p) for p in (user_data.get("party") or [])]

            # ✅ FILTRO ATUALIZADO: Pega capturados, remove quem está na party e ignora EXT
            caught_all = [
                str(c) for c in user_data.get("caught", []) 
                if str(c) not in party_ids
            ]
            # Remove duplicados mantendo a ordem
            caught_all = list(dict.fromkeys(caught_all))

            # Lógica de Paginação (Mantém igual)
            PAGE_SIZE = 36  # Grid 6x6
            total_pages = max(1, ((len(caught_all) + PAGE_SIZE - 1) // PAGE_SIZE)) if caught_all else 1
            page_now = int(st.session_state.get("hub_box_page", 1))
            page_now = max(1, min(total_pages, page_now))
            st.session_state["hub_box_page"] = page_now

            # Controles de Navegação da BOX
            cpg1, cpg2, cpg3 = st.columns([1, 2, 1])
            with cpg1:
                if st.button("⬅️", key="hub_box_prev", disabled=(page_now <= 1)):
                    st.session_state["hub_box_page"] = page_now - 1
                    st.rerun()
            with cpg2:
                st.markdown(f"<div class='hub-sub' style='text-align:center'>Página <b>{page_now}</b> / {total_pages}</div>", unsafe_allow_html=True)
            with cpg3:
                if st.button("➡️", key="hub_box_next", disabled=(page_now >= total_pages)):
                    st.session_state["hub_box_page"] = page_now + 1
                    st.rerun()

            if not caught_all:
                st.info("Você ainda não marcou nenhum Pokémon como capturado.")
            else:
                # Delimitação dos Pokémons da página atual
                start = (page_now - 1) * PAGE_SIZE
                end = start + PAGE_SIZE
                page_ids = caught_all[start:end]

                # Renderização do Grid 6 colunas
                grid_cols = 6
                for r in range(0, len(page_ids), grid_cols):
                    cols = st.columns(grid_cols)
                    for col, pid in zip(cols, page_ids[r : r + grid_cols]):
                        with col:
                            sprite = _get_sprite(pid)
                            name = _get_pokemon_name(pid)
                            
                            # Usamos HTML puro para o Slot para evitar que o Streamlit quebre o layout
                            st.markdown(f'''
                                <div class="box-slot-grass">
                                    <img src="{sprite}" style="width: 55px; image-rendering: pixelated;">
                                </div>
                            ''', unsafe_allow_html=True)
                            
                            if st.button(name, key=f"hub_box_{pid}_{r}", use_container_width=True):
                                _open_box_context(pid)
                                st.rerun()

            # 
            # Menu de contexto (Aparece ao clicar em um Pokémon da BOX)
            ctx_pid = st.session_state.get("hub_context_pid")
            if ctx_pid:
                st.markdown("<div class='hub-divider'></div>", unsafe_allow_html=True)
                pname_ctx = _get_pokemon_name(ctx_pid)
                st.markdown(f"**Menu:** {pname_ctx}")
                c1, c2, c3, c4 = st.columns([1, 1, 1, 1])
                with c1:
                    if st.button("➡️ Mover p/ equipe", key="hub_ctx_move"):
                        _move_box_to_party(ctx_pid)
                with c2:
                    if st.button("📄 Ver ficha", key="hub_ctx_view"):
                        _select_pid(ctx_pid)
                        st.rerun()
                with c3:
                    if st.button("✖️ Fechar", key="hub_ctx_close"):
                        st.session_state["hub_context_pid"] = None
                        st.rerun()
                with c4:
                    if st.button("🗑️ Excluir", key="hub_ctx_delete"):
                        st.session_state["hub_context_action"] = "confirm_delete"

                if st.session_state.get("hub_context_action") == "confirm_delete":
                    st.warning("Excluir removerá este Pokémon da BOX, equipe, wishlist e status salvos.")
                    cd1, cd2 = st.columns([1, 1])
                    with cd1:
                        if st.button("✅ Confirmar exclusão", key="hub_ctx_delete_yes"):
                            _delete_pokemon_from_hub(ctx_pid)
                    with cd2:
                        if st.button("Cancelar", key="hub_ctx_delete_no"):
                            st.session_state["hub_context_action"] = None
                            st.rerun()

            st.markdown('</div>', unsafe_allow_html=True) # Fecha grass-box

            # --- 👇 NOVO BLOCO: ADICIONAR EXTERNO ---
            st.markdown('<div style="margin-top: 10px;"></div>', unsafe_allow_html=True)
            with st.expander("➕ Adicionar Visitante (Fora da Dex)"):
                st.caption("Adicione um Pokémon que não está na Pokédex do sistema.")
                ext_name = st.text_input("Nome (ex: MissingNo)", key="hub_add_ext")
                
                if st.button("📥 Adicionar à Box"):
                    if ext_name.strip():
                        # Cria ID único com prefixo EXT:
                        new_ext_id = f"EXT:{ext_name.strip()}"
                        
                        # Garante a lista
                        user_data.setdefault("caught", [])
                        
                        if new_ext_id not in user_data["caught"]:
                            user_data["caught"].append(new_ext_id)
                            save_data_cloud(trainer_name, user_data)
                            st.success(f"**{ext_name}** adicionado com sucesso!")
                            st.rerun()
                        else:
                            st.warning("Este Pokémon já está na sua lista.")
            # ----------------------------------------

        # ---------
        # PARTY (equipe ativa)
        # ---------
        # ---------
        # PARTY (equipe ativa)
        # ---------
        with col_right:
            # Container principal da Equipe
            st.markdown('<div class="team-box">', unsafe_allow_html=True)

            # Busca a lista de IDs na party
            party = [str(p) for p in (user_data.get("party") or [])]
            
            if not party:
                st.info("Sua equipe está vazia. Use a BOX para mover um Pokémon.")
            else:
                n = len(party)
        
                # Layout adaptativo conforme a quantidade de Pokémon
                if n <= 2:
                    cols_n = 1
                elif n <= 4:
                    cols_n = 2
                else:
                    cols_n = 2  # 2x4 compacto
        
                # Renderização da grade da equipe
                for r in range(0, n, cols_n):
                    cols = st.columns(cols_n)
                    for col, pid in zip(cols, party[r:r+cols_n]):
                        with col:
                            sprite = _get_sprite(pid)
                            name = _get_pokemon_name(pid)
                            typ = _get_pokemon_type(pid)
                            npv = _get_pokemon_np(pid)
                
                            st.markdown(f'''
                                <div class="gba-party-slot">
                                    <img src="{sprite}" style="width: 85px; image-rendering: pixelated; margin-bottom: 8px;">
                                    <div style="text-align: center; width: 100%;">
                                        <div style="color: white; font-family: 'Press Start 2P'; font-size: 10px; margin-bottom: 5px; text-shadow: 2px 2px 0px #000;">{name}</div>
                                        <div style="color: #ffd166; font-size: 9px; font-weight: bold;">{typ} • NP {npv}</div>
                                    </div>
                                </div>
                            ''', unsafe_allow_html=True)
                
                            b1, b2 = st.columns(2)
                            with b1:
                                if st.button("Abrir", key=f"hub_party_open_{pid}", use_container_width=True):
                                    _select_pid(pid)
                                    st.rerun()
                            with b2:
                                if st.button("Remover", key=f"hub_party_rm_{pid}", use_container_width=True):
                                    _remove_from_party(pid)
        
            st.markdown('</div>', unsafe_allow_html=True)  # Fecha team-box
        

    # ==========================
    # TAB: Lista de interesses
    # ==========================
    with t_wish:
        st.subheader("🌟 Lista de Interesses")
        st.caption("Marque aqui o que você quer caçar / o que quer registrar como shiny.")
        colw1, colw2 = st.columns(2)

        with colw1:
            st.markdown("#### ✅ Wishlist")
            wish = user_data.get("wishlist", [])
            new_item = st.text_input("Adicionar Pokémon (nome)", key="hub_wish_add")
            if st.button("Adicionar", key="hub_wish_add_btn"):
                if new_item:
                    wish.append(new_item.strip())
                    user_data["wishlist"] = list(dict.fromkeys(wish))
                    save_data_cloud(trainer_name, user_data)
                    st.rerun()
            if wish:
                for i, it in enumerate(wish):
                    c1, c2 = st.columns([0.8, 0.2])
                    with c1:
                        st.write(f"- {it}")
                    with c2:
                        if st.button("Remover", key=f"hub_wish_rm_{i}"):
                            user_data["wishlist"] = [x for x in wish if x != it]
                            save_data_cloud(trainer_name, user_data)
                            st.rerun()
            else:
                st.info("Sua wishlist está vazia.")

        with colw2:
            st.markdown("#### ✨ Shinies")
            st.caption("IDs (da sua Pokédex) marcados como shiny para imagem.")
            caught_all = [str(c) for c in user_data.get("caught", []) if not str(c).startswith("EXT:")]
            if caught_all:
                shiny_sel = st.multiselect(
                    "Marcar como shiny (capturados)",
                    options=caught_all,
                    default=[str(x) for x in user_data.get("shinies", []) if str(x) in caught_all],
                    key="hub_shiny_multi",
                )
                if set(shiny_sel) != set([str(x) for x in user_data.get("shinies", [])]):
                    user_data["shinies"] = shiny_sel
                    save_data_cloud(trainer_name, user_data)
                    st.rerun()
            else:
                st.info("Você precisa capturar Pokémons para marcar shinies.")

    # ==========================
    # TAB: vistos
    # ==========================
    with t_seen:
        total = len(df) if df is not None else 0
        vistos = len(user_data.get("seen", []))
        st.markdown("### Progresso da Pokédex")
        st.progress(min(vistos / total, 1.0) if total else 0.0)
        st.write(f"**{vistos}** de **{total}** Pokémons registrados.")

    # ==========================
    # TAB: MEU TREINADOR
    # ==========================
    with t_trainer:
        st.subheader("📸 Meu Treinador")
        user_data.setdefault("trainer_profile", {})
        profile = user_data["trainer_profile"]

        if "show_trainer_uploader" not in st.session_state:
            st.session_state["show_trainer_uploader"] = False
        if "show_trainer_badges" not in st.session_state:
            st.session_state["show_trainer_badges"] = False

        storage_path = profile.get("photo_storage_path")
        photo_bytes = download_storage_bytes(storage_path) if storage_path else None
        if (not photo_bytes) and profile.get("photo_thumb_b64"):
            try:
                photo_bytes = base64.b64decode(profile["photo_thumb_b64"])
            except Exception:
                photo_bytes = None

        chosen_avatar, chosen_path = get_selected_trainer_avatar(user_data)
        badge_dir = os.path.join("Assets", "insignias")
        badge_files = []
        if os.path.isdir(badge_dir):
            badge_files = sorted(
                f
                for f in os.listdir(badge_dir)
                if f.lower().endswith((".png", ".jpg", ".jpeg", ".webp"))
            )
        badge_options = [os.path.splitext(f)[0] for f in badge_files]
        badge_paths = {
            os.path.splitext(f)[0]: os.path.join(badge_dir, f) for f in badge_files
        }
        selected_badges = profile.get("badges", [])

        top_col_left, top_col_right = st.columns([2, 1])
        with top_col_left:
            st.markdown("#### 🖼️ Foto enviada")
            if photo_bytes:
                st.image(photo_bytes, width=220)
            else:
                st.info("Nenhuma foto enviada.")

            st.markdown("#### 🎭 Avatar escolhido")
            if chosen_avatar and chosen_path:
                st.image(chosen_path, width=160, caption=chosen_avatar)
            else:
                st.info("Ainda não há avatar selecionado.")

        with top_col_right:
            st.markdown("#### 🏅 Insígnias")
            if selected_badges:
                badge_cols = st.columns(min(4, len(selected_badges)))
                for idx, badge_key in enumerate(selected_badges):
                    badge_path = badge_paths.get(badge_key)
                    if not badge_path:
                        continue
                    with badge_cols[idx % len(badge_cols)]:
                        st.image(badge_path, width=70)
            else:
                st.info("Nenhuma insígnia adicionada.")

        action_cols = st.columns(2)
        with action_cols[0]:
            if st.button("📤 Enviar foto e construir avatar"):
                st.session_state["show_trainer_uploader"] = True
        with action_cols[1]:
            if st.button("🏅 Adicionar insígnias"):
                st.session_state["show_trainer_badges"] = True

        if st.session_state["show_trainer_badges"]:
            st.markdown("---")
            st.subheader("🏅 Minhas insígnias")
            if not badge_options:
                st.info("Nenhuma insígnia encontrada em Assets/insignias.")
            else:
                badge_pick = st.multiselect(
                    "Selecione suas insígnias",
                    options=badge_options,
                    default=selected_badges,
                    key="trainer_badge_select",
                )
                if st.button("💾 Salvar insígnias", key="trainer_badge_save"):
                    profile["badges"] = badge_pick
                    save_data_cloud(trainer_name, user_data)
                    st.success("Insígnias atualizadas!")
                    st.rerun()

        if st.session_state["show_trainer_uploader"]:
            st.markdown("---")
            st.subheader("📷 Enviar foto e sugerir avatar")
            upload_limit = 3
            today = datetime.now().strftime("%Y-%m-%d")
            upload_state = profile.get("photo_uploads", {})
            last_upload_date = upload_state.get("date")
            upload_count = int(upload_state.get("count", 0) or 0) if last_upload_date == today else 0
            remaining_uploads = max(0, upload_limit - upload_count)
            upload_limit_reached = upload_count >= upload_limit

            if upload_limit_reached:
                st.warning("Você atingiu o limite de 3 uploads de foto hoje. Tente novamente amanhã.")
            else:
                st.caption(f"Uploads restantes hoje: {remaining_uploads} de {upload_limit}.")

            uploaded_photo = st.file_uploader(
                "Envie uma foto para recortar e buscar o seu avatar",
                type=["png", "jpg", "jpeg", "webp"],
                key="trainer_photo_upload",
                disabled=upload_limit_reached,
            )

            cropped_image = None
            if uploaded_photo is not None:
                try:
                    raw_bytes = uploaded_photo.getvalue()
                    raw_image = Image.open(io.BytesIO(raw_bytes))
                    cropped_image = crop_center_square(raw_image)
                    st.image(cropped_image, caption="Pré-visualização (recorte central)", width=260)
                except Exception:
                    st.error("Não foi possível ler essa imagem. Tente outro arquivo.")

            c_photo_actions = st.columns([1, 1, 2])

            with c_photo_actions[0]:
                if st.button("💾 Salvar foto", disabled=cropped_image is None or upload_limit_reached):
                    if upload_limit_reached:
                        st.warning("Limite diário atingido. Aguarde até amanhã para enviar outra foto.")
                        st.stop()
                    # --- gera PNG cheio (para o Storage) ---
                    buffer_full = io.BytesIO()
                    cropped_image.convert("RGB").save(buffer_full, format="PNG")

                    # --- gera thumb pequeno (para UI rápida / cabe no Sheets) ---
                    thumb = cropped_image.copy()
                    thumb.thumbnail((96, 96), Image.Resampling.LANCZOS)
                    buffer_thumb = io.BytesIO()
                    thumb.convert("RGB").save(buffer_thumb, format="PNG")

                    # --- sobe a foto cheia pro Firebase Storage ---
                    db, bucket = init_firebase()
                    storage_path = f"trainer_photos/{safe_doc_id(trainer_name)}/profile.png"
                    upload_png_to_storage(bucket, buffer_full.getvalue(), storage_path)

                    # --- salva só o path + thumb no user_data (Sheets não estoura) ---
                    profile["photo_storage_path"] = storage_path
                    profile["photo_thumb_b64"] = base64.b64encode(buffer_thumb.getvalue()).decode("utf-8")

                    # nunca mais salvar a foto cheia em base64 no Sheets
                    profile.pop("photo_b64", None)

                    profile["photo_updated_at"] = str(datetime.now())
                    profile["photo_uploads"] = {
                        "date": today,
                        "count": upload_count + 1,
                    }
                    save_data_cloud(trainer_name, user_data)

                    st.success("Foto salva! Agora vamos sugerir avatares.")
                    st.rerun()

            with c_photo_actions[1]:
                if st.button("🗑️ Remover foto"):
                    # apaga do Storage também
                    db, bucket = init_firebase()
                    old_path = profile.get("photo_storage_path")
                    if old_path:
                        try:
                            bucket.blob(old_path).delete()
                        except Exception:
                            pass

                    profile.pop("photo_storage_path", None)
                    profile.pop("photo_b64", None)
                    profile.pop("photo_thumb_b64", None)
                    profile.pop("photo_updated_at", None)

                    save_data_cloud(trainer_name, user_data)
                    st.success("Foto removida.")
                    st.rerun()

            st.markdown("---")
            st.subheader("🎭 Avatares sugeridos")

            catalog = build_trainer_avatar_catalog()
            index_entries = build_trainer_avatar_index()

            if not catalog or not index_entries:
                st.warning("Nenhum avatar encontrado na pasta trainer.")
            else:
                random_bases = []

                available_bases = sorted(catalog.keys())
                random_suggestions_key = "trainer_random_avatar_bases"
                random_suggestions_pool_key = "trainer_random_avatar_pool"
                expected_total = min(10, len(available_bases))

                if (
                    st.session_state.get(random_suggestions_pool_key) != available_bases
                    or len(st.session_state.get(random_suggestions_key, [])) != expected_total
                ):
                    st.session_state[random_suggestions_pool_key] = available_bases
                    st.session_state[random_suggestions_key] = random.sample(available_bases, k=expected_total)
                random_bases = st.session_state.get(random_suggestions_key, [])
                chosen_avatar = profile.get("avatar_choice")

                def render_avatar_suggestions(section_id: str, title: str, bases: list[str], best_defaults: dict[str, str]):
                    if not bases:
                        return
                    st.markdown(f"**{title}**")
                    cols = st.columns(min(5, len(bases)))
                    for idx, base in enumerate(bases):
                        items = catalog.get(base, [])
                        if not items:
                            continue
                        names = [item["name"] for item in items]
                        default_name = chosen_avatar if chosen_avatar in names else best_defaults.get(base)
                        default_idx = names.index(default_name) if default_name in names else 0
                        with cols[idx]:
                            st.markdown(f"**{base.title()}**")
                            selected_skin = st.selectbox(
                                "Skin",
                                names,
                                index=default_idx,
                                key=f"trainer_skin_{section_id}_{base}",
                            )
                            sel_path = next(
                                (i["path"] for i in items if i["name"] == selected_skin),
                                items[0]["path"],
                            )
                            st.image(sel_path, width=120)

                            if st.button("✅ Escolher", key=f"trainer_pick_{section_id}_{base}"):
                                profile["avatar_choice"] = selected_skin
                                profile["avatar_base"] = base

                                try:
                                    _, bucket = init_firebase()
                                    avatar_upload = upload_avatar_choice_to_storage(bucket, trainer_name, selected_skin)
                                    if avatar_upload:
                                        profile["avatar_storage_path"] = avatar_upload.get("storage_path")
                                        profile["avatar_url"] = avatar_upload.get("url")
                                except Exception:
                                    pass

                                save_data_cloud(trainer_name, user_data)
                                st.success(f"Avatar selecionado: {selected_skin}.")
                                st.rerun()


                render_avatar_suggestions(
                    "random",
                    "10 sorteados aleatoriamente",
                    random_bases,
                    {},
                )

            st.markdown("---")
            st.subheader("🧵 Skins do personagem escolhido")
            chosen_avatar, chosen_path = get_selected_trainer_avatar(user_data)
            if chosen_avatar and chosen_path:
                chosen_base = profile.get("avatar_base") or _trainer_avatar_base(chosen_avatar)
                base_items = catalog.get(chosen_base, [])
                base_names = [item["name"] for item in base_items]
                current_idx = base_names.index(chosen_avatar) if chosen_avatar in base_names else 0
                st.image(chosen_path, width=140, caption=f"Atual: {chosen_avatar}")
                new_skin = st.selectbox(
                    "Trocar skin (mesmo personagem)",
                    base_names,
                    index=current_idx,
                    key="trainer_skin_swap",
                )
                if st.button("🔁 Atualizar skin"):
                    profile["avatar_choice"] = new_skin
                    profile["avatar_base"] = chosen_base

                    try:
                        _, bucket = init_firebase()
                        avatar_upload = upload_avatar_choice_to_storage(bucket, trainer_name, new_skin)
                        if avatar_upload:
                            profile["avatar_storage_path"] = avatar_upload.get("storage_path")
                            profile["avatar_url"] = avatar_upload.get("url")
                    except Exception:
                        pass

                    save_data_cloud(trainer_name, user_data)
                    st.success("Skin atualizada!")
                    st.rerun()
            else:
                st.info("Escolha um avatar nas sugestões acima para liberar as skins.")

    # ==========================
    # CRIAÇÃO DE 
    # ==========================

elif page == "Criação Guiada de Fichas":
    st.title("🧩 Criação Guiada de Fichas")
    
    # CORREÇÃO: Inicializa o last_page se não existir
    if "last_page" not in st.session_state:
        st.session_state["last_page"] = ""

    # Só reseta o menu se estivermos REALMENTE vindo de outra página
    if st.session_state["last_page"] != "Criação Guiada de Fichas":
        if st.session_state.get("cg_force_guided"):
            st.session_state["cg_view"] = "guided"
            st.session_state["cg_force_guided"] = False
        else:
            st.session_state["cg_view"] = "menu"
        
        # IMPORTANTE: Atualiza o last_page para evitar o loop no próximo rerun
        st.session_state["last_page"] = "Criação Guiada de Fichas"

    if "cg_view" not in st.session_state:
        st.session_state["cg_view"] = "menu"

    # menu interno
    if st.session_state["cg_view"] == "menu":
        choice = st.radio(
            "Escolha o que você quer fazer ao abrir a criação de ficha:",
            ["Criação Guiada", "Criação de Golpes"],
            horizontal=True
        )
        if choice == "Criação Guiada":
            st.session_state["cg_view"] = "guided"
            st.rerun()
        else:
            st.session_state["cg_view"] = "moves"
            st.rerun()

    # ==========================
    # A) CRIAÇÃO DE GOLPES (Golpe Builder MM3e)
    # ==========================
    if st.session_state["cg_view"] == "moves":
        trainer_name_gb = st.session_state.get("trainer_name", "Treinador")
        return_to_view_gb = st.session_state.get("cg_return_to")
        if return_to_view_gb != "guided":
            return_to_view_gb = None

        if _GB_AVAILABLE:
            _gb_render_builder(
                excel_path=_resolve_asset_path("golpes_pokemon_MM_reescritos.xlsx"),
                state_key_prefix="cg_gb_main",
                return_to_view=return_to_view_gb,
                trainer_name=trainer_name_gb,
            )
        else:
            st.subheader("⚔️ Criação de Golpes")
            st.warning("Módulo golpe_builder_ui não encontrado. Usando criador legado.")
            return_to_view = return_to_view_gb
            render_move_creator(
                excel_path=_resolve_asset_path("golpes_pokemon_MM_reescritos.xlsx"),
                state_key_prefix="cg_moves_ui",
                return_to_view=return_to_view,
            )

        # botão voltar (fallback se render_golpe_builder não mostrar o próprio)
        if not _GB_AVAILABLE:
            if st.button("⬅️ Voltar para a ficha", key="btn_back_to_sheet"):
                st.session_state["cg_view"] = st.session_state.get("cg_return_to", "menu")
                st.rerun()


    # ==========================
    # B) CRIAÇÃO GUIADA (FICHA) - OTIMIZADA
    # ==========================
    if st.session_state["cg_view"] == "guided":
        st.subheader("🧬 Criação Guiada")
        with st.expander("📄 Importar ficha M&M (PDF)", expanded=False):
            mm_pdf = st.file_uploader("Envie a ficha (.pdf) exportada", type=["pdf"], key="mm_pdf_upl")
    
            if mm_pdf and st.button("Importar ficha M&M", type="primary", key="btn_import_mm_pdf"):
                pdf_bytes = mm_pdf.read()
    
                try:
                    # 1) tenta primeiro o PDF gerado pelo próprio site (padrão Blastoise)
                    try:
                        sheet_payload = parse_sheet_pdf(pdf_bytes)
                        # garante id no pokemon (alguns PDFs antigos podem não ter)
                        sheet_payload.setdefault("pokemon", {})
                        if "id" not in sheet_payload["pokemon"]:
                            fallback_name = sheet_payload["pokemon"].get("name", "pokemon")
                            sheet_payload["pokemon"]["id"] = normalize_sheet_pokemon_id(f"pdf_{safe_doc_id(fallback_name)}", fallback_name)
                    except Exception:
                        # 2) fallback: Hero Lab / M&M
                        sheet_payload = import_mm_pdf_to_sheet_payload(pdf_bytes)

    
                    db, bucket = init_firebase()
                    trainer_name = st.session_state.get("trainer_name", "Treinador")
    
                    sheet_id, storage_path = save_sheet_with_pdf(
                        db, bucket, trainer_name, sheet_payload, pdf_bytes=pdf_bytes
                    )
    
                    st.success(f"Importado e salvo! sheet_id={sheet_id}")
                    st.json(sheet_payload)  # debug
                except Exception as e:
                    st.error(f"Falha ao importar: {e}")
    
        # 1. Inicialização segura: Só executa se o draft não existir
        if "cg_draft" not in st.session_state:
            cg_init()
            
        nome_salvo = st.session_state.get("cg_draft", {}).get("pname", "")
    
        # 2. Input de Nome ÚNICO (Evita DuplicateElementKey)
        pname = st.text_input(
            "Digite o nome do Pokémon (ex: Blastoise)", 
            value=nome_salvo,  # <--- Aqui está a mágica. Se for a 1ª vez, entra ""
            placeholder="Ex: Blastoise", 
            key="cg_pname"
        )

        pname_norm = _norm(pname) if pname else ""
        last_pname_norm = st.session_state.get("cg_active_pokemon_norm", "")
        if pname_norm and last_pname_norm and pname_norm != last_pname_norm:
            cg_reset_for_new_pokemon(pname)
        if pname_norm:
            st.session_state["cg_active_pokemon_norm"] = pname_norm
        else:
            st.session_state["cg_active_pokemon_norm"] = ""

        # Tenta pegar o NP salvo. Se não existir, assume 0.
        np_salvo = st.session_state.get("cg_np", 0)
        # Garante que é um inteiro (caso tenha vindo None ou string por algum erro estranho)
        if np_salvo is None: 
            np_salvo = 0
    
        # 3. Processamento e Criação da Ficha (TUDO deve estar dentro deste IF)
        if pname:
            # ✅ LAYOUT: Sugestões da Pokédex logo abaixo da busca
            if len(pname) >= 2:
                matches = df[df["Nome"].str.lower().str.contains(pname.lower(), na=False)].head(10)
                if not matches.empty:
                    st.caption("Sugestões encontradas na sua Pokédex:")
                    st.write(matches[["Nº", "Nome"]])
                    
            raw_name = pname.strip().lower()

            is_nidoran_generic = raw_name in ["nidoran", "nidoran♀", "nidoran♂", "nidoran-f", "nidoran-m"]
            is_lycanroc_any = "lycanroc" in raw_name
            is_lycanroc_generic = raw_name == "lycanroc"

            if is_nidoran_generic:
                choice = st.radio(
                    "Qual Nidoran?",
                    ["Nidoran ♀", "Nidoran ♂"],
                    horizontal=True,
                    key="nidoran_choice"
                )
                poke_query = "nidoran-f" if "♀" in choice else "nidoran-m"
                pname = "Nidoran ♀" if "♀" in choice else "Nidoran ♂"

            elif is_lycanroc_any:
                # Nome genérico ("lycanroc") → mostra radio; nome com forma já salva → detecta forma
                if is_lycanroc_generic:
                    lyc_choice = st.radio(
                        "Qual forma do Lycanroc?",
                        ["Midday (Dia)", "Midnight (Noite)", "Dusk (Crepúsculo)"],
                        horizontal=True,
                        key="lycanroc_choice",
                    )
                else:
                    # Forma já conhecida (ex: "Lycanroc (Midnight)") — detecta automaticamente
                    if "midnight" in raw_name or "noite" in raw_name:
                        lyc_choice = "Midnight (Noite)"
                    elif "dusk" in raw_name or "crepúscul" in raw_name:
                        lyc_choice = "Dusk (Crepúsculo)"
                    else:
                        lyc_choice = "Midday (Dia)"

                if "Noite" in lyc_choice:
                    poke_query = "lycanroc-midnight"
                    pname = "Lycanroc (Midnight)"
                elif "Crepúsculo" in lyc_choice:
                    poke_query = "lycanroc-dusk"
                    pname = "Lycanroc (Dusk)"
                else:
                    poke_query = "lycanroc-midday"
                    pname = "Lycanroc (Midday)"

            else:
                # Usa a função auxiliar apenas aqui
                poke_query = to_pokeapi_name(pname)

            # Atualiza o rascunho com o nome específico da forma
            st.session_state["cg_draft"]["pname"] = pname
    
            # Busca ID no Excel
            pid = resolve_pokemon_pid(df, pname)


            # Linha da Pokédex para recursos contextuais (ex.: Viabilidade)
            row = pd.DataFrame()
            try:
                pid_norm = _norm_pid(pid)
                if pid_norm:
                    row = df[df["Nº"].apply(_norm_pid) == pid_norm]

                # fallback por nome quando o PID não resolve (formas/aliases)
                if row.empty and "Nome" in df.columns:
                    nomes = df["Nome"].astype(str)
                    row = df[nomes.apply(normalize_text) == normalize_text(pname)]

                if row.empty and "Nome" in df.columns:
                    api_key = to_pokeapi_name(pname)
                    if api_key:
                        row = df[nomes.apply(lambda x: to_pokeapi_name(str(x))) == api_key]
            except Exception:
                row = pd.DataFrame()

    
            # Busca dados na API
            with st.spinner("Buscando dados do Pokémon online (stats + ability + tipos)..."):
                pjson = pokeapi_get_pokemon(poke_query)
                base_stats = pokeapi_parse_stats(pjson)
                types = pokeapi_parse_types(pjson)
                abilities = pokeapi_parse_abilities(pjson)

                imported_name = st.session_state.get("cg_imported_name")
                imported_types = (
                    st.session_state.get("cg_imported_types")
                    if imported_name and _norm(imported_name) == _norm(pname)
                    else None
                )
                imported_abilities = (
                    st.session_state.get("cg_imported_abilities")
                    if imported_name and _norm(imported_name) == _norm(pname)
                    else None
                )
                if imported_types:
                    types = list(imported_types)
                if imported_abilities:
                    abilities = sorted(set(abilities + list(imported_abilities)))
                
                # Escolha de Habilidades
                saved_abilities = st.session_state.get("cg_abilities")
                if not isinstance(saved_abilities, list):
                    saved_abilities = None
                chosen_abilities = st.multiselect(
                    "Escolha a(s) habilidade(s) (pode mais de uma):",
                    options=abilities,
                    default=[
                        a
                        for a in (
                            saved_abilities
                            or (imported_abilities if imported_abilities else (abilities[:1] if abilities else []))
                        )
                        if a in abilities
                    ],
                )
                
                if not chosen_abilities:
                    chosen_abilities = abilities
                st.session_state["cg_abilities"] = chosen_abilities
    
            # 3) NP / PP
            np_sugerido = get_np_for_pokemon(df, pid, fallback_np=6)
            # --- LÓGICA DE RECUPERAÇÃO DO NP (CORREÇÃO) ---
            # Tenta pegar do rascunho salvo (draft) e do widget atual
            val_draft = st.session_state.get("cg_draft", {}).get("np", 0)
            val_widget = st.session_state.get("cg_np", 0)
            np_salvo_final = max(int(val_draft or 0), int(val_widget or 0))
            
            np_ = st.number_input(
                "NP do seu Pokémon (o jogador informa)", 
                min_value=0, 
                value=int(np_salvo_final), # <--- Carrega o valor salvo ou 0
                step=1, 
                key="cg_np", 
                on_change=_cg_sync_from_np
            )
            # --- SALVAMENTO FORÇADO ---
            # Salva no rascunho imediatamente para não perder ao trocar de aba
            if "cg_draft" in st.session_state:
                st.session_state["cg_draft"]["np"] = np_
            pp_total = calc_pp_budget(np_)
    
            pp_spent_moves = sum((m.get("pp_cost") or 0) for m in st.session_state.get("cg_moves", []))
            pp_spent_moves += sum(_move_accuracy_pp(m) for m in st.session_state.get("cg_moves", []))
    
            tabs = st.tabs(
                [
                    "1️⃣ Visão Geral",
                    "2️⃣ Abilities e Defesas",
                    "3️⃣ Skills e Advantages",
                    "4️⃣ Golpes",
                    "5️⃣ Revisão e Exportação",
                ]
            )
            
            pp_abilities = 0
            pp_defenses = 0
            pp_skills = 0
            pp_advantages = 0
            pp_moves = pp_spent_moves
    
            # 4) Atributos base
            PL = int(np_)
            cap = 2 * PL
            
            atk = int(base_stats.get("attack", 10))
            spatk = int(base_stats.get("special-attack", 10))
            spe = int(base_stats.get("speed", 10))
            def_ = int(base_stats.get("defense", 10))
            spdef = int(base_stats.get("special-defense", 10))
            
            int_base = max(0, (spatk - 10) // 10)
            stgr_base = max(0, (atk - 10) // 10)
            
            den_td = max(1, def_ + spe)
            thg_base = round((def_ / den_td) * cap)
            dodge_base = cap - thg_base
            
            den_wf = max(1, spdef + def_)
            will_base = round((spdef / den_wf) * cap)
            fort_base = cap - will_base
    
            with tabs[0]:
                st.markdown(
                    f"""
                    <div class="cg-card">
                        <div class="cg-title">Visão Geral</div>
                        <span class="cg-pill">NP {np_}</span>
                        <span class="cg-pill">PP Total {pp_total}</span>
                        <span class="cg-pill">Tipos: {', '.join(types)}</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                st.markdown(
                    f"**Pokémon:** {pname}  \n"
                    f"**Abilities escolhidas:** {', '.join(chosen_abilities)}"
                )
                st.info("Use as abas para preencher cada etapa. O total de PP gastos é somado automaticamente no final.")

                # ── Anotações da ficha ──────────────────────────────────────
                _notas_saved = st.session_state.get("cg_draft", {}).get("notas", "")
                _notas_new = st.text_area(
                    "📝 Anotações / Notas",
                    value=_notas_saved,
                    height=90,
                    key="cg_notas",
                    placeholder="Estratégia, personalidade, histórico do Pokémon, observações...",
                    help="Notas livres salvas junto da ficha.",
                )
                if _notas_new != _notas_saved:
                    st.session_state["cg_draft"]["notas"] = _notas_new

                # --------------------------
                # Preset por Viabilidade (Arquétipos)
                # --------------------------
                viab_text = ""
                try:
                    if isinstance(row, pd.DataFrame) and (not row.empty) and ("Viabilidade" in row.columns):
                        viab_text = str(row.iloc[0].get("Viabilidade", "") or "").strip()
                except Exception:
                    viab_text = ""

                if viab_text and viab_text.strip() and viab_text.strip().lower() not in {"nan", "sem dados.", "sem dados"}:
                    with st.expander("📌 Arquétipo de Viabilidade (1 clique)", expanded=True):
                        st.caption("Isso vira um *preset* para facilitar a criação: sugere **golpes + habilidade** e você ajusta depois. Não altera suas fórmulas nem regras de PP.")
                        archetypes = _viab_parse_archetypes(viab_text)
                        if not archetypes:
                            st.write(viab_text)
                        else:
                            labels = [a.get("label", "Arquétipo") for a in archetypes]
                            sel_idx = st.selectbox(
                                "Escolha um arquétipo (da sua Pokédex):",
                                options=list(range(len(labels))),
                                format_func=lambda i: labels[int(i)],
                                key="cg_viab_sel_idx",
                            )
                            arch = archetypes[int(sel_idx)]

                            st.markdown(f"**Estratégia:** `{arch.get('code') or '—'}`")
                            if arch.get("ability"):
                                st.markdown(f"**Habilidade sugerida:** `{arch['ability']}`")
                            if arch.get("partners"):
                                st.markdown("**Parceiros sugeridos:** " + ", ".join(arch["partners"]))

                            # Com key fixa, o Streamlit prioriza session_state e ignora `value` em trocas
                            # de arquétipo; sincronizamos explicitamente o conteúdo antes de renderizar.
                            st.session_state["cg_viab_raw_view"] = str(arch.get("raw", "") or "")
                            st.text_area(
                                "Texto da Viabilidade (referência)",
                                value=st.session_state["cg_viab_raw_view"],
                                height=120,
                                key="cg_viab_raw_view",
                                disabled=True,
                            )

                            rank_default = st.slider(
                                "Rank padrão para adicionar golpes desse arquétipo",
                                1, 20,
                                min(10, max(1, int(np_ or 1))),
                                key="cg_viab_rank_default",
                            )

                            excel_path_viab = _resolve_asset_path("golpes_pokemon_MM_reescritos.xlsx")

                            colA, colB = st.columns([1, 1])
                            with colA:
                                core_n = 3
                                b1, b2 = st.columns(2)
                                with b1:
                                    apply_full = st.button(
                                        "✅ Adicionar kit completo",
                                        key="cg_viab_apply_full_btn",
                                        use_container_width=True,
                                    )
                                with b2:
                                    apply_core = st.button(
                                        f"⚡ Só núcleo ({core_n} golpes)",
                                        key="cg_viab_apply_core_btn",
                                        use_container_width=True,
                                    )

                                mode = "full" if apply_full else ("core" if apply_core else None)
                                if mode:
                                    added, missing, abil = _viab_apply_to_session(
                                        arch,
                                        int(rank_default),
                                        excel_path_viab,
                                        mode=mode,
                                        core_n=core_n,
                                    )

                                    # tenta setar habilidade sugerida (se ela existir no Pokémon)
                                    if abil:
                                        abil_norm = _norm(abil)
                                        matched_abil = None
                                        for a in (abilities or []):
                                            if _norm(a) == abil_norm:
                                                matched_abil = a
                                                break
                                        if matched_abil:
                                            cur = st.session_state.get("cg_abilities") or []
                                            if matched_abil not in cur:
                                                cur = [matched_abil] + cur
                                            st.session_state["cg_abilities"] = cur

                                    # guarda para revisão
                                    st.session_state["cg_viab_selected"] = arch
                                    st.session_state["cg_viab_mode"] = (
                                        "Kit completo" if mode == "full" else f"Núcleo ({core_n})"
                                    )

                                    msg = f"Golpes adicionados: {len(added)}"
                                    if added:
                                        msg += " — " + ", ".join(added[:8]) + ("…" if len(added) > 8 else "")
                                    st.success(msg)

                                    if missing:
                                        st.warning(
                                            "Não encontrei no banco: "
                                            + ", ".join(missing[:12])
                                            + ("…" if len(missing) > 12 else "")
                                            + ". Você pode criar esses golpes na aba **Criação de Golpes**."
                                        )
                                    st.rerun()
                            with colB:
                                st.caption("Depois, vá na aba **4️⃣ Golpes** para ajustar Rank/PP/Acerto em **Lista & ajustes**.")

    
            with tabs[1]:
                st.markdown("### 📊 Atributos (auto + editável)")
                cap = 2 * int(st.session_state.get("cg_np", 0) or 0)
                _cg_init_defenses_if_missing(dodge_base, fort_base)
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    stgr_suggested_max = max(0, int(stgr_base))
                    stgr_current = int(st.session_state.get("cg_stgr", stgr_suggested_max))
                    if stgr_current > stgr_suggested_max:
                        st.session_state["cg_stgr"] = stgr_suggested_max
                        stgr_current = stgr_suggested_max
                    stgr = st.number_input(
                        "Stgr (Força)",
                        value=stgr_current,
                        min_value=0, max_value=stgr_suggested_max, key="cg_stgr",
                    )

                    int_suggested_max = max(0, int(int_base))
                    int_current = int(st.session_state.get("cg_int", int_suggested_max))
                    if int_current > int_suggested_max:
                        st.session_state["cg_int"] = int_suggested_max
                        int_current = int_suggested_max
                    intellect = st.number_input(
                        "Int (Intelecto)",
                        value=int_current,
                        min_value=0, max_value=int_suggested_max, key="cg_int",
                    )
                
                with col2:
                    dodge_value = int(st.session_state.get("cg_dodge", dodge_base))
                    dodge_min = min(dodge_value, max(0, int(dodge_base) - 2))
                    dodge_max = max(dodge_value, min(99, int(dodge_base) + 2))
                    dodge = st.number_input(
                        "Dodge", key="cg_dodge",
                        value=dodge_value, min_value=dodge_min, max_value=dodge_max,
                        on_change=_cg_sync_from_dodge,
                    )
                    parry = st.number_input(
                        "Parry", key="cg_parry",
                        value=int(st.session_state.get("cg_parry", st.session_state.get("cg_dodge", dodge_base))),
                        min_value=0, max_value=99, disabled=True,
                    )
            
                with col3:
                    thg = st.number_input(
                        "Thg (Toughness)", key="cg_thg",
                        value=int(st.session_state.get("cg_thg", max(0, cap - int(st.session_state.get("cg_dodge", dodge_base))))),
                        min_value=0, max_value=99, disabled=True,
                    )
                    fortitude = st.number_input(
                        "Fortitude", key="cg_fortitude",
                        value=int(st.session_state.get("cg_fortitude", fort_base)),
                        min_value=min(int(st.session_state.get("cg_fortitude", fort_base)), max(0, int(fort_base) - 2)),
                        max_value=max(int(st.session_state.get("cg_fortitude", fort_base)), min(99, int(fort_base) + 2)),
                        on_change=_cg_sync_from_fortitude,
                    )
                    will = st.number_input(
                        "Will", key="cg_will",
                        value=int(st.session_state.get("cg_will", max(0, cap - int(st.session_state.get("cg_fortitude", fort_base))))),
                        min_value=0, max_value=99, disabled=True,
                    )
    
                st.markdown("### 💰 PP automático")
                pp_stgr = int(stgr) * 2
                pp_int  = int(intellect) * 2
                pp_dodge = int(dodge) * 1
                pp_parry = int(parry) * 1
                pp_thg   = int(thg) * 1
                pp_will  = int(will) * 1
                pp_fort  = int(fortitude) * 1
                
                pp_abilities_auto = pp_stgr + pp_int
                pp_defenses_auto  = pp_dodge + pp_parry + pp_thg + pp_will + pp_fort
                
                st.write(f"**Abilities:** {pp_abilities_auto} PP (STGR {pp_stgr} + INT {pp_int})")
                st.write(f"**Defesas:** {pp_defenses_auto} PP (Dodge {pp_dodge} + Parry {pp_parry} + Thg {pp_thg} + Will {pp_will} + Fort {pp_fort})")
                
                pp_abilities = pp_abilities_auto
                pp_defenses  = pp_defenses_auto
    
            with tabs[2]:
                st.markdown("### 🧠 Skills (M&M 3e)")
                if "cg_skills" not in st.session_state:
                    st.session_state["cg_skills"] = {k: 0 for k in SKILLS_MM3}
                    st.session_state["cg_skill_custom"] = []
    
                cols_sk = st.columns(3)
                total_skill_ranks = 0
                for i, sk in enumerate(SKILLS_MM3):
                    with cols_sk[i % 3]:
                        v_sk = st.number_input(sk, min_value=0, max_value=40, value=int(st.session_state["cg_skills"].get(sk, 0)), step=1, key=f"cg_skill_{sk}")
                        st.session_state["cg_skills"][sk] = int(v_sk)
                        total_skill_ranks += int(v_sk)
    
                st.divider()
                st.markdown("### Skills extras")
                add_name = st.text_input("Nome da skill extra", key="cg_skill_add_name")
                if st.button("➕ Adicionar skill extra", key="cg_skill_add_btn"):
                    if add_name.strip():
                        st.session_state["cg_skill_custom"].append({"name": add_name.strip(), "ranks": 0})
    
                for idx, row_sk in enumerate(list(st.session_state["cg_skill_custom"])):
                    c1, c2, c3 = st.columns([6, 2, 2])
                    with c1: st.write(row_sk["name"])
                    with c2:
                        rv = st.number_input("Ranks", min_value=0, max_value=40, value=int(row_sk.get("ranks", 0)), key=f"cg_skill_custom_rank_{idx}")
                        row_sk["ranks"] = int(rv)
                        total_skill_ranks += int(rv)
                    with c3:
                        if st.button("❌", key=f"cg_skill_custom_del_{idx}"):
                            st.session_state["cg_skill_custom"].pop(idx)
                            st.rerun()
    
                pp_skills = total_skill_ranks / 2
                st.info(f"Total de ranks: **{total_skill_ranks}** → PP em Skills: **{pp_skills}**")
    
                st.markdown("### ⭐ Advantages (sugestões)")
                st.session_state.setdefault("cg_advantages_custom", [])
                adv_suggestions = suggest_advantages(pjson=pjson, base_stats=base_stats, types=types, abilities=abilities)
                if not adv_suggestions:
                    st.info("Nenhuma vantagem sugerida automaticamente.")

                st.markdown("### Advantages manuais")
                adv_manual_name = st.text_input("Digite a advantage manual", key="cg_advantage_add_name")
                if st.button("➕ Adicionar advantage manual", key="cg_advantage_add_btn"):
                    adv_manual_name = adv_manual_name.strip()
                    if adv_manual_name and adv_manual_name not in st.session_state["cg_advantages_custom"]:
                        st.session_state["cg_advantages_custom"].append(adv_manual_name)

                for idx, adv_name in enumerate(list(st.session_state["cg_advantages_custom"])):
                    c1, c2 = st.columns([8, 2])
                    with c1:
                        st.write(adv_name)
                    with c2:
                        if st.button("❌", key=f"cg_advantage_custom_del_{idx}"):
                            st.session_state["cg_advantages_custom"].pop(idx)
                            st.rerun()

                labels = [a.label() for a in adv_suggestions]
                notes_map = {a.label(): (a.note or "") for a in adv_suggestions}
                saved_adv = st.session_state.get("cg_advantages") or []
                all_advantages = sorted(set(labels + list(saved_adv) + list(st.session_state["cg_advantages_custom"])))
                chosen_labels = st.multiselect(
                    "Selecione advantages:",
                    options=all_advantages,
                    default=[lab for lab in saved_adv if lab in all_advantages],
                )
                chosen_adv = chosen_labels
                for lab in chosen_labels:
                    if notes_map.get(lab):
                        st.caption(f"• {lab}: {notes_map[lab]}")
                st.session_state["cg_advantages"] = chosen_adv
    
                pp_advantages = len(chosen_adv)
                st.info(f"Advantages escolhidas: **{pp_advantages} PP**")
    
            with tabs[3]:
                st.markdown("### ⚔️ Golpes & Poderes (fluxo guiado)")
                st.caption(
                    "Adicione golpes prontos, monte um poder **passo a passo** (sem precisar ler o texto do M&M) "
                    "ou use o criador completo. Suas fórmulas e regras de PP não mudam — apenas organizamos o processo."
                )

                excel_path = _resolve_asset_path("golpes_pokemon_MM_reescritos.xlsx")
                try:
                    db_moves_guided = load_move_db(excel_path)
                except Exception as e:
                    st.error(f"Não consegui carregar o Excel de golpes: {e}")
                    db_moves_guided = None

                # garante lista de golpes na sessão
                st.session_state.setdefault("cg_moves", st.session_state.get("cg_draft", {}).get("moves", []))
                if "cg_draft" in st.session_state:
                    st.session_state["cg_draft"]["moves"] = st.session_state["cg_moves"]

                # orçamento unificado com o total gasto (mesmo limite de PP total)
                pp_cap_moves = int(pp_total)
                acerto_pp_total = sum(_move_accuracy_pp(m) for m in st.session_state.get("cg_moves", []))
                pp_spent_moves_live = sum((m.get("pp_cost") or 0) for m in st.session_state.get("cg_moves", [])) + acerto_pp_total
                pp_spent_total_live = float(pp_abilities) + float(pp_defenses) + float(pp_skills) + float(pp_advantages) + float(pp_spent_moves_live)

                cA, cB, cC = st.columns(3)
                cA.metric("PP total gasto", int(pp_spent_total_live))
                cB.metric("Limite total", pp_cap_moves)
                cC.metric("Slots de golpes", len(st.session_state.get("cg_moves", [])))
                st.progress(min(1.0, float(pp_spent_total_live) / float(max(1, pp_cap_moves))))

                if pp_spent_total_live > pp_cap_moves:
                    st.error("Você estourou o limite de PP total. Ajuste golpes, acerto ou outros gastos.")

                sub_tabs = st.tabs([
                    "⚡ Adicionar rápido",
                    "🛠️ Criador completo (avançado)",
                    "📦 Lista & ajustes",
                ])

                # --------------------------
                # (A) Adicionar rápido
                # --------------------------
                with sub_tabs[0]:
                    st.markdown("#### 1) Sugestões pelo Pokémon (Bulbapedia)")
                    st.caption("Lista de golpes que o Pokémon pode aprender. Eu comparo com o banco e deixo 1 clique para adicionar.")

                    rank_default = st.slider(
                        "Rank padrão para adicionar",
                        1, 20,
                        min(10, max(1, int(np_ or 1))),
                        key="cg_quick_rank_default"
                    )

                    if db_moves_guided is not None and isinstance(pjson, dict):
                        api_moves = _pokeapi_parse_move_names(pjson)
                        matched = []
                        for nm in api_moves[:120]:  # limita para não ficar pesado
                            mv = _try_match_move_in_db(db_moves_guided, nm)
                            if mv:
                                matched.append(mv)

                        # remove duplicados por nome
                        seen = set()
                        uniq = []
                        for mv in matched:
                            k = _norm(mv.name)
                            if k not in seen:
                                seen.add(k)
                                uniq.append(mv)

                        if uniq:
                            # tenta priorizar STAB (mesmo tipo do Pokémon)
                            stab = [m for m in uniq if _norm(getattr(m, "tipo", "")) in {_norm(t) for t in (types or [])}]
                            other = [m for m in uniq if m not in stab]
                            uniq = stab + other

                            default_pick = [m.name for m in uniq[:8]]
                            pick = st.multiselect(
                                "Selecione e adicione (sugestões)",
                                options=[m.name for m in uniq[:60]],
                                default=default_pick[:4],
                                key="cg_quick_pick"
                            )

                            col1, col2 = st.columns(2)
                            with col1:
                                if st.button("➕ Adicionar selecionados", key="cg_quick_add_selected"):
                                    existing = {_norm(m.get("name", "")) for m in st.session_state.get("cg_moves", [])}
                                    added = 0
                                    for nm in pick:
                                        mv = db_moves_guided.get_by_name(nm)
                                        if not mv:
                                            continue
                                        if _norm(mv.name) in existing:
                                            continue

                                        pp_auto = None
                                        try:
                                            tmp = mv.pp_cost(int(rank_default))
                                            if isinstance(tmp, tuple):
                                                pp_auto = tmp[0]
                                        except Exception:
                                            pp_auto = None

                                        st.session_state["cg_moves"].append(
                                            _cg_confirm_move(
                                                mv,
                                                int(rank_default),
                                                pp_override=(int(pp_auto) if pp_auto is not None else None),
                                                accuracy=_default_accuracy_from_raw(mv),
                                            )
                                        )
                                        added += 1

                                    st.success(f"Adicionados: {added} golpe(s).")
                                    st.rerun()
                            with col2:
                                st.caption("Dica: para personalizar Reaction, Área, Extended, Resist — use a busca abaixo ou **Lista & ajustes**.")
                        else:
                            st.info("Não encontrei golpes do learnset no banco atual. Use a busca abaixo ou o assistente.")

                    st.divider()
                    st.markdown("#### 2) Buscar no banco (rápido)")
                    q = st.text_input("Digite parte do nome (ex.: thunder, protect)", key="cg_quick_search")
                    if db_moves_guided is None:
                        st.stop()

                    if q and len(q.strip()) >= 2:
                        results = db_moves_guided.search_by_name_prefix(q.strip())[:20]
                        if not results:
                            sugg = db_moves_guided.suggest_by_description(q.strip(), top_k=8)
                            results = [mv for (mv, _s) in sugg]

                        # ✅ BLINDADO: remove duplicados por nome (evita StreamlitDuplicateElementKey)
                        if results:
                            seen = set()
                            uniq = []
                            for mv in results:
                                k = _norm(getattr(mv, "name", "") or "")
                                if not k or k in seen:
                                    continue
                                seen.add(k)
                                uniq.append(mv)
                            results = uniq                       

                        if results:
                            existing = {_norm(m.get("name", "")) for m in st.session_state.get("cg_moves", [])}
                            for mv in results:
                                key_base = f"{_norm(mv.name)}_{rank_default}"
                                with st.container(border=True):
                                    hc1, hc2 = st.columns([7, 3])
                                    with hc1:
                                        st.write(f"**{mv.name}** — {getattr(mv,'tipo','—')} / {getattr(mv,'categoria','—')}")
                                        bullets = _summarize_build(getattr(mv, "build", ""))
                                        if bullets:
                                            st.caption(" • ".join(bullets))
                                    with hc2:
                                        if _norm(mv.name) in existing:
                                            st.button("✅ Já na ficha", disabled=True, key=f"cg_quick_already_{key_base}")

                                    if _norm(mv.name) not in existing:
                                        with st.expander("⚙️ Personalizar antes de adicionar", expanded=False):
                                            pp_auto = None
                                            try:
                                                tmp = mv.pp_cost(int(rank_default))
                                                if isinstance(tmp, tuple):
                                                    pp_auto = tmp[0]
                                            except Exception:
                                                pp_auto = None

                                            pc1, pc2, pc3 = st.columns(3)
                                            with pc1:
                                                cust_rank = st.number_input("Rank", min_value=1, max_value=30, value=int(rank_default), step=1, key=f"cg_q_rank_{key_base}")
                                                cust_pp_val = int(pp_auto) if pp_auto is not None else 1
                                                try:
                                                    tmp2 = mv.pp_cost(int(cust_rank))
                                                    if isinstance(tmp2, tuple) and tmp2[0] is not None:
                                                        cust_pp_val = int(tmp2[0])
                                                except Exception:
                                                    pass
                                                cust_pp = st.number_input("PP", min_value=1, value=cust_pp_val, step=1, key=f"cg_q_pp_{key_base}")
                                            with pc2:
                                                cust_is_reaction = st.checkbox("⚡ É uma Reaction", value=False, key=f"cg_q_react_{key_base}")
                                                cust_ranged = st.checkbox("🎯 Ranged", value=bool(getattr(mv, "ranged", False)), key=f"cg_q_ranged_{key_base}")
                                                cust_perception = st.checkbox("👁️ Perception Area", value=bool(getattr(mv, "perception_area", False)), key=f"cg_q_perc_{key_base}")
                                            with pc3:
                                                cust_area = st.selectbox("Área", ["—", "Burst", "Cone", "Line", "Cloud"], index=0, key=f"cg_q_area_{key_base}")
                                                cust_area_ext = st.checkbox("Área Extended (+1r)", value=False, key=f"cg_q_areaext_{key_base}")
                                                cust_resist = st.selectbox("Resiste via", ["— (padrão)", "Thg", "Dodge", "Will", "Fortitude"], index=0, key=f"cg_q_resist_{key_base}")

                                            if st.button("➕ Adicionar com personalização", key=f"cg_quick_add_{key_base}", type="primary"):
                                                move_entry = _cg_confirm_move(mv, int(cust_rank), pp_override=int(cust_pp))
                                                move_entry["meta"]["is_reaction"] = bool(cust_is_reaction)
                                                move_entry["meta"]["ranged"] = bool(cust_ranged)
                                                move_entry["meta"]["perception_area"] = bool(cust_perception)
                                                if cust_area != "—":
                                                    move_entry["meta"]["area_type"] = cust_area
                                                    move_entry["meta"]["area_extended"] = bool(cust_area_ext)
                                                if cust_resist != "— (padrão)":
                                                    move_entry["meta"]["resist_stat"] = cust_resist
                                                # Aplica extras na build
                                                build_extra = move_entry.get("build", "")
                                                if cust_area != "—":
                                                    area_rank = 2 if cust_area_ext else 1
                                                    if f"[Area:" not in build_extra:
                                                        build_extra = build_extra.rstrip("; ") + f" [Area: {cust_area} {area_rank}]"
                                                if cust_is_reaction and "[Reaction]" not in build_extra:
                                                    build_extra = build_extra.rstrip("; ") + " [Reaction]"
                                                if cust_resist not in ("— (padrão)", "Thg") and "(Resisted by" not in build_extra:
                                                    build_extra = build_extra.rstrip("; ") + f" (Resisted by {cust_resist})"
                                                move_entry["build"] = build_extra.strip()
                                                st.session_state["cg_moves"].append(move_entry)
                                                st.rerun()
                        else:
                            st.info("Nada encontrado. Tente outra palavra ou use o Assistente para montar o efeito.")
                    else:
                        st.caption("Informe pelo menos 2 caracteres para buscar.")

                # --------------------------
                with sub_tabs[1]:
                    if _GB_AVAILABLE:
                        st.caption(
                            "Golpe Builder completo (MM3e): busca, modo simples/avançado, preview em tempo real, "
                            "templates e histórico. Golpes confirmados vão direto para sua lista acima."
                        )
                        trainer_name_inline = st.session_state.get("trainer_name", "Treinador")
                        _gb_render_builder(
                            excel_path=_resolve_asset_path("golpes_pokemon_MM_reescritos.xlsx"),
                            state_key_prefix="cg_gb_inline",
                            return_to_view=None,
                            trainer_name=trainer_name_inline,
                        )
                    else:
                        st.caption("Modo avançado: use a tela completa de criação e edição de golpes.")
                        show_full = st.checkbox(
                            "Abrir o criador completo aqui (avançado)",
                            value=False, key="cg_show_full_creator_inline"
                        )
                        if show_full:
                            render_move_creator(
                                excel_path=_resolve_asset_path("golpes_pokemon_MM_reescritos.xlsx"),
                                state_key_prefix="cg_moves_inline",
                                return_to_view=None,
                            )
                        else:
                            st.info("Marque a caixa acima para abrir o criador completo nesta aba.")
                        if st.button("Abrir pelo menu Criação de Golpes", key="cg_go_creator_menu"):
                            st.session_state["cg_view"] = "moves"
                            st.session_state["cg_return_to"] = "guided"
                            st.rerun()

                # --------------------------
                # (D) Lista & ajustes
                # --------------------------
                with sub_tabs[2]:
                    st.markdown("#### 📦 Golpes confirmados nesta ficha")
                
                    # CSS local (não interfere no resto do app)
                    st.markdown("""
                    <style>
                      .mv-kpi { opacity: .92; font-size: 0.92rem; }
                      .mv-muted { opacity: .75; }
                      .mv-chips { display:flex; flex-wrap:wrap; gap:6px; margin-top:6px; }
                      .mv-chip {
                        display:inline-flex; align-items:center; gap:6px;
                        padding: 2px 10px; border-radius: 999px;
                        border: 1px solid rgba(255,255,255,0.14);
                        background: rgba(255,255,255,0.06);
                        font-size: 0.82rem;
                      }
                      .mv-warn {
                        border-left: 4px solid rgba(255, 180, 0, 0.9);
                        padding-left: 10px;
                      }
                    </style>
                    """, unsafe_allow_html=True)
                
                    if not st.session_state.get("cg_moves"):
                        st.info("Nenhum golpe confirmado ainda. Use as abas acima para adicionar e ajustar.")
                    else:
                        # Stats atuais (mesma lógica que você já usa)
                        stats_fallback = st.session_state.get("cg_draft", {}).get("stats", {})
                        stats_now = {
                            "stgr": int(st.session_state.get("cg_stgr", stats_fallback.get("stgr", 0)) or 0),
                            "int": int(st.session_state.get("cg_int", stats_fallback.get("int", 0)) or 0),
                            "dodge": int(st.session_state.get("cg_dodge", stats_fallback.get("dodge", 0)) or 0),
                            "parry": int(st.session_state.get("cg_parry", stats_fallback.get("parry", 0)) or 0),
                            "thg": int(st.session_state.get("cg_thg", stats_fallback.get("thg", 0)) or 0),
                            "fortitude": int(st.session_state.get("cg_fortitude", stats_fallback.get("fortitude", 0)) or 0),
                            "will": int(st.session_state.get("cg_will", stats_fallback.get("will", 0)) or 0),
                        }
                        np_value = int(st.session_state.get("cg_np", 0) or 0)
                        target_total = 2 * int(np_value)
                
                        # -------- Toolbar: busca / filtros / ordenação / compacto --------
                        all_moves = st.session_state["cg_moves"]
                
                        # prepara lista com índices reais (pra editar/remover sem bagunçar)
                        entries = []
                        all_tags_set = set()
                        for real_idx, m in enumerate(all_moves):
                            ui_id = _mv_ui_id(m)
                            name = str(m.get("name") or "Golpe")
                            tags = _mv_tags_from_move(m)
                            all_tags_set.update(tags)
                
                            meta = m.setdefault("meta", {})
                            accuracy = int(m.get("accuracy", 0) or 0)
                            pp_here = int(m.get("pp_cost") or 0) if m.get("pp_cost") is not None else 0
                            base_rank = int(m.get("rank", 0) or 0)
                            based_label, stat_val = _move_stat_value(m, stats_now)
                            final_rank = base_rank + int(stat_val)
                            mod_acerto = accuracy + final_rank
                
                            entries.append({
                                "real_idx": real_idx,
                                "ui_id": ui_id,
                                "name": name,
                                "tags": tags,
                                "accuracy": accuracy,
                                "pp": pp_here,
                                "base_rank": base_rank,
                                "final_rank": final_rank,
                                "mod_acerto": mod_acerto,
                                "based_label": based_label,
                                "stat_val": int(stat_val),
                            })
                
                        tcol1, tcol2, tcol3, tcol4 = st.columns([4, 3, 3, 2], vertical_alignment="center")
                        with tcol1:
                            q = st.text_input("Buscar golpe", value="", placeholder="Digite parte do nome…", key="mv_list_q")
                        with tcol2:
                            tag_opts = sorted(all_tags_set)
                            tag_pick = st.multiselect("Filtrar por tags", options=tag_opts, default=[], key="mv_list_tags")
                        with tcol3:
                            sort_mode = st.selectbox(
                                "Ordenar",
                                ["Nome (A→Z)", "PP (maior→menor)", "Rank final (maior→menor)", "Acerto (maior→menor)", "Mod. acerto (maior→menor)"],
                                index=0,
                                key="mv_list_sort",
                            )
                        with tcol4:
                            compact = st.toggle("Compacto", value=False, key="mv_list_compact")
                
                        # aplica busca/filtro
                        qn = (q or "").strip().lower()
                        if qn:
                            entries = [e for e in entries if qn in e["name"].lower()]
                
                        if tag_pick:
                            need = set(tag_pick)
                            entries = [e for e in entries if need.issubset(set(e["tags"]))]
                
                        # ordena
                        entries = sorted(
                            entries,
                            key=lambda e: _mv_sort_key(sort_mode, e["name"], e["pp"], e["final_rank"], e["accuracy"], e["mod_acerto"])
                        )
                
                        # -------- Render dos cards --------
                        for order_i, e in enumerate(entries, start=1):
                            m_gv = all_moves[e["real_idx"]]
                            ui_id = e["ui_id"]
                            meta = m_gv.setdefault("meta", {})
                
                            # Descrição: custom salva na ficha > banco/Excel > derivada da build
                            _custom_desc = m_gv.get("descricao") or m_gv.get("meta", {}).get("descricao")
                            desc_text = _custom_desc or _mv_desc_for(db_moves_guided, m_gv.get("name"))
                            # Preview: usa descrição customizada ou derivada de build se não houver desc_text
                            if desc_text:
                                preview = _mv_preview(desc_text, 200)
                            elif m_gv.get("build"):
                                _bul = _summarize_build(m_gv.get("build", ""))
                                preview = " • ".join(_bul[:5]) if _bul else "Criado via Criador Completo."
                            else:
                                preview = "Descrição não disponível."

                            # limites
                            acc_limit = _move_accuracy_limit(m_gv, np_value, stats_now)
                            current_acc = int(m_gv.get("accuracy", 0) or 0)
                            base_rank = int(m_gv.get("rank", 0) or 0)
                            based_label, stat_val = _move_stat_value(m_gv, stats_now)
                            final_rank = base_rank + int(stat_val)
                            mod_acerto = current_acc + final_rank

                            # card
                            with st.container(border=True):
                                left, right = st.columns([7, 3], vertical_alignment="top")

                                with left:
                                    st.markdown(f"### {order_i}. {m_gv.get('name','Golpe')}")
                                    # Chips: tipo do golpe + tags
                                    _tipo_chip = f'<span class="mv-chip">🏷️ {m_gv["tipo"]}</span>' if m_gv.get("tipo") else ""
                                    _tag_chips = "".join([f'<span class="mv-chip">{t}</span>' for t in e["tags"]])
                                    if _tipo_chip or _tag_chips:
                                        st.markdown(f'<div class="mv-chips">{_tipo_chip}{_tag_chips}</div>', unsafe_allow_html=True)

                                    # linha de KPIs (compacta, sem poluir)
                                    pp_here = m_gv.get("pp_cost")
                                    pp_show = "—" if pp_here is None else str(pp_here)

                                    if based_label == "—":
                                        rank_label = f"Rank final {final_rank}"
                                    else:
                                        rank_label = f"Rank final {final_rank} (Base {base_rank} + {based_label} {stat_val})"

                                    st.markdown(
                                        f"<div class='mv-kpi'>"
                                        f"<b>{rank_label}</b> &nbsp;•&nbsp; "
                                        f"<b>PP:</b> {pp_show} &nbsp;•&nbsp; "
                                        f"<b>Acerto:</b> {current_acc} &nbsp;•&nbsp; "
                                        f"<b>Mod. acerto:</b> {mod_acerto} <span class='mv-muted'>(2×NP = {target_total})</span>"
                                        f"</div>",
                                        unsafe_allow_html=True,
                                    )

                                    # avisos de leitura
                                    warn_lines = []
                                    if current_acc > int(acc_limit):
                                        warn_lines.append(f"Acerto acima do sugerido ({current_acc} > {acc_limit}).")
                                    if mod_acerto > target_total:
                                        warn_lines.append(f"Mod. acerto acima de 2×NP ({mod_acerto} > {target_total}).")

                                    if warn_lines:
                                        st.markdown(
                                            "<div class='mv-warn mv-muted'>⚠️ " + " ".join(warn_lines) + "</div>",
                                            unsafe_allow_html=True,
                                        )

                                    if not compact:
                                        st.caption(preview)

                                    # Detalhes
                                    with st.expander("📘 Detalhes", expanded=False):
                                        st.markdown("**Descrição**")
                                        if desc_text:
                                            st.write(desc_text)
                                        elif m_gv.get("build"):
                                            st.caption("_(criado via Criador Completo — descrição derivada da build)_")
                                            _bul2 = _summarize_build(m_gv.get("build", ""))
                                            if _bul2:
                                                st.write(" • ".join(_bul2[:8]))
                                        else:
                                            st.caption("Descrição não disponível.")
                                        if m_gv.get("build"):
                                            st.markdown("**Build MM3e**")
                                            st.code(m_gv["build"], language="text")

                                    # ── Editar nome / build / descrição ──────────────────
                                    with st.expander("✏️ Editar golpe", expanded=False):
                                        _en = st.text_input(
                                            "Nome", value=m_gv.get("name", ""),
                                            key=f"mv_edit_name_{ui_id}",
                                        )
                                        _eb = st.text_area(
                                            "Build string (MM3e)",
                                            value=m_gv.get("build", ""),
                                            height=80,
                                            key=f"mv_edit_build_{ui_id}",
                                            help="Cole ou edite a build MM3e diretamente aqui.",
                                        )
                                        _ed = st.text_area(
                                            "Descrição personalizada",
                                            value=m_gv.get("descricao", ""),
                                            height=60,
                                            key=f"mv_edit_desc_{ui_id}",
                                            help="Se preenchida, substitui a descrição do banco/Excel.",
                                        )
                                        if st.button("💾 Salvar edição", key=f"mv_edit_save_{ui_id}", type="primary"):
                                            if _en.strip():
                                                m_gv["name"] = _en.strip()
                                            m_gv["build"] = _eb.strip()
                                            m_gv["descricao"] = _ed.strip()
                                            st.rerun()
                
                                with right:
                                    # Ajustes (mantém suas funções)
                                    cur_rank = int(m_gv.get("rank", 1) or 1)
                                    new_rank = st.number_input(
                                        "Rank",
                                        min_value=1,
                                        max_value=50,
                                        value=int(cur_rank),
                                        step=1,
                                        key=f"mv_rank_{ui_id}",
                                    )
                                    if st.button("Definir rank", key=f"mv_set_rank_{ui_id}", use_container_width=True):
                                        m_gv["rank"] = int(new_rank)
                                        pp_recalc, _why = _cg_recalculate_pp(m_gv, int(new_rank), db_moves_guided)
                                        if pp_recalc is not None:
                                            m_gv["pp_cost"] = int(pp_recalc)
                
                                        # atualiza build renderizado quando possível (mesmo comportamento atual)
                                        if db_moves_guided:
                                            try:
                                                mv_db = db_moves_guided.get_by_name(str(m_gv.get("name") or ""))
                                            except Exception:
                                                mv_db = None
                                            if mv_db:
                                                try:
                                                    m_gv["build"] = mv_db.render_build(int(new_rank))
                                                except Exception:
                                                    pass
                                        st.rerun()
                
                                    # Acerto (mantém limite sugerido)
                                    acc_limit = _move_accuracy_limit(m_gv, np_value, stats_now)
                                    current_acc = int(m_gv.get("accuracy", 0) or 0)
                                    safe_max = max(int(acc_limit), int(current_acc))
                                    new_acc = st.number_input(
                                        "Acerto",
                                        min_value=0,
                                        max_value=int(safe_max),
                                        value=int(current_acc),
                                        step=1,
                                        key=f"mv_acc_{ui_id}",
                                    )
                                    st.caption(f"Limite sugerido: {acc_limit}")
                                    if st.button("Definir acerto", key=f"mv_set_acc_{ui_id}", use_container_width=True):
                                        m_gv["accuracy"] = int(new_acc)
                                        st.rerun()
                
                                    # --- Propriedades do golpe ---
                                    _meta_gv = m_gv.setdefault("meta", {})

                                    # Reaction
                                    is_reaction = st.checkbox(
                                        "⚡ Reaction",
                                        value=bool(_meta_gv.get("is_reaction", False)),
                                        key=f"mv_reaction_{ui_id}",
                                        help="Golpe usado como reação (não consome a ação do turno)",
                                    )
                                    if is_reaction != bool(_meta_gv.get("is_reaction", False)):
                                        _meta_gv["is_reaction"] = bool(is_reaction)

                                    # Ranged
                                    is_ranged = st.checkbox(
                                        "🎯 Ranged",
                                        value=bool(_meta_gv.get("ranged", False)),
                                        key=f"mv_ranged_{ui_id}",
                                    )
                                    if is_ranged != bool(_meta_gv.get("ranged", False)):
                                        _meta_gv["ranged"] = bool(is_ranged)

                                    # Perception Area
                                    area_checked = st.checkbox(
                                        "👁️ Perception Area",
                                        value=bool(_meta_gv.get("perception_area", False)),
                                        key=f"mv_area_{ui_id}",
                                    )
                                    if area_checked != bool(_meta_gv.get("perception_area", False)):
                                        _meta_gv["perception_area"] = bool(area_checked)

                                    # Tipo de área
                                    _area_opts = ["—", "Burst", "Cone", "Line", "Cloud"]
                                    _cur_area = str(_meta_gv.get("area_type") or "—")
                                    _cur_area_idx = _area_opts.index(_cur_area) if _cur_area in _area_opts else 0
                                    new_area_type = st.selectbox(
                                        "Tipo de área",
                                        _area_opts,
                                        index=_cur_area_idx,
                                        key=f"mv_areatype_{ui_id}",
                                    )
                                    if new_area_type != _cur_area:
                                        _meta_gv["area_type"] = new_area_type if new_area_type != "—" else None

                                    # Extended Area
                                    area_ext = st.checkbox(
                                        "Área Extended (+1r)",
                                        value=bool(_meta_gv.get("area_extended", False)),
                                        key=f"mv_areaext_{ui_id}",
                                        help="Compra 1 rank extra de área (Extended Area)",
                                        disabled=(new_area_type == "—"),
                                    )
                                    if area_ext != bool(_meta_gv.get("area_extended", False)):
                                        _meta_gv["area_extended"] = bool(area_ext)

                                    # Resist stat override
                                    _resist_opts = ["— (padrão)", "Thg", "Dodge", "Will", "Fortitude"]
                                    _cur_resist = str(_meta_gv.get("resist_stat") or "— (padrão)")
                                    _cur_resist_idx = _resist_opts.index(_cur_resist) if _cur_resist in _resist_opts else 0
                                    new_resist = st.selectbox(
                                        "Resiste via",
                                        _resist_opts,
                                        index=_cur_resist_idx,
                                        key=f"mv_resist_{ui_id}",
                                        help="Substitui a defesa usada para resistir este golpe",
                                    )
                                    if new_resist != _cur_resist:
                                        _meta_gv["resist_stat"] = new_resist if new_resist != "— (padrão)" else None

                                    # Tipo do golpe (pokémon type)
                                    _tipo_opts = ["—", "Normal", "Fire", "Water", "Electric", "Grass",
                                                  "Ice", "Fighting", "Poison", "Ground", "Flying",
                                                  "Psychic", "Bug", "Rock", "Ghost", "Dragon",
                                                  "Dark", "Steel", "Fairy"]
                                    _cur_tipo = str(m_gv.get("tipo") or "—")
                                    _cur_tipo_idx = _tipo_opts.index(_cur_tipo) if _cur_tipo in _tipo_opts else 0
                                    new_tipo = st.selectbox(
                                        "Tipo do golpe",
                                        _tipo_opts,
                                        index=_cur_tipo_idx,
                                        key=f"mv_tipo_{ui_id}",
                                        help="Tipo Pokémon do golpe (exibido como chip no card)",
                                    )
                                    if new_tipo != _cur_tipo:
                                        m_gv["tipo"] = new_tipo if new_tipo != "—" else None

                                    # Se o golpe está sem PP, mantém fix manual
                                    if m_gv.get("pp_cost") is None:
                                        pp_fix = st.number_input(
                                            "PP do golpe",
                                            min_value=1,
                                            value=1,
                                            step=1,
                                            key=f"mv_fix_pp_{ui_id}",
                                        )
                                        if st.button("Definir PP", key=f"mv_fix_pp_btn_{ui_id}", use_container_width=True):
                                            m_gv["pp_cost"] = int(pp_fix)
                                            st.rerun()
                
                                    if st.button("🗑️ Remover", key=f"mv_rm_{ui_id}", use_container_width=True):
                                        st.session_state["cg_moves"].pop(e["real_idx"])
                                        st.rerun()
                
                        # Totais (mantém seu cálculo)
                        acerto_pp_total = sum(_move_accuracy_pp(m) for m in st.session_state.get("cg_moves", []))
                        pp_spent_moves_live = sum((m.get("pp_cost") or 0) for m in st.session_state.get("cg_moves", [])) + acerto_pp_total
                        st.info(f"PP de Acerto (soma): {acerto_pp_total}")
                        st.info(f"PP gastos em Golpes (incluindo acerto): **{int(pp_spent_moves_live)}** / **{pp_cap_moves}**")
                
                    # atualiza para a aba de revisão (mantém seu fluxo)
                    pp_moves = pp_spent_moves_live if st.session_state.get("cg_moves") else 0


                # atualiza para a aba de revisão
                pp_moves = pp_spent_moves_live
            with tabs[4]:
                st.markdown("### 🧾 Revisão de PP")
                # Arquétipo aplicado (opcional)
                arch = st.session_state.get("cg_viab_selected")
                if isinstance(arch, dict):
                    st.markdown("### 🧩 Arquétipo aplicado (Viabilidade)")
                    st.write(f"**Estratégia:** {arch.get('code') or '—'}")
                    if arch.get("ability"):
                        st.write(f"**Habilidade sugerida:** {arch.get('ability')}")
                    if arch.get("partners"):
                        st.write("**Parceiros sugeridos:** " + ", ".join(arch.get("partners") or []))
                
                pp_spent_total = float(pp_abilities) + float(pp_defenses) + float(pp_skills) + float(pp_advantages) + float(pp_moves)
                st.markdown(f"""
                    <div class="cg-card">
                        <div class="cg-title">Resumo de PP</div>
                        <div>Abilities: {pp_abilities} | Defesas: {pp_defenses} | Skills: {pp_skills} | Adv: {pp_advantages} | Golpes: {pp_moves}</div>
                        <hr/><strong>Total gasto:</strong> {pp_spent_total} / {pp_total}
                    </div>
                """, unsafe_allow_html=True)
    
                if pp_spent_total > pp_total: st.warning("PP total ultrapassado.")
                else: st.success("PP total dentro do limite. ✅")
    
                pdf_bytes = build_sheet_pdf(
                    pname=pname, np_=np_, types=types, abilities=chosen_abilities,
                    stats={"stgr": int(stgr), "intellect": int(intellect), "dodge": int(dodge), "parry": int(parry), "fortitude": int(fortitude), "will": int(will)},
                    chosen_adv=chosen_adv, moves=st.session_state.get("cg_moves", [])
                )
    
                st.download_button("⬇️ Exportar PDF", data=pdf_bytes, file_name=f"ficha_{pname}_{np_}.pdf", mime="application/pdf")
    
                if st.button("☁️ Salvar ficha na Nuvem", key="btn_save_sheet_cloud"):
                    db_fs, bkt_fs = init_firebase()
                    skills_payload = []
                    for sk_n, sk_r in st.session_state.get("cg_skills", {}).items():
                        if int(sk_r) > 0:
                            skills_payload.append({"name": sk_n, "ranks": int(sk_r)})
                    for row_sk in st.session_state.get("cg_skill_custom", []):
                        if row_sk.get("name") and int(row_sk.get("ranks", 0)) > 0:
                            skills_payload.append({"name": row_sk["name"], "ranks": int(row_sk["ranks"])})
                
                    # =========================
                    # FIX: resolve ID real antes de salvar (evita salvar como 0)
                    # =========================
                    pid_save = 0
                
                    try:
                        # tenta resolver pelo nome (padrão EXT), ex: "EXT:Zoroark"
                        resolved = resolve_pokemon_pid(df, pname)
                        if resolved is not None and str(resolved).isdigit():
                            pid_save = int(resolved)
                    except Exception:
                        pid_save = 0
                
                    # fallback: usa o pid atual se for numérico
                    if pid_save == 0:
                        try:
                            if pid is not None and str(pid).isdigit():
                                pid_save = int(pid)
                        except Exception:
                            pid_save = 0
                
                    payload_fs = {
                        "pokemon": {"id": int(pid_save), "name": pname, "types": types, "abilities": chosen_abilities},
                        "np": int(np_), "pp_budget_total": int(pp_total), "pp_spent_total": float(pp_spent_total),
                        "stats": {"stgr": int(stgr), "int": int(intellect), "dodge": int(dodge), "parry": int(parry), "thg": int(thg), "fortitude": int(fortitude), "will": int(will)},
                        "advantages": chosen_adv,
                        "skills": skills_payload,
                        "moves": st.session_state.get("cg_moves", [])
                    }
                
                    sid_fs, _ = save_sheet_with_pdf(
                        db=db_fs,
                        bucket=bkt_fs,
                        trainer_name=trainer_name,
                        sheet_payload=payload_fs,
                        pdf_bytes=pdf_bytes,
                        sheet_id=st.session_state.get("cg_edit_sheet_id")
                    )
                    st.success(f"✅ Salva! ID: {sid_fs}")
                    st.session_state["cg_edit_sheet_id"] = None
    
    
            if st.button("⬅️ Voltar"):
                st.session_state["cg_view"] = "menu"
                st.rerun()
    
        else:
            # Caso não tenha nome digitado
            st.info("💡 Digite o nome do Pokémon acima para começar a gerar a ficha.")
            if st.button("⬅️ Voltar ao Menu", key="btn_back_empty"):
                st.session_state["cg_view"] = "menu"
                st.rerun()


# =================
# MINHAS FICHAS
# =================
elif page == "Minhas Fichas":
    st.title("📚 Minhas Fichas")
    st.caption("Veja e gerencie as fichas salvas na nuvem.")

    with st.expander("📥 Importar ficha via PDF", expanded=False):
        st.caption("Envie um PDF exportado aqui para recriar a ficha automaticamente.")
        pdf_upload = st.file_uploader("Selecionar PDF da ficha", type=["pdf"], key="sheet_pdf_import")
        if pdf_upload and st.button("📥 Importar PDF", key="sheet_pdf_import_btn"):
            try:
                imported_sheet = parse_sheet_pdf(pdf_upload.getvalue())
                apply_imported_sheet_to_session(imported_sheet)
                st.session_state["cg_force_guided"] = True
                st.session_state["nav_to"] = "Criação Guiada de Fichas"
                st.success("Ficha importada! Abrindo na criação guiada...")
                st.rerun()
            except Exception as exc:
                st.error(f"Não consegui importar o PDF: {exc}")

    db, bucket = init_firebase()
    sheets = list_sheets(db, trainer_name)
    party_ids_set = {str(p) for p in (user_data.get("party") or [])}
    
    hub_pokemon_ids = []
    _seen_hub_ids = set()
    for raw_pid in (user_data.get("caught") or []):
        pid = _normalize_hub_pid(raw_pid)
        if not pid or pid in _seen_hub_ids:
            continue
        _seen_hub_ids.add(pid)
        hub_pokemon_ids.append(pid)
    
    def _sheet_pokemon_name(pid_value: str) -> str:
        pid = str(pid_value).strip()
        if not pid:
            return "Pokémon inválido"

        # Visitante fora da dex
        if pid.startswith("EXT:"):
            return pid.replace("EXT:", "", 1).strip() or "Pokémon Externo"
        if pid.startswith("PID:"):
            pid = pid.replace("PID:", "", 1).strip()

        # Busca no df da Pokédex usando _norm_pid para normalizar corretamente
        try:
            pid_norm = _norm_pid(pid)
            hit = df[df["Nº"].apply(_norm_pid) == pid_norm]
            if not hit.empty:
                name = str(hit.iloc[0].get("Nome") or "").strip()
                if name:
                    return name
        except Exception:
            pass

        return "Pokémon desconhecido"
    
    def _hub_pid_label(pid_value: str) -> str:
        pid_norm = str(pid_value).strip()
        pname_local = _sheet_pokemon_name(pid_norm)
        location_tag = "Equipe" if pid_norm in party_ids_set else "Box"
        return f"{pname_local} — {location_tag}"
        # se quiser aparecer SÓ o nome, sem "— Box/Equipe":
        # return pname_local

    if not sheets:
        st.info("Você ainda não tem fichas salvas.")
    else:
        for sheet in sheets:
            pokemon = sheet.get("pokemon", {})
            pname = pokemon.get("name", "Pokémon")
            pid = pokemon.get("id", "—")
            np_ = sheet.get("np", "—")
            updated_at = sheet.get("updated_at", "—")
            created_at = sheet.get("created_at", "—")
            sheet_id = sheet.get("_sheet_id")
            pdf_meta = sheet.get("pdf") or {}
            storage_path = pdf_meta.get("storage_path")

            with st.expander(f"🧾 {pname} (ID {pid}) — NP {np_}"):
                action_col1, action_col2 = st.columns(2)
                with action_col1:
                    if st.button("✏️ Editar ficha", key=f"edit_sheet_{sheet_id}"):
                        apply_sheet_to_session(sheet, sheet_id=sheet_id)
                        st.session_state["cg_force_guided"] = True
                        st.session_state["nav_to"] = "Criação Guiada de Fichas"
                        st.rerun()
                with action_col2:
                    if st.button("🗑️ Excluir ficha", key=f"delete_sheet_{sheet_id}"):
                        delete_sheet(db, bucket, trainer_name, sheet_id, storage_path=storage_path)
                        st.success("Ficha excluída com sucesso.")
                        st.rerun()
                st.write(f"**Atualizada em:** {updated_at}")
                st.write(f"**Criada em:** {created_at}")
                st.write(f"**PP Total:** {sheet.get('pp_budget_total', '—')}")

                linked_pid = _normalize_hub_pid(sheet.get("linked_pid", ""))

                if linked_pid:
                    st.write(f"**Associada a:** {_hub_pid_label(linked_pid)}")
                else:
                    st.caption("Sem Pokémon associado nesta ficha.")

                if hub_pokemon_ids:
                    assoc_options = [""] + hub_pokemon_ids
                    current_assoc = linked_pid if linked_pid in hub_pokemon_ids else ""
                    selected_assoc = st.selectbox(
                        "Associar ficha a um Pokémon do Trainer Hub",
                        options=assoc_options,
                        index=assoc_options.index(current_assoc),
                        format_func=lambda opt: "❌ Sem associação" if not opt else _hub_pid_label(opt),
                        key=f"sheet_assoc_select_{sheet_id}",
                    )

                    if st.button("💾 Salvar associação", key=f"sheet_assoc_save_{sheet_id}"):
                        payload = {"linked_pid": _normalize_hub_pid(selected_assoc) or None}
                        save_sheet_to_firestore(db, trainer_name, payload, sheet_id=sheet_id)
                        st.success("Associação atualizada.")
                        st.rerun()
                else:
                    st.caption("Capture Pokémon no Trainer Hub para habilitar associações.")

                moves = sheet.get("moves") or []
                if moves:
                    st.markdown("**Golpes:**")
                    for i, m in enumerate(moves, start=1):
                        st.write(f"{i}. {m.get('name', 'Golpe')} (Rank {m.get('rank', '—')})")
                else:
                    st.caption("Sem golpes registrados.")


                if storage_path:
                    if st.button("📄 Baixar PDF", key=f"download_pdf_{sheet.get('_sheet_id')}"):
                        blob = bucket.blob(storage_path)
                        pdf_bytes = blob.download_as_bytes()
                        st.download_button(
                            "Clique para baixar o PDF",
                            data=pdf_bytes,
                            file_name=f"ficha_{pname}_{pid}.pdf",
                            mime="application/pdf",
                            key=f"download_pdf_btn_{sheet.get('_sheet_id')}",
                        )


# PVP ARENA

elif page == "PvP – Arena Tática":
    st.title("⚔️ PvP – Arena Tática")
    st.caption(
        "Base multiplayer: criar/abrir arena, entrar por código, espectadores "
    )

    db, bucket = init_firebase()
    view = st.session_state.get("pvp_view", "lobby")
    rid = st.session_state.get("active_room_id")  # ✅ sempre existe ou None
    room = get_room(db, rid) if rid else None     # ✅ evita NameError
    role = get_role(room, trainer_name) if room else "spectator"
    is_player = role in ["owner", "challenger"]
    battle_sheets_map = {}
    if is_player:
        try:
            for sh in list_sheets(db, trainer_name) or []:
                p = (sh.get("pokemon") or {})
                pid = p.get("id")
                if pid is not None:
                    battle_sheets_map[str(pid)] = sh
    
                lpid = sh.get("linked_pid")
                if lpid:
                    battle_sheets_map[str(lpid)] = sh
        except Exception:
            battle_sheets_map = {}



    # =========================
    # VIEW: BATTLE (CÓDIGO CONSOLIDADO E CORRIGIDO)
    # =========================
    if view == "battle":
        if not rid or not room:
            st.session_state["pvp_view"] = "lobby"
            st.rerun()
    
        ensure_pvp_sync_listener(db, rid)
        sync_watchdog(db, rid)
        battle_watchdog(db, rid)

            
        # --- AQUI: INICIA O SISTEMA DE SYNC AUTOMÁTICO ---
        # Isso cria a thread que fica "dormindo" até o Firebase avisar de uma mudança.
        # -------------------------------------------------

        if "last_click_processed" not in st.session_state:
            st.session_state["last_click_processed"] = None
        if "placing_trainer" not in st.session_state:
            st.session_state["placing_trainer"] = None

        # --- 1. SINCRONIZAÇÃO DE DADOS ---
        current_party = user_data.get("party") or []
        db.collection("rooms").document(rid).collection("public_state").document("players").set(
            {trainer_name: current_party}, merge=True
        )

        if "stats" in user_data:
            # Prepara um dicionário estruturado para o merge funcionar direito
            # Estrutura: { "NomeTreinador": { "ID_Pokemon": { "stats": {...} } } }
            nested_update = {}
            
            for pid in current_party:
                is_shiny = pid in user_data.get("shinies", [])
                hub_stats = user_data["stats"].get(pid, {})
                saved_form = user_data.get("forms", {}).get(str(pid))
                if hub_stats:
                    # Se o treinador ainda não está no dicionário, cria
                    if trainer_name not in nested_update:
                        nested_update[trainer_name] = {}
                    
                    clean_stats = {
                        "dodge": int(hub_stats.get("dodge", 0)),
                        "parry": int(hub_stats.get("parry", 0)),
                        "will": int(hub_stats.get("will", 0)),
                        "fort": int(hub_stats.get("fort", 0)),
                        "thg": int(hub_stats.get("thg", 0)),
                    }
                    # Adiciona os dados do Pokémon
                    nested_update[trainer_name][str(pid)] = {
                        "stats": hub_stats,
                        "shiny": is_shiny,
                        "form": saved_form, # ✅ ENVIANDO PARA O BANCO
                        "updatedAt": str(datetime.now())
                    }
            
            if nested_update:
                # Agora o .set(..., merge=True) vai entender a estrutura aninhada corretamente!
                db.collection("rooms").document(rid).collection("public_state").document("party_states").set(nested_update, merge=True)

        # --- 2. CARREGAMENTO DO ESTADO ---
        state = get_state(db, rid)
        seed = state.get("seed")
        no_water_state = bool(state.get("noWater", False))
        
        all_pieces = state.get("pieces") or []
        seen_pids = state.get("seen") or []
        field_effects = state.get("effects") or []

        # --- 3. HELPERS LOCAIS ---
        ps_doc = db.collection("rooms").document(rid).collection("public_state").document("party_states").get()
        party_states_data = ps_doc.to_dict() or {}

        # --- FUNÇÃO DE LEITURA INTELIGENTE (CORRIGIDA) ---
        def get_poke_data(t_name, p_id):
            user_dict = party_states_data.get(t_name, {})
            p_data = user_dict.get(str(p_id), {})
            
            hp = p_data.get("hp", 6)
            cond = p_data.get("cond", [])
            stats = p_data.get("stats", {})
            shiny_status = p_data.get("shiny", False)
            saved_form = p_data.get("form", None)  # ✅ LER DO BANCO
            
            if t_name == trainer_name:
                # CORREÇÃO: Converte para string, remove espaços e trata vazio como '0'
                # Além disso, ignora a chave 'notes' que é texto puro.
                stats_is_bad = not stats or all(
                    int(str(v).strip() or 0) == 0 
                    for k, v in stats.items() if k != "notes"
                )
                
                if stats_is_bad:
                    if "stats" in user_data:
                        local_s = user_data["stats"].get(str(p_id)) or user_data["stats"].get(p_id)
                        if local_s:
                            stats = local_s
            
            return hp, cond, stats, shiny_status, saved_form
        
        def get_poke_display_name(pid):
            row = df[df['Nº'].astype(str) == str(pid)]
            if not row.empty: return row.iloc[0]['Nome']
            return str(pid)

        # Definição da Função de Renderização da Coluna (DEFINIDA ANTES DE USAR)
        # Definição da Função de Renderização da Coluna (MELHORADA)
        def render_player_column(p_name, p_label, is_me):

            # --- Painel visual (conceito) ---
            with st.container():
                st.markdown("<div class='pvp-team-panel-marker'></div>", unsafe_allow_html=True)

                if is_me:
                    photo_data = get_trainer_photo_thumb(user_data)
                    if photo_data:
                        st.markdown(
                            f"""### <span class='pvp-team-title'><img src='{photo_data}' class='pvp-mini'/> {p_label}</span>""",
                            unsafe_allow_html=True,
                        )
                    else:
                        st.markdown(f"### {p_label}")
                else:
                    st.markdown(f"### {p_label}")

                # Busca party e estado público
# Busca party e estado público
            p_doc_data = db.collection("rooms").document(rid).collection("public_state").document("players").get().to_dict() or {}
            party_list = p_doc_data.get(p_name, [])[:8] 
            
            # Variáveis de Estado de Ação
            moving_piece_id = st.session_state.get("moving_piece_id")
            placing_pid = st.session_state.get("placing_pid")
            placing_trainer = st.session_state.get("placing_trainer")
            
            state = get_state(db, rid)
            all_pieces = state.get("pieces") or []
            seen_pids = state.get("seen") or []
            
            if not party_list:
                st.caption("Aguardando...")
                return
        
            trainer_piece = next(
                (p for p in all_pieces if p.get("owner") == p_name and p.get("kind") == "trainer"),
                None,
            )
            p_pieces_on_board = [
                p for p in all_pieces if p.get("owner") == p_name and p.get("kind") != "trainer"
            ]

            if is_me:
                with st.container():
                    st.markdown("<div class='pvp-avatar-card-marker'></div>", unsafe_allow_html=True)
                    st.markdown("#### 🙂 Meu Avatar")
                avatar_choice, avatar_path = get_selected_trainer_avatar(user_data)
                if avatar_path:
                    st.image(avatar_path, width=96)
                else:
                    st.caption("Escolha um avatar na aba Meu Treinador.")

                is_busy = (moving_piece_id is not None) or (placing_pid is not None) or bool(placing_trainer)

                if placing_trainer:
                    st.info("Clique no mapa para posicionar seu avatar.")
                    if st.button("🔙 Cancelar avatar", key="cancel_place_trainer"):
                        st.session_state["placing_trainer"] = None
                        request_rerun("cancel_place_trainer")
                else:
                    if trainer_piece:
                        c_avatar_1, c_avatar_2, c_avatar_3 = st.columns(3)
                        with c_avatar_1:
                            if st.button("🚶 Mover", key="move_trainer", disabled=is_busy):
                                st.session_state["moving_piece_id"] = trainer_piece.get("id")
                                st.session_state["arena_pause_until"] = time.time() + 0.15
                                st.session_state["placing_pid"] = None
                                st.session_state["placing_trainer"] = None
                                request_rerun("move_select", force=True)
                        with c_avatar_2:
                            trainer_revealed = trainer_piece.get("revealed", True)
                            if st.button("👁️" if trainer_revealed else "✅", key="toggle_trainer"):
                                trainer_piece["revealed"] = not trainer_revealed
                                upsert_piece(db, rid, trainer_piece)
                                request_rerun("toggle_trainer")
                        with c_avatar_3:
                            if st.button("❌", key="remove_trainer"):
                                delete_piece(db, rid, trainer_piece.get("id"))
                                if st.session_state.get("moving_piece_id") == trainer_piece.get("id"):
                                    st.session_state["moving_piece_id"] = None
                                request_rerun("remove_trainer")
                    else:
                        if st.button("📍 Colocar avatar", key="place_trainer", disabled=not avatar_choice or is_busy):
                            st.session_state["placing_trainer"] = True
                            st.session_state["arena_pause_until"] = time.time() + 1.2
                            st.session_state["placing_pid"] = None
                            st.session_state["moving_piece_id"] = None
                            request_rerun("place_trainer_start", force=True)
            else:
                with st.container():
                    st.markdown("<div class='pvp-avatar-card-marker'></div>", unsafe_allow_html=True)
                    st.markdown("#### 🙂 Avatar")

                opponent_avatar = trainer_piece.get("avatar") if trainer_piece else None
                opponent_avatar_path = None
                if opponent_avatar:
                    avatar_lookup = build_trainer_avatar_lookup()
                    avatar_info = avatar_lookup.get(opponent_avatar) or {}
                    opponent_avatar_path = avatar_info.get("path")

                if opponent_avatar_path:
                    st.image(opponent_avatar_path, width=96)
                else:
                    st.caption("Avatar do oponente ainda não foi revelado no mapa.")

            # ==========================
            # UI compacta: lista + detalhes do selecionado
            # ==========================
            sel_key = f"pvp_sel_{rid}_{p_name}"
            if sel_key not in st.session_state:
                st.session_state[sel_key] = str(party_list[0]) if party_list else None
            selected_pid = st.session_state.get(sel_key)

            def _hp_color(hp_val: int) -> str:
                try:
                    v = int(hp_val)
                except Exception:
                    v = 0
                if v <= 0:
                    return "#64748b"  # slate
                if v <= 2:
                    return "#ef4444"  # red
                if v <= 4:
                    return "#f59e0b"  # amber
                return "#22c55e"      # green

            def _hp_bar_html(hp_val: int, max_hp: int = 6) -> str:
                try:
                    v = int(hp_val)
                except Exception:
                    v = 0
                v = max(0, min(max_hp, v))
                pct = int(round((v / max_hp) * 100)) if max_hp else 0
                col = _hp_color(v)
                return f"""<div style='display:flex;align-items:center;gap:10px;margin-top:4px;'>
                    <div style='flex:1;height:10px;background:rgba(148,163,184,0.25);border:1px solid rgba(148,163,184,0.22);border-radius:999px;overflow:hidden;'>
                      <div style='width:{pct}%;height:100%;background:{col};'></div>
                    </div>
                    <div style='min-width:48px;text-align:right;font-family:monospace;color:#e2e8f0;'>{v}/{max_hp}</div>
                </div>"""

            def _badges_html(conds: list[str], limit: int = 3) -> str:
                conds = conds or []
                show = conds[:limit]
                extra = max(0, len(conds) - len(show))
                chips = "".join([
                    f"<span class='pvp-badge' style='font-size:11px;padding:2px 8px;border-radius:999px;margin-right:6px;display:inline-block;'>{c}</span>"
                    for c in show
                ])
                if extra:
                    chips += f"<span class='pvp-badge' style='font-size:11px;padding:2px 8px;border-radius:999px;opacity:0.85;'>+{extra}</span>"
                return chips


            for i, pid in enumerate(party_list):
                # 1. Agora recuperamos 5 valores (incluindo p_form)
                cur_hp, cur_cond, cur_stats, is_shiny, p_form = get_poke_data(p_name, pid)
                
                is_on_map = any(str(p["pid"]) == str(pid) for p in p_pieces_on_board)
                p_obj = next((p for p in p_pieces_on_board if str(p["pid"]) == str(pid)), None)
                already_seen = str(pid) in seen_pids
                
                # Ícone de HP
                if cur_hp >= 5: hpi = "💚"
                elif cur_hp >= 3: hpi = "🟡"
                elif cur_hp >= 1: hpi = "🔴"
                else: hpi = "💀"
                
                # 2. Lógica da imagem com prioridade para a Forma
                if p_form:
                     sprite_url = pokemon_pid_to_image(f"EXT:{p_form}", mode="sprite", shiny=is_shiny)
                else:
                     sprite_url = pokemon_pid_to_image(pid, mode="sprite", shiny=is_shiny)
        
                # Checa se ESTE Pokémon específico está realizando uma ação
                is_moving_this = (p_obj and moving_piece_id == p_obj.get("id"))
                is_placing_this = (placing_pid == pid)

                # Estilo da Borda: Amarelo se movendo, Azul se colocando, Padrão caso contrário
                border_color = "#FFCC00" if is_moving_this else ("#38bdf8" if is_placing_this else None)
                
                # Container Visual
                with st.container():
                    card_cls = "pvp-mon-card-marker" + (" pvp-selected" if (str(pid) == str(selected_pid)) else "")
                    st.markdown(f"<div class='{card_cls}'></div>", unsafe_allow_html=True)
                    if is_me:
                        # --- Linha compacta (sempre visível) ---
                        is_busy = (moving_piece_id is not None) or (placing_pid is not None) or bool(placing_trainer)
                        is_selected = (str(pid) == str(selected_pid))
                        p_real_name = get_poke_display_name(pid)
                        loc_lbl = "No campo" if is_on_map else "Mochila"

                        # Destaque sutil se estiver em ação
                        if is_moving_this or is_placing_this:
                            border_color = "#FFCC00" if is_moving_this else "#38bdf8"
                            st.markdown(
                                f"<div style='height:2px;background:{border_color};border-radius:999px;margin:4px 0 10px 0;'></div>",
                                unsafe_allow_html=True,
                            )

                        c_row_img, c_row_mid, c_row_act = st.columns([0.85, 2.85, 1.30])

                        with c_row_img:
                            if cur_hp == 0:
                                st.markdown(
                                    f'<img src="{sprite_url}" style="width:56px;filter:grayscale(100%);opacity:0.65;">',
                                    unsafe_allow_html=True,
                                )
                                st.caption("FAINTED")
                            else:
                                st.image(sprite_url, width=56)

                        with c_row_mid:
                            st.markdown(
                                f"**{p_real_name}** <span style='color:#94a3b8;font-size:12px;margin-left:8px;'>{loc_lbl}</span>",
                                unsafe_allow_html=True,
                            )
                            if cur_cond:
                                st.markdown(_badges_html(cur_cond, limit=3), unsafe_allow_html=True)
                            else:
                                st.markdown(
                                    "<span style='color:#94a3b8;font-size:12px;'>Sem status negativos.</span>",
                                    unsafe_allow_html=True,
                                )
                            st.markdown(_hp_bar_html(cur_hp, 6), unsafe_allow_html=True)

                        with c_row_act:
                            a1, a2 = st.columns(2)
                            with a1:
                                if st.button(
                                    "🔍",
                                    key=f"sel_{p_name}_{pid}_{i}",
                                    help="Detalhes",
                                    use_container_width=True,
                                ):
                                    st.session_state[sel_key] = str(pid)
                                    st.rerun()

                            with a2:
                                main_icon = "🚶" if is_on_map else "📍"
                                main_help = "Mover" if is_on_map else "Colocar no campo"
                                disabled_main = is_busy and not (is_moving_this or is_placing_this)
                                if st.button(
                                    main_icon,
                                    key=f"main_{p_name}_{pid}_{i}",
                                    help=main_help,
                                    disabled=disabled_main or (cur_hp <= 0),
                                    use_container_width=True,
                                ):
                                    if is_on_map and p_obj:
                                        st.session_state["moving_piece_id"] = p_obj.get("id")
                                        st.session_state["arena_pause_until"] = time.time() + 1.2
                                        st.session_state["placing_pid"] = None
                                        st.session_state["placing_trainer"] = None
                                        st.rerun()
                                    else:
                                        st.session_state["placing_pid"] = pid
                                        st.session_state["arena_pause_until"] = time.time() + 1.2
                                        st.session_state["placing_effect"] = None
                                        st.session_state["moving_piece_id"] = None
                                        st.session_state["placing_trainer"] = None
                                        st.rerun()

                        # --- Detalhes (apenas do selecionado) ---
                        if is_selected:
                            st.markdown(
                                "<div style='height:8px;'></div>",
                                unsafe_allow_html=True,
                            )

                            if is_moving_this:
                                st.info("📍 Selecione um quadrado vazio no mapa.")
                                if st.button("🔙 Cancelar mover", key=f"cncl_move_{p_name}_{pid}_{i}"):
                                    st.session_state["moving_piece_id"] = None
                                    st.rerun()

                            elif is_placing_this:
                                st.info("📍 Clique no mapa para posicionar.")
                                if st.button("🔙 Cancelar posicionamento", key=f"cncl_place_{p_name}_{pid}_{i}"):
                                    st.session_state["placing_pid"] = None
                                    st.rerun()

                            else:
                                # Controles do Pokémon no campo (revelar/remover)
                                if is_on_map and p_obj:
                                    c_vis, c_del = st.columns(2)
                                    with c_vis:
                                        is_rev = p_obj.get("revealed", True)
                                        if st.button(
                                            "👁️" if is_rev else "✅",
                                            key=f"v_{p_name}_{pid}_{i}_detail",
                                            help="Revelar/Esconder",
                                            use_container_width=True,
                                        ):
                                            p_obj["revealed"] = not is_rev
                                            upsert_piece(db, rid, p_obj)
                                            if p_obj["revealed"]:
                                                mark_pid_seen(db, rid, pid)
                                            st.rerun()
                                    with c_del:
                                        if st.button(
                                            "❌",
                                            key=f"r_{p_name}_{pid}_{i}_detail",
                                            help="Remover do Mapa",
                                            use_container_width=True,
                                        ):
                                            delete_piece(db, rid, p_obj.get("id"))
                                            add_public_event(db, rid, "pokemon_removed", p_name, {"pid": pid})
                                            st.session_state["moving_piece_id"] = None
                                            st.rerun()

                                # Editor de HP / Status (mesmas chaves e callback)
                                st.slider(
                                    "HP",
                                    0,
                                    6,
                                    value=int(cur_hp),
                                    key=f"hp_{p_name}_{pid}_{i}",
                                    label_visibility="collapsed",
                                    on_change=update_poke_state_callback,
                                    args=(db, rid, p_name, pid, i),
                                )

                                st.multiselect(
                                    "Status",
                                    ["⚡", "❄️", "🔥", "💤", "☠️", "💓"],
                                    default=cur_cond,
                                    key=f"cond_{p_name}_{pid}_{i}",
                                    label_visibility="collapsed",
                                    on_change=update_poke_state_callback,
                                    args=(db, rid, p_name, pid, i),
                                )


                            st.markdown("<hr style='opacity:0.15;margin:10px 0;'>", unsafe_allow_html=True)

                    else:
                        # ==========================
                        # VISÃO DO OPONENTE (Corrigida)
                        # ==========================
                        piece_obj = next((p for p in p_pieces_on_board if str(p["pid"]) == str(pid)), None)
                        is_revealed = piece_obj.get("revealed", True) if piece_obj else False
                        
                        # Se estiver no campo e revelado, OU se já foi visto antes (está na bench mas conhecido)
                        show_full = (piece_obj and is_revealed) or already_seen
                        
                        status_txt = "(Mochila)" if not piece_obj else ("(Escondido)" if not is_revealed else "")
        
                        if show_full:
                            p_real_name = get_poke_display_name(pid)
                            c1, c2 = st.columns([1, 2])
                            
                            with c1: 
                                # ✅ CORREÇÃO 1: Imagem em escala de cinza se HP for 0
                                # Antes usava st.image direto, que sempre é colorido.
                                if cur_hp == 0:
                                    st.markdown(f'<img src="{sprite_url}" style="width:50px; filter:grayscale(100%); opacity:0.6;">', unsafe_allow_html=True)
                                else:
                                    st.image(sprite_url, width=50)
        
                            with c2:
                                st.markdown(f"**{p_real_name}**")
                                st.caption(f"{hpi} HP: {cur_hp}/6 {status_txt}")
                                
                                # ✅ CORREÇÃO 2: Exibir ícones de status para o oponente
                                if cur_cond:
                                    # Junta os ícones em uma string (ex: "🔥 ☠️")
                                    cond_str = " ".join(cur_cond)
                                    st.markdown(f"Status: {cond_str}")
        
                                if cur_hp == 0: 
                                    st.caption("**FAINTED**")
                        else:
                            # Pokémon desconhecido / escondido
                            c1, c2 = st.columns([1, 2])
                            with c1: st.image("https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg", width=40)
                            with c2: st.caption(f"??? {status_txt}")
        
        # --- 4. PREPARAÇÃO DE TIMES E VARIÁVEIS (UNIFICADO) ---
        owner_name = (room.get("owner") or {}).get("name", "Host")
        challengers = room.get("challengers") or []
        challenger_names = [c.get("name") for c in challengers]
        
        # Lista total de quem está na arena para a lógica de cores e calculadora
        all_players = [owner_name] + challenger_names

        # 1. MECÂNICA DE INTERFACE: Define quem é "Você" e as etiquetas (do código SEU)
        if trainer_name == owner_name:
            p1_name = owner_name
            p1_label = f"🎒 {owner_name} (Você)"
            viewer_is_p1 = True
            p2_name = challenger_names[0] if challenger_names else None
            p2_label = f"🆚 {p2_name}" if p2_name else "🆚 Aguardando..."
        elif trainer_name in challenger_names:
            p1_name = trainer_name
            p1_label = f"🎒 {trainer_name} (Você)"
            viewer_is_p1 = True 
            p2_name = owner_name
            p2_label = f"🆚 {owner_name}"
        else:
            p1_name = owner_name
            p1_label = f"🔴 {owner_name}"
            viewer_is_p1 = False
            p2_name = challenger_names[0] if challenger_names else None
            p2_label = f"🔵 {p2_name}" if p2_name else "🔵 Aguardando..."

        # 2. MECÂNICA DE MAPA E CALCULADORA: Processa peças e visibilidade (do código MEU)
        pieces_to_draw = []
        player_pieces_map = {name: [] for name in all_players}

        for p in all_pieces:
            if p.get("kind") == "trainer":
                if p.get("owner") == trainer_name:
                    pieces_to_draw.append(p)
                elif p.get("revealed", True):
                    pieces_to_draw.append(p)
                if p.get("owner") in player_pieces_map:
                    player_pieces_map[p.get("owner")].append(p)
                continue
            # 1. Captura a forma (5ª variável) do banco de dados
            hp_check, _, _, _, p_form = get_poke_data(p.get("owner"), p.get("pid"))
            
            p["status"] = "fainted" if hp_check == 0 else "active"
            
            # ✅ O TRUQUE: Se tiver forma salva, substituímos o PID temporariamente
            # Isso força o renderizador do mapa a baixar o sprite da forma correta (ex: EXT:lycanroc-midnight)
            if p_form:
                p["pid"] = f"EXT:{p_form}"

            # Lógica de Visibilidade: Dono vê tudo, outros veem apenas revelados 
            if p.get("owner") == trainer_name: 
                pieces_to_draw.append(p)
            elif p.get("revealed", True): 
                pieces_to_draw.append(p)
            
            # Popula o mapa para a calculadora encontrar os alvos corretamente
            if p.get("owner") in player_pieces_map:
                player_pieces_map[p.get("owner")].append(p)

        theme_key = room.get("theme", "cave_water")
        grid = int(state.get("gridSize") or room.get("gridSize") or 10)

        # --- 5. INTERFACE DO TOPO ---
        last_events = list_public_events(db, rid, limit=1)
        last_dice = next((e for e in last_events if e.get("type") == "dice"), None)

        top = st.columns([1, 1, 1, 1, 4])
        with top[0]:
            if st.button("⬅️ Lobby"):
                st.session_state["pvp_view"] = "lobby"
                st.rerun()
        with top[1]:
            if st.button("🔄 Atualizar"): st.rerun()
        with top[2]:
            if st.button("🎲 d20", disabled=not is_player): roll_die(db, rid, trainer_name, sides=20); st.rerun()
        with top[3]:
            if st.button("🎲 d6", disabled=not is_player): roll_die(db, rid, trainer_name, sides=6); st.rerun()
        with top[4]:
            if last_dice:
                pl = last_dice.get("payload", {})
                dice_line = f"🎲 {last_dice.get('by')}: <strong>{pl.get('result')}</strong> (d{pl.get('sides')})"
                dice_bg = "rgba(34, 197, 94, 0.18)"
                dice_border = "rgba(34, 197, 94, 0.5)"
            else:
                dice_line = "🎲 Aguardando rolagem..."
                dice_bg = "rgba(148, 163, 184, 0.15)"
                dice_border = "rgba(148, 163, 184, 0.35)"
            st.markdown(
                f"""
                <div style="
                    min-height: 42px;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    margin-bottom: 8px;
                    padding: 6px 10px;
                    border-radius: 10px;
                    border: 1px solid {dice_border};
                    background: {dice_bg};
                    font-weight: 600;">
                    {dice_line}
                </div>
                """,
                unsafe_allow_html=True,
            )
            # ✅ PEDIDO: Mostrar Código da Sala aqui em cima
            st.markdown(f"""
    <div style='display: flex; align-items: center; gap: 10px;'>
        <span style='font-family: "Press Start 2P"; font-size: 18px;'>🏟️ Arena:</span>
        <span style='font-family: "Press Start 2P"; font-size: 24px; color: #FFCC00; background: #333; padding: 5px 10px; border-radius: 5px;'>{rid}</span>
    </div>
    """, unsafe_allow_html=True) 

        # ==========================================
        # 🧮 6. CALCULADORA DE COMBATE
        # ==========================================
        battle_ref = db.collection("rooms").document(rid).collection("public_state").document("battle")
        battle_doc = battle_ref.get()
        
        # ✅ BLINDAGEM: update() falha se o doc não existe. Então criamos o doc na hora.
        if not battle_doc.exists:
            b_data = {"status": "idle", "logs": [], "initiative": {}}
            battle_ref.set(
                {
                    **b_data,
                    "createdAt": firestore.SERVER_TIMESTAMP,
        "updatedAt": firestore.SERVER_TIMESTAMP,
                    "updatedAt": firestore.SERVER_TIMESTAMP,
                },
                merge=True,
            )
        else:
            b_data = battle_doc.to_dict() or {}
            b_data.setdefault("status", "idle")
            b_data.setdefault("logs", [])
            b_data.setdefault("initiative", {})
        # =========================
        # 🎛️ HUD / Barra de Status
        # =========================
        b_status = (b_data.get("status") or "idle")
        b_attacker = b_data.get("attacker") or "—"
        sync_badge = "Atualizando…" if st.session_state.get("pvp_sync_pending") else "Sincronizado ✓"

        # CSS leve para cards/badges (aplicado só na Arena PvP)
        st.markdown("""
        <style>
        .pvp-statusbar{
            display:flex; gap:10px; align-items:center; flex-wrap:wrap;
            padding:10px 12px; border-radius:14px;
            border:1px solid rgba(148,163,184,0.35);
            background: rgba(15,23,42,0.55);
            box-shadow: inset 0 0 18px rgba(15,23,42,0.35);
            margin: 6px 0 10px 0;
        }
        .pvp-pill{
            display:inline-flex; gap:8px; align-items:center;
            padding:6px 10px; border-radius:999px;
            border:1px solid rgba(148,163,184,0.35);
            background: rgba(2,6,23,0.35);
            font-weight:600; font-size:12px;
        }
        .pvp-badge{
            display:inline-flex; align-items:center; justify-content:center;
            padding:2px 8px; border-radius:999px;
            background: rgba(56,189,248,0.18);
            border: 1px solid rgba(56,189,248,0.35);
            font-size: 12px; font-weight: 700;
            margin-right: 6px;
        }
        
        /* =========================
           PvP / Arena - Team Panel (conceito)
           ========================= */
        [data-testid="stVerticalBlock"]:has(.pvp-team-panel-marker){
            border-radius: 18px;
            padding: 14px 14px 10px 14px !important;
            border: 2px solid rgba(148,163,184,0.35);
            background: radial-gradient(1200px 500px at 20% 0%, rgba(56,189,248,0.12), transparent 60%),
                        rgba(15,23,42,0.58);
            box-shadow: inset 0 0 18px rgba(15,23,42,0.40);
        }
        .pvp-team-title{
            display:flex; align-items:center; gap:10px;
            font-weight:800;
            letter-spacing: 0.5px;
        }
        .pvp-team-title .pvp-mini{
            width:32px;height:32px;border-radius:10px;object-fit:cover;
            border:1px solid rgba(226,232,240,0.28);
            box-shadow: 0 6px 14px rgba(2,6,23,0.35);
        }

        /* Card de Avatar (topo) */
        [data-testid="stVerticalBlock"]:has(.pvp-avatar-card-marker){
            margin-top: 6px;
            border-radius: 18px;
            padding: 12px 12px 12px 12px !important;
            border: 1px solid rgba(148,163,184,0.28);
            background: rgba(2,6,23,0.28);
            box-shadow: inset 0 0 14px rgba(2,6,23,0.22);
        }

        /* Card de Pokémon (cada linha) */
        [data-testid="stVerticalBlock"]:has(.pvp-mon-card-marker){
            border-radius: 18px;
            padding: 12px 12px 10px 12px !important;
            border: 1px solid rgba(148,163,184,0.22);
            background: rgba(2,6,23,0.22);
            box-shadow: 0 10px 20px rgba(2,6,23,0.18);
            margin-bottom: 10px;
        }
        [data-testid="stVerticalBlock"]:has(.pvp-mon-card-marker.pvp-selected){
            border: 2px solid rgba(56,189,248,0.55);
            background: radial-gradient(900px 380px at 15% 0%, rgba(56,189,248,0.12), transparent 60%),
                        rgba(2,6,23,0.22);
        }

        /* Botões dentro dos cards (aproxima o conceito) */
        [data-testid="stVerticalBlock"]:has(.pvp-mon-card-marker) [data-testid="stButton"] button,
        [data-testid="stVerticalBlock"]:has(.pvp-avatar-card-marker) [data-testid="stButton"] button{
            border-radius: 14px !important;
            border: 1px solid rgba(148,163,184,0.28) !important;
            background: linear-gradient(180deg, rgba(56,189,248,0.55) 0%, rgba(14,165,233,0.35) 100%) !important;
            box-shadow: 0 8px 18px rgba(2,6,23,0.22);
            color: #e2e8f0 !important;
            font-weight: 800 !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            line-height: 1 !important;
        }
        [data-testid="stVerticalBlock"]:has(.pvp-mon-card-marker) [data-testid="stButton"] button p,
        [data-testid="stVerticalBlock"]:has(.pvp-avatar-card-marker) [data-testid="stButton"] button p{
            margin: 0 !important;
            width: 100%;
            line-height: 1 !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
        }
        [data-testid="stVerticalBlock"]:has(.pvp-mon-card-marker) [data-testid="stButton"] button:hover,
        [data-testid="stVerticalBlock"]:has(.pvp-avatar-card-marker) [data-testid="stButton"] button:hover{
            filter: brightness(1.08);
            transform: translateY(-1px);
        }

        /* Botões "ícone" (🔍 / 📍 / 🚶 / 👁️ / ❌) */
        [data-testid="stVerticalBlock"]:has(.pvp-mon-card-marker) [data-testid="stButton"] button:has(span){
            min-height: 40px;
        }

        /* Remove padding extra do bloco interno quando usamos marcadores */
        [data-testid="stVerticalBlock"] > div:has(.pvp-team-panel-marker),
        [data-testid="stVerticalBlock"] > div:has(.pvp-avatar-card-marker),
        [data-testid="stVerticalBlock"] > div:has(.pvp-mon-card-marker){
            padding: 0 !important;
            margin: 0 !important;
        }

        </style>
        """, unsafe_allow_html=True)

        st.markdown(
            f"""
            <div class="pvp-statusbar">
              <span class="pvp-pill">🧩 Sala <b>{rid}</b></span>
              <span class="pvp-pill">🎚️ Fase: <b>{b_status}</b></span>
              <span class="pvp-pill">⚔️ Agindo: <b>{b_attacker}</b></span>
              <span class="pvp-pill">📡 {sync_badge}</span>
            </div>
            """,
            unsafe_allow_html=True
        )

        tab_arena, tab_combate, tab_inic, tab_fichas, tab_log = st.tabs(["🗺️ Arena", "⚔️ Combate", "🧭 Iniciativa", "📋 Fichas", "📜 Log"])
        import math

        with tab_combate:
            with st.expander("⚔️ Calculadora de Combate", expanded=(b_status != "idle")):
            
                # [FASE 0] IDLE
                if b_status == "idle":
                    if is_player:
                        if st.button("Nova Batalha (Atacar)"):
                            battle_ref.set({"status": "setup", "attacker": trainer_name, "attack_move": None, "logs": []})
                            st.rerun()
                    else:
                        st.caption("Aguardando combate...")
            
                # [FASE 1] CONFIGURAR ATAQUE
                elif b_status == "setup":
                    # CORREÇÃO: Parêntese fechado corretamente aqui
                    st.caption(f"**Atacante:** {b_data.get('attacker')}")
                
                    if b_data.get("attacker") == trainer_name:
                        attacker_pid = None
                        attacker_sheet = None
                        attacker_stats = {}
                        if current_party:
                            attacker_pid = st.selectbox(
                                "Seu Pokémon (Atacante)",
                                options=current_party,
                                format_func=get_poke_display_name,
                                key=f"atk_self_{rid}",
                            )
                            attacker_sheet = battle_sheets_map.get(str(attacker_pid))
                            if attacker_sheet:
                                attacker_stats = attacker_sheet.get("stats") or {}
                            else:
                                _, _, attacker_stats, _, _ = get_poke_data(trainer_name, attacker_pid)
                        else:
                            st.info("Sua party está vazia para selecionar o atacante.")
                        # Busca peças de TODOS os outros jogadores que não são você
                        # Busca peças de TODOS os outros jogadores (somente Pokémon com pid)
                        target_options = {}
                        for p_owner, p_pieces in (player_pieces_map or {}).items():
                            if p_owner == trainer_name:
                                continue
                        
                            for p in (p_pieces or []):
                                piece_id = p.get("id")
                                pid = p.get("pid")
                                kind = (p.get("kind") or "").lower()
                        
                                # ignora peças inválidas / não-Pokémon
                                if not piece_id:
                                    continue
                                if kind == "trainer":
                                    continue
                                if pid in (None, "", 0, "0"):
                                    continue
                        
                                label = f"{get_poke_display_name(pid)} ({p_owner})"
                                target_options[piece_id] = label
                    
                        c_atk1, c_atk2, c_atk3 = st.columns(3)
                    
                        # CORREÇÃO: Indentação alinhada com o bloco acima (4 espaços dentro do IF)
                        with c_atk1:
                            target_id = st.selectbox("Alvo", options=list(target_options.keys()), 
                                        format_func=lambda x: target_options[x],
                                        key=f"atk_target_{rid}") if target_options else None
                    
                        with c_atk2:
                            attack_mode = st.radio(
                                "Modo",
                                ["Normal", "Área"],
                                horizontal=True,
                                key=f"atk_mode_{rid}"
                            )
                        
                        # Se for Área
                        if attack_mode == "Área":
                            st.info("Ataque em Área: Dodge (CD 10 + Nível) reduz dano pela metade.")
                            lvl_effect = st.number_input("Nível do Efeito / Dano", min_value=1, value=1)
                            is_eff_area = st.checkbox("É Efeito? (Affliction)", key=f"area_eff_{rid}")
                        
                            if st.button("🚀 Lançar Área"):
                                if not target_id:
                                    st.warning("Selecione um alvo.")
                                    st.stop()
                        
                                # Busca a peça alvo de forma blindada
                                t_p = next((p for p in (all_pieces or []) if p.get("id") == target_id), None)
                                if not t_p:
                                    st.error("Alvo inválido (não encontrado no mapa).")
                                    st.stop()
                        
                                t_owner = t_p.get("owner")
                                t_pid = t_p.get("pid")
                        
                                # Se não tiver pid, não é Pokémon (ou está incompleto)
                                if not t_pid:
                                    st.error("Esse alvo não é um Pokémon (sem pid).")
                                    st.stop()
                        
                                battle_ref.update({
                                    "status": "aoe_defense",
                                    "target_id": target_id,
                                    "target_owner": t_owner,
                                    "target_pid": t_pid,
                                    "aoe_dc": int(lvl_effect) + 10,
                                    "dmg_base": int(lvl_effect),
                                    "is_effect": bool(is_eff_area),
                                    "logs": [f"{trainer_name} lançou Área (Nv {lvl_effect}). Defensor rola Dodge (CD {int(lvl_effect)+10})."]
                                })
                                st.rerun()

                        else:
                            # Normal
                            with c_atk3:
                                atk_type = st.selectbox("Alcance", ["Distância (Dodge)", "Corpo-a-corpo (Parry)"])
                        

                            move_payload = None
                            selected_accuracy = None
                            selected_damage = None
                            if attacker_sheet:
                                moves = attacker_sheet.get("moves") or []
                            else:
                                moves = []

                            if moves:
                                def _move_label(idx: int) -> str:
                                    mv = moves[idx]
                                    name = mv.get("name", "Golpe")
                                    rank = int(mv.get("rank", 0) or 0)
                                    based, stat_val = _move_stat_value(mv, attacker_stats)
                                    damage = rank + stat_val
                                    acc = int(mv.get("accuracy", 0) or 0)
                                    return f"{name}. Acerto: {acc}. Dano {damage}"

                                move_options = ["manual"] + list(range(len(moves)))
                                move_choice = st.selectbox(
                                    "Golpe (opcional)",
                                    options=move_options,
                                    format_func=lambda x: "Manual (sem golpe)" if x == "manual" else _move_label(x),
                                    key=f"atk_move_{rid}",
                                )

                                if move_choice != "manual":
                                    selected_move = moves[int(move_choice)]
                                    selected_accuracy = int(selected_move.get("accuracy", 0) or 0)
                                    rank = int(selected_move.get("rank", 0) or 0)
                                    based, stat_val = _move_stat_value(selected_move, attacker_stats)
                                    selected_damage = rank + stat_val
                                    move_payload = {
                                        "name": selected_move.get("name", "Golpe"),
                                        "accuracy": selected_accuracy,
                                        "damage": selected_damage,
                                        "rank": rank,
                                        "based_stat": based,
                                        "stat_value": stat_val,
                                    }
                            else:
                                st.caption("Sem golpes disponíveis para este Pokémon.")

                            ac1, ac2 = st.columns([2, 1])
                            with ac1:
                                atk_mod = st.number_input("Acerto (Modificador)", value=int(selected_accuracy or 0), step=1)
                            with ac2:
                                if selected_accuracy is not None:
                                    st.markdown(f"**Acerto sugerido:** {selected_accuracy}")
                                else:
                                    st.caption("Sem acerto sugerido.")
                        
                            if st.button("⚔️ Rolar Ataque"):
                                if not target_id:
                                    st.warning("Selecione um alvo.")
                                    st.stop()

                                # Busca a peça alvo de forma blindada
                                t_p = next((p for p in (all_pieces or []) if p.get("id") == target_id), None)
                                if not t_p:
                                    st.error("Alvo inválido (não encontrado no mapa).")
                                    st.stop()

                                t_owner = t_p.get("owner")
                                t_pid = t_p.get("pid")

                                # Se não tiver pid, não é Pokémon (ou está incompleto)
                                if not t_pid:
                                    st.error("Esse alvo não é um Pokémon (sem pid).")
                                    st.stop()

                                # Pega stats do alvo
                                _, _, t_stats, _, _ = get_poke_data(t_owner, t_pid)
                                dodge = int((t_stats or {}).get("dodge", 0))
                                parry = int((t_stats or {}).get("parry", 0))

                                defense_val = dodge if ("Distância" in (atk_type or "")) else parry
                                needed = defense_val + 10

                                d20 = random.randint(1, 20)
                                total_atk = int(atk_mod) + int(d20)

                                # RAW M&M (simplificado)
                                if d20 == 1:
                                    hit = False
                                    crit_bonus = 0
                                elif d20 == 20:
                                    # 20 natural: acerta automaticamente e consideramos "Increased Effect" (+5)
                                    hit = True
                                    crit_bonus = 5
                                else:
                                    hit = (total_atk >= needed)
                                    crit_bonus = 0

                                result_msg = "ACERTOU! ✅" if hit else "ERROU! ❌"
                                crit_txt = " (CRÍTICO +5)" if crit_bonus else ""

                                battle_ref.update({
                                    "status": "hit_confirmed" if hit else "missed",
                                    "attacker": trainer_name,
                                    "attacker_pid": attacker_pid,
                                    "target_id": target_id,
                                    "target_owner": t_owner,
                                    "target_pid": t_pid,
                                    "attack_move": move_payload,
                                    "attack_range": atk_type,
                                    "atk_mod": int(atk_mod),
                                    "d20": int(d20),
                                    "defense_val": int(defense_val),
                                    "needed": int(needed),
                                    "total_atk": int(total_atk),
                                    "crit_bonus": int(crit_bonus),
                                    "logs": [
                                        f"{trainer_name} rolou {d20}+{atk_mod}=**{total_atk}** (vs Def {needed} [{defense_val}+10]){crit_txt}... {result_msg}"
                                    ],
                                })
                                st.rerun()

                    else:
                        st.info(f"Aguardando {b_data.get('attacker')}...")

                # [FASE 1.5] DEFESA DE ÁREA
                elif b_status == "aoe_defense":
                    st.info(b_data["logs"][-1])

                    if b_data.get("target_owner") == trainer_name:
                        st.markdown("### 🛡️ Defesa (Área)")
                        st.caption("Regra-base: Dodge (CD 10 + Nível) reduz o Rank pela metade. (Aqui você pode escolher qualquer defesa.)")

                        c1, c2, c3, c4 = st.columns(4)
                        chosen = None
                        if c1.button("Dodge"): chosen = "dodge"
                        if c2.button("Parry"): chosen = "parry"
                        if c3.button("Fort"): chosen = "fort"
                        if c4.button("Will"): chosen = "will"
                        if st.button("THG (Toughness)"): chosen = "thg"

                        if chosen:
                            d20 = random.randint(1, 20)
                            _, _, t_stats, _, _ = get_poke_data(trainer_name, b_data.get("target_pid"))
                            stat_val = int((t_stats or {}).get(chosen, 0))

                            total_roll = d20 + stat_val
                            dc = int(b_data.get("aoe_dc", 10))
                            base_rank = int(b_data.get("dmg_base", 0))

                            if total_roll >= dc:
                                final_rank = max(1, math.floor(base_rank / 2))
                                msg = f"Sucesso! ({total_roll} vs {dc}) com {chosen.upper()}. Rank reduzido: {base_rank} -> {final_rank}."
                            else:
                                final_rank = base_rank
                                msg = f"Falha! ({total_roll} vs {dc}) com {chosen.upper()}. Rank total: {final_rank}."

                            battle_ref.update({
                                "status": "waiting_defense",
                                "dmg_base": int(final_rank),
                                "logs": firestore.ArrayUnion([msg + " Escolha a resistência agora."])
                            })
                            st.rerun()
                    else:
                        st.warning("Aguardando defesa de área...")

                # [FASE 2] CONFIRMAÇÃO DE ACERTO / DEFINIR RANK
                elif b_status == "hit_confirmed":
                    st.success(b_data["logs"][-1])
                    if b_data.get("attacker") == trainer_name:
                        dmg_input = st.number_input("Rank do Dano / Efeito", min_value=0, value=int(b_data.get("attack_move",{}).get("damage",0) if b_data.get("attack_move") else 0), step=1)
                        is_eff_check = st.checkbox("É efeito? (Affliction)", value=False)

                        if st.button("Confirmar Rank"):
                            battle_ref.update({
                                "status": "waiting_defense",
                                "dmg_base": int(dmg_input),
                                "is_effect": bool(is_eff_check),
                                "logs": firestore.ArrayUnion([f"Rank/Dano: {dmg_input} ({'Efeito' if is_eff_check else 'Dano'}). Aguardando resistência..."])
                            })
                            st.rerun()
                    else:
                        st.info("Aguardando atacante definir o dano...")

                elif b_status == "missed":
                    st.error(b_data["logs"][-1])
                    if b_data.get("attacker") == trainer_name:
                        if st.button("Encerrar"):
                            battle_ref.update({"status": "idle", "logs": []})
                            st.rerun()

                # [FASE 3] RESISTÊNCIA FINAL
                elif b_status == "waiting_defense":
                    is_eff = b_data.get("is_effect", False)
                    base_val = 10 if is_eff else 15
                    rank = int(b_data.get("dmg_base", 0))
                    dc_total = base_val + rank + int(b_data.get("crit_bonus", 0))
                
                    st.info(f"Resistir contra: **CD {dc_total}** ({base_val} + {rank})")
                
                    if b_data.get("target_owner") == trainer_name:
                        st.markdown("### 🛡️ Resistir com:")
                        c1, c2, c3, c4 = st.columns(4)
                        res_type = None
                        if c1.button("Dodge"): res_type = "dodge"
                        if c2.button("Parry"): res_type = "parry"
                        if c3.button("Fort"): res_type = "fort"
                        if c4.button("Will"): res_type = "will"
                        if st.button("THG (Toughness)"): res_type = "thg"

                        if res_type:
                            def_die = random.randint(1, 20)
                            _, _, t_stats, _, _ = get_poke_data(trainer_name, b_data.get('target_pid'))
                            stat_val = int(t_stats.get(res_type, 0))
                        
                            check_total = def_die + stat_val
                        
                            # Cálculo M&M: Falha = CD - Check
                            diff = dc_total - check_total
                        
                            if diff <= 0:
                                bars_lost = 0
                                res_msg = "SUCESSO (Nenhum dano)"
                            else:
                                # 1 grau a cada 5 pontos
                                bars_lost = math.ceil(diff / 5)
                                res_msg = f"FALHA por {diff}"
                        
                            final_msg = f"🛡️ Defensor rolou {def_die} + {stat_val} = **{check_total}** ({res_type.upper()}). {res_msg}. Perdeu **{bars_lost}** barras."
                        
                            battle_ref.update({
                                "status": "finished",
                                "final_bars": bars_lost,
                                "logs": firestore.ArrayUnion([final_msg])
                            })
                            st.rerun()
                    else:
                        st.warning("Aguardando defesa...")

                # [FASE 4] FIM / SECUNDÁRIO
                elif b_status == "finished":
                    st.markdown(f"## 🩸 Resultado: -{b_data.get('final_bars')} Barras")
                    for log in b_data.get("logs", []): st.text(log)
                
                    if b_data.get("attacker") == trainer_name:
                        c_end1, c_end2 = st.columns(2)
                        if c_end1.button("Encerrar Combate"):
                            battle_ref.update({"status": "idle", "logs": []})
                            st.rerun()
                    
                        target_name = get_poke_display_name(b_data.get('target_pid'))
                        if c_end2.button(f"⚡ Efeito Secundário em {target_name}"):
                            battle_ref.update({
                                "status": "hit_confirmed", 
                                "is_effect": False, 
                                "logs": [f"⚡ Efeito Secundário ativado em {target_name}!"]
                            })
                            st.rerun()
                    else:
                        st.info("Aguardando encerramento do atacante...")



            # =========================
        with tab_arena:
            # 7. LAYOUT DAS COLUNAS (EQUILIBRADO PARA 4 JOGADORES)
            # =========================
            if seed is None:
                st.warning("Sem mapa.")
                st.stop()

            # Proporção [1.8, 4.6, 1.8] dá leve redução do mapa e mantém fôlego nas laterais
            opponents_list = sorted(list(set([p for p in all_players if p != trainer_name])))
            if opponents_list:
                c_me, c_map, c_opps = st.columns([1.9, 5.2, 1.9])
            else:
                c_me, c_map = st.columns([2.2, 6.0])
                c_opps = None

            with c_me:
                render_player_column(trainer_name, "🎒 Sua Equipe", is_me=True)

            with c_map:
                st.markdown(f"### 🗺️ Arena (Sala {rid})")
            
                # --- ALERTA VISUAL DE AÇÃO NO MAPA ---
                #if st.session_state.get("moving_piece_id"):
                #    st.warning("🏃 MODO MOVIMENTO: Clique em um quadrado vazio para mover o Pokémon.", icon="📍")
                #elif st.session_state.get("placing_pid"):
                #    st.info("📍 MODO POSICIONAMENTO: Clique no mapa para colocar o Pokémon.", icon="⬇️")
                #elif st.session_state.get("placing_effect"):
                #    eff_icon = st.session_state.get("placing_effect")
                #    st.info(f"✨ MODO TERRENO: Clique para adicionar {eff_icon}.", icon="✨")
                # -------------------------------------

                            # Ferramentas de Campo (UX melhorada)
                moving_piece_id = st.session_state.get("moving_piece_id")
                placing_pid = st.session_state.get("placing_pid")
                placing_trainer = st.session_state.get("placing_trainer")
                placing_eff = st.session_state.get("placing_effect")

                # Banner do "modo" atual + botão de cancelar
                if moving_piece_id or placing_pid or placing_trainer or placing_eff:
                    bmsg, bbtn = st.columns([6, 1])
                    with bmsg:
                        if moving_piece_id:
                            st.warning("🏃 MODO MOVIMENTO: clique em um quadrado vazio para mover.", icon="📍")
                        elif placing_pid:
                            st.info("📍 MODO POSICIONAMENTO: clique no mapa para colocar o Pokémon.", icon="⬇️")
                        elif placing_trainer:
                            st.info("🧍 MODO AVATAR: clique no mapa para posicionar seu treinador.", icon="⬇️")
                        elif placing_eff == "__erase__":
                            st.warning("🧽 MODO BORRACHA: clique em um quadrado para REMOVER o efeito.", icon="🧽")
                        elif placing_eff:
                            st.info(f"✨ MODO TERRENO: clique para adicionar {placing_eff}.", icon="✨")
                    with bbtn:
                        if st.button("✖", key=f"cancel_action_{rid}", help="Cancelar ação atual", use_container_width=True):
                            st.session_state["moving_piece_id"] = None
                            st.session_state["placing_pid"] = None
                            st.session_state["placing_trainer"] = None
                            st.session_state["placing_effect"] = None
                            st.rerun()

                # Ferramentas de Campo
                with st.expander("🛠️ Itens e Terrenos", expanded=False):
                    if is_player:
                        st.checkbox(
                            "🖌️ Pincel (não desarmar ao clicar)",
                            value=bool(st.session_state.get(f"effect_brush_{rid}", True)),
                            key=f"effect_brush_{rid}",
                            help="Se ligado, você pode clicar várias vezes no mapa sem precisar re-selecionar o ícone.",
                        )

                        top_tools = st.columns([1.1, 1.1, 2.8])
                        with top_tools[0]:
                            erase_curr = (st.session_state.get("placing_effect") == "__erase__")
                            if st.button(
                                "🧽 Apagar",
                                key=f"effect_erase_{rid}",
                                type=("primary" if erase_curr else "secondary"),
                                use_container_width=True,
                                help="Ativa um modo para remover o efeito clicando no mapa.",
                            ):
                                st.session_state["placing_effect"] = None if erase_curr else "__erase__"
                                st.session_state["moving_piece_id"] = None
                                st.session_state["placing_pid"] = None
                                st.session_state["placing_trainer"] = None
                                request_rerun("remove_trainer")
                        with top_tools[1]:
                            if st.button(
                                "🧼 Limpar Tudo",
                                key=f"effect_clear_{rid}",
                                use_container_width=True,
                                help="Remove todos os efeitos do mapa.",
                            ):
                                
                                st.rerun()
                        with top_tools[2]:
                            st.caption("Selecione um ícone e clique no mapa. (Pincel mantém o modo ativo.)")

                        effects_map = {"Fogo":"🔥", "Gelo":"🧊", "Água":"💧", "Rocha":"🪨", "Nuvem":"☁️", "Sol":"☀️", "Grama":"🍃", "Raio":"⚡"}
                        curr = st.session_state.get("placing_effect")

                        # Contagem dos efeitos atuais (opcional, ajuda a visualizar densidade)
                        counts = {}
                        for e in (field_effects or []):
                            ic = (e or {}).get("icon")
                            if ic:
                                counts[ic] = counts.get(ic, 0) + 1

                        cols = st.columns(8)
                        for i, (k, v) in enumerate(effects_map.items()):
                            btn_type = "primary" if curr == v else "secondary"
                            label = f"{v}" + (f" {counts[v]}" if counts.get(v) else "")
                            if cols[i].button(label, key=f"ef_{rid}_{k}", type=btn_type, use_container_width=True):
                                st.session_state["placing_effect"] = None if curr == v else v
                                # Se ativar efeito, cancela outras ações para evitar bugs
                                st.session_state["moving_piece_id"] = None
                                st.session_state["placing_pid"] = None
                                st.session_state["placing_trainer"] = None
                                request_rerun("remove_trainer")

                # Ajustes de visualização do mapa (zoom automático por tamanho para manter proporção)
                toolbar = st.columns([1.15, 2.35, 1.0])
                with toolbar[0]:
                    show_grid = st.checkbox("Grade Tática", value=bool(st.session_state.get(f"grid_{rid}", True)), key=f"grid_{rid}")

                # Menor mapa => maior zoom, com limite para não ficar desproporcional.
                auto_zoom_by_grid = {
                    6: 1.15,
                    8: 1.00,
                    10: 0.90,
                    12: 0.82,
                }
                map_zoom = float(auto_zoom_by_grid.get(int(grid), 0.90))
                st.session_state[f"map_zoom_{rid}"] = map_zoom

                with toolbar[1]:
                    st.markdown("**Zoom do mapa:** automático")
                with toolbar[2]:
                    st.caption(f"{int(map_zoom * 100)}%")

                # ... (Restante do código de renderização do mapa permanece igual) ...
                # --- assinatura SOMENTE do que muda o VISUAL do mapa ---
                def _compact_piece(p: dict) -> dict:
                    # mantenha só o que altera o desenho
                    return {
                        "id": p.get("id"),
                        "row": int(p.get("row", -1) or -1),
                        "col": int(p.get("col", -1) or -1),
                        "kind": p.get("kind"),
                        "pid": str(p.get("pid", "")),
                        "shiny": bool(p.get("shiny", False)),
                        "owner": p.get("owner"),
                        "avatar": p.get("avatar"),
                    }
                
                def _compact_eff(e: dict) -> dict:
                    return {
                        "row": int(e.get("row", -1) or -1),
                        "col": int(e.get("col", -1) or -1),
                        "icon": e.get("icon"),
                    }
                
                map_signature = json.dumps({
                    "seed": int(seed or 0),
                    "theme": str(theme_key or ""),
                    "noWater": bool(no_water_state),
                    "grid": bool(show_grid),
                
                    # IMPORTANTE: não use `updatedAt` aqui
                    "pieces": sorted([_compact_piece(p) for p in (pieces_to_draw or [])], key=lambda x: (x["id"] or "", x["row"], x["col"])),
                    "effects": sorted([_compact_eff(e) for e in (field_effects or [])], key=lambda x: (x["row"], x["col"], str(x["icon"]))),
                }, sort_keys=True, default=str)

                if st.session_state.get("map_cache_sig") != map_signature:
                    st.session_state["map_cache_sig"] = map_signature
                    st.session_state["map_cache_img"] = render_map_with_pieces(
                        grid,
                        theme_key,
                        seed,
                        no_water_state,
                        pieces_to_draw,
                        trainer_name,
                        room,
                        effects=field_effects,
                        show_grid=show_grid,
                    )
                img = st.session_state.get("map_cache_img")

                # Zoom do mapa (ajuda a encaixar melhor em 100%/50% de zoom do navegador)
                map_zoom = float(st.session_state.get(f"map_zoom_{rid}", 0.90) or 0.90)
                tile_px = max(1, int(TILE_SIZE * map_zoom))
                st.session_state[f"_tile_px_{rid}"] = tile_px

                img_to_show = img
                
                # Cache do resize por (sala, assinatura do mapa, zoom)
                scale_key = (rid, st.session_state.get("map_cache_sig"), float(map_zoom))
                
                if img_to_show and map_zoom != 1.0:
                    if st.session_state.get("map_scaled_key") != scale_key:
                        st.session_state["map_scaled_key"] = scale_key
                        try:
                            st.session_state["map_scaled_img"] = img_to_show.resize(
                                (
                                    max(1, int(img_to_show.width * map_zoom)),
                                    max(1, int(img_to_show.height * map_zoom)),
                                ),
                                resample=Image.NEAREST,
                            )
                        except Exception:
                            # Mantém o seu fallback original
                            st.session_state["map_scaled_img"] = img_to_show
                            tile_px = TILE_SIZE
                            st.session_state[f"_tile_px_{rid}"] = tile_px
                
                    img_to_show = st.session_state.get("map_scaled_img") or img_to_show


                with st.container():
                    sig = st.session_state.get("map_cache_sig", "") or ""
                    sig_short = hashlib.md5(sig.encode("utf-8")).hexdigest()[:10]
                    zoom_tag = int(float(map_zoom) * 100)
                    
                    # Key muda quando o visual do mapa muda (assinatura) ou quando muda o zoom
                    map_widget_key = f"map_{rid}_{sig_short}_{zoom_tag}"
                    
                    click = streamlit_image_coordinates(img_to_show, key=map_widget_key)
            if c_opps is not None:
                with c_opps:
                    st.markdown("### 🆚 Oponentes")
                    opponents = opponents_list

                    if not opponents:
                        st.caption("Aguardando...")
                    else:
                        for idx, opp_name in enumerate(opponents):
                            icons = ["🔴", "🟡", "🌸"]
                            icon = icons[idx] if idx < len(icons) else "⚪"

                            with st.expander(f"{icon} {opp_name}", expanded=True):
                                render_player_column(opp_name, "", is_me=False)


            # =========================
            # 8. LÓGICA DE CLIQUE
            # =========================
            if click and "x" in click and "y" in click:
                tile_px_eff = int(st.session_state.get(f"_tile_px_{rid}", TILE_SIZE) or TILE_SIZE)
                col = int(click["x"] // tile_px_eff)
                row = int(click["y"] // tile_px_eff)
                if 0 <= row < grid and 0 <= col < grid:
                    ppid = st.session_state.get("placing_pid")
                    peff = st.session_state.get("placing_effect")
                    moving_piece_id = st.session_state.get("moving_piece_id")
                    placing_trainer = st.session_state.get("placing_trainer")
                    if peff:
                        curr = state.get("effects") or []
                        brush_on = bool(st.session_state.get(f"effect_brush_{rid}", True))

                        if peff == "__erase__":
                            # Remove efeito (se existir) no quadrado clicado
                            new = [
                                e for e in curr
                                if not (int((e or {}).get("row", -1)) == row and int((e or {}).get("col", -1)) == col)
                            ]
                            if len(new) != len(curr):
                                db.collection("rooms").document(rid).collection("public_state").document("state").update({
                                    "effects": new,
                                    "updatedAt": firestore.SERVER_TIMESTAMP,
                                })
                                add_public_event(db, rid, "effect_removed", trainer_name, {"to": [row, col]})

                            if not brush_on:
                                st.session_state["placing_effect"] = None
                            request_rerun("effect_erase")

                        # Aplica/atualiza efeito
                        new = [
                            e for e in curr
                            if not (int((e or {}).get("row", -1)) == row and int((e or {}).get("col", -1)) == col)
                        ]
                        new.append({"icon": peff, "row": row, "col": col, "id": str(uuid.uuid4())[:8]})
                        db.collection("rooms").document(rid).collection("public_state").document("state").update({
                            "effects": new,
                            "updatedAt": firestore.SERVER_TIMESTAMP,
                        })
                        add_public_event(db, rid, "effect", trainer_name, {"icon": peff, "to": [row, col]})

                        if not brush_on:
                            st.session_state["placing_effect"] = None
                        request_rerun("effect_place")

                    elif ppid:
                        new_id = str(uuid.uuid4())[:8]
                        # Stats já estão no banco, não precisa passar aqui
                        am_i_shiny = ppid in user_data.get("shinies", [])
                        new_piece = {
                            "id": new_id, 
                            "pid": ppid, 
                            "owner": trainer_name, 
                            "row": row, 
                            "col": col, 
                            "revealed": True, 
                            "status": "active",
                            "shiny": am_i_shiny # <--- SALVA NA PEÇA
                        }
                        upsert_piece(db, rid, new_piece)
                        mark_pid_seen(db, rid, ppid)
                        add_public_event(db, rid, "piece_placed", trainer_name, {"pid": ppid, "to": [row, col]})
                        st.session_state.pop("placing_pid", None)
                        request_rerun("piece_place")
                    elif placing_trainer:
                        s_now = get_state(db, rid)
                        all_p = s_now.get("pieces") or []
                        avatar_choice = user_data.get("trainer_profile", {}).get("avatar_choice")
                        existing_trainer = next(
                            (p for p in all_p if p.get("owner") == trainer_name and p.get("kind") == "trainer"),
                            None,
                        )
                        if avatar_choice:
                            if existing_trainer:
                                existing_trainer["row"] = row
                                existing_trainer["col"] = col
                                existing_trainer["avatar"] = avatar_choice
                                existing_trainer["revealed"] = True
                                upsert_piece(db, rid, existing_trainer)
                            else:
                                new_id = str(uuid.uuid4())[:8]
                                new_piece = {
                                    "id": new_id,
                                    "kind": "trainer",
                                    "avatar": avatar_choice,
                                    "owner": trainer_name,
                                    "row": row,
                                    "col": col,
                                    "revealed": True,
                                    "status": "active",
                                }
                                upsert_piece(db, rid, new_piece)
                            add_public_event(db, rid, "trainer_placed", trainer_name, {"to": [row, col]})
                        st.session_state["placing_trainer"] = None
                        request_rerun("trainer_place")
                    elif moving_piece_id and is_player:
                        s_now = get_state(db, rid)
                        all_p = s_now.get("pieces") or []
                        mover = next((p for p in all_p if p["id"] == moving_piece_id), None)
                        if mover:
                            # 1. Guarda a posição antiga para o Log
                            old_pos = [mover["row"], mover["col"]]

                            # 2. Atualiza para a nova posição
                            mover["row"] = row
                            mover["col"] = col

                            # 3. Salva a peça no Firebase
                            upsert_piece(db, rid, mover)

                            # 4. Registra o movimento publicamente NO LOG
                            add_public_event(db, rid, "move", trainer_name, {
                                "pid": mover["pid"],
                                "from": old_pos,
                                "to": [row, col]
                            })



                            # 5. Limpa a seleção e recarrega
                            st.session_state["moving_piece_id"] = None
                            request_rerun("move_commit")


        with tab_inic:
            st.markdown("### 🧭 Iniciativa")
            st.caption(
                "Speed vem da Pokédex (get_poke_data). Mod Speed vem da tabela. "
                "O dono da sala pode rolar todos; cada jogador pode rolar apenas os próprios Pokémon."
            )
        
            st.markdown("""
            <style>
              .init-wrap{
                border: 1px solid rgba(255,255,255,.14);
                background: rgba(255,255,255,.035);
                border-radius: 16px;
                padding: 12px;
              }
              .init-row{
                border: 1px solid rgba(255,255,255,.10);
                background: rgba(0,0,0,.18);
                border-radius: 14px;
                padding: 8px 10px;
                margin: 8px 0;
              }
              .init-title{
                font-weight: 900;
                font-size: .95rem;
                line-height: 1.05;
              }
              .init-sub{
                opacity:.8;
                font-size: .78rem;
                margin-top: 1px;
              }
              .init-badge{
                display:inline-flex;
                align-items:center;
                padding: 4px 8px;
                border-radius: 999px;
                border: 1px solid rgba(255,255,255,.16);
                background: rgba(255,255,255,.06);
                font-weight: 900;
                font-size: .78rem;
                line-height: 1;
                white-space: nowrap;
              }
              .init-badge-strong{ background: rgba(59,130,246,.12); border-color: rgba(59,130,246,.22); }
              .init-badge-ok{ background: rgba(34,197,94,.10); border-color: rgba(34,197,94,.22); }
              .init-badge-warn{ background: rgba(245,158,11,.10); border-color: rgba(245,158,11,.22); }
              .init-divider{ height: 1px; background: rgba(255,255,255,.10); margin: 10px 0; }
              .init-head{
                display:flex; align-items:center; justify-content:space-between; gap: 10px;
                margin-bottom: 6px;
              }
            </style>
            """, unsafe_allow_html=True)
        
            def _to_int(v, default=0):
                try:
                    return int(float(v))
                except Exception:
                    return int(default)
        
            def _speed_to_initiative_mod(speed_value: int) -> int:
                if speed_value >= 121: return 8
                if speed_value >= 101: return 4
                if speed_value >= 81:  return 2
                if speed_value >= 71:  return 1
                if speed_value >= 61:  return 0
                if speed_value >= 41:  return -1
                if speed_value >= 1:   return -4
                return 0
            
            def _get_speed_value(pid: str, display_name: str, stats_dict: dict) -> int:
                # 1) tenta vir do stats_dict (se algum dia você passar a salvar speed lá)
                v = _extract_speed_from_stats(stats_dict)
                if v > 0:
                    return v
            
                # 2) tenta pela PokeAPI usando o ID, se for numérico
                try:
                    pid_s = str(pid).strip()
                    if pid_s.isdigit():
                        pjson = pokeapi_get_pokemon(pid_s)
                        base = pokeapi_parse_stats(pjson)
                        v2 = _to_int(base.get("speed", 0), 0)
                        if v2 > 0:
                            return v2
                except Exception:
                    pass
            
                # 3) tenta pela PokeAPI usando o nome (limpando prefixos tipo "EXT:")
                try:
                    name = str(display_name or "").strip()
                    name = name.replace("EXT:", "").strip()
                    # se tiver sufixos seus tipo " - Delta", corta aqui:
                    name = name.split(" - ")[0].strip()
                    if name:
                        pjson = pokeapi_get_pokemon(name)
                        base = pokeapi_parse_stats(pjson)
                        return _to_int(base.get("speed", 0), 0)
                except Exception:
                    return 0
            
                return 0


            
            def _extract_speed_from_stats(stats_dict: dict) -> int:
                if not isinstance(stats_dict, dict):
                    return 0
                for k in ["speed", "spe", "spd", "velocidade", "vel"]:
                    if k in stats_dict:
                        return _to_int(stats_dict.get(k), 0)
                return 0
        
            # Estado da sala (peças em campo)
            s_now = get_state(db, rid)
            pieces = s_now.get("pieces") or []
        
            # Armazenamento sincronizado da iniciativa (fica salvo na sala via battle_ref)
            init_store = (b_data.get("initiative") or {})  # dict[str, dict]
        
            # Monta linhas (Pokémon + (opcional) Treinador)
            rows = []
            for p in pieces:
                p_kind = str(p.get("kind") or "piece")
                if p_kind not in {"trainer", "piece"}:
                    continue
        
                key = f"{p_kind}:{p.get('id')}"
                owner = str(p.get("owner") or "")
        
                if p_kind == "trainer":
                    # Sem sprite de pokémon
                    label = f"Treinador"
                    display = f"🧍 {owner}" if owner else "🧍 Treinador"
                    pid = ""
                    speed_val = 0
                    speed_mod = 0
                    sprite_url = None
                else:
                    pid = str(p.get("pid") or "")
                    display = get_poke_display_name(pid)
                    label = f"Pokémon"
                    _, _, poke_stats, _, _ = get_poke_data(owner, pid)
                    speed_val = _get_speed_value(pid, display, poke_stats)
                    speed_mod = _speed_to_initiative_mod(speed_val)
                    sprite_url = pokemon_pid_to_image(pid, mode="sprite", shiny=(str(pid) in {str(x).strip() for x in (user_data.get("shinies") or [])} if isinstance(user_data, dict) else set()))
        
                saved = init_store.get(key) or {}
                d20_roll = _to_int(saved.get("d20_roll"), 0)
                bonus_input = _to_int(saved.get("bonus_input"), 0)
        
                # iniciativa final (Pokémon: d20 + mod + ajuste; Treinador: só ajuste por enquanto)
                if p_kind == "piece" and d20_roll > 0:
                    final_init = d20_roll + speed_mod + bonus_input
                else:
                    final_init = bonus_input if p_kind == "trainer" else 0
        
                rows.append({
                    "key": key,
                    "owner": owner,
                    "kind": "Avatar" if p_kind == "trainer" else "Pokémon",
                    "pid": pid,
                    "display": display,
                    "label": label,
                    "sprite": sprite_url,
                    "speed": speed_val,
                    "mod_speed": speed_mod,
                    "d20": d20_roll,
                    "bonus": bonus_input,
                    "initiative": final_init
                })
        
            if not rows:
                st.info("Sem peças em campo para registrar iniciativa.")
            else:
                st.markdown('<div class="init-wrap">', unsafe_allow_html=True)
        
                # Ações de rolagem (permissões)
                pokemon_rows = [r for r in rows if r["kind"] == "Pokémon"]
                my_pokemon_rows = [r for r in pokemon_rows if r["owner"] == trainer_name]
        
                st.markdown('<div class="init-head">', unsafe_allow_html=True)
                st.markdown("**Controles de rolagem**", unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
        
                c1, c2, c3 = st.columns([2, 2, 2])
        
                # Dono: rolar todos
                with c1:
                    if role == "owner":
                        if st.button("🎲 Rolar todos (Pokémon em campo)", use_container_width=True):
                            out = dict(init_store)
                            for rec in pokemon_rows:
                                k = rec["key"]
                                prev = out.get(k) or {}
                                rolled = random.randint(1, 20)
                                bonus_val = _to_int(prev.get("bonus_input"), _to_int(rec.get("bonus"), 0))
                                speed_mod = _to_int(rec.get("mod_speed"), 0)
                                out[k] = {
                                    "d20_roll": rolled,
                                    "speed": _to_int(rec.get("speed"), 0),
                                    "speed_mod": speed_mod,
                                    "bonus_input": bonus_val,
                                    "initiative": rolled + speed_mod + bonus_val,
                                    "note": "",
                                }
                            battle_ref.update({
                                "initiative": out,
                                "updatedAt": firestore.SERVER_TIMESTAMP,
                            })
                            st.success("Iniciativas roladas ✅")
                            st.rerun()
                    else:
                        st.caption("Apenas o dono rola todos.")
        
                # Todos: rolar selecionado (dono rola qualquer; outros rolam apenas os próprios)
                with c2:
                    if role == "owner":
                        pool = pokemon_rows
                        help_txt = "Selecione qualquer Pokémon em campo."
                    else:
                        pool = my_pokemon_rows
                        help_txt = "Selecione um dos seus Pokémon em campo."
        
                    options = {r["key"]: f"{r['display']} • {r['owner']}" for r in pool}
                    sel = st.selectbox(
                        "Rolar selecionado",
                        options=list(options.keys()) if options else [],
                        format_func=lambda k: options.get(k, k),
                        key=f"init_sel_{rid}_{trainer_name}",
                    )
                    st.caption(help_txt)
        
                with c3:
                    can_roll_sel = bool(sel)
                    if st.button("🎯 Rolar selecionado", disabled=not can_roll_sel, use_container_width=True):
                        rec = next((x for x in pokemon_rows if x["key"] == sel), None)
                        if rec:
                            # valida permissão
                            if role != "owner" and rec["owner"] != trainer_name:
                                st.error("Você só pode rolar iniciativa dos seus próprios Pokémon.")
                            else:
                                out = dict(init_store)
                                prev = out.get(sel) or {}
                                rolled = random.randint(1, 20)
                                speed_mod = _to_int(rec.get("mod_speed"), 0)
                                bonus_val = _to_int(prev.get("bonus_input"), _to_int(rec.get("bonus"), 0))
                                out[sel] = {
                                    "d20_roll": rolled,
                                    "speed": _to_int(rec.get("speed"), 0),
                                    "speed_mod": speed_mod,
                                    "bonus_input": bonus_val,
                                    "initiative": rolled + speed_mod + bonus_val,
                                    "note": "",
                                }
                                battle_ref.update({
                                    "initiative": out,
                                    "updatedAt": firestore.SERVER_TIMESTAMP,
                                })
                                st.success("Rolado ✅")
                                st.rerun()
        
                st.markdown('<div class="init-divider"></div>', unsafe_allow_html=True)
                st.caption("Tabela Speed → Mod: 1-40=-4, 41-60=-1, 61-70=0, 71-80=1, 81-100=2, 101-120=4, 121+=8.")
        
                # Editor clean por linha (sincroniza no botão salvar)
                out = dict(init_store)
        
                for rec in rows:
                    k = rec["key"]
                    owner = rec["owner"]
                    kind = rec["kind"]
        
                    can_edit = (role == "owner") or (owner == trainer_name)
        
                    # carrega do store (pra refletir sincronizado)
                    saved = out.get(k) or {}
                    d20 = _to_int(saved.get("d20_roll"), _to_int(rec.get("d20"), 0))
                    spd = _to_int(rec.get("speed"), 0)
                    spd_mod = _to_int(rec.get("mod_speed"), 0)
                    bonus_prev = _to_int(saved.get("bonus_input"), _to_int(rec.get("bonus"), 0))
        
                    # calc
                    if kind == "Pokémon" and d20 > 0:
                        init_final = d20 + spd_mod + bonus_prev
                    else:
                        init_final = bonus_prev if kind == "Avatar" else 0
        
                    st.markdown('<div class="init-row">', unsafe_allow_html=True)
                    cA, cB, cC, cD, cE, cF, cG = st.columns([0.9, 2.7, 1.1, 1.1, 0.9, 1.5, 1.2])
        
                    # Sprite
                    with cA:
                        if kind == "Pokémon" and rec.get("sprite"):
                            st.image(rec["sprite"], width=38)
                        else:
                            st.markdown('<span class="init-badge">🧍</span>', unsafe_allow_html=True)
        
                    # Nome + dono
                    with cB:
                        st.markdown(f'<div class="init-title">{html.escape(rec["display"])}</div>', unsafe_allow_html=True)
                        st.markdown(f'<div class="init-sub">{html.escape(owner) if owner else "—"} • {html.escape(kind)}</div>', unsafe_allow_html=True)
        
                    # Speed
                    with cC:
                        st.markdown('<span class="init-badge">Speed</span>', unsafe_allow_html=True)
                        st.markdown(f'<span class="init-badge init-badge-ok">{spd}</span>', unsafe_allow_html=True)
        
                    # Mod Speed
                    with cD:
                        st.markdown('<span class="init-badge">Mod</span>', unsafe_allow_html=True)
                        st.markdown(f'<span class="init-badge init-badge-warn">{spd_mod:+d}</span>', unsafe_allow_html=True)
        
                    # d20
                    with cE:
                        st.markdown('<span class="init-badge">d20</span>', unsafe_allow_html=True)
                        st.markdown(f'<span class="init-badge init-badge-strong">{d20}</span>', unsafe_allow_html=True)
        
                    # Ajuste do jogador (editável)
                    with cF:
                        new_bonus = st.number_input(
                            "Ajuste",
                            value=int(bonus_prev),
                            step=1,
                            min_value=-99,
                            max_value=99,
                            key=f"init_bonus_{rid}_{k}",
                            disabled=not can_edit,
                            label_visibility="collapsed",
                        )
        
                    # Iniciativa final
                    with cG:
                        st.markdown('<span class="init-badge">Final</span>', unsafe_allow_html=True)
                        st.markdown(f'<span class="init-badge init-badge-strong">{init_final}</span>', unsafe_allow_html=True)
        
                    # Atualiza out em memória
                    out[k] = {
                        "d20_roll": d20,
                        "speed": spd,
                        "speed_mod": spd_mod,
                        "bonus_input": _to_int(new_bonus, 0),
                        "initiative": (d20 + spd_mod + _to_int(new_bonus, 0)) if (kind == "Pokémon" and d20 > 0) else (_to_int(new_bonus, 0) if kind == "Avatar" else 0),
                        "note": "",
                    }
                    st.markdown('</div>', unsafe_allow_html=True)
        
                st.markdown('<div class="init-divider"></div>', unsafe_allow_html=True)
        
                # Ordem automática enumerada
                st.markdown("#### 🏁 Ordem automática")
                order_rows = []
                for rec in rows:
                    saved = out.get(rec["key"]) or {}
                    order_rows.append({
                        "owner": rec["owner"],
                        "kind": rec["kind"],
                        "item": (f"{rec['display']} • {rec['owner']}" if rec["owner"] else rec["display"]),
                        "initiative": _to_int(saved.get("initiative"), 0),
                    })
        
                import pandas as _pd
                df_order = _pd.DataFrame(order_rows).sort_values(by=["initiative", "item"], ascending=[False, True]).reset_index(drop=True)
                df_order.insert(0, "ordem", df_order.index + 1)
        
                st.dataframe(
                    df_order[["ordem", "owner", "kind", "item", "initiative"]],
                    hide_index=True,
                    use_container_width=True,
                    column_config={
                        "ordem": st.column_config.NumberColumn("Ordem", format="%d"),
                        "owner": st.column_config.TextColumn("Dono"),
                        "kind": st.column_config.TextColumn("Tipo"),
                        "item": st.column_config.TextColumn("Em campo"),
                        "initiative": st.column_config.NumberColumn("Iniciativa", format="%d"),
                    },
                )
        
                # Salvar (sincroniza)
                if st.button("💾 Salvar iniciativa", use_container_width=True):
                    battle_ref.update({
                        "initiative": out,
                        "updatedAt": firestore.SERVER_TIMESTAMP,
                    })
                    st.success("Iniciativa salva e sincronizada ✅")
        
                st.markdown('</div>', unsafe_allow_html=True)

        with tab_fichas:
            import html, re
        
            st.markdown("### 📋 Fichas")
            st.caption("Cards à esquerda (favoritos + resumo). Ficha completa à direita (imagem grande + stats + golpes).")
        
            # =========================
            # CSS (apenas para o HTML dos cards/painel)
            # =========================
            st.markdown("""
            <style>
              .pvp-card{
                border-radius: 16px;
                padding: 12px 12px 10px 12px;
                border: 1px solid rgba(255,255,255,.14);
                background: rgba(255,255,255,.04);
                margin-bottom: 10px;
                overflow: hidden;
              }
              .pvp-card-selected{
                box-shadow: 0 0 0 2px rgba(59,130,246,.35) inset, 0 10px 30px rgba(0,0,0,.25);
                border-color: rgba(59,130,246,.45);
              }
              .pvp-card-head{
                display:flex;
                gap:10px;
                align-items:center;
                min-width: 0;
              }
              .pvp-card-title{
                font-weight: 900;
                line-height: 1.15;
                font-size: .98rem;
                margin: 0;
              }
              .pvp-card-sub{
                opacity: .85;
                font-size: .82rem;
                margin-top: 2px;
              }
              .pvp-pill-row{
                display:flex;
                flex-wrap: wrap;
                gap:6px;
                margin-top: 6px;
              }
              .pvp-pill{
                display:inline-flex;
                align-items:center;
                gap:6px;
                padding: 3px 8px;
                border-radius: 999px;
                border: 1px solid rgba(255,255,255,.18);
                background: rgba(0,0,0,.18);
                font-weight: 900;
                font-size: .72rem;
                line-height: 1;
                white-space: nowrap;
              }
              .pvp-pill-acc{ background: rgba(34,197,94,.10); border-color: rgba(34,197,94,.25); }
              .pvp-pill-rk { background: rgba(245,158,11,.10); border-color: rgba(245,158,11,.25); }
              .pvp-pill-area{ background: rgba(168,85,247,.10); border-color: rgba(168,85,247,.25); }
        
              .pvp-divider{ height:1px; background: rgba(255,255,255,.12); margin: 10px 0; }
        
              .pvp-move-row{
                display:flex;
                align-items:center;
                justify-content:space-between;
                gap: 8px;
                padding: 6px 8px;
                border-radius: 12px;
                border: 1px solid rgba(255,255,255,.12);
                background: rgba(0,0,0,.18);
                margin-top: 6px;
              }
              .pvp-move-name{
                font-weight: 900;
                font-size: .86rem;
                line-height: 1.1;
                opacity: .96;
                overflow:hidden;
                text-overflow: ellipsis;
                white-space: nowrap;
                max-width: 240px;
                min-width: 0;
              }
              .pvp-move-badges{
                display:flex;
                gap:6px;
                flex-wrap: wrap;
                justify-content:flex-end;
              }
        
              .pvp-open{
                display:block;
                width:100%;
                text-align:center;
                padding: 10px 12px;
                border-radius: 12px;
                border: 1px solid rgba(255,255,255,.16);
                background: rgba(56,189,248,.20);
                font-weight: 900;
                color: #e2e8f0;
                text-decoration: none;
                margin-top: 10px;
              }
              .pvp-open:hover{ filter: brightness(1.08); }
        
              .pvp-panel{
                border-radius: 18px;
                border: 1px solid rgba(255,255,255,.14);
                background: rgba(255,255,255,.04);
                padding: 14px;
                overflow: hidden;
              }
              .pvp-panel-title{ font-weight: 950; font-size: 1.05rem; line-height: 1.15; margin: 0; }
              .pvp-panel-sub{ opacity: .85; margin-top: 3px; }
        
              .pvp-stat-grid{
                display:grid;
                grid-template-columns: repeat(3, minmax(0,1fr));
                gap: 8px;
                margin-top: 10px;
              }
              .pvp-stat{
                border: 1px solid rgba(255,255,255,.12);
                background: rgba(0,0,0,.18);
                border-radius: 14px;
                padding: 10px;
              }
              .pvp-stat-k{ font-size: .75rem; opacity: .78; font-weight: 900; margin-bottom: 4px; }
              .pvp-stat-v{ font-size: 1.05rem; font-weight: 950; line-height: 1; }
        
              .pvp-chips{
                display:flex;
                flex-wrap: wrap;
                gap:6px;
                margin-top: 6px;
              }
              .pvp-chip{
                display:inline-flex;
                align-items:center;
                padding: 4px 8px;
                border-radius: 999px;
                border: 1px solid rgba(255,255,255,.14);
                background: rgba(0,0,0,.16);
                font-weight: 900;
                font-size: .75rem;
              }
            </style>
            """, unsafe_allow_html=True)
        
            # =========================
            # Helpers
            # =========================
            def _hex_to_rgba(hex_color: str, a: float) -> str:
                hc = (hex_color or "").lstrip("#").strip()
                if len(hc) != 6:
                    return f"rgba(100,116,139,{a})"
                r = int(hc[0:2], 16)
                g = int(hc[2:4], 16)
                b = int(hc[4:6], 16)
                return f"rgba({r},{g},{b},{a})"
        
            def _type_bg_style(ptypes: list[str]) -> str:
                t1 = _type_color(ptypes[0]) if ptypes else "#64748b"
                t2 = _type_color(ptypes[1]) if (ptypes and len(ptypes) > 1) else t1
                bg1 = _hex_to_rgba(t1, 0.20)
                bg2 = _hex_to_rgba(t2, 0.20)
                bd  = _hex_to_rgba(t1, 0.45)
                return f"background: linear-gradient(135deg, {bg1}, {bg2}); border-color: {bd};"
        
            def _norm_name(s: str) -> str:
                try:
                    return normalize_text(str(s or ""))
                except Exception:
                    return str(s or "").strip().lower()
        
            def _safe_pid(v) -> str:
                s = str(v or "").strip()
                if not s:
                    return ""
                s = s.lstrip("#").strip()
                if s.upper().startswith("EXT:"):
                    return s
                try:
                    return str(int(float(s)))
                except Exception:
                    return s
        
            def _resolve_base_pid(pid: str, pname: str = "") -> str:
                """
                Se vier EXT:Hydreigon, tenta achar o Nº no df pelo Nome.
                Caso já seja numérico, retorna normal.
                """
                pid = _safe_pid(pid)
                if not pid:
                    return ""
                if pid.isdigit():
                    return pid
        
                # EXT:...
                if pid.upper().startswith("EXT:"):
                    name = pid.split(":", 1)[1].strip()
                else:
                    name = pname.strip() if pname else pid
        
                key = _norm_name(name)
        
                try:
                    # df tem colunas "Nome" e "Nº"
                    df2 = df.copy()
                    df2["_k"] = df2["Nome"].astype(str).apply(_norm_name)
                    row = df2[df2["_k"] == key]
                    if not row.empty:
                        return str(row.iloc[0]["Nº"])
                except Exception:
                    pass
        
                return ""  # não achou base
        
            def _img_fallback(pid: str, shiny: bool, prefer_art: bool = True) -> str:
                pid = _safe_pid(pid)
                pokeball = "https://upload.wikimedia.org/wikipedia/commons/5/53/Pok%C3%A9_Ball_icon.svg"
                url = None
                if prefer_art:
                    try:
                        url = pokemon_pid_to_image(pid, mode="artwork", shiny=shiny)
                    except Exception:
                        url = None
                if not url:
                    try:
                        url = pokemon_pid_to_image(pid, mode="sprite", shiny=shiny)
                    except Exception:
                        url = None
                return url or pokeball
        
            def _mv_is_area(mv: dict) -> bool:
                mv = mv or {}
                meta = mv.get("meta") or {}
                if meta.get("perception_area") is True:
                    return True
                if meta.get("is_area") is True or meta.get("area") is True:
                    return True
                build = str(mv.get("build") or "").lower()
                if "área" in build or "area" in build or "aoe" in build:
                    return True
                return False
        
            def _mv_summary(mv: dict, stats: dict) -> tuple[int,int,str,int,bool]:
                mv = mv or {}
                stats = stats or {}
                base_rank = int(mv.get("rank") or mv.get("Rank") or 0)
                acc_base  = int(mv.get("accuracy") or mv.get("Accuracy") or mv.get("acerto") or 0)
                based_label, stat_val = _move_stat_value(mv, stats)
                stat_val = int(stat_val or 0)
                rank_total = int(base_rank) + int(stat_val)
                return rank_total, acc_base, str(based_label), stat_val, _mv_is_area(mv)
        
            def _fav_moves_for_pid(pid: str) -> list[str]:
                fav = (user_data.get("favorite_moves") or {}) if isinstance(user_data, dict) else {}
                pid_s = str(pid or "").strip()
                if pid_s in fav and isinstance(fav[pid_s], list):
                    return fav[pid_s]
                if pid_s.isdigit():
                    pid2 = (pid_s.lstrip("0") or "0")
                    if pid2 in fav and isinstance(fav[pid2], list):
                        return fav[pid2]
                return []
        
            def _auto_thg_from_pokedex(pid: str, display_name: str, np_: int) -> int:
                cap = max(0, 2 * int(np_ or 0))
                try:
                    if str(pid).strip().isdigit():
                        pjson = pokeapi_get_pokemon(str(pid).strip())
                    else:
                        name = str(display_name or "").replace("EXT:", "").split(" - ")[0].strip()
                        pjson = pokeapi_get_pokemon(name)
        
                    base = pokeapi_parse_stats(pjson)
                    def_ = int(base.get("defense", 10) or 10)
                    spe  = int(base.get("speed", 10) or 10)
        
                    den = max(1, def_ + spe)
                    thg_base = round((def_ / den) * cap)
                    return int(max(0, thg_base))
                except Exception:
                    return 0
        
            # =========================
            # Query param: seleção do card (HTML link)
            # =========================
            try:
                qp_sheet = st.query_params.get("sheet", "")
            except Exception:
                qp_sheet = ""
        
            if qp_sheet:
                st.session_state["pvp_sheet_selected_pid"] = _safe_pid(qp_sheet)
        
            # =========================
            # Resolve fichas da party (aceita EXT)
            # =========================
            party_ids_raw = []
            _seen = set()
            for item in (current_party or []):
                raw = _safe_pid(item)
                if not raw or raw in _seen:
                    continue
                _seen.add(raw)
                party_ids_raw.append(raw)
            
            sheets_to_show = []
            missing_ids = []
            
            for raw_pid in party_ids_raw:
                # raw_pid pode ser "87" ou "EXT:Hydreigon"
                base_pid = raw_pid if raw_pid.isdigit() else str(_resolve_base_pid(raw_pid, raw_pid) or raw_pid)
            
                # tenta por base_pid
                sh = (battle_sheets_map or {}).get(str(base_pid))
            
                # se ainda não achou, tenta pelo raw_pid direto (EXT:...)
                if not sh:
                    sh = (battle_sheets_map or {}).get(str(raw_pid))
            
                # normaliza "00087" -> "87"
                if not sh and base_pid and str(base_pid).isdigit():
                    pid2 = (str(base_pid).lstrip("0") or "0")
                    sh = (battle_sheets_map or {}).get(pid2)
            
                if sh:
                    sh = dict(sh)
                    sh["_party_pid_raw"] = raw_pid
                    sh["_party_pid_base"] = base_pid
                    sheets_to_show.append(sh)
                else:
                    missing_ids.append(raw_pid)
            
            if missing_ids:
                st.caption("Sem ficha salva (ou não pertence ao seu trainer) para: " + ", ".join(missing_ids))
            
            if not sheets_to_show:
                st.info("Não encontrei fichas salvas para a sua party. Vá em **Minhas Fichas** / **Criação Guiada de Fichas** e salve uma ficha para ela aparecer aqui.")
            else:
                # default: seleciona o primeiro
                if not st.session_state.get("pvp_sheet_selected_pid"):
                    p0 = (sheets_to_show[0].get("pokemon") or {}).get("id")
                    st.session_state["pvp_sheet_selected_pid"] = str(p0) if p0 is not None else ""
            
                selected_pid = _safe_pid(st.session_state.get("pvp_sheet_selected_pid") or "")

        
                # tenta achar sheet pelo id salvo
                selected_sheet = None
                for sh in sheets_to_show:
                    spid = _safe_pid((sh.get("pokemon") or {}).get("id") or "")
                    if spid == selected_pid:
                        selected_sheet = sh
                        break
                if selected_sheet is None:
                    selected_sheet = sheets_to_show[0]
        
                _shinies = {str(x).strip() for x in (user_data.get("shinies") or [])} if isinstance(user_data, dict) else set()
        
                left, right = st.columns([0.60, 0.40], gap="large")
        
                # =========================
                # LEFT: cards em HTML puro (moldura e cor funcionam)
                # =========================
                with left:
                    st.markdown("#### Cards")
        
                    cols_per_row = 2 if len(sheets_to_show) > 1 else 1
                    card_cols = st.columns(cols_per_row, gap="small") if cols_per_row > 1 else [st.container()]
        
                    for i_sh, sh in enumerate(sheets_to_show):
                        with card_cols[i_sh % cols_per_row]:
                            pkm = sh.get("pokemon") or {}
                            pid_sheet = _safe_pid(pkm.get("id") or "")
                            pname = str(pkm.get("name") or "").strip() or "Pokémon"
                            ptypes = pkm.get("types") or []
                            np_ = sh.get("np", pkm.get("np", "—"))
        
                            # pid raw (pode ser EXT) para imagem / visual
                            pid_raw = _safe_pid(sh.get("_party_pid_raw") or pid_sheet)
                            shiny = str(pid_sheet) in _shinies
        
                            sprite_url = _img_fallback(pid_raw, shiny=shiny, prefer_art=False)
        
                            moves = sh.get("moves") or []
                            if isinstance(moves, dict):
                                moves = list(moves.values())
                            if not isinstance(moves, list):
                                moves = []
        
                            stats = sh.get("stats") or {}
        
                            fav_names = _fav_moves_for_pid(pid_sheet)
                            mv_by_name = { _norm_name(m.get("name") or m.get("Nome") or m.get("nome") or ""): m for m in moves if isinstance(m, dict) }
        
                            fav_moves = []
                            for fn in (fav_names or []):
                                key = _norm_name(fn)
                                if key in mv_by_name:
                                    fav_moves.append(mv_by_name[key])
        
                            if not fav_moves:
                                fav_moves = [m for m in moves[:3] if isinstance(m, dict)]
        
                            is_selected = (pid_sheet == selected_pid)
                            card_class = "pvp-card pvp-card-selected" if is_selected else "pvp-card"
        
                            pills_types = "".join([f'<span class="pvp-pill">{html.escape(str(t))}</span>' for t in (ptypes or [])]) if ptypes else ""
        
                            # movimentos (HTML)
                            mv_html = ""
                            for mv in fav_moves[:5]:
                                mv_name = str(mv.get("name") or mv.get("Nome") or mv.get("nome") or "Golpe").strip()
                                rk_total, acc_base, based_label, stat_val, is_area = _mv_summary(mv, stats)
                                base_rank = int(mv.get("rank") or mv.get("Rank") or 0)
                                if based_label in ("Stgr", "Int") and int(stat_val) != 0:
                                    rk_break = f"(R{base_rank}+{int(stat_val)} {based_label})"
                                else:
                                    rk_break = f"(R{base_rank})"
                                area_txt = "Área" if is_area else "Alvo"
        
                                mv_html += dedent(f"""\
                                <div class="pvp-move-row">
                                  <div class="pvp-move-name">{html.escape(mv_name)}</div>
                                  <div class="pvp-move-badges">
                                    <span class="pvp-pill pvp-pill-acc">A+{int(acc_base)}</span>
                                    <span class="pvp-pill pvp-pill-rk">R{int(rk_total)}</span>
                                    <span class="pvp-pill pvp-pill-area">{html.escape(area_txt)}</span>
                                  </div>
                                </div>
                                <div style="opacity:.78;font-weight:900;font-size:.73rem;margin-top:3px;">
                                  {html.escape(rk_break)}
                                </div>
                                """)
        
                            if not mv_html:
                                mv_html = "<div style='opacity:.75;font-weight:900;'>Sem golpes nesta ficha.</div>"
        
                            # link seleciona por query param (mantém molde e cor!)
                            # usa o PID DA FICHA (numérico) como "sheet"
                            base_url = "?"
                            open_href = f"{base_url}sheet={html.escape(pid_sheet)}"
        
                            fav_label = "Favoritos" if fav_names else "Golpes (sem favoritos)"

                            _tpl = dedent("""
                            <div class="{card_class}" style="{bg_style}">
                              <div class="pvp-card-head">
                                <img src="{sprite_url}"
                                     style="width:72px;height:72px;object-fit:contain;border-radius:12px;border:1px solid rgba(255,255,255,12);
                                            background:rgba(0,0,0,12);padding:6px;">
                                <div style="flex:1;min-width:0;">
                                  <div class="pvp-card-title">{pname}</div>
                                  <div class="pvp-card-sub">#{pid_sheet} &bull; NP {np_}</div>
                                  <div class="pvp-pill-row">{pills_types}</div>
                                </div>
                              </div>
                            
                              <div class="pvp-divider"></div>
                              <div style="font-weight:950;">{fav_label}</div>
                            
                              {mv_html}
                            
                              <a class="pvp-open" href="{open_href}" target="_self">Abrir ficha &rarr;</a>
                            </div>
                            """).strip()
                            
                            st.markdown(
                                _tpl.format(
                                    card_class=card_class,
                                    bg_style=_type_bg_style(ptypes),
                                    sprite_url=html.escape(sprite_url),
                                    pname=html.escape(pname),
                                    pid_sheet=html.escape(pid_sheet),
                                    np_=html.escape(str(np_)),
                                    pills_types=pills_types,   # já é HTML montado
                                    fav_label=html.escape(fav_label),
                                    mv_html=mv_html,           # já é HTML montado
                                    open_href=open_href
                                ),
                                unsafe_allow_html=True
                            )

        
                # =========================
                # RIGHT: painel ficha completa (fallback imagem ok)
                # =========================
                with right:
                    st.markdown("#### Ficha completa")
        
                    sh = selected_sheet or {}
                    pkm = sh.get("pokemon") or {}
                    pid = _safe_pid(pkm.get("id") or "")
                    pname = str(pkm.get("name") or "").strip() or "Pokémon"
                    ptypes = pkm.get("types") or []
        
                    try:
                        np_ = int(sh.get("np", pkm.get("np", 0)) or 0)
                    except Exception:
                        np_ = 0
        
                    stats = sh.get("stats") or {}
                    stgr = int(stats.get("stgr", 0) or 0)
                    intelect = int(stats.get("int", 0) or 0)
                    dodge = int(stats.get("dodge", 0) or 0)
                    parry = int(stats.get("parry", 0) or 0)
                    fort = int(stats.get("fortitude", 0) or 0)
                    will = int(stats.get("will", 0) or 0)
        
                    thg = int(stats.get("thg", 0) or 0)
                    if thg <= 0:
                        thg = _auto_thg_from_pokedex(pid, pname, np_)
        
                    cap = 2 * int(np_ or 0)
                    if int(stats.get("dodge", 0) or 0) <= 0 and cap > 0 and thg > 0:
                        dodge = max(0, cap - thg)
        
                    shiny = str(pid) in _shinies
        
                    # tenta artwork do "pid normal"; se falhar, sprite; se falhar, pokebola
                    art_url = _img_fallback(pid, shiny=shiny, prefer_art=True)
        
                    skills = sh.get("skills") or []
                    if isinstance(skills, dict):
                        skills = list(skills.values())
                    if not isinstance(skills, list):
                        skills = []
        
                    advantages = sh.get("advantages") or []
                    if isinstance(advantages, str):
                        advantages = [advantages]
                    if not isinstance(advantages, list):
                        advantages = []
        
                    moves = sh.get("moves") or []
                    if isinstance(moves, dict):
                        moves = list(moves.values())
                    if not isinstance(moves, list):
                        moves = []
        
                    st.markdown(f'<div class="pvp-panel" style="{_type_bg_style(ptypes)}">', unsafe_allow_html=True)
        
                    cA, cB = st.columns([1, 1.1])
                    with cA:
                        st.image(art_url, use_container_width=True)
        
                    with cB:
                        st.markdown(f'<div class="pvp-panel-title">{html.escape(pname)}</div>', unsafe_allow_html=True)
                        st.markdown(
                            f'<div class="pvp-panel-sub">#{html.escape(pid) if pid else "—"} • <b>NP {np_}</b></div>',
                            unsafe_allow_html=True
                        )
        
                        if ptypes:
                            st.markdown(
                                '<div class="pvp-chips">' +
                                ''.join([f'<span class="pvp-chip">{html.escape(str(t))}</span>' for t in ptypes]) +
                                '</div>',
                                unsafe_allow_html=True
                            )
        
                        st.markdown('<div class="pvp-divider"></div>', unsafe_allow_html=True)
        
                        st.markdown(
                            '<div class="pvp-stat-grid">'
                            f'  <div class="pvp-stat"><div class="pvp-stat-k">Stgr</div><div class="pvp-stat-v">{stgr}</div></div>'
                            f'  <div class="pvp-stat"><div class="pvp-stat-k">Int</div><div class="pvp-stat-v">{intelect}</div></div>'
                            f'  <div class="pvp-stat"><div class="pvp-stat-k">Thg</div><div class="pvp-stat-v">{thg}</div></div>'
                            f'  <div class="pvp-stat"><div class="pvp-stat-k">Dodge</div><div class="pvp-stat-v">{dodge}</div></div>'
                            f'  <div class="pvp-stat"><div class="pvp-stat-k">Parry</div><div class="pvp-stat-v">{parry}</div></div>'
                            f'  <div class="pvp-stat"><div class="pvp-stat-k">Fort</div><div class="pvp-stat-v">{fort}</div></div>'
                            f'  <div class="pvp-stat"><div class="pvp-stat-k">Will</div><div class="pvp-stat-v">{will}</div></div>'
                            '</div>',
                            unsafe_allow_html=True
                        )
        
                    st.markdown('<div class="pvp-divider"></div>', unsafe_allow_html=True)
        
                    st.markdown("**Skills**")
                    if skills:
                        chips = []
                        for sk in skills:
                            if not isinstance(sk, dict):
                                continue
                            n = str(sk.get("name") or sk.get("nome") or "").strip()
                            r = int(sk.get("ranks") or sk.get("rank") or 0)
                            if n and r > 0:
                                chips.append(f"{n} R{r}")
                        if chips:
                            st.markdown('<div class="pvp-chips">' + ''.join([f'<span class="pvp-chip">{html.escape(x)}</span>' for x in chips]) + '</div>', unsafe_allow_html=True)
                        else:
                            st.caption("Sem skills registradas.")
                    else:
                        st.caption("Sem skills registradas.")
        
                    st.markdown("**Advantages**")
                    adv = [str(a).strip() for a in advantages if str(a).strip()]
                    if adv:
                        st.markdown('<div class="pvp-chips">' + ''.join([f'<span class="pvp-chip">{html.escape(a)}</span>' for a in adv]) + '</div>', unsafe_allow_html=True)
                    else:
                        st.caption("Sem advantages registradas.")
        
                    st.markdown('<div class="pvp-divider"></div>', unsafe_allow_html=True)
        
                    st.markdown("**Golpes**")
                    if not moves:
                        st.caption("Sem golpes nesta ficha.")
                    else:
                        for j, mv in enumerate(moves):
                            if not isinstance(mv, dict):
                                continue
                            mv_name = str(mv.get("name") or mv.get("Nome") or mv.get("nome") or "Golpe").strip()
                            rk_total, acc_base, based_label, stat_val, is_area = _mv_summary(mv, stats)
                            base_rank = int(mv.get("rank") or mv.get("Rank") or 0)
        
                            if based_label in ("Stgr", "Int") and int(stat_val) != 0:
                                rk_break = f"R{base_rank}+{int(stat_val)} {based_label}"
                            else:
                                rk_break = f"R{base_rank}"
        
                            meta = mv.get("meta") or {}
                            ranged = meta.get("ranged") is True
                            tags = []
                            tags.append("Área" if is_area else "Alvo")
                            if ranged:
                                tags.append("Ranged")
        
                            header = f"{mv_name} • A+{int(acc_base)} • R{int(rk_total)} ({rk_break})"
                            with st.expander(header, expanded=False):
                                if tags:
                                    st.markdown('<div class="pvp-chips">' + ''.join([f'<span class="pvp-chip">{html.escape(t)}</span>' for t in tags]) + '</div>', unsafe_allow_html=True)
        
                                desc = mv.get("description") or mv.get("desc")
                                build = mv.get("build")
                                if isinstance(desc, str) and desc.strip():
                                    st.write(desc.strip())
                                elif isinstance(build, str) and build.strip():
                                    st.code(build.strip(), language="text")
                                else:
                                    st.caption("Descrição não disponível.")
        
                                st.text_input(
                                    "Notas (local)",
                                    key=f"pvp_sheet_notes_{rid}_{trainer_name}_{pid}_{j}",
                                    placeholder="Anotações…",
                                )
        
                    st.markdown("</div>", unsafe_allow_html=True)  # panel end


        
        
        with tab_log:
            st.markdown("### 📜 Log")
            st.caption("Eventos públicos em tempo real (movimentos, efeitos, rolagens, etc.).")
            render_public_log_fragment(db, rid, title=None, height=320, show_divider=False, limit=60)

            st.markdown("### ⚔️ Logs do combate")
            logs = list(reversed(b_data.get("logs") or []))[:60]
            if not logs:
                st.caption("Sem logs de combate ainda.")
            else:
                for line in logs:
                    st.markdown(f"- {line}")

        st.stop()
    elif view == "lobby":
            stop_pvp_sync_listener()
            # --- MAPA DE NOMES (Para exibição amigável) ---
            THEME_NAMES = {
                "cave_water": "Caverna (com água)",
                "forest": "Floresta (padrão)",
                "mountain_slopes": "Montanha (padrão)",
                "plains": "Pradaria",
                "dirt": "Terra Batida",
                "river": "Rio",
                "sea_coast": "Costa Marítima",
                "center_lake": "Lago Central",

                # --- BIOMAS (novos) ---
                "biome_grass": "Campos / Rotas gramadas",
                "biome_forest": "Floresta (densidade)",
                "biome_meadow": "Meadow / Campo florido",
                "biome_desert": "Deserto / Árido",
                "biome_mountain": "Montanha / Rochoso (bioma)",
                "biome_snow": "Neve / Gelo",
                "biome_water": "Água (rio/lago/mar) (bioma)",
                "biome_cave": "Caverna / Dungeon (bioma)",
                "biome_mix": "Mix (rotas variadas) (bioma)",

                # --- BIOMAS nativos do BiomeGenerator ---
                "grasslands": "Grasslands / Campos",
                "deepforest": "Deep Forest / Floresta densa",
                "desert": "Desert / Árido",
                "beach": "Beach / Costa",
                "snowlands": "Snowlands / Neve",
                "cave": "Cave / Caverna",
                "mines": "Mines / Rochoso",
                "temple": "Temple / Ruínas",
                "seafloor": "Seafloor / Fundo do mar",
                "interior": "Interior / Arena",
                "lake": "Lake / Lago",
                "river": "River / Rio",
            }
    
            # ==========================================
            # 1. PAINEL SUPERIOR (CRIAR / LISTAR / ENTRAR)
            # ==========================================
            
            # --- Criar nova arena ---
            st.subheader("➕ Criar nova arena")
            c1, c2, c3 = st.columns([1, 1, 2])
            with c1:
                grid = st.selectbox("Tamanho do grid", [6, 8, 10, 12], index=0)
            with c2:
                    # Usa os mesmos biomas disponíveis no BiomeGenerator.
                    BIOME_LABELS = {
                        "grasslands": "Grasslands / Campos",
                        "deepforest": "Deep Forest / Floresta densa",
                        "desert": "Desert / Árido",
                        "beach": "Beach / Costa",
                        "snowlands": "Snowlands / Neve",
                        "cave": "Cave / Caverna",
                        "mines": "Mines / Rochoso",
                        "temple": "Temple / Ruínas",
                        "seafloor": "Seafloor / Fundo do mar",
                        "interior": "Interior / Arena",
                        "lake": "Lake / Lago",
                        "river": "River / Rio",
                    }
                    selector_options = [
                        f"{BIOME_LABELS.get(b, b)} ({b})" for b in BIOME_CONFIG.keys()
                    ]
                    selected_option = st.selectbox("Tema / Bioma", selector_options, index=0)
                    theme = selected_option.rsplit("(", 1)[-1].rstrip(")")
            with c3:
                st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
                if st.button("🆕 Criar arena", type="primary"):
                    rid, err = create_room(db, trainer_name, grid, theme, max_active=5)
                    if err:
                        st.error(err)
                    else:
                        st.success(f"Arena criada! Código: **{rid}**")
                        st.session_state["active_room_id"] = rid
                        sync_me_into_room(db, rid, trainer_name, role="owner")
                        st.rerun()
        
            st.markdown("---")
        
            # --- Minhas arenas ---
            st.subheader("📌 Minhas arenas")
            my_rooms = list_my_rooms(db, trainer_name)
            
            if not my_rooms:
                st.info("Você ainda não tem arenas ativas. Crie uma acima.")
            else:
                map_choice = {} 
                for rid in my_rooms[:20]:
                    info = get_room(db, rid)
                    if info:
                        gs = info.get("gridSize", "?")
                        th_key = info.get("theme", "cave_water")
                        th_nice = THEME_NAMES.get(th_key, th_key)
                        chal = (info.get("challenger") or {})
                        chal_name = chal.get("name") if isinstance(chal, dict) else (chal or "Ninguém")
                        
                        label = f"{th_nice} vs {chal_name} ({gs}x{gs}) [ID: {rid}]"
                        map_choice[label] = rid
        
                if map_choice:
                    chosen_label = st.selectbox("Selecionar Arena", list(map_choice.keys()))
                    chosen_rid = map_choice[chosen_label]
        
                    b1, b2 = st.columns([1, 4])
                    with b1:
                        if st.button("📂 Abrir Selecionada"):
                            st.session_state["active_room_id"] = chosen_rid
                            st.rerun()
                    with b2:
                        if st.button("🗄️ Arquivar"):
                            remove_room_from_user(db, trainer_name, chosen_rid)
                            if st.session_state.get("active_room_id") == chosen_rid:
                                st.session_state.pop("active_room_id", None)
                            st.rerun()
        
            st.markdown("---")
        
            # --- Entrar por código ---
            st.subheader("🔑 Entrar por código")
            cc1, cc2, cc3 = st.columns([2, 1, 1])
            with cc1:
                code = st.text_input("Código da arena (roomId)", value="")
            with cc2:
                st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
                if st.button("🥊 Desafiante"):
                    if code.strip():
                        res = join_room_as_challenger(db, code.strip(), trainer_name)
                        if res == "OK":
                            st.session_state["active_room_id"] = code.strip()
                            sync_me_into_room(db, code.strip(), trainer_name, role="challenger")
                            st.rerun()
                        elif res == "ALREADY_OWNER":
                            st.warning("Você é o dono desta sala.")
                            st.session_state["active_room_id"] = code.strip()
                            sync_me_into_room(db, code.strip(), trainer_name, role="owner")
                            st.rerun()
                        else:
                            st.error(res)
                            pass
            with cc3:
                # Adiciona um espaço para alinhar com o input que tem label
                st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
                if st.button("👀 Espectador"):
                    if code.strip():
                        res = join_room_as_spectator(db, code.strip(), trainer_name)
                        if res in ["OK", "PLAYER"]:
                            st.session_state["active_room_id"] = code.strip()
                            sync_me_into_room(db, code.strip(), trainer_name, role="spectator")
                            st.rerun()
                        else:
                            st.error(res)
                            pass
    
            st.markdown("---")
    
            # ==========================================
            # 2. PAINEL INFERIOR (PRÉ-VISUALIZAÇÃO DA ARENA ATIVA)
            # ==========================================
            rid = st.session_state.get("active_room_id")
            st.subheader("🎮 Arena Ativa (Pré-visualização)")
            
            # Botões de Navegação Básica
            c_nav1, c_nav2 = st.columns([1, 5])
            with c_nav1:
                if st.button("🔄 Atualizar"): st.rerun()
            with c_nav2:
                 if st.button("❌ Fechar Prévia"): 
                    st.session_state["active_room_id"] = None
                    st.rerun()
    
            if not rid:
                st.info("Nenhuma arena selecionada. Abra uma na lista acima.")
            else:
                room = get_room(db, rid)
                if not room:
                    st.error("Arena não encontrada.")
                else:
                    owner = (room.get("owner") or {}).get("name")
                    chal = room.get("challenger") or {}
                    chal_name = chal.get("name") if isinstance(chal, dict) else (chal or "Aguardando...")
                    
                    st.info(f"📍 **Arena {rid}** | {room.get('theme')} | {owner} vs {chal_name}")
    
                    state_ref = db.collection("rooms").document(rid).collection("public_state").document("state")
                    state_doc = state_ref.get()
                    state = state_doc.to_dict() if state_doc.exists else {}
    
                    grid = int(room.get("gridSize") or 6)
                    theme_key = room.get("theme") or "cave_water"
                    seed = state.get("seed")
                    no_water_state = bool(state.get("noWater", False))
                    all_pieces = state.get("pieces") or []
                    pieces = visible_pieces_for(room, trainer_name, all_pieces)
                    
                    role = get_role(room, trainer_name)
                    is_player = role in ["owner", "challenger"]
                    
                    no_water = st.checkbox("🚫 Gerar sem água", value=no_water_state, disabled=not is_player)

                    # ✅ Se a sala já tem seed mas ainda não tem mapUrl no Firestore, publique o mapa (1 vez por sessão).
                    try:
                        if seed is not None and not state.get("mapUrl") and is_player:
                            key_flag = f"pvp_map_published_{rid}_{seed}_{int(grid)}_{theme_key}_{int(bool(no_water_state))}"
                            if not st.session_state.get(key_flag):
                                _, bucket = init_firebase()
                                ensure_room_map_published(
                                    db, bucket,
                                    rid=str(rid), grid=int(grid), theme_key=str(theme_key),
                                    seed=int(seed), no_water=bool(no_water_state),
                                    show_grid=True
                                )
                                st.session_state[key_flag] = True
                    except Exception:
                        pass

                    if seed is not None and not state.get("mapUrl") and is_player:
                        if st.button("📤 Publicar mapa no site", help="Gera e envia o mapa base ao Firebase Storage e grava mapUrl no public_state/state."):
                            try:
                                _, bucket = init_firebase()
                                ensure_room_map_published(
                                    db, bucket,
                                    rid=str(rid), grid=int(grid), theme_key=str(theme_key),
                                    seed=int(seed), no_water=bool(no_water_state),
                                    show_grid=True
                                )
                                st.success("Mapa publicado! Recarregue o battle-site.")
                            except Exception as e:
                                st.error(f"Falha ao publicar mapa: {e}")
                            st.rerun()

                    
                    if seed is None:
                        if st.button("🗺️ Gerar mapa (pixel art)", disabled=not is_player):
                            seed = generate_biome_seed()
                            state_ref.set({
                                "gridSize": grid, "theme": theme_key, "seed": seed,
                                "tilesPacked": None, "noWater": bool(no_water),
                                "updatedAt": firestore.SERVER_TIMESTAMP,
                            }, merge=True)

                            # ✅ publica mapa base no Storage + JSON estruturado (mapUrl/mapDataUrl)
                            try:
                                _, bucket = init_firebase()
                                ensure_room_map_published(
                                    db, bucket,
                                    rid=str(rid), grid=int(grid), theme_key=str(theme_key),
                                    seed=int(seed), no_water=bool(no_water),
                                    show_grid=True
                                )
                            except Exception:
                                pass

                            # ✅ redireciona para o battle-site (o outro site consome o mapa)
                            redirect_to_battle_site(str(rid), trainer_name)

                    else:
                        show_grid = st.checkbox("Mostrar grade tática", value=False, key=f"show_grid_preview_{rid}")
                        img = render_map_with_pieces(grid, theme_key, seed, no_water_state, pieces, trainer_name, room, show_grid=show_grid)
                        st.image(img, caption="Prévia do Campo")
                        
                        if st.button("⚔️ IR PARA O CAMPO DE BATALHA", type="primary"):
                            redirect_to_battle_site(str(rid), trainer_name)
                            
                            st.rerun()
                                                    
                        if st.button("🔁 Regerar Mapa", disabled=not is_player):
                             seed = generate_biome_seed()
                             state_ref.set({"seed": seed, "tilesPacked": None, "noWater": bool(no_water)}, merge=True)
                             st.rerun()
    
            
    
    

    

elif page == "Compendium de Ga'Al":
    render_compendium_page()

elif page == "Gerador de Encontros":
    render_encounter_generator_page()

elif page == "Mochila":
    if "backpack" not in user_data:
        user_data["backpack"] = {
            "money": 0, 
            "medicine": [], 
            "pokeballs": [], 
            "tms": [], 
            "key_items": []
        }

    # Cabeçalho com Dinheiro (AGORA IDENTADO CORRETAMENTE)
    st.markdown(f'<div class="money-display">💰 Dinheiro: ₽ {user_data["backpack"]["money"]}</div>', unsafe_allow_html=True)    

    col_bag, col_content = st.columns([1, 2.5])
    
    with col_bag:
        # Tenta carregar a imagem local; se não existir, usa o link reserva
        try:
            st.image("mochila.png", width=150, caption="MINHA MOCHILA")
        except:
            st.image("https://raw.githubusercontent.com/PokeAPI/sprites/master/sprites/items/main-stats/adventure-guide.png", width=150)
        
        # Ajuste de Saldo
        new_money = st.number_input("Ajustar saldo", value=int(user_data["backpack"]["money"]), step=100)
        if new_money != user_data["backpack"]["money"]:
            user_data["backpack"]["money"] = new_money
            save_data_cloud(trainer_name, user_data) 
        
        if st.button("🧹 Limpar itens vazios"):
            for k in ["medicine", "pokeballs", "tms", "key_items"]:
                user_data["backpack"][k] = [i for i in user_data["backpack"][k] if i["name"] and i.get("qty", 0) > 0]
            save_data_cloud(trainer_name, user_data) 
            st.rerun()
    
    with col_content:
            tabs = st.tabs(["💊 Medicamentos", "🔴 Poké Bolas", "💿 TMs", "🔑 Itens-chave"])
            cfg = [(tabs[0], "medicine", True), (tabs[1], "pokeballs", True), 
                   (tabs[2], "tms", True), (tabs[3], "key_items", False)]
    
            for tab, key, use_img in cfg:
                with tab:
                    current_list = user_data["backpack"].get(key, [])
                    
                    # Interface de edição 
                    updated_items = []
                    # Exibimos a lista atual + 1 linha sempre vazia no final
                    display_count = len(current_list) + 1
                    
                    for i in range(display_count):
                        item_data = current_list[i] if i < len(current_list) else {"name": "", "qty": 0}
                        res = render_item_row(key, i, item_data, show_image=use_img)
                        
                        if res["name"]: # Se o usuário digitou algo 
                            updated_items.append(res)
    
                    # Se a lista atualizada for maior que a original, salvamos 
                    # automaticamente para gerar a nova linha vazia 
                    if len(updated_items) > len(current_list):
                        user_data["backpack"][key] = updated_items
                        save_data_cloud(trainer_name, user_data)
                        st.rerun()
    
                    if st.button(f"💾 Confirmar {key.title()}", key=f"sv_{key}"):
                        user_data["backpack"][key] = updated_items
                        save_data_cloud(trainer_name, user_data)
                        st.success("Mochila sincronizada com sucesso.")    
    
    
    
    
    
    
